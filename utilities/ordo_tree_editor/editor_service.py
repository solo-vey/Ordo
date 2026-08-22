from __future__ import annotations
import mimetypes
import urllib.parse
import fnmatch

import argparse
import copy
import contextvars
import base64
import hashlib
import io
import json
import sys
import os
import re
import shutil
import shlex
import subprocess
import tempfile
import threading
import time
import traceback
import uuid
import urllib.error
import urllib.request
from urllib.parse import urlparse, parse_qs
import webbrowser
import zipfile
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from datetime import datetime, timezone

import yaml

UTILITY_ROOT = Path(__file__).resolve().parent

INTEGRATED_COMPILER_ROOT = UTILITY_ROOT / "integrated_compiler"
INTEGRATED_COMPILER_VERSION = "ordo-runtime-semantic-compiler/0.7.15.5-r3dev-profile-adapter"

def _run_integrated_compile(package_root: Path, program_path: Path) -> tuple[dict[str, Any], dict[str, Any]]:
    """Compile source YAML into a Runtime Semantic Plan using the bundled compiler.

    The compiler remains an internal module/CLI boundary so the same code can be used
    by Editor and headless CI. No playbook-specific semantics are implemented here.
    """
    compiler = INTEGRATED_COMPILER_ROOT / "compile_runtime_semantic_plan_v7.py"
    validator = INTEGRATED_COMPILER_ROOT / "validate_runtime_semantic_plan_v7.py"
    if not compiler.is_file() or not validator.is_file():
        raise ValueError("Integrated compiler runtime is missing from the Editor package.")
    with tempfile.TemporaryDirectory(prefix="ordo-integrated-compile-") as td:
        out = Path(td) / "runtime_semantic_plan.json"
        env = dict(os.environ)
        env["PYTHONPATH"] = str(INTEGRATED_COMPILER_ROOT) + (os.pathsep + env["PYTHONPATH"] if env.get("PYTHONPATH") else "")
        cp = subprocess.run([sys.executable, str(compiler), str(program_path), "-o", str(out)], cwd=str(package_root), env=env, capture_output=True, text=True, timeout=120)
        if cp.returncode != 0 or not out.is_file():
            detail = (cp.stderr or cp.stdout or "compiler failed").strip()[-4000:]
            raise ValueError(f"Integrated playbook compilation failed: {detail}")
        try:
            plan = json.loads(out.read_text(encoding="utf-8"))
        except Exception as error:
            raise ValueError(f"Integrated compiler produced invalid JSON: {error}") from error
        vp = subprocess.run([sys.executable, str(validator), str(out)], cwd=str(package_root), env=env, capture_output=True, text=True, timeout=120)
        try:
            validation = json.loads(vp.stdout or "{}")
        except Exception:
            validation = {"status":"FAIL", "detail": (vp.stderr or vp.stdout or "validator failed").strip()[-4000:]}
        if vp.returncode != 0 or validation.get("status") != "PASS":
            raise ValueError("Integrated runtime-plan validation failed: " + json.dumps(validation, ensure_ascii=False))
        compile_summary = {}
        try:
            compile_summary = json.loads((cp.stdout or "{}").splitlines()[-1])
        except Exception:
            compile_summary = {"status":"PASS"}
        report = {
            "mode":"integrated", "compiler_version": plan.get("compiler_version") or INTEGRATED_COMPILER_VERSION,
            "compile": compile_summary, "validation": validation,
            "stages":[
                {"id":"load_source","status":"PASS"}, {"id":"compile_runtime_plan","status":"PASS"},
                {"id":"validate_runtime_plan","status":"PASS"}, {"id":"verify_compatibility","status":"PASS"},
            ],
        }
        return plan, report

_SOURCE_RESOURCE_PATH_RE = re.compile(r"^[A-Za-z0-9_.-]+(?:/[A-Za-z0-9_.-]+)+\.(?:yaml|yml|json|md|py)$", re.IGNORECASE)

def _source_resource_references(value: Any) -> list[str]:
    out: set[str] = set()
    def walk(v: Any):
        if isinstance(v, dict):
            for x in v.values(): walk(x)
        elif isinstance(v, list):
            for x in v: walk(x)
        elif isinstance(v, str):
            text=v.strip()
            if "://" not in text and _SOURCE_RESOURCE_PATH_RE.fullmatch(text): out.add(text)
    walk(value)
    return sorted(out)

def _wrap_yaml_as_source_package(filename: str, raw: bytes) -> tuple[str, bytes]:
    if len(raw) > 12 * 1024 * 1024:
        raise ValueError("Playbook YAML is larger than the 12 MB local editor limit.")
    try:
        doc = yaml.safe_load(raw.decode("utf-8-sig"))
    except (UnicodeDecodeError, yaml.YAMLError) as error:
        raise ValueError(f"Playbook YAML could not be parsed: {error}") from error
    if not _looks_like_playbook(doc):
        raise ValueError("The selected YAML is not an Ordo playbook source.")
    external_refs=_source_resource_references(doc)
    if external_refs:
        preview=", ".join(external_refs[:5])
        more=f" (+{len(external_refs)-5} more)" if len(external_refs)>5 else ""
        raise ValueError(f"This playbook YAML references external package resources: {preview}{more}. Upload a source ZIP containing the YAML and its referenced resources.")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("program.ordo.yaml", raw)
    return filename, buf.getvalue()
if str(UTILITY_ROOT) not in sys.path:
    sys.path.insert(0, str(UTILITY_ROOT))
try:
    from .alpha20_runtime import apply_state_patch_atomic, canonicalize_runtime_state, legacy_updates_to_state_patch, normalize_gate_failure, validate_state_patch
except ImportError:
    from alpha20_runtime import apply_state_patch_atomic, canonicalize_runtime_state, legacy_updates_to_state_patch, normalize_gate_failure, validate_state_patch
try:
    from .ordo_yaml_semantics import declared_writes as _shared_declared_writes, declared_routes as _shared_declared_routes, routes as _shared_generic_routes
except ImportError:
    from ordo_yaml_semantics import declared_writes as _shared_declared_writes, declared_routes as _shared_declared_routes, routes as _shared_generic_routes


OPENAI_MODELS = ("gpt-5.6-sol", "gpt-5.6-terra", "gpt-5.6-luna")
PROVIDERS = ("openai", "mlx", "custom")
DEFAULT_MLX_BASE_URL = "http://127.0.0.1:8080/v1"
DEFAULT_CUSTOM_BASE_URL = "http://ml03.ligazakon.net:8555/v1"
LIVE_SESSIONS: dict[str, dict[str, Any]] = {}
PROVIDER_CAPABILITY_CACHE: dict[str, dict[str, Any]] = {}
PROVIDER_CAPABILITY_CACHE_PATH = Path.home() / ".ordo_tree_editor" / "provider_capabilities.json"

def _provider_capability_cache_key(provider: str, base_url: str, model: str, api_style: str) -> str:
    return "|".join([
        str(provider or "").strip().lower(),
        str(base_url or "").strip().rstrip("/"),
        str(model or "").strip(),
        str(api_style or "").strip().lower(),
    ])

def _load_provider_capability_cache() -> None:
    try:
        if not PROVIDER_CAPABILITY_CACHE_PATH.is_file():
            return
        raw=json.loads(PROVIDER_CAPABILITY_CACHE_PATH.read_text(encoding="utf-8"))
        if isinstance(raw,dict):
            PROVIDER_CAPABILITY_CACHE.clear()
            PROVIDER_CAPABILITY_CACHE.update({str(k):copy.deepcopy(v) for k,v in raw.items() if isinstance(v,dict)})
    except Exception:
        # Capability evidence is an optimization/safety input, never a startup blocker.
        return

def _persist_provider_capability_cache() -> None:
    try:
        PROVIDER_CAPABILITY_CACHE_PATH.parent.mkdir(parents=True,exist_ok=True)
        tmp=PROVIDER_CAPABILITY_CACHE_PATH.with_suffix(".tmp")
        tmp.write_text(json.dumps(PROVIDER_CAPABILITY_CACHE,ensure_ascii=False,indent=2),encoding="utf-8")
        tmp.replace(PROVIDER_CAPABILITY_CACHE_PATH)
    except Exception:
        return

def _remember_provider_capability(profile: dict[str, Any]) -> None:
    if not isinstance(profile,dict):
        return
    provider=str(profile.get("provider") or "").strip().lower()
    base_url=str(profile.get("base_url") or "").strip().rstrip("/")
    model=str(profile.get("model") or "").strip()
    api_style=str(profile.get("api_style") or "").strip().lower()
    if not provider or not base_url or not model or not api_style:
        return
    PROVIDER_CAPABILITY_CACHE[_provider_capability_cache_key(provider,base_url,model,api_style)]=copy.deepcopy(profile)
    _persist_provider_capability_cache()

def _cached_provider_capability(provider: str, base_url: str, model: str, api_style: str) -> dict[str,Any] | None:
    value=PROVIDER_CAPABILITY_CACHE.get(_provider_capability_cache_key(provider,base_url,model,api_style))
    return copy.deepcopy(value) if isinstance(value,dict) else None

_load_provider_capability_cache()

RUN_ARTIFACT_REGISTRY: dict[tuple[str,str,str], dict[str, dict[str, Any]]] = {}
REPLAY_PACKAGES: dict[str, dict[str, Any]] = {}
EXECUTE_RUNS: dict[str, dict[str, Any]] = {}
EXECUTE_RUNS_LOCK = threading.RLock()

LIVE_RUNTIME: dict[str, Any] = {
    "enabled": False,
    "provider": "openai",
    "model": None,
    "api_key": None,
    "base_url": "https://api.openai.com/v1",
}
EDITOR_STARTUP: dict[str, Any] = {"gitlab_root": ""}

def _resolve_startup_runtime_config(*, provider: str | None = None, model: str | None = None, base_url: str | None = None, api_key: str | None = None) -> dict[str, Any]:
    """Resolve application startup defaults without binding them to one provider/model."""
    resolved_provider=str(provider or os.environ.get("ORDO_MODEL_PROVIDER") or "openai").strip().lower()
    if resolved_provider not in PROVIDERS:
        raise ValueError(f"Unsupported startup model provider: {resolved_provider}")
    resolved_model=str(model or os.environ.get("ORDO_MODEL_NAME") or os.environ.get("OPENAI_MODEL") or ("gpt-5.6-terra" if resolved_provider=="openai" else "")).strip()
    resolved_key=api_key or os.environ.get("ORDO_MODEL_API_KEY") or (os.environ.get("OPENAI_API_KEY") if resolved_provider=="openai" else None)
    if resolved_provider=="openai":
        resolved_base="https://api.openai.com/v1"
        if resolved_model and resolved_model not in OPENAI_MODELS:
            raise ValueError(f"Startup OpenAI model must be one of: {', '.join(OPENAI_MODELS)}")
        enabled=bool(resolved_model and resolved_key)
    elif resolved_provider=="mlx":
        resolved_base=_normalize_base_url(base_url or os.environ.get("ORDO_MODEL_BASE_URL"), DEFAULT_MLX_BASE_URL)
        enabled=bool(resolved_model)
    else:
        resolved_base=_normalize_base_url(base_url or os.environ.get("ORDO_MODEL_BASE_URL"), DEFAULT_CUSTOM_BASE_URL)
        enabled=bool(resolved_model)
    return {"enabled":enabled,"provider":resolved_provider,"api_key":resolved_key,"model":resolved_model or None,"base_url":resolved_base}

def _parse_gitlab_tree_url(root_url: str) -> dict[str, str]:
    raw=str(root_url or "").strip().rstrip("/")
    if not raw:
        raise ValueError("GitLab root URL is not configured.")
    parsed=urlparse(raw)
    if parsed.scheme not in {"http","https"} or not parsed.netloc:
        raise ValueError("GitLab root must be an absolute http(s) URL.")
    path=urllib.parse.unquote(parsed.path).strip("/")
    marker="/-/tree/"
    if marker not in "/"+path:
        raise ValueError("GitLab root must be a repository tree URL containing /-/tree/<ref>/<path>.")
    project_part, tree_part=("/"+path).split(marker,1)
    project=project_part.strip("/")
    parts=[x for x in tree_part.split("/") if x]
    if not project or not parts:
        raise ValueError("GitLab tree URL must include project and ref.")
    return {"origin":f"{parsed.scheme}://{parsed.netloc}","project":project,"ref":parts[0],"path":"/".join(parts[1:])}

def _gitlab_api_url(spec: dict[str,str], suffix: str, params: dict[str,Any] | None=None) -> str:
    project=urllib.parse.quote(spec["project"],safe="")
    url=f"{spec['origin']}/api/v4/projects/{project}/{suffix.lstrip('/')}"
    if params:
        url += "?" + urllib.parse.urlencode(params)
    return url

def _gitlab_http_json(url: str) -> Any:
    req=urllib.request.Request(url,headers={"Accept":"application/json","User-Agent":"Ordo-Tree-Editor"})
    try:
        with urllib.request.urlopen(req,timeout=15) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        raise ValueError(f"GitLab request failed with HTTP {error.code}.") from error
    except (urllib.error.URLError,TimeoutError,json.JSONDecodeError) as error:
        raise ValueError(f"GitLab request failed: {error}") from error

def _gitlab_repository_tree(spec: dict[str,str], path: str) -> list[dict[str,Any]]:
    rows=[]
    page=1
    while True:
        data=_gitlab_http_json(_gitlab_api_url(spec,"repository/tree",{"path":path,"ref":spec["ref"],"per_page":100,"page":page}))
        if not isinstance(data,list):
            raise ValueError("GitLab repository tree response is not a list.")
        batch=[x for x in data if isinstance(x,dict)]
        rows.extend(batch)
        if len(batch)<100:
            break
        page += 1
        if page>1000:
            raise ValueError("GitLab repository tree pagination exceeded the safety limit.")
    return rows

def _gitlab_validate_directory_path(spec: dict[str,str], path: str) -> str:
    requested=str(path or "").strip().strip("/")
    root=(spec.get("path") or "").strip("/")
    if root and not (requested==root or requested.startswith(root+"/")):
        raise ValueError("Requested GitLab directory is outside the configured root.")
    return requested

def _gitlab_directory_listing(spec: dict[str,str], path: str) -> dict[str,Any]:
    """Read exactly one repository directory; child directories remain lazy stubs."""
    path=_gitlab_validate_directory_path(spec,path)
    entries=sorted(_gitlab_repository_tree(spec,path),key=lambda x:(str(x.get("type") or ""),str(x.get("name") or "").lower()))
    readme=None
    archives=[]
    children=[]
    for item in entries:
        item_type=str(item.get("type") or "")
        name=str(item.get("name") or "")
        item_path=str(item.get("path") or "")
        if item_type=="blob":
            lower=name.lower()
            if lower=="readme.md":
                readme={"filename":name,"path":item_path}
            elif lower.endswith(".zip"):
                archives.append({"filename":name,"path":item_path})
            continue
        if item_type=="tree":
            children.append({"name":name,"path":item_path,"readme":None,"archives":[],"children":[],"loaded":False})
    return {
        "name": path.rstrip("/").split("/")[-1] if path else spec.get("project","root").split("/")[-1],
        "path": path,
        "readme": readme,
        "archives": archives,
        "children": children,
        "loaded": True,
    }

def _gitlab_playbook_catalog(root_url: str) -> dict[str,Any]:
    spec=_parse_gitlab_tree_url(root_url)
    root=_gitlab_directory_listing(spec,spec["path"])
    # Lazy contract: only the configured root is read here. Descendants are
    # fetched independently as their disclosure controls are opened in the UI.
    return {
        "status":"passed",
        "root_url":root_url,
        "spec":spec,
        "root":root,
        "directories":root.get("children") or [],
        "directory_count":len(root.get("children") or []),
        "archive_count":len(root.get("archives") or []),
        "lazy":True,
    }

def _gitlab_directory_payload(root_url: str, directory_path: str) -> dict[str,Any]:
    spec=_parse_gitlab_tree_url(root_url)
    node=_gitlab_directory_listing(spec,directory_path)
    return {"status":"passed","root_url":root_url,"directory":node,"lazy":True}

def _gitlab_download_archive(root_url: str, archive_path: str) -> bytes:
    spec=_parse_gitlab_tree_url(root_url)
    path=str(archive_path or "").strip().lstrip("/")
    prefix=(spec.get("path") or "").strip("/")
    if not path.lower().endswith(".zip") or (prefix and not (path==prefix or path.startswith(prefix+"/"))):
        raise ValueError("Requested GitLab archive is outside the configured root or is not a ZIP file.")
    encoded_path=urllib.parse.quote(path,safe="")
    url=_gitlab_api_url(spec,f"repository/files/{encoded_path}/raw",{"ref":spec["ref"]})
    req=urllib.request.Request(url,headers={"Accept":"application/zip","User-Agent":"Ordo-Tree-Editor"})
    try:
        with urllib.request.urlopen(req,timeout=30) as response:
            raw=response.read(60*1024*1024+1)
    except urllib.error.HTTPError as error:
        raise ValueError(f"GitLab archive download failed with HTTP {error.code}.") from error
    except (urllib.error.URLError,TimeoutError) as error:
        raise ValueError(f"GitLab archive download failed: {error}") from error
    if len(raw)>60*1024*1024:
        raise ValueError("GitLab playbook archive exceeds the 60 MB Editor limit.")
    return raw


def _gitlab_read_text_file(root_url: str, file_path: str) -> str:
    spec=_parse_gitlab_tree_url(root_url)
    path=str(file_path or "").strip().lstrip("/")
    prefix=(spec.get("path") or "").strip("/")
    if not path.lower().endswith(".md") or (prefix and not (path==prefix or path.startswith(prefix+"/"))):
        raise ValueError("Requested GitLab Markdown file is outside the configured root or is not Markdown.")
    encoded_path=urllib.parse.quote(path,safe="")
    url=_gitlab_api_url(spec,f"repository/files/{encoded_path}/raw",{"ref":spec["ref"]})
    req=urllib.request.Request(url,headers={"Accept":"text/markdown,text/plain","User-Agent":"Ordo-Tree-Editor"})
    try:
        with urllib.request.urlopen(req,timeout=20) as response:
            raw=response.read(2*1024*1024+1)
    except urllib.error.HTTPError as error:
        raise ValueError(f"GitLab README download failed with HTTP {error.code}.") from error
    except (urllib.error.URLError,TimeoutError) as error:
        raise ValueError(f"GitLab README download failed: {error}") from error
    if len(raw)>2*1024*1024:
        raise ValueError("GitLab README exceeds the 2 MB Editor limit.")
    return raw.decode("utf-8",errors="replace")
PLAYBOOK_PACKAGE: dict[str, Any] = {
    "id": None,
    "filename": None,
    "source_name": None,
    "source": None,
    "resources": {},
    "manifest": [],
    "compiled_plan": None,
    "compiled_plan_status": {"available": False, "valid": False, "reason": "not_loaded"},
    "semantic_plan": None,
    "semantic_plan_status": {"available": False, "valid": False, "reason": "not_loaded"},
}

# Loaded playbook packages are retained by immutable content id so concurrent browser
# sessions/tabs cannot invalidate an in-progress run by loading another package.
PLAYBOOK_PACKAGES: dict[str, dict[str, Any]] = {}
_ACTIVE_PLAYBOOK_PACKAGE: contextvars.ContextVar[dict[str, Any] | None] = contextvars.ContextVar("ordo_active_playbook_package", default=None)
_ACTIVE_RUN_CONTEXT: contextvars.ContextVar[dict[str, str] | None] = contextvars.ContextVar("ordo_active_run_context", default=None)

def _active_playbook_package() -> dict[str, Any]:
    return _ACTIVE_PLAYBOOK_PACKAGE.get() or PLAYBOOK_PACKAGE


REPOSITORY_ROOT = UTILITY_ROOT.parents[1]
CLI_ROOT = REPOSITORY_ROOT / "cli"
if not CLI_ROOT.is_dir():
    CLI_ROOT = REPOSITORY_ROOT / "cli_embedded" / "ordo_pkg"
if str(CLI_ROOT) not in sys.path:
    sys.path.insert(0, str(CLI_ROOT))



NODE_TEMPLATES = {
    "question": {
        "id": "N_NEW_QUESTION",
        "question": "Describe the decision this node must collect.",
        "answer_type": "free_text",
        "allow_unmatched_input": True,
        "on_answer": {"continue": {"next": "N_NEXT"}},
        "allowed_from": [],
    },
    "gate": {
        "id": "N_NEW_GATE",
        "question": "Run or review the required gate.",
        "answer_type": "enum",
        "allow_unmatched_input": True,
        "allowed_answers": ["pass", "revise"],
        "on_answer": {
            "pass": {"next": "N_NEXT"},
            "revise": {"next": "N_REPAIR"},
        },
        "allowed_from": [],
    },
    "materialization": {
        "id": "N_NEW_MATERIALIZATION",
        "question": "Materialize the reviewed output artifact.",
        "answer_type": "confirmation",
        "allow_unmatched_input": True,
        "on_answer": {"confirmed": {"next": "N_NEXT"}},
        "allowed_from": [],
    },
    "terminal": {
        "id": "N_NEW_TERMINAL",
        "question": "Confirm the terminal outcome.",
        "answer_type": "confirmation",
        "allow_unmatched_input": True,
        "terminal": True,
        "allowed_from": [],
    },
}


def parse_yaml(text: str) -> dict[str, Any]:
    data = yaml.safe_load(text)
    if not isinstance(data, dict):
        raise ValueError("The YAML root must be a mapping.")
    return data


def dump_yaml(source: dict[str, Any]) -> str:
    return yaml.safe_dump(source, allow_unicode=True, sort_keys=False)


def dump_value_yaml(value: Any) -> str:
    rendered = yaml.safe_dump(value, allow_unicode=True, sort_keys=False)
    if rendered.endswith("\n...\n"):
        rendered = rendered[:-5]
    return rendered.strip()


def node_sections(node: dict[str, Any]) -> list[dict[str, str]]:
    return [{"key": str(key), "value_yaml": dump_value_yaml(value)} for key, value in node.items()]


def _records(source: dict[str, Any], collection: str) -> list[dict[str, Any]]:
    records = source.get(collection)
    if not isinstance(records, list):
        raise ValueError(f"The loaded source does not contain a {collection} list.")
    return [record for record in records if isinstance(record, dict)]


def _replace_target_references(value: Any, old_id: str, new_id: str) -> Any:
    if isinstance(value, dict):
        if set(value) and all(isinstance(child, str) for child in value.values()):
            return {key: new_id if child == old_id else child for key, child in value.items()}
        return {
            key: new_id if key in {"next", "to", "on_pass", "on_fail", "pass_to", "fail_to"} and child == old_id else _replace_target_references(child, old_id, new_id)
            for key, child in value.items()
        }
    if isinstance(value, list):
        return [_replace_target_references(child, old_id, new_id) for child in value]
    return value


def replace_record(source: dict[str, Any], collection: str, old_id: str, replacement: dict[str, Any]) -> dict[str, Any]:
    """Replace a node or executable gate while preserving the rest of the source."""
    if collection not in {"nodes", "gates"}:
        raise ValueError("Only nodes and gates can be edited as records.")
    new_id = replacement.get("id")
    if not isinstance(new_id, str) or not new_id.strip():
        raise ValueError("The replacement record must declare a non-empty id.")
    records = source.get(collection)
    if not isinstance(records, list):
        raise ValueError(f"The loaded source does not contain a {collection} list.")
    matches = [index for index, record in enumerate(records) if isinstance(record, dict) and record.get("id") == old_id]
    if len(matches) != 1:
        raise ValueError(f"Cannot identify exactly one {collection[:-1]} with id {old_id!r}.")
    all_records = [
        record
        for records_key in ("nodes", "gates")
        for record in source.get(records_key, [])
        if isinstance(record, dict)
    ]
    if new_id != old_id and any(record.get("id") == new_id for record in all_records):
        raise ValueError(f"A node or gate with id {new_id!r} already exists.")
    records[matches[0]] = replacement
    if new_id != old_id:
        for record in all_records:
            for contract_key in ("allowed_from", "allowed_to"):
                if isinstance(record.get(contract_key), list):
                    record[contract_key] = [new_id if value == old_id else value for value in record[contract_key]]
            navigation = record.get("navigation_contract")
            if isinstance(navigation, dict):
                for contract_key in ("allowed_from", "allowed_to"):
                    if isinstance(navigation.get(contract_key), list):
                        navigation[contract_key] = [new_id if value == old_id else value for value in navigation[contract_key]]
            for key, value in list(record.items()):
                record[key] = _replace_target_references(value, old_id, new_id)
        for container_key in ("graph_contract", "playbook"):
            container = source.get(container_key)
            if isinstance(container, dict) and container.get("entry_node") == old_id:
                container["entry_node"] = new_id
    return source


def replace_node(source: dict[str, Any], old_id: str, replacement: dict[str, Any]) -> dict[str, Any]:
    return replace_record(source, "nodes", old_id, replacement)


def replace_record_sections(source: dict[str, Any], collection: str, old_id: str, sections: dict[str, str]) -> dict[str, Any]:
    replacement: dict[str, Any] = {}
    for key, value_yaml in sections.items():
        if not isinstance(key, str) or not key.strip() or not isinstance(value_yaml, str):
            raise ValueError("Each record section must have a non-empty key and YAML value.")
        replacement[key] = yaml.safe_load(value_yaml)
    return replace_record(source, collection, old_id, replacement)


def replace_node_sections(source: dict[str, Any], old_id: str, sections: dict[str, str]) -> dict[str, Any]:
    return replace_record_sections(source, "nodes", old_id, sections)


def _targets(value: Any) -> list[str]:
    targets: list[str] = []
    if isinstance(value, dict):
        for key, item in value.items():
            if key == "next" and isinstance(item, str):
                targets.append(item)
            else:
                targets.extend(_targets(item))
    elif isinstance(value, list):
        for item in value:
            targets.extend(_targets(item))
    return targets


def _node_targets(node: dict[str, Any]) -> list[str]:
    """Return graph targets from the canonical and ARF prototype forms."""
    targets = _targets(node.get("on_answer", {}))
    transitions = node.get("transitions", {})
    if isinstance(transitions, dict):
        targets.extend(
            target
            for target in transitions.values()
            if isinstance(target, str) and not target.startswith("$")
        )
    elif isinstance(transitions, list):
        targets.extend(
            item["to"] for item in transitions
            if isinstance(item, dict) and isinstance(item.get("to"), str) and not item["to"].startswith("$")
        )
    navigation = node.get("navigation_contract", {})
    if isinstance(navigation, dict) and isinstance(navigation.get("allowed_to"), list):
        targets.extend(target for target in navigation["allowed_to"] if isinstance(target, str) and not target.startswith("$"))
    return targets


def _node_edges(node: dict[str, Any]) -> list[dict[str, str]]:
    edges: list[dict[str, str]] = []
    explicit_targets: set[str] = set()
    if isinstance(node.get("next"), str) and not node["next"].startswith("$"):
        edges.append({"target": node["next"], "storage": "next", "key": "next"})
    transitions = node.get("transitions", {})
    if isinstance(transitions, dict):
        edges.extend(
            {"target": target, "storage": "transitions", "key": str(key)}
            for key, target in transitions.items()
            if isinstance(target, str) and not target.startswith("$")
        )
        explicit_targets.update(edge["target"] for edge in edges)
    elif isinstance(transitions, list):
        explicit_targets = set()
        for index, transition in enumerate(transitions):
            if not isinstance(transition, dict) or not isinstance(transition.get("to"), str):
                continue
            target = transition["to"]
            if target.startswith("$"):
                continue
            explicit_targets.add(target)
            label = transition.get("id") or transition.get("when") or transition.get("outcome") or f"transition_{index + 1}"
            edges.append({"target": target, "storage": "transitions_list", "key": str(label), "index": str(index)})
    on_answer = node.get("on_answer", {})
    if isinstance(on_answer, dict):
        if isinstance(on_answer.get("next"), str) and not on_answer["next"].startswith("$"):
            edges.append({"target": on_answer["next"], "storage": "on_answer_next", "key": "next"})
        for outcome, route in on_answer.items():
            if outcome == "next":
                continue
            targets = _targets(route)
            edges.extend({"target": target, "storage": "on_answer", "key": str(outcome)} for target in targets)
    # ``navigation_contract.allowed_to`` is canonical routing authority when a
    # record intentionally omits an executable transition list.  Do not render
    # it in addition to explicit transitions: that would duplicate the same
    # control edge and misleadingly imply two runtime routes.
    if not edges:
        navigation = node.get("navigation_contract")
        if isinstance(navigation, dict):
            for target in navigation.get("allowed_to", []):
                if isinstance(target, str) and target and not target.startswith("$"):
                    edges.append({"target": target, "storage": "navigation_allowed_to", "key": target})
    for edge in edges:
        edge.setdefault("edge_type", "control_flow")
    return edges


def _route_targets(value: Any) -> list[str]:
    if isinstance(value, str):
        return [value]
    return _targets(value)


# Release-1 graph-projection safety: these are failure dispositions/actions, not
# graph entities unless the playbook explicitly declares an entity with exactly
# that ID.  The renderer must never synthesize them as terminal nodes.
_RESERVED_GRAPH_DISPOSITIONS = {
    "block", "blocked", "fail", "failure", "retry", "stop", "halt",
    "abort", "reject", "rejected", "skip", "continue", "pass",
}


def _is_reserved_graph_disposition(value: Any) -> bool:
    return isinstance(value, str) and value.strip().lower() in _RESERVED_GRAPH_DISPOSITIONS


def _gate_edges(gate: dict[str, Any], known_entity_ids: set[str] | None = None) -> list[dict[str, str]]:
    known_entity_ids = known_entity_ids or set()
    result: list[dict[str, str]] = []
    for key in ("on_pass", "on_fail", "pass_to", "fail_to"):
        for target in _route_targets(gate.get(key)):
            if _is_reserved_graph_disposition(target) and target not in known_entity_ids:
                continue
            result.append({"target": target, "storage": "gate_route", "key": key, "edge_type": "control_flow"})
    return result


def _terminal_records(source: dict[str, Any]) -> list[dict[str, Any]]:
    records = source.get("terminals", [])
    if not isinstance(records, list):
        return []
    result = []
    for item in records:
        if isinstance(item, str) and item.strip():
            result.append({"id": item})
        elif isinstance(item, dict) and isinstance(item.get("id"), str) and item["id"].strip():
            result.append(item)
    return result


def _external_terminal_ids(source: dict[str, Any], edges: list[dict[str, str]], known_ids: set[str]) -> list[str]:
    declared: list[str] = []
    contract = source.get("graph_contract", {})
    if isinstance(contract, dict):
        declared.extend(value for value in contract.get("external_terminal_targets", []) if isinstance(value, str))
    routed = [
        edge["target"] for edge in edges
        if edge.get("edge_type") == "control_flow"
        and edge["target"] not in known_ids
        and not edge["target"].startswith("$")
        and not _is_reserved_graph_disposition(edge["target"])
    ]
    return list(dict.fromkeys(
        identifier for identifier in [*declared, *routed]
        if identifier not in known_ids and not _is_reserved_graph_disposition(identifier)
    ))



def _first_display_text(record: dict[str, Any], fields: tuple[str, ...]) -> str:
    """Return the first meaningful scalar display value, skipping ID duplicates."""
    record_id = str(record.get("id") or "").strip()
    for field in fields:
        value = record.get(field)
        if value is None or isinstance(value, (dict, list, tuple, set)):
            continue
        text = str(value).strip()
        if not text or text == record_id:
            continue
        return text
    return record_id or "Unnamed element"


def _node_display_label(node: dict[str, Any]) -> str:
    return _first_display_text(node, (
        "title", "question", "purpose", "description", "summary",
        "prompt", "instruction", "propose", "proposal",
        "display_name", "name", "label", "action", "type", "kind",
    ))


def _gate_display_label(gate: dict[str, Any]) -> str:
    return _first_display_text(gate, (
        "title", "purpose", "condition", "description", "summary",
        "question", "prompt", "instruction", "propose", "proposal",
        "display_name", "name", "label", "specification", "validator",
        "method", "type", "kind",
    ))


def _terminal_display_label(terminal: dict[str, Any]) -> str:
    return _first_display_text(terminal, (
        "title", "purpose", "description", "summary", "label", "name",
    ))


def _projection_declared_outputs(nodes: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Derive visual output entities only from explicit artifact/output declarations."""
    by_path: dict[str, dict[str, Any]] = {}
    for node in nodes:
        nid=str(node.get("id") or "")
        candidates=[]
        if isinstance(node.get("output"),str): candidates.append((node["output"],"output"))
        rem=node.get("rematerialization")
        if isinstance(rem,dict) and isinstance(rem.get("output"),str): candidates.append((rem["output"],"rematerialization.output"))
        pkg=node.get("package")
        if isinstance(pkg,dict) and isinstance(pkg.get("path"),str): candidates.append((pkg["path"],"package.path"))
        art=node.get("artifact")
        if isinstance(art,dict) and isinstance(art.get("expected_path"),str): candidates.append((art["expected_path"],"artifact.expected_path"))
        for path,source_field in candidates:
            path=str(path).strip()
            if not path: continue
            rec=by_path.setdefault(path,{"id":"OUT::"+path,"path":path,"producers":[],"declarations":[]})
            if nid and nid not in rec["producers"]: rec["producers"].append(nid)
            rec["declarations"].append({"node_id":nid,"source_field":source_field})
    return list(by_path.values())




def _normalize_output_reference(value: Any) -> str:
    if not isinstance(value, str):
        return ""
    text = value.strip().replace("\\", "/")
    while text.startswith("./"):
        text = text[2:]
    return text


def _append_projection_reference(refs: list[dict[str, Any]], value: Any, source_field: str, priority: int) -> None:
    if isinstance(value, str):
        normalized = _normalize_output_reference(value)
        if normalized:
            refs.append({"value": normalized, "source_field": source_field, "priority": priority})
        return
    if isinstance(value, list):
        for index, item in enumerate(value):
            _append_projection_reference(refs, item, f"{source_field}[{index}]", priority)


def _producer_output_references(node: dict[str, Any]) -> list[dict[str, Any]]:
    """Formal output/materialization references used only by the editor projection.

    No prose, ID similarity, or action-name inference is allowed here.  These
    references are exact values already carried by structured Ordo fields.
    """
    refs: list[dict[str, Any]] = []
    _append_projection_reference(refs, node.get("output"), "output", 0)
    remat = node.get("rematerialization")
    if isinstance(remat, dict):
        _append_projection_reference(refs, remat.get("output"), "rematerialization.output", 0)
    artifact = node.get("artifact")
    if isinstance(artifact, dict):
        _append_projection_reference(refs, artifact.get("expected_path"), "artifact.expected_path", 1)
        _append_projection_reference(refs, artifact.get("state_path"), "artifact.state_path", 2)
    artifacts = node.get("artifacts")
    if isinstance(artifacts, list):
        for index, item in enumerate(artifacts):
            if not isinstance(item, dict):
                continue
            _append_projection_reference(refs, item.get("expected_path"), f"artifacts[{index}].expected_path", 1)
            _append_projection_reference(refs, item.get("state_path"), f"artifacts[{index}].state_path", 2)
    package = node.get("package")
    if isinstance(package, dict):
        _append_projection_reference(refs, package.get("path"), "package.path", 1)
    return refs


def _declared_output_records(source: dict[str, Any]) -> list[dict[str, Any]]:
    raw = source.get("outputs", [])
    if not isinstance(raw, list):
        return []
    records: list[dict[str, Any]] = []
    for item in raw:
        if isinstance(item, str) and item.strip():
            records.append({"id": item.strip()})
        elif isinstance(item, dict) and isinstance(item.get("id"), str) and item["id"].strip():
            records.append(item)
    return records


def _declared_output_references(record: dict[str, Any]) -> list[dict[str, Any]]:
    refs: list[dict[str, Any]] = []
    # Exact structured identity/reference values only.  ID participates only as
    # an exact formal value; there is deliberately no token/name similarity.
    _append_projection_reference(refs, record.get("output"), "output", 0)
    _append_projection_reference(refs, record.get("path"), "path", 0)
    _append_projection_reference(refs, record.get("expected_path"), "expected_path", 1)
    _append_projection_reference(refs, record.get("state_path"), "state_path", 2)
    artifact = record.get("artifact")
    if isinstance(artifact, dict):
        _append_projection_reference(refs, artifact.get("expected_path"), "artifact.expected_path", 1)
        _append_projection_reference(refs, artifact.get("state_path"), "artifact.state_path", 2)
    package = record.get("package")
    if isinstance(package, dict):
        _append_projection_reference(refs, package.get("path"), "package.path", 1)
    _append_projection_reference(refs, record.get("id"), "id", 3)
    return refs


def _declared_output_display_path(record: dict[str, Any]) -> str:
    refs = [r for r in _declared_output_references(record) if r.get("source_field") != "id"]
    refs.sort(key=lambda item: (int(item.get("priority", 99)), str(item.get("source_field") or "")))
    return str(refs[0]["value"]) if refs else ""


def _projection_resolve_declared_outputs(
    source: dict[str, Any],
    nodes: list[dict[str, Any]],
    execution_terminal_ids: set[str] | None = None,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], set[str]]:
    """Resolve declared outputs to producers using exact structured references.

    Returns (view entities, traceability diagnostics, claimed materialization
    references).  This function never creates control-flow edges and never
    mutates source/runtime semantics.
    """
    execution_terminal_ids = execution_terminal_ids or set()
    node_refs: dict[str, list[dict[str, Any]]] = {
        str(node.get("id")): _producer_output_references(node)
        for node in nodes if isinstance(node.get("id"), str)
    }
    entities: list[dict[str, Any]] = []
    diagnostics: list[dict[str, Any]] = []
    claimed_values: set[str] = set()

    for record in _declared_output_records(source):
        output_id = str(record.get("id") or "").strip()
        if not output_id or output_id in execution_terminal_ids:
            # Explicit execution-terminal authority wins over view-only output
            # projection.  This preserves existing routing semantics.
            continue
        output_refs = _declared_output_references(record)
        output_values = {str(item["value"]) for item in output_refs}
        matches: list[dict[str, Any]] = []
        for node_id, refs in node_refs.items():
            for ref in refs:
                if str(ref["value"]) not in output_values:
                    continue
                matches.append({
                    "node_id": node_id,
                    "value": str(ref["value"]),
                    "producer_field": str(ref["source_field"]),
                    "producer_priority": int(ref["priority"]),
                    "output_fields": [str(orow["source_field"]) for orow in output_refs if str(orow["value"]) == str(ref["value"])],
                })
        producer_best: dict[str, dict[str, Any]] = {}
        for match in matches:
            current = producer_best.get(match["node_id"])
            if current is None or match["producer_priority"] < current["producer_priority"]:
                producer_best[match["node_id"]] = match
        selected: list[dict[str, Any]] = []
        if producer_best:
            best_priority = min(item["producer_priority"] for item in producer_best.values())
            selected = [item for item in producer_best.values() if item["producer_priority"] == best_priority]

        producer_ids: list[str] = []
        status = "WARNING"
        reason = "unresolved_producer"
        matched_reference = None
        candidates = sorted(producer_best)
        if len(selected) == 1:
            producer_ids = [selected[0]["node_id"]]
            status = "PASS"
            reason = "resolved_formal_reference"
            matched_reference = selected[0]
            claimed_values.add(str(selected[0]["value"]))
        elif len(selected) > 1:
            status = "FAIL"
            reason = "ambiguous_producer"
            candidates = sorted(item["node_id"] for item in selected)

        path = _declared_output_display_path(record)
        # Older editor payloads used a bare ``outputs: [{id: ...}]`` entry as
        # an external terminal declaration.  Preserve that representation
        # without conflating modern output contracts that carry metadata.
        legacy_terminal = set(record) == {"id"}
        entities.append({
            "id": output_id,
            "path": path,
            "output_type": record.get("type") or record.get("kind") or "output",
            "producers": producer_ids,
            "traceability_status": status,
            "traceability_reason": reason,
            "record": record,
            "declared": True,
            "legacy_terminal": legacy_terminal,
        })
        diagnostics.append({
            "check": "DECLARED_OUTPUT_PRODUCER_TRACEABILITY",
            "output_id": output_id,
            "status": status,
            "reason": reason,
            "producer_nodes": producer_ids,
            "candidates": candidates,
            "matched_reference": matched_reference,
            "declared_references": output_refs,
        })

    return entities, diagnostics, claimed_values

def _projection_state_writes(node: dict[str, Any]) -> set[str]:
    writes={str(x) for x in (node.get("writes") or []) if isinstance(x,str)}
    def walk(obj: Any):
        if isinstance(obj,dict):
            update=obj.get("update_state")
            if isinstance(update,dict):
                for k in update:
                    if isinstance(k,str): writes.add(k)
            for v in obj.values(): walk(v)
        elif isinstance(obj,list):
            for v in obj: walk(v)
    walk(node)
    return writes


def _projection_state_inputs(record: dict[str, Any]) -> list[str]:
    vals=[]
    for key in ("inputs","reads","required_inputs"):
        v=record.get(key)
        if isinstance(v,list): vals.extend(str(x) for x in v if isinstance(x,str))
    return list(dict.fromkeys(vals))


def _projection_write_covers(write_path: str, read_path: str) -> bool:
    return write_path==read_path or read_path.startswith(write_path+'.') or write_path.startswith(read_path+'.')


def _projection_control_adjacency(nodes: list[dict[str, Any]], gates: list[dict[str, Any]], declared_ids: set[str]):
    edges=[{"source":str(n["id"]),**e} for n in nodes for e in _node_edges(n)]
    edges += [{"source":str(g["id"]),**e} for g in gates for e in _gate_edges(g,declared_ids)]
    out: dict[str,set[str]]={x:set() for x in declared_ids}
    inc: dict[str,set[str]]={x:set() for x in declared_ids}
    for e in edges:
        if e.get("edge_type")!="control_flow": continue
        a,b=e["source"],e["target"]
        if a in out and b in declared_ids: out[a].add(b); inc[b].add(a)
    return edges,out,inc


def _projection_dominators(entry: str|None, ids: set[str], out: dict[str,set[str]], inc: dict[str,set[str]]) -> dict[str,set[str]]:
    if not entry or entry not in ids: return {x:{x} for x in ids}
    reachable=set(); q=[entry]
    while q:
        x=q.pop(0)
        if x in reachable: continue
        reachable.add(x); q.extend(y for y in out.get(x,set()) if y not in reachable)
    dom={n:({n} if n==entry else set(reachable)) for n in reachable}
    changed=True
    while changed:
        changed=False
        for n in reachable:
            if n==entry: continue
            ps=[p for p in inc.get(n,set()) if p in reachable]
            common=set.intersection(*(dom[p] for p in ps)) if ps else set()
            new={n}|common
            if new!=dom[n]: dom[n]=new; changed=True
    return dom


def _projection_dependency_edges(source: dict[str, Any], nodes: list[dict[str, Any]], gates: list[dict[str, Any]], declared_ids: set[str], out: dict[str,set[str]], inc: dict[str,set[str]]) -> tuple[list[dict[str,Any]],list[dict[str,Any]]]:
    """Deterministic producer/state→gate projection; unresolved ambiguity is diagnostic."""
    entry=_entry_node_id(source); dom=_projection_dominators(entry,declared_ids,out,inc)
    producer_records=[*nodes,*gates]
    writes={str(n["id"]):_projection_state_writes(n) for n in producer_records}
    gc=source.get("graph_contract") if isinstance(source.get("graph_contract"),dict) else {}
    optional={(str(x.get("consumer")),str(x.get("path"))) for x in (gc.get("explicit_optional_dependencies") or []) if isinstance(x,dict) and x.get("consumer") and x.get("path")}
    deps=[]; unresolved=[]
    for gate in gates:
        gid=str(gate["id"])
        for path in _projection_state_inputs(gate):
            if (gid,path) in optional:
                continue
            candidates=[]
            for nid,wpaths in writes.items():
                if any(_projection_write_covers(w,path) for w in wpaths): candidates.append(nid)
            dominating=[x for x in candidates if x in dom.get(gid,set())]
            # Prefer the deepest dominating producer (dominated by the others).
            chosen=None
            if dominating:
                ranked=sorted(dominating,key=lambda x:len(dom.get(x,set())),reverse=True)
                chosen=ranked[0]
            elif len(candidates)==1:
                # Candidate may be recovery-only; don't invent a dependency as guaranteed.
                unresolved.append({"gate":gid,"path":path,"candidates":candidates,"reason":"producer_does_not_dominate_gate"})
            elif len(candidates)>1:
                unresolved.append({"gate":gid,"path":path,"candidates":sorted(candidates),"reason":"ambiguous_non_dominating_producers"})
            else:
                unresolved.append({"gate":gid,"path":path,"candidates":[],"reason":"no_declared_producer"})
            if chosen:
                deps.append({"source":chosen,"target":gid,"edge_type":"validation_dependency","relation_type":"validation_dependency","state_path":path,"inference":"dominating_state_producer"})
    # dedupe exact relation/path
    seen=set(); outdeps=[]
    for e in deps:
        key=(e["source"],e["target"],e.get("state_path"))
        if key not in seen: seen.add(key);outdeps.append(e)
    return outdeps,unresolved


_INSPECTOR_RESOURCE_PATH_RE = re.compile(
    r"(?<![A-Za-z0-9_.-])([A-Za-z0-9_.-]+(?:/[A-Za-z0-9_.-]+)+\.(?:py|md|markdown|json|ya?ml|txt|html?|css|js|mjs|cjs|ts|tsx|jsx|xml|csv|sql|sh|toml|ini|cfg))(?![A-Za-z0-9_.-])",
    re.IGNORECASE,
)


def _resolve_package_resource(resources: dict[str, Any], path_value: str) -> tuple[str | None, Any]:
    """Resolve a package-relative resource path without guessing across ambiguities."""
    text = resources.get(path_value)
    if text is not None:
        return path_value, text
    matches = [(str(k), v) for k, v in resources.items() if str(k).endswith(path_value) or path_value.endswith(str(k))]
    if len(matches) == 1:
        return matches[0]
    return None, None


def _generic_record_resource_references(record: dict[str, Any], resources: dict[str, Any]) -> list[dict[str, Any]]:
    """Discover package resources anywhere in a source record.

    Structured scalar values that look like package paths are retained even when
    unresolved, because an explicit contract reference is useful diagnostic data.
    Paths merely mentioned inside prose are included only when they resolve to an
    actual package resource, preventing examples/output paths from polluting the
    inspector.
    """
    discovered: list[dict[str, Any]] = []

    def add(role: str, path_value: str, *, explicit: bool) -> None:
        path_value = str(path_value or "").strip().strip('`"\'')
        if not path_value or "://" in path_value:
            return
        resolved_path, _ = _resolve_package_resource(resources, path_value)
        if explicit or resolved_path:
            discovered.append({"role": role, "path": path_value})

    def walk(value: Any, role: str) -> None:
        if isinstance(value, dict):
            for key, child in value.items():
                walk(child, f"{role}.{key}" if role else str(key))
            return
        if isinstance(value, list):
            for index, child in enumerate(value):
                walk(child, f"{role}[{index}]")
            return
        if not isinstance(value, str):
            return
        text = value.strip()
        # A scalar whose entire value is a package-like path is an explicit ref.
        full = _INSPECTOR_RESOURCE_PATH_RE.fullmatch(text)
        if full:
            add(role, full.group(1), explicit=True)
            return
        # In prose/commands only surface paths that actually exist in the package.
        for match in _INSPECTOR_RESOURCE_PATH_RE.finditer(value):
            add(role, match.group(1), explicit=False)

    walk(record, "")
    return discovered


def _template_inspector_payload(package: dict[str, Any], source: dict[str, Any], node_id: str) -> dict[str, Any]:
    source_records = [r for r in [*(source.get("nodes") or []), *(source.get("gates") or [])] if isinstance(r, dict)]
    record = next((r for r in source_records if str(r.get("id") or "") == node_id), None)
    output_path = node_id[len("OUT::"):] if node_id.startswith("OUT::") else ""
    producer_records = []
    declared_output_record = next((item for item in _declared_output_records(source) if str(item.get("id") or "") == node_id), None)
    declared_output_view = None
    if record is None and output_path:
        declared = next((item for item in _projection_declared_outputs([r for r in (source.get("nodes") or []) if isinstance(r, dict)]) if str(item.get("path") or "") == output_path), None)
        producer_ids = list(declared.get("producers") or []) if isinstance(declared, dict) else []
        producer_records = [r for r in source_records if str(r.get("id") or "") in producer_ids]
        if producer_records:
            record = producer_records[0]
    elif record is None and declared_output_record is not None:
        projected, _, _ = _projection_resolve_declared_outputs(source, [r for r in (source.get("nodes") or []) if isinstance(r, dict)], set())
        declared_output_view = next((item for item in projected if str(item.get("id") or "") == node_id), None)
        producer_ids = list(declared_output_view.get("producers") or []) if isinstance(declared_output_view, dict) else []
        producer_records = [r for r in source_records if str(r.get("id") or "") in producer_ids]
        output_path = str((declared_output_view or {}).get("path") or "")
        if producer_records:
            record = producer_records[0]
    if record is None and declared_output_record is not None:
        return {
            "status":"passed", "node_id":node_id, "action":None,
            "purpose":declared_output_record.get("purpose") or declared_output_record.get("title"),
            "output":output_path or None, "template":None, "bindings":None, "parameters":{}, "references":[],
            "record":{"declared_output":declared_output_record,"traceability":declared_output_view},
            "entity_type":"declared_output", "producers":[],
        }
    if record is None:
        raise ValueError("Selected element has no inspectable source/producer contract.")
    resources = package.get("resources") if isinstance(package.get("resources"), dict) else {}
    refs: list[dict[str, Any]] = []
    records_for_refs = producer_records or [record]
    for rec in records_for_refs:
        rec_id = str(rec.get("id") or "")
        role_prefix = f"{rec_id}." if producer_records else ""
        # Preserve the historical canonical roles first.
        for key in ("template","bindings","validator","specification","schema","resource","reference"):
            value = rec.get(key)
            if isinstance(value, str) and value.strip(): refs.append({"role":role_prefix+key,"path":value.strip(),"producer_node":rec_id})
        rec_remat = rec.get("rematerialization") if isinstance(rec.get("rematerialization"), dict) else {}
        for key in ("template","bindings","validator","specification","schema","resource","reference"):
            value = rec_remat.get(key)
            if isinstance(value, str) and value.strip(): refs.append({"role":role_prefix+f"rematerialization.{key}","path":value.strip(),"producer_node":rec_id})
        # Generic discovery covers allowed_tools, knowledge refs and future fields.
        for item in _generic_record_resource_references(rec, resources):
            refs.append({"role":role_prefix+item["role"],"path":item["path"],"producer_node":rec_id})

    # Dedupe identical path references while retaining all semantic origins.
    grouped: dict[str, dict[str, Any]] = {}
    for item in refs:
        path_value = str(item.get("path") or "").strip()
        if not path_value:
            continue
        bucket = grouped.setdefault(path_value, {"path":path_value, "roles":[], "producer_nodes":[]})
        role = str(item.get("role") or "reference")
        producer = str(item.get("producer_node") or "")
        if role not in bucket["roles"]:
            bucket["roles"].append(role)
        if producer and producer not in bucket["producer_nodes"]:
            bucket["producer_nodes"].append(producer)

    enriched=[]
    for path_value, item in grouped.items():
        resolved_path, text = _resolve_package_resource(resources, path_value)
        roles = item["roles"]
        enriched.append({
            "role": " · ".join(roles),
            "roles": roles,
            "path": path_value,
            "producer_node": item["producer_nodes"][0] if item["producer_nodes"] else "",
            "producer_nodes": item["producer_nodes"],
            "resolved_path": resolved_path,
            "available": text is not None,
            "text": text if isinstance(text,str) else None,
        })
    enriched.sort(key=lambda x: (not x["available"], x["path"].lower()))
    primary_template = next((x for x in enriched if any(r == "template" or r.endswith(".template") for r in x.get("roles",[])) and x["available"]), None)
    primary_bindings = next((x for x in enriched if any(r == "bindings" or r.endswith(".bindings") for r in x.get("roles",[])) and x["available"]), None)
    remat = record.get("rematerialization") if isinstance(record.get("rematerialization"), dict) else {}
    output = output_path or record.get("output") or remat.get("output")
    parameters={k:record.get(k) for k in ("inputs","derive_before_generate","rendering_contract","update_state","artifact","preconditions") if record.get(k) is not None}
    producer_summaries=[{"id":str(r.get("id") or ""),"action":r.get("action"),"purpose":r.get("purpose") or r.get("title"),"output":r.get("output") or ((r.get("rematerialization") or {}).get("output") if isinstance(r.get("rematerialization"),dict) else None)} for r in (producer_records or [record])]
    if declared_output_record is not None:
        raw_record={"declared_output":declared_output_record,"traceability":declared_output_view,"producer_records":producer_records}
        entity_type="declared_output"
    elif output_path:
        raw_record={"derived_output":{"id":node_id,"path":output_path,"producers":producer_summaries},"producer_records":producer_records}
        entity_type="output"
    else:
        raw_record=record
        entity_type="source"
    return {
        "status":"passed","node_id":node_id,"action":record.get("action"),"purpose":record.get("purpose") or record.get("title"),
        "output":output,"template":primary_template,"bindings":primary_bindings,"parameters":parameters,"references":enriched,"record":raw_record,
        "entity_type":entity_type,"producers":producer_summaries
    }

def graph_view(source: dict[str, Any], resources: dict[str, Any] | None = None) -> dict[str, Any]:
    nodes = [node for node in source.get("nodes", []) if isinstance(node, dict) and node.get("id")]
    gates = [gate for gate in source.get("gates", []) if isinstance(gate, dict) and gate.get("id")]
    terminal_records = _terminal_records(source)
    declared_ids = {str(record["id"]) for record in [*nodes, *gates, *terminal_records]}
    declared_ids.update(
        str(item["id"]) for item in source.get("outputs", [])
        if isinstance(item, dict) and isinstance(item.get("id"), str)
    )

    suppressed_reserved_routes: list[dict[str, str]] = []
    for gate in gates:
        for key in ("on_pass", "on_fail", "pass_to", "fail_to"):
            for target in _route_targets(gate.get(key)):
                if _is_reserved_graph_disposition(target) and target not in declared_ids:
                    suppressed_reserved_routes.append({"source": str(gate["id"]), "key": key, "value": target})

    navigation_permissions = [
        {"source": str(node["id"]), "target": str(target)}
        for node in nodes
        for target in ((node.get("navigation_contract") or {}).get("allowed_to") or [])
        if isinstance(node.get("navigation_contract"), dict) and isinstance(target, str)
    ]

    edges, control_out, control_in = _projection_control_adjacency(nodes,gates,declared_ids)
    validation_dependencies, unresolved_dependencies = _projection_dependency_edges(source,nodes,gates,declared_ids,control_out,control_in)
    edges.extend(validation_dependencies)
    known_ids = {str(record["id"]) for record in [*nodes, *gates, *terminal_records]}
    # Determine execution terminals from explicit graph/routing authority before
    # projecting declared outputs.  source.outputs alone is not an execution edge.
    terminals = _external_terminal_ids(source, edges, known_ids)
    declared_output_entities, output_traceability, claimed_output_values = _projection_resolve_declared_outputs(
        source, nodes, set(terminals) | {str(item["id"]) for item in terminal_records}
    )
    derived_outputs = [
        item for item in _projection_declared_outputs(nodes)
        if _normalize_output_reference(item.get("path")) not in claimed_output_values
    ]
    declared_outputs = [*declared_output_entities, *derived_outputs]
    for out_entity in declared_outputs:
        relation = "declares_output" if out_entity.get("declared") else "produces_output"
        for producer in out_entity.get("producers",[]):
            edges.append({
                "source":producer, "target":out_entity["id"],
                "edge_type":relation, "relation_type":relation,
                "artifact_path":out_entity.get("path") or "",
                "label":"declares output" if relation == "declares_output" else "produces",
            })
        # Historical derived-output enablement remains view-only. Declared outputs
        # intentionally link only to their formal producer to avoid branch-like noise.
        if not out_entity.get("declared"):
            for producer in out_entity.get("producers",[]):
                for pred in control_in.get(producer,set()):
                    if any(str(g.get("id"))==pred for g in gates):
                        edges.append({"source":pred,"target":out_entity["id"],"edge_type":"enables_output","relation_type":"enables_output","artifact_path":out_entity.get("path") or ""})
    for edge in edges:
        edge.setdefault("relation_type",edge.get("edge_type","control_flow"))

    # Reachability is computed from execution edges only.  It is diagnostic in
    # Release 1; we do not mutate playbook semantics or delete source entities.
    graph_ids = known_ids | set(terminals)
    outgoing: dict[str, list[str]] = {identifier: [] for identifier in graph_ids}
    for edge in edges:
        if edge.get("edge_type") == "control_flow" and edge["source"] in outgoing and edge["target"] in graph_ids:
            outgoing[edge["source"]].append(edge["target"])
    entry = _entry_node_id(source)
    reachable: set[str] = set()
    queue = [entry] if entry in graph_ids else []
    while queue:
        current = queue.pop(0)
        if current in reachable:
            continue
        reachable.add(current)
        queue.extend(target for target in outgoing.get(current, []) if target not in reachable)
    unreachable_terminals = sorted(identifier for identifier in _declared_terminal_ids(source) if identifier in graph_ids and identifier not in reachable)
    package_resources = resources if isinstance(resources, dict) else {}
    def record_reference_paths(record: dict[str, Any]) -> list[str]:
        if not package_resources:
            return []
        refs=[]
        for item in _generic_record_resource_references(record, package_resources):
            path_value=str(item.get("path") or "").strip()
            resolved_path,_=_resolve_package_resource(package_resources,path_value)
            if resolved_path and resolved_path not in refs:
                refs.append(resolved_path)
        return refs

    return {
        "nodes": [
            {
                "id": str(node["id"]),
                "element_type": "node",
                "entity_type": "execution_node",
                "collection": "nodes",
                "label": _node_display_label(node),
                "answer_type": node.get("answer_type") or node.get("kind", "unspecified"),
                "terminal": node.get("terminal") is True,
                "allowed_from": node.get("allowed_from", []),
                "record_yaml": dump_yaml(node),
                "sections": node_sections(node),
                "resource_references": record_reference_paths(node),
            }
            for node in nodes
        ] + [
            {
                "id": str(terminal["id"]),
                "element_type": "terminal",
                "entity_type": "terminal",
                "collection": None,
                "label": _terminal_display_label(terminal),
                "answer_type": "terminal",
                "terminal": True,
                "allowed_from": terminal.get("allowed_from", []),
                "record_yaml": dump_yaml(terminal),
                "sections": [],
            }
            for terminal in terminal_records
        ] + [
            {
                "id": str(gate["id"]),
                "element_type": "gate",
                "entity_type": "gate",
                "collection": "gates",
                "label": _gate_display_label(gate),
                "answer_type": str(gate.get("method") or "gate"),
                "terminal": False,
                "allowed_from": [],
                "record_yaml": dump_yaml(gate),
                "sections": node_sections(gate),
                "resource_references": record_reference_paths(gate),
            }
            for gate in gates
        ] + [
            {
                "id": terminal_id,
                "element_type": "terminal",
                "entity_type": "terminal",
                "collection": None,
                "label": terminal_id,
                "answer_type": "external terminal",
                "terminal": True,
                "allowed_from": [],
                "record_yaml": "",
                "sections": [],
            }
            for terminal_id in terminals
        ] + [
            {
                "id": out_entity["id"],
                # A bare legacy top-level output ID historically represented a
                # terminal destination.  Rich declared outputs remain their
                # own view-only entity, with producer traceability.
                "element_type": "terminal" if out_entity.get("legacy_terminal") else "output",
                "entity_type": "terminal" if out_entity.get("legacy_terminal") else ("declared_output" if out_entity.get("declared") else "output"),
                "collection": "declared_outputs" if out_entity.get("declared") else "derived_outputs",
                "label": (out_entity.get("path") or out_entity["id"]),
                "path": out_entity.get("path") or "",
                "answer_type": (f"declared output · {str(out_entity.get('traceability_status') or '').lower()}" if out_entity.get("declared") else "materialized output"),
                "terminal": bool(out_entity.get("legacy_terminal")),
                "producers": out_entity.get("producers",[]),
                "traceability_status": out_entity.get("traceability_status"),
                "traceability_reason": out_entity.get("traceability_reason"),
                "output_type": out_entity.get("output_type"),
                "record_yaml": dump_yaml(out_entity.get("record")) if out_entity.get("declared") and isinstance(out_entity.get("record"), dict) else "",
                "sections": node_sections(out_entity.get("record")) if out_entity.get("declared") and isinstance(out_entity.get("record"), dict) else [],
            }
            for out_entity in declared_outputs
        ],
        # Alpha.20 clients infer a missing relation as control-flow. Preserve
        # the established pre-Alpha.20 contract for older compact payloads;
        # canonical graph-contract payloads retain typed control edges.
        "edges": [
            {key: value for key, value in edge.items() if key not in {"edge_type", "relation_type"}}
            if edge.get("edge_type") == "control_flow" and (
                not isinstance(source.get("graph_contract"), dict)
                or (source.get("ordo") or {}).get("version") == "0.12"
            ) else edge
            for edge in edges
        ],
        "projection_validation": {
            "status": "PASS" if not unresolved_dependencies and not [row for row in output_traceability if row.get("status") == "FAIL"] and not unreachable_terminals else "FAIL",
            "checks": {
                "synthetic_execution_nodes": 0,
                "unresolved_validation_dependencies": len(unresolved_dependencies),
                "orphan_outputs": len([row for row in output_traceability if row.get("status") == "WARNING"]),
                "ambiguous_outputs": len([row for row in output_traceability if row.get("status") == "FAIL"]),
                "unreachable_terminals": len(unreachable_terminals),
                "non_control_relations_are_typed": all(e.get("relation_type") in {"control_flow","validation_dependency","state_dependency","produces_output","declares_output","enables_output","traceability_reference"} for e in edges),
            },
            "DECLARED_OUTPUT_PRODUCER_TRACEABILITY": {
                "status": "FAIL" if any(row.get("status") == "FAIL" for row in output_traceability) else ("WARNING" if any(row.get("status") == "WARNING" for row in output_traceability) else "PASS"),
                "resolved": len([row for row in output_traceability if row.get("status") == "PASS"]),
                "unresolved": len([row for row in output_traceability if row.get("status") == "WARNING"]),
                "ambiguous": len([row for row in output_traceability if row.get("status") == "FAIL"]),
            },
        },
        "projection_diagnostics": {
            "profile": "release2_typed_semantic_projection",
            "suppressed_reserved_routes": suppressed_reserved_routes,
            "navigation_permissions_not_rendered_as_control_flow": navigation_permissions,
            "unreachable_terminals": unreachable_terminals,
            "unresolved_validation_dependencies": unresolved_dependencies,
            "orphan_outputs": [row["output_id"] for row in output_traceability if row.get("status") == "WARNING"],
            "ambiguous_outputs": [row["output_id"] for row in output_traceability if row.get("status") == "FAIL"],
            "declared_output_producer_traceability": output_traceability,
            "synthetic_execution_nodes": [],
            "runtime_semantics_mutated": False,
        },
    }


def _declared_terminal_ids(source: dict[str, Any]) -> set[str]:
    terminal_ids = {str(item["id"]) for item in _terminal_records(source)}
    for node in source.get("nodes", []):
        if isinstance(node, dict) and node.get("terminal") is True and isinstance(node.get("id"), str):
            terminal_ids.add(node["id"])
    for output in source.get("outputs", []):
        if isinstance(output, dict) and isinstance(output.get("id"), str):
            terminal_ids.add(output["id"])
        elif isinstance(output, str) and output.strip():
            terminal_ids.add(output)
    contract = source.get("graph_contract", {})
    if isinstance(contract, dict):
        terminal_ids.update(value for value in contract.get("external_terminal_targets", []) if isinstance(value, str))
    return terminal_ids


def _entry_node_id(source: dict[str, Any]) -> str | None:
    contract = source.get("graph_contract", {})
    if isinstance(contract, dict) and isinstance(contract.get("entry_node"), str):
        return contract["entry_node"]
    playbook = source.get("playbook", {})
    if isinstance(playbook, dict) and isinstance(playbook.get("entry_node"), str):
        return playbook["entry_node"]
    return None


def _declared_dynamic_route_edges(record: dict[str, Any]) -> list[dict[str, str]]:
    """Return declared runtime routing possibilities as structural-only edges.

    These edges are used by editor-local validation for reachability/dead-end/path
    checks. The runtime expression itself (for example ``$plan.selected_node``) is
    deliberately not evaluated.
    """
    routes = record.get("declared_dynamic_routes")
    if not isinstance(routes, dict):
        return []

    edges: list[dict[str, str]] = []
    seen: set[str] = set()
    for route_name, route in routes.items():
        targets: list[str] = []
        if isinstance(route, str):
            targets = [route]
        else:
            targets = _targets(route)
        for target in targets:
            if not isinstance(target, str) or not target or target.startswith("$") or target in seen:
                continue
            seen.add(target)
            edges.append({
                "target": target,
                "storage": "declared_dynamic_routes",
                "key": str(route_name),
            })
    return edges


def _structural_model(source: dict[str, Any]) -> dict[str, Any]:
    nodes = [item for item in source.get("nodes", []) if isinstance(item, dict)] if isinstance(source.get("nodes", []), list) else []
    gates = [item for item in source.get("gates", []) if isinstance(item, dict)] if isinstance(source.get("gates", []), list) else []
    terminal_records = _terminal_records(source)
    outputs = source.get("outputs", []) if isinstance(source.get("outputs", []), list) else []

    records = [("node", item) for item in nodes] + [("gate", item) for item in gates] + [("terminal", item) for item in terminal_records]
    ids = [str(item.get("id")) for _, item in records if isinstance(item.get("id"), str) and item.get("id")]
    executable_ids = {
        str(item["id"])
        for _, item in [("node", item) for item in nodes] + [("gate", item) for item in gates]
        if isinstance(item.get("id"), str) and item.get("id")
    }
    terminal_ids = _declared_terminal_ids(source)
    known_targets = set(ids) | terminal_ids
    edges = [
        {"source": str(node["id"]), **edge}
        for node in nodes if isinstance(node.get("id"), str) and node.get("id")
        for edge in [*_node_edges(node), *_declared_dynamic_route_edges(node)]
    ] + [
        {"source": str(gate["id"]), **edge}
        for gate in gates if isinstance(gate.get("id"), str) and gate.get("id")
        for edge in [*_gate_edges(gate), *_declared_dynamic_route_edges(gate)]
    ]
    adjacency = {identifier: [] for identifier in known_targets | executable_ids}
    reverse = {identifier: [] for identifier in known_targets | executable_ids}
    for edge in edges:
        adjacency.setdefault(edge["source"], []).append(edge["target"])
        reverse.setdefault(edge["target"], []).append(edge["source"])
    return {
        "nodes": nodes,
        "gates": gates,
        "records": records,
        "ids": ids,
        "executable_ids": executable_ids,
        "terminal_ids": terminal_ids,
        "known_targets": known_targets,
        "edges": edges,
        "adjacency": adjacency,
        "reverse": reverse,
        "entry": _entry_node_id(source),
        "outputs": outputs,
    }


def _result(check_id: str, name: str, findings: list[dict[str, Any]], pass_summary: str) -> dict[str, Any]:
    if any(item.get("severity") == "error" for item in findings):
        status = "error"
    elif any(item.get("severity") == "warning" for item in findings):
        status = "warning"
    else:
        status = "passed"
    return {
        "id": check_id,
        "name": name,
        "status": status,
        "summary": pass_summary if not findings else f"{len(findings)} finding(s).",
        "findings": findings,
    }


def _reachable(start: str | None, adjacency: dict[str, list[str]], allowed: set[str]) -> set[str]:
    if not start or start not in allowed:
        return set()
    seen: set[str] = set()
    stack = [start]
    while stack:
        current = stack.pop()
        if current in seen or current not in allowed:
            continue
        seen.add(current)
        stack.extend(target for target in adjacency.get(current, []) if target in allowed and target not in seen)
    return seen


def _cycle_components(executable_ids: set[str], adjacency: dict[str, list[str]]) -> list[list[str]]:
    index = 0
    indices: dict[str, int] = {}
    lowlink: dict[str, int] = {}
    stack: list[str] = []
    on_stack: set[str] = set()
    components: list[list[str]] = []

    def visit(node: str) -> None:
        nonlocal index
        indices[node] = index
        lowlink[node] = index
        index += 1
        stack.append(node)
        on_stack.add(node)
        for target in adjacency.get(node, []):
            if target not in executable_ids:
                continue
            if target not in indices:
                visit(target)
                lowlink[node] = min(lowlink[node], lowlink[target])
            elif target in on_stack:
                lowlink[node] = min(lowlink[node], indices[target])
        if lowlink[node] == indices[node]:
            component: list[str] = []
            while stack:
                member = stack.pop()
                on_stack.remove(member)
                component.append(member)
                if member == node:
                    break
            if len(component) > 1 or node in adjacency.get(node, []):
                components.append(sorted(component))

    for node in sorted(executable_ids):
        if node not in indices:
            visit(node)
    return components


def validate_source(source: dict[str, Any]) -> dict[str, Any]:
    """Run editor-local structural checks only; this is not full Ordo validation."""
    model = _structural_model(source)
    checks: list[dict[str, Any]] = []

    checks.append(_result(
        "yaml_root",
        "YAML structure",
        [],
        "YAML root is a mapping and can be structurally inspected.",
    ))

    id_findings: list[dict[str, Any]] = []
    missing_id_records = [kind for kind, item in model["records"] if not isinstance(item.get("id"), str) or not item.get("id", "").strip()]
    if missing_id_records:
        id_findings.append({"severity": "error", "code": "MISSING_ID", "message": f"{len(missing_id_records)} graph record(s) have no non-empty id."})
    seen: set[str] = set()
    duplicates: set[str] = set()
    for identifier in model["ids"]:
        if identifier in seen:
            duplicates.add(identifier)
        seen.add(identifier)
    if duplicates:
        id_findings.append({"severity": "error", "code": "DUPLICATE_ID", "message": "Duplicate graph IDs: " + ", ".join(sorted(duplicates))})
    checks.append(_result("unique_ids", "Graph element IDs", id_findings, f"{len(model['ids'])} graph element IDs are present and unique."))

    entry_findings: list[dict[str, Any]] = []
    entry = model["entry"]
    if not entry:
        entry_findings.append({"severity": "error", "code": "ENTRY_MISSING", "message": "No explicit entry_node is declared in graph_contract or playbook."})
    elif entry not in model["executable_ids"]:
        entry_findings.append({"severity": "error", "code": "ENTRY_UNKNOWN", "message": f"Entry node {entry} does not reference an existing node or gate."})
    checks.append(_result("entry_point", "Entry point", entry_findings, f"Entry point {entry} exists." if entry else "Entry point exists."))

    target_findings: list[dict[str, Any]] = []
    for edge in model["edges"]:
        if edge["target"] not in model["known_targets"]:
            target_findings.append({
                "severity": "error",
                "code": "DANGLING_TARGET",
                "message": f"{edge['source']} -> {edge['target']} references an undeclared target.",
                "element": edge["source"],
            })
    checks.append(_result("transition_targets", "Transition targets", target_findings, f"All {len(model['edges'])} transition target(s) resolve to declared graph elements or terminals."))

    allowed = model["known_targets"] | model["executable_ids"]
    reachable = _reachable(entry, model["adjacency"], allowed)
    unreachable_exec = sorted(model["executable_ids"] - reachable)
    reachability_findings = [
        {"severity": "warning", "code": "UNREACHABLE_ELEMENT", "message": f"{identifier} is not reachable from the entry point.", "element": identifier}
        for identifier in unreachable_exec
    ]
    checks.append(_result("reachability", "Reachability from entry", reachability_findings, f"All {len(model['executable_ids'])} executable graph elements are reachable from the entry point."))

    reachable_terminals = sorted(model["terminal_ids"] & reachable)
    unreachable_terminals = sorted(model["terminal_ids"] - reachable)
    terminal_findings: list[dict[str, Any]] = []
    if not model["terminal_ids"]:
        terminal_findings.append({"severity": "error", "code": "NO_TERMINALS", "message": "No terminal/output targets are declared."})
    elif not reachable_terminals:
        terminal_findings.append({"severity": "error", "code": "NO_REACHABLE_TERMINAL", "message": "No declared terminal/output is reachable from the entry point."})
    terminal_findings.extend(
        {"severity": "warning", "code": "UNREACHABLE_TERMINAL", "message": f"Terminal/output {identifier} is not reachable from the entry point.", "element": identifier}
        for identifier in unreachable_terminals
    )
    checks.append(_result("terminal_reachability", "Terminal reachability", terminal_findings, f"All {len(model['terminal_ids'])} declared terminal/output target(s) are reachable."))

    terminal_like = set(model["terminal_ids"])
    dead_ends = sorted(
        identifier for identifier in (model["executable_ids"] & reachable)
        if identifier not in terminal_like and not [target for target in model["adjacency"].get(identifier, []) if target in allowed]
    )
    dead_end_findings = [
        {"severity": "error", "code": "DEAD_END", "message": f"{identifier} is reachable but has no outgoing transition and is not terminal.", "element": identifier}
        for identifier in dead_ends
    ]
    checks.append(_result("dead_ends", "Dead-end branches", dead_end_findings, "No reachable non-terminal node or gate ends without an outgoing transition."))

    can_reach_terminal: set[str] = set()
    stack = list(model["terminal_ids"] & allowed)
    while stack:
        current = stack.pop()
        if current in can_reach_terminal:
            continue
        can_reach_terminal.add(current)
        stack.extend(source_id for source_id in model["reverse"].get(current, []) if source_id in allowed)
    trapped = sorted((model["executable_ids"] & reachable) - can_reach_terminal)
    closure_findings = [
        {"severity": "error", "code": "NO_TERMINAL_PATH", "message": f"{identifier} is reachable from entry but no path from it reaches a declared terminal/output.", "element": identifier}
        for identifier in trapped
    ]
    checks.append(_result("terminal_paths", "Paths to terminal outcomes", closure_findings, "Every reachable executable element has at least one path to a declared terminal/output."))

    cycles = _cycle_components(model["executable_ids"], model["adjacency"])
    cycle_findings = [
        {"severity": "info", "code": "CYCLE_DETECTED", "message": "Cycle detected: " + " -> ".join(component), "elements": component}
        for component in cycles
    ]
    checks.append(_result("cycles", "Cycle discovery", cycle_findings, "No executable graph cycles detected."))

    output_trace_findings: list[dict[str, Any]] = []
    for row in graph_view(source).get("projection_diagnostics", {}).get("declared_output_producer_traceability", []):
        status = row.get("status")
        if status == "WARNING":
            output_trace_findings.append({
                "severity": "warning",
                "code": "DECLARED_OUTPUT_PRODUCER_UNRESOLVED",
                "message": f"Declared output {row.get('output_id')} has no uniquely resolvable formal producer reference.",
                "element": row.get("output_id"),
            })
        elif status == "FAIL":
            candidates = ", ".join(row.get("candidates") or [])
            output_trace_findings.append({
                "severity": "error",
                "code": "DECLARED_OUTPUT_PRODUCER_AMBIGUOUS",
                "message": f"Declared output {row.get('output_id')} matches multiple producer nodes: {candidates}.",
                "element": row.get("output_id"),
            })
    checks.append(_result(
        "declared_output_producer_traceability",
        "DECLARED_OUTPUT_PRODUCER_TRACEABILITY",
        output_trace_findings,
        "Every declared output with formal artifact/output references resolves to a unique producer; semantic links remain view-only.",
    ))

    errors = sum(1 for check in checks for item in check["findings"] if item.get("severity") == "error")
    warnings = sum(1 for check in checks for item in check["findings"] if item.get("severity") == "warning")
    infos = sum(1 for check in checks for item in check["findings"] if item.get("severity") == "info")
    issues = []
    for check in checks:
        for finding in check["findings"]:
            issue = dict(finding)
            if issue.get("code") == "DANGLING_TARGET":
                issue["code"] = "GRAPH_TARGET_MISSING"
            issue["check"] = check["id"]
            issues.append(issue)
    return {
        "scope": "editor_structural_validation",
        "status": "failed" if errors else ("warning" if warnings else "passed"),
        "summary": {"checks": len(checks), "errors": errors, "warnings": warnings, "info": infos},
        "checks": checks,
        # Compatibility alias retained for integrations that consume a flat
        # graph-finding list rather than grouped validation checks.
        "issues": issues,
        "note": "Structural validation only. Full Ordo playbook validation must be performed by the Ordo validation/playbook tooling.",
    }


def tree_module_manifest_path() -> Path:
    candidates = (
        REPOSITORY_ROOT / "packages" / "ordo_applied_project_factory" / "source" / "tree_module_library" / "manifest.yaml",
        REPOSITORY_ROOT / "source" / "tree_module_library" / "manifest.yaml",
    )
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    raise FileNotFoundError("Cannot find an ARF tree-module library manifest.")



def _humanize_replay_id(record_id: str) -> str:
    base = record_id[2:] if len(record_id) > 2 and record_id[1:2] == "_" else record_id
    words = [part for part in base.split("_") if part]
    if not words:
        return record_id
    return " ".join(word.upper() if len(word) <= 3 and word.isupper() else word.lower() for word in words).capitalize()


def _replay_display_label(record: dict[str, Any] | None, record_id: str, kind: str) -> str:
    if not isinstance(record, dict):
        return _humanize_replay_id(record_id)
    fields = (
        ("title", "purpose", "condition", "description", "summary", "question", "prompt", "instruction", "propose", "proposal", "display_name", "name", "label", "specification", "validator", "method", "type", "kind")
        if kind == "gate"
        else ("title", "question", "purpose", "description", "summary", "prompt", "instruction", "propose", "proposal", "display_name", "name", "label", "action", "type", "kind")
    )
    for key in fields:
        value = record.get(key)
        if isinstance(value, str) and value.strip() and value.strip() != record_id:
            return value.strip()
    return _humanize_replay_id(record_id)


def _replay_record_maps(source: dict[str, Any]) -> tuple[dict[str, dict[str, Any]], dict[str, str]]:
    records: dict[str, dict[str, Any]] = {}
    kinds: dict[str, str] = {}
    for collection, kind in (("nodes", "node"), ("gates", "gate"), ("outputs", "output"), ("terminals", "terminal")):
        values = source.get(collection, [])
        if isinstance(values, list):
            for record in values:
                if isinstance(record, dict) and isinstance(record.get("id"), str):
                    records[record["id"]] = record
                    kinds[record["id"]] = kind
    return records, kinds


def _decision_target_node(decision: dict[str, Any], path: list[str], kinds: dict[str, str]) -> str | None:
    checkpoint = decision.get("checkpoint_id")
    if isinstance(checkpoint, str) and checkpoint in path:
        index = path.index(checkpoint)
        if kinds.get(checkpoint) == "node":
            return checkpoint
        for candidate in reversed(path[:index]):
            if kinds.get(candidate) == "node":
                return candidate
    decision_id = str(decision.get("decision_id", ""))
    dtokens = {part for part in (decision_id[2:] if decision_id.startswith("D_") else decision_id).split("_") if len(part) > 2 and part not in {"IMPLEMENTATION"}}
    best: tuple[int, str] | None = None
    for node_id in path:
        if kinds.get(node_id) != "node":
            continue
        ntokens = set((node_id[2:] if node_id.startswith("N_") else node_id).split("_"))
        score = len(dtokens & ntokens)
        if score and (best is None or score > best[0]):
            best = (score, node_id)
    return best[1] if best else None


def _interaction_events(trace: dict[str, Any]) -> list[dict[str, Any]]:
    value = trace.get("interaction_trace")
    if not isinstance(value, list):
        return []
    return [item for item in value if isinstance(item, dict)]


def _interaction_text(event: dict[str, Any], *keys: str) -> str | None:
    for key in keys:
        value = event.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return None


def build_replay_view(trace: dict[str, Any], source: dict[str, Any] | None) -> dict[str, Any]:
    source = source if isinstance(source, dict) else {}
    records, kinds = _replay_record_maps(source)
    path = [item for item in trace.get("node_path", []) if isinstance(item, str)]

    interactions = _interaction_events(trace)
    interactions_by_node: dict[str, list[dict[str, Any]]] = {}
    for event in interactions:
        node_id = event.get("node_id")
        if isinstance(node_id, str) and node_id:
            interactions_by_node.setdefault(node_id, []).append(event)
    for events in interactions_by_node.values():
        events.sort(key=lambda item: item.get("sequence") if isinstance(item.get("sequence"), int) else 10**9)

    decisions_by_node: dict[str, list[dict[str, Any]]] = {}
    unattached: list[dict[str, Any]] = []
    for decision in trace.get("accepted_decisions", []):
        if not isinstance(decision, dict):
            continue
        target = _decision_target_node(decision, path, kinds)
        if target:
            decisions_by_node.setdefault(target, []).append(decision)
        else:
            unattached.append(decision)
    checkpoints = set(item for item in trace.get("checkpoints", []) if isinstance(item, str))
    impl_verification = trace.get("implementation_verification") if isinstance(trace.get("implementation_verification"), dict) else {}
    blocked_gate = impl_verification.get("gate") if str(impl_verification.get("status", "")).lower() not in {"passed", "pass", "ok", "success"} else None
    steps = []
    for index, record_id in enumerate(path):
        record = records.get(record_id, {})
        kind = kinds.get(record_id, "gate" if record_id.startswith("G_") else "node")
        node_interactions = interactions_by_node.get(record_id, [])

        prompt = None
        if node_interactions:
            # The exported interaction trace is factual runtime/reconstruction evidence and
            # therefore has precedence over wording reconstructed from the source playbook.
            prompt = _interaction_text(node_interactions[0], "question_or_prompt_shown", "question", "prompt")
        if not prompt:
            if kind == "gate":
                prompt = _replay_display_label(record, record_id, "gate")
            else:
                for key in ("title", "question", "purpose", "description", "summary", "prompt", "instruction", "action"):
                    value = record.get(key) if isinstance(record, dict) else None
                    if isinstance(value, str) and value.strip():
                        prompt = value.strip(); break
                prompt = prompt or _replay_display_label(record, record_id, kind)

        label = _replay_display_label(record, record_id, kind)
        # When source metadata is absent, a captured prompt is a more useful display label
        # than repeating the technical ID. Keep the ID separately as the factual graph key.
        if label == record_id and prompt and prompt != record_id:
            label = prompt

        gate_status = None
        if kind == "gate":
            gate_status = "blocked" if record_id == blocked_gate else ("passed" if record_id in checkpoints else "visited")
        steps.append({
            "index": index + 1,
            "id": record_id,
            "kind": kind,
            "label": label,
            "prompt": prompt,
            "answer_type": record.get("answer_type") if isinstance(record, dict) else None,
            "gate_status": gate_status,
            "decisions": decisions_by_node.get(record_id, []),
            "interactions": [
                {
                    "sequence": event.get("sequence"),
                    "question": _interaction_text(event, "question_or_prompt_shown", "question", "prompt"),
                    "analyst_response": _interaction_text(event, "analyst_answer_verbatim", "analyst_response", "answer"),
                    "transition_to": event.get("transition_to"),
                    "capture_mode": event.get("capture_mode"),
                    "timestamp": event.get("timestamp"),
                    "timestamp_status": event.get("timestamp_status"),
                }
                for event in node_interactions
            ],
        })
    playbook = trace.get("playbook") if isinstance(trace.get("playbook"), dict) else {}
    interaction_contract = trace.get("interaction_trace_contract") if isinstance(trace.get("interaction_trace_contract"), dict) else {}
    return {
        "run_id": trace.get("run_id"),
        "playbook": playbook,
        "export_timestamp": trace.get("export_timestamp"),
        "final_active_node": trace.get("final_active_node"),
        "readiness_status": trace.get("readiness_status") if isinstance(trace.get("readiness_status"), dict) else {},
        "steps": steps,
        "unattached_decisions": unattached,
        "source_available": bool(records),
        "interaction_trace_available": bool(interactions),
        "interaction_trace_schema_version": trace.get("interaction_trace_schema_version"),
        "interaction_trace_status": interaction_contract.get("status"),
        "verbatim_answers_available": bool(interaction_contract.get("verbatim_answers_available")) or any(
            item.get("analyst_response") for step in steps for item in step.get("interactions", [])
        ),
        "verbatim_chat_available": False,
    }

def build_debug_reproduction_view(debug: dict[str, Any]) -> dict[str, Any]:
    run = debug.get("run") if isinstance(debug.get("run"), dict) else {}
    calls = debug.get("calls") if isinstance(debug.get("calls"), list) else []
    history = run.get("history") if isinstance(run.get("history"), list) else []
    answers_by_node: dict[str, list[str]] = {}
    for item in history:
        if not isinstance(item, dict) or item.get("role") != "analyst":
            continue
        node_id = str(item.get("node_id") or "").strip()
        text = str(item.get("text") or "").strip()
        if node_id and text:
            answers_by_node.setdefault(node_id, []).append(text)
    replay_calls=[]
    for call in calls:
        if not isinstance(call, dict):
            continue
        replay_calls.append({
            "index": call.get("index"),
            "current_id": call.get("current_id"),
            "element_kind": call.get("element_kind"),
            "phase": call.get("phase"),
            "llm_call_skipped": bool((call.get("runtime") or {}).get("llm_call_skipped")),
            "parsed_result": copy.deepcopy(((call.get("output") or {}).get("parsed_result"))),
            "state_before": copy.deepcopy(((call.get("runtime") or {}).get("state_before")) or {}),
            "state_after": copy.deepcopy(((call.get("runtime") or {}).get("state_after")) or {}),
            "selected_route_key": (call.get("runtime") or {}).get("selected_route_key"),
            "next_id": (call.get("runtime") or {}).get("next_id"),
        })
    outcome = run.get("outcome") if isinstance(run.get("outcome"), dict) else {}
    checkpoint = str(outcome.get("nodeId") or run.get("current_id") or "").strip()
    path=[x for x in (run.get("path") or []) if isinstance(x,str)]
    if not checkpoint and path: checkpoint=path[-1]
    return {
        "kind":"debug_reproduction",
        "schema_version":"ordo.debug_reproduction.v1",
        "run_id": run.get("run_id"),
        "path": path,
        "suggested_checkpoint": checkpoint,
        "halt_reason": outcome.get("reason"),
        "answers_by_node": answers_by_node,
        "recorded_calls": replay_calls,
        "total_recorded_calls": len(replay_calls),
        "total_recorded_answers": sum(len(v) for v in answers_by_node.values()),
        "steps": [],
        "source_available": False,
        "interaction_trace_available": bool(answers_by_node),
        "verbatim_answers_available": bool(answers_by_node),
        "verbatim_chat_available": True,
    }

def _strict_replay_provenance_errors(view: dict[str, Any]) -> list[str]:
    if not isinstance(view, dict) or not bool(view.get("strict_replay_provenance")):
        return []
    errors=[]
    if not str(view.get("source_run_id") or "").strip(): errors.append("source_run_id is required for strict replay provenance")
    if len(str(view.get("source_sha256") or "")) != 64: errors.append("source_sha256 is required for strict replay provenance")
    recorded=view.get("recorded_against") if isinstance(view.get("recorded_against"),dict) else {}
    for key in ("compiler","editor","playbook","semantic_plan_sha256"):
        if not str(recorded.get(key) or "").strip(): errors.append(f"recorded_against.{key} is required for strict replay provenance")
    return errors

def _with_replay_provenance(view: dict[str, Any], filename: str, raw: bytes) -> dict[str, Any]:
    result=copy.deepcopy(view) if isinstance(view,dict) else {}
    trace=result.get("source_trace") if isinstance(result.get("source_trace"),dict) else {}
    raw_meta={}
    try:
        candidate=json.loads(raw.decode("utf-8")) if filename.lower().endswith(".json") else {}
        if isinstance(candidate,dict): raw_meta=candidate
    except Exception:
        raw_meta={}
    # Preserve package-authored component provenance, then enrich provider identity from trace.
    recorded_against=copy.deepcopy(result.get("recorded_against")) if isinstance(result.get("recorded_against"),dict) else {}
    if isinstance(raw_meta.get("recorded_against"),dict):
        recorded_against.update(copy.deepcopy(raw_meta["recorded_against"]))
    for key in ("provider","base_url","model"):
        value=(trace.get(key) if isinstance(trace,dict) else None) or raw_meta.get(key)
        if value is not None: recorded_against[key]=value
    run_meta=(trace.get("run") if isinstance(trace,dict) else None) or raw_meta.get("run") or {}
    result["source_run_id"]=str(result.get("source_run_id") or (run_meta or {}).get("run_id") or raw_meta.get("run_id") or "")
    result["source_sha256"]=hashlib.sha256(raw).hexdigest()
    result["source_filename"]=filename
    result["recorded_against"]=recorded_against
    errors=_strict_replay_provenance_errors(result)
    if errors:
        raise ValueError("Strict replay provenance invalid: " + "; ".join(errors))
    return result

def parse_replay_package(filename: str, raw: bytes) -> dict[str, Any]:
    lower = filename.lower()
    trace: dict[str, Any] | None = None
    source: dict[str, Any] | None = None
    if lower.endswith(".json"):
        loaded = json.loads(raw.decode("utf-8"))
        if not isinstance(loaded, dict):
            raise ValueError("Replay JSON root must be an object.")
        if isinstance(loaded.get("recorded_calls"), list) and (str(loaded.get("kind") or "") == "debug_reproduction" or str(loaded.get("schema_version") or "").startswith("ordo.debug_reproduction.")):
            return _with_replay_provenance(loaded, filename, raw)
        if isinstance(loaded.get("calls"), list) and isinstance(loaded.get("run"), dict):
            return _with_replay_provenance(build_debug_reproduction_view(loaded), filename, raw)
        trace = loaded
    else:
        try:
            archive = zipfile.ZipFile(io.BytesIO(raw))
        except zipfile.BadZipFile as error:
            raise ValueError("Replay package must be a ZIP or run_trace.json file.") from error
        names = archive.namelist()
        reproduction_name = next((name for name in names if Path(name).name == "reproduction.json"), None)
        if reproduction_name:
            loaded = json.loads(archive.read(reproduction_name).decode("utf-8"))
            if isinstance(loaded, dict) and isinstance(loaded.get("recorded_calls"), list):
                return _with_replay_provenance(loaded, filename, raw)
        debug_name = next((name for name in names if "debug-run-summary" in Path(name).name.lower() and name.lower().endswith(".json")), None)
        if debug_name:
            loaded = json.loads(archive.read(debug_name).decode("utf-8"))
            if isinstance(loaded, dict) and isinstance(loaded.get("calls"), list) and isinstance(loaded.get("run"), dict):
                return _with_replay_provenance(build_debug_reproduction_view(loaded), filename, raw)
        trace_name = next((name for name in names if Path(name).name == "run_trace.json"), None)
        if not trace_name:
            raise ValueError("Replay ZIP does not contain run_trace.json or reproduction.json.")
        loaded = json.loads(archive.read(trace_name).decode("utf-8"))
        if not isinstance(loaded, dict):
            raise ValueError("run_trace.json root must be an object.")
        trace = loaded
        if not isinstance(trace.get("interaction_trace"), list):
            interaction_name = next((name for name in names if Path(name).name == "interaction_trace.json"), None)
            if interaction_name:
                try:
                    interaction_candidate = json.loads(archive.read(interaction_name).decode("utf-8"))
                    if isinstance(interaction_candidate, list):
                        trace["interaction_trace"] = interaction_candidate
                except (json.JSONDecodeError, UnicodeDecodeError):
                    pass
        preferred = "playbook/source/program.ordo.yaml"
        yaml_name = preferred if preferred in names else next((name for name in names if Path(name).name in {"program.ordo.yaml", "program.ordo.yml"}), None)
        if yaml_name:
            try:
                candidate = yaml.safe_load(archive.read(yaml_name).decode("utf-8"))
                if isinstance(candidate, dict):
                    source = candidate
            except yaml.YAMLError:
                source = None
    return _with_replay_provenance(build_replay_view(trace or {}, source), filename, raw)


def _replay_auto_answers(replay: dict[str, Any]) -> dict[str, list[str]]:
    answers: dict[str, list[str]] = {}
    if isinstance(replay.get("answers_by_node"), dict):
        for node_id, values in replay["answers_by_node"].items():
            if isinstance(values, list):
                cleaned=[str(v).strip() for v in values if str(v).strip()]
                if cleaned: answers[str(node_id)]=cleaned
    for step in replay.get("steps") or []:
        if not isinstance(step, dict) or not step.get("id"): continue
        node_id=str(step["id"]); values=answers.setdefault(node_id,[])
        for interaction in step.get("interactions") or []:
            if not isinstance(interaction, dict): continue
            value=str(interaction.get("analyst_response") or "").strip()
            if value and value not in values: values.append(value)
        if not values: answers.pop(node_id,None)
    return answers

def _register_replay_package(filename: str, raw: bytes) -> dict[str, Any]:
    replay=parse_replay_package(filename,raw)
    replay_id=hashlib.sha256(raw).hexdigest()[:16]
    REPLAY_PACKAGES[replay_id]={
        "replay_id":replay_id,"filename":filename,"replay":copy.deepcopy(replay),
        "answers_by_node":_replay_auto_answers(replay),
    }
    return {"replay_id":replay_id,"replay":replay,"auto_answer_count":sum(len(v) for v in REPLAY_PACKAGES[replay_id]["answers_by_node"].values())}


_TEXT_RESOURCE_EXTENSIONS = {
    ".yaml", ".yml", ".json", ".md", ".txt", ".csv", ".tsv", ".xml", ".html", ".htm",
    ".jinja", ".jinja2", ".j2", ".tmpl", ".template", ".prompt", ".schema", ".py", ".js", ".ts",
}

def _safe_zip_name(name: str) -> bool:
    path = Path(name)
    return bool(name and not path.is_absolute() and ".." not in path.parts and not name.endswith("/"))

def _looks_like_playbook(source: Any) -> bool:
    if not isinstance(source, dict):
        return False
    return isinstance(source.get("nodes"), list) or isinstance(source.get("gates"), list) or _entry_id(source) is not None


R3_PACKAGE_CAPABILITIES = frozenset({
    "typed_state_patch",
    "model_gate_check_results",
    "artifact_materialization",
    "deterministic_validator",
    "revisit_context",
    "recovery_router",
    "analyst_auto_answers",
    "live_evidence_export",
    "guided_replay",
    "package_session_isolation",
    "package_tool_execution",
})

def _validate_r3_package_manifest_v2(archive: zipfile.ZipFile, infos: list[zipfile.ZipInfo], source_name: str, source_raw: bytes, semantic_names: list[str], semantic_plan: dict[str, Any] | None) -> dict[str, Any]:
    candidates=[]
    for info in infos:
        if Path(info.filename).name != "manifest.json":
            continue
        try:
            obj=json.loads(archive.read(info.filename).decode("utf-8-sig"))
        except (UnicodeDecodeError,json.JSONDecodeError):
            continue
        if isinstance(obj,dict) and str(obj.get("manifest_version") or "") == "2.0":
            candidates.append((info.filename,obj))
    if not candidates:
        return {"available":False,"valid":False,"reason":"not_present"}
    if len(candidates)>1:
        raise ValueError("Package contains multiple manifest v2 files.")
    manifest_name, manifest=candidates[0]
    required={"manifest_version","package_id","package_version","playbook_contract_version","semantic_plan_format_version","entry_node","terminal_nodes","capabilities","integrity","files"}
    missing=sorted(required-set(manifest))
    if missing:
        raise ValueError(f"Package manifest v2 missing required fields: {', '.join(missing)}")
    capabilities=manifest.get("capabilities")
    if not isinstance(capabilities,dict):
        raise ValueError("Package manifest v2 capabilities must be an object.")
    required_caps=capabilities.get("required")
    optional_caps=capabilities.get("optional")
    if not isinstance(required_caps,list) or not all(isinstance(x,str) and x for x in required_caps):
        raise ValueError("Package manifest v2 capabilities.required must be a string array.")
    if not isinstance(optional_caps,list) or not all(isinstance(x,str) and x for x in optional_caps):
        raise ValueError("Package manifest v2 capabilities.optional must be a string array.")
    if len(required_caps)!=len(set(required_caps)) or len(optional_caps)!=len(set(optional_caps)):
        raise ValueError("Package manifest v2 capabilities contain duplicates.")
    overlap=sorted(set(required_caps)&set(optional_caps))
    if overlap:
        raise ValueError(f"Package manifest v2 capability appears as both required and optional: {overlap[0]}")
    unsupported=sorted(set(required_caps)-R3_PACKAGE_CAPABILITIES)
    if unsupported:
        raise ValueError(f"Package requires unsupported runtime capability: {unsupported[0]}")
    integrity=manifest.get("integrity")
    if not isinstance(integrity,dict) or integrity.get("hash_algorithm")!="sha256" or integrity.get("coverage")!="all_files_except_manifest":
        raise ValueError("Package manifest v2 integrity contract is unsupported.")
    source_sha=str(integrity.get("source_sha256") or "").lower()
    actual_source_sha=hashlib.sha256(source_raw).hexdigest()
    if source_sha != actual_source_sha:
        raise ValueError(f"Package manifest v2 source SHA mismatch: declared={source_sha} actual={actual_source_sha}")
    if manifest.get("entry_node") != _entry_id(yaml.safe_load(source_raw.decode("utf-8-sig"))):
        raise ValueError("Package manifest v2 entry_node does not match playbook source.")
    files=manifest.get("files")
    if not isinstance(files,dict) or not files:
        raise ValueError("Package manifest v2 files must contain hashes for package content.")
    expected_paths={info.filename for info in infos if info.filename != manifest_name}
    if set(files) != expected_paths:
        missing_paths=sorted(expected_paths-set(files)); extra_paths=sorted(set(files)-expected_paths)
        detail=[]
        if missing_paths: detail.append(f"missing={missing_paths[:3]}")
        if extra_paths: detail.append(f"extra={extra_paths[:3]}")
        raise ValueError("Package manifest v2 file coverage mismatch: " + " ".join(detail))
    for name,declared in files.items():
        if not isinstance(declared,str) or len(declared)!=64:
            raise ValueError(f"Package manifest v2 invalid SHA256 for {name}.")
        actual=hashlib.sha256(archive.read(name)).hexdigest()
        if declared.lower()!=actual:
            raise ValueError(f"Package manifest v2 file SHA mismatch: {name}")
    declared_plan_sha=integrity.get("semantic_plan_sha256")
    if declared_plan_sha is not None:
        if len(semantic_names)!=1 or not isinstance(semantic_plan,dict):
            raise ValueError("Package manifest v2 declares semantic_plan_sha256 but no usable semantic plan is present.")
        actual=hashlib.sha256(archive.read(semantic_names[0])).hexdigest()
        if str(declared_plan_sha).lower()!=actual:
            raise ValueError("Package manifest v2 semantic plan SHA mismatch.")
        if str(manifest.get("semantic_plan_format_version") or "") != str(semantic_plan.get("format_version") or ""):
            raise ValueError("Package manifest v2 semantic_plan_format_version mismatch.")
    return {
        "available":True,"valid":True,"reason":"ok","path":manifest_name,
        "manifest_version":"2.0","package_id":manifest.get("package_id"),"package_version":manifest.get("package_version"),
        "required_capabilities":required_caps,"optional_capabilities":optional_caps,
    }


_RUNTIME_SEMANTIC_PLAN_RELATIVE_PATH = "compiled/runtime_semantic_plan.json"
_CANONICAL_ORDO_SOURCE_RELATIVE_PATHS = (
    "playbook/source/program.ordo.yaml",
    "playbook/source/program.ordo.yml",
    "program.ordo.yaml",
    "program.ordo.yml",
)


def _single_enclosing_zip_root(names: list[str]) -> str | None:
    """Return the sole archive wrapper directory, if every file is below one root.

    This is package-structure discovery only. It does not grant semantic authority to
    arbitrary nested directories; it merely lets canonical package-relative paths be
    expressed below a conventional ZIP wrapper directory.
    """
    if not names or any("/" not in name.strip("/") for name in names):
        return None
    roots = {name.strip("/").split("/", 1)[0] for name in names if name.strip("/")}
    return next(iter(roots)) if len(roots) == 1 else None


def _top_level_runtime_plan_package_roots(names: list[str]) -> list[str]:
    """Find malformed competing wrapper roots that each look like a package root.

    A top-level directory only becomes a candidate here when it contains both a
    canonical Ordo source location and the canonical runtime-plan location. This is
    used solely to fail explicitly on ambiguous multi-package ZIPs; it is never used
    as a broad nested-plan discovery mechanism.
    """
    name_set = set(names)
    roots: list[str] = []
    for name in names:
        suffix = "/" + _RUNTIME_SEMANTIC_PLAN_RELATIVE_PATH
        if not name.endswith(suffix):
            continue
        prefix = name[: -len(suffix)]
        if not prefix or "/" in prefix:
            continue
        if any(f"{prefix}/{rel}" in name_set for rel in _CANONICAL_ORDO_SOURCE_RELATIVE_PATHS):
            roots.append(prefix)
    return sorted(set(roots))


def _resolve_runtime_semantic_plan_authority(names: list[str]) -> dict[str, Any]:
    """Resolve active Runtime Semantic Plan from package-authoritative locations.

    Canonical Ordo does not define a package-wide basename namespace. Therefore a
    file named runtime_semantic_plan.json only has execution authority when it is at
    the canonical package-relative path (archive root) or the equivalent path under
    the ZIP's single enclosing wrapper directory.
    """
    name_set = set(names)
    named_resources = sorted(
        name for name in names if Path(name).name == "runtime_semantic_plan.json"
    )

    authoritative: list[tuple[str, str]] = []
    if _RUNTIME_SEMANTIC_PLAN_RELATIVE_PATH in name_set:
        authoritative.append((_RUNTIME_SEMANTIC_PLAN_RELATIVE_PATH, "canonical_package_path"))

    enclosing_root = _single_enclosing_zip_root(names)
    if enclosing_root:
        rooted = f"{enclosing_root}/{_RUNTIME_SEMANTIC_PLAN_RELATIVE_PATH}"
        if rooted in name_set:
            authoritative.append((rooted, "package_root_canonical_path"))

    # A malformed archive containing multiple sibling package roots must not be
    # resolved by ZIP ordering. Report explicit semantic-authority ambiguity.
    competing_roots = _top_level_runtime_plan_package_roots(names)
    if not authoritative and len(competing_roots) > 1:
        paths = [f"{root}/{_RUNTIME_SEMANTIC_PLAN_RELATIVE_PATH}" for root in competing_roots]
        return {
            "path": None,
            "authority": "ambiguous",
            "reason": "ambiguous_authoritative_runtime_plans",
            "paths": paths,
            "ignored": sorted(name for name in named_resources if name not in paths),
        }

    # These two candidates can only coexist in a structurally malformed archive
    # (a root-level canonical plan plus a wrapped canonical package). Never pick by
    # ordering if that situation is encountered.
    unique_authoritative = []
    seen = set()
    for item in authoritative:
        if item[0] not in seen:
            unique_authoritative.append(item)
            seen.add(item[0])
    if len(unique_authoritative) > 1:
        paths = [path for path, _ in unique_authoritative]
        return {
            "path": None,
            "authority": "ambiguous",
            "reason": "ambiguous_authoritative_runtime_plans",
            "paths": paths,
            "ignored": sorted(name for name in named_resources if name not in paths),
        }

    if unique_authoritative:
        path, authority = unique_authoritative[0]
        return {
            "path": path,
            "authority": authority,
            "reason": "authoritative_plan_present",
            "paths": [path],
            "ignored": sorted(name for name in named_resources if name != path),
        }

    return {
        "path": None,
        "authority": None,
        "reason": "not_present",
        "paths": [],
        "ignored": named_resources,
    }


def parse_playbook_package(filename: str, raw: bytes) -> dict[str, Any]:
    original_filename = filename
    original_raw = raw
    input_kind = "zip"
    if filename.lower().endswith((".yaml", ".yml")):
        input_kind = "yaml"
        _, raw = _wrap_yaml_as_source_package(filename, raw)
    elif not filename.lower().endswith(".zip"):
        raise ValueError("Playbook source must be a YAML/YML file or ZIP package.")
    if len(raw) > 60 * 1024 * 1024:
        raise ValueError("Playbook package is larger than the 60 MB local editor limit.")
    try:
        archive = zipfile.ZipFile(io.BytesIO(raw))
    except zipfile.BadZipFile as error:
        raise ValueError("Playbook package is not a valid ZIP archive.") from error
    infos = [info for info in archive.infolist() if not info.is_dir()]
    if len(infos) > 1500:
        raise ValueError("Playbook package contains too many files (limit: 1500).")
    unsafe = [info.filename for info in infos if not _safe_zip_name(info.filename)]
    if unsafe:
        raise ValueError(f"Playbook package contains unsafe paths: {unsafe[0]}")
    total_uncompressed = sum(info.file_size for info in infos)
    if total_uncompressed > 120 * 1024 * 1024:
        raise ValueError("Playbook package expands beyond the 120 MB local editor limit.")

    names = [info.filename for info in infos]
    preferred_names = [
        "playbook/source/program.ordo.yaml", "playbook/source/program.ordo.yml",
        "program.ordo.yaml", "program.ordo.yml",
    ]
    yaml_names = [name for name in names if Path(name).suffix.lower() in {".yaml", ".yml"}]
    ordered = []
    for preferred in preferred_names:
        ordered.extend([name for name in yaml_names if name == preferred or name.endswith("/" + preferred)])
    ordered.extend([name for name in yaml_names if Path(name).name in {"program.ordo.yaml", "program.ordo.yml"} and name not in ordered])
    ordered.extend([name for name in yaml_names if name not in ordered])

    source = None
    source_name = None
    source_raw: bytes | None = None
    yaml_errors = []
    for name in ordered:
        try:
            candidate_raw = archive.read(name)
            candidate = yaml.safe_load(candidate_raw.decode("utf-8-sig"))
        except (UnicodeDecodeError, yaml.YAMLError) as error:
            yaml_errors.append(f"{name}: {error}")
            continue
        if _looks_like_playbook(candidate):
            source, source_name, source_raw = candidate, name, candidate_raw
            break
    if source is None or source_name is None:
        detail = f" Last YAML error: {yaml_errors[-1]}" if yaml_errors else ""
        raise ValueError("Could not locate an Ordo playbook YAML in the ZIP." + detail)

    resources: dict[str, str] = {}
    manifest: list[dict[str, Any]] = []
    total_text = 0
    for info in infos:
        name = info.filename
        ext = Path(name).suffix.lower()
        item = {"path": name, "size": info.file_size, "text": False}
        if ext in _TEXT_RESOURCE_EXTENSIONS and info.file_size <= 2 * 1024 * 1024 and total_text <= 24 * 1024 * 1024:
            try:
                text = archive.read(name).decode("utf-8-sig")
                resources[name] = text
                total_text += len(text.encode("utf-8"))
                item["text"] = True
            except UnicodeDecodeError:
                pass
        manifest.append(item)

    compiled_plan = None
    semantic_plan = None
    compiled_status: dict[str, Any] = {"available": False, "valid": False, "reason": "not_present"}
    semantic_status: dict[str, Any] = {"available": False, "valid": False, "reason": "not_present"}

    # Runtime Semantic Plan authority is package-structural, never basename-based.
    # Nested tests/fixtures/examples may legitimately contain same-named artifacts.
    runtime_plan_resolution = _resolve_runtime_semantic_plan_authority(names)
    if runtime_plan_resolution.get("reason") == "ambiguous_authoritative_runtime_plans":
        paths = runtime_plan_resolution.get("paths") or []
        raise ValueError(
            "Runtime Semantic Plan authority is ambiguous: "
            "ambiguous_authoritative_runtime_plans. paths=" + json.dumps(paths, ensure_ascii=False)
        )
    semantic_names = [runtime_plan_resolution["path"]] if runtime_plan_resolution.get("path") else []
    ignored_runtime_plan_named_resources = list(runtime_plan_resolution.get("ignored") or [])

    preparation_report: dict[str, Any] = {"mode":"precompiled", "stages":[{"id":"load_source","status":"PASS"}]}
    semantic_generated = False
    if not semantic_names:
        # Source-first R3 path: extract the already path-safe package, compile internally,
        # and validate before the package can become executable.
        with tempfile.TemporaryDirectory(prefix="ordo-source-package-") as td:
            root = Path(td)
            archive.extractall(root)
            program_path = root / source_name
            try:
                semantic_plan, preparation_report = _run_integrated_compile(root, program_path)
            except ValueError as error:
                # Release-2 compatibility: old ZIPs that predate graph_contract may still
                # use the historical YAML execution path. New canonical source inputs are
                # fail-closed and must compile successfully before execution.
                legacy_zip = input_kind == "zip" and not isinstance(source.get("graph_contract"), dict)
                if legacy_zip:
                    semantic_plan = None
                    preparation_report = {
                        "mode":"legacy_yaml_compat",
                        "stages":[{"id":"load_source","status":"PASS"},{"id":"compile_runtime_plan","status":"SKIPPED"}],
                        "warning":"Legacy Release-2 source does not declare graph_contract; integrated compilation was not required for compatibility.",
                    }
                else:
                    raise ValueError(str(error)) from error
        if semantic_plan is not None:
            semantic_generated = True
            semantic_status = {
                "available": True, "valid": True, "reason":"integrated_compile_ok",
                "path":"internal://runtime_semantic_plan.json", "generated":True,
                "authority":"integrated_source_compile",
                "source_sha256": hashlib.sha256(source_raw or b"").hexdigest(),
                "format": semantic_plan.get("format"), "format_version": semantic_plan.get("format_version"),
                "structural_status": str((semantic_plan.get("validation") or {}).get("structural_status") or "PASS").upper(),
                "semantic_status": str((semantic_plan.get("validation") or {}).get("semantic_status") or "PASS").upper(),
                "compiler_version": semantic_plan.get("compiler_version"),
            }
    if not semantic_generated and len(semantic_names) == 1:
        plan_name = semantic_names[0]
        semantic_status = {
            "available": True, "valid": False, "reason": "invalid", "path": plan_name,
            "authority": runtime_plan_resolution.get("authority") or "canonical_package_path",
        }
        try:
            candidate = json.loads(archive.read(plan_name).decode("utf-8-sig"))
            actual_sha = hashlib.sha256(source_raw or b"").hexdigest()
            declared_sha = str(((candidate.get("source") or {}).get("sha256") if isinstance(candidate, dict) else "") or "").lower()
            validation = candidate.get("validation") if isinstance(candidate, dict) else None
            structural = str((validation or {}).get("structural_status") or "pending").upper() if isinstance(validation, dict) else "PENDING"
            semantic = str((validation or {}).get("semantic_status") or "pending").upper() if isinstance(validation, dict) else "PENDING"
            if not isinstance(candidate, dict):
                semantic_status["reason"] = "root_must_be_object"
            elif candidate.get("format") != "ordo.runtime_semantic_plan":
                semantic_status.update({"reason": "unsupported_format", "format": candidate.get("format")})
            elif not str(candidate.get("format_version") or "").startswith("1."):
                semantic_status.update({"reason": "unsupported_format_version", "format_version": candidate.get("format_version")})
            elif not isinstance(candidate.get("elements"), dict):
                semantic_status["reason"] = "elements_must_be_object"
            elif structural == "FAIL":
                semantic_status.update({"reason": "structural_validation_failed", "structural_status": structural, "semantic_status": semantic})
            elif any(isinstance(i, dict) and i.get("severity") == "error" and i.get("code") in set((validation or {}).get("blocking_issue_codes") or ["CLASSIFICATION_CONFLICT", "UNKNOWN_ELEMENT_KIND"]) for i in ((validation or {}).get("compilation_issues") or [])):
                semantic_status.update({"reason": "blocking_semantic_execution_issue", "structural_status": structural, "semantic_status": semantic})
            elif not declared_sha:
                semantic_status.update({"reason": "source_sha256_missing", "actual_source_sha256": actual_sha})
            elif declared_sha != actual_sha:
                semantic_status.update({"reason": "source_sha256_mismatch", "declared_source_sha256": declared_sha, "actual_source_sha256": actual_sha})
            else:
                semantic_plan = candidate
                semantic_status.update({
                    "valid": True,
                    "reason": "ok" if semantic != "FAIL" else "usable_with_semantic_issues",
                    "source_sha256": actual_sha,
                    "format": candidate.get("format"),
                    "format_version": candidate.get("format_version"),
                    "structural_status": structural,
                    "semantic_status": semantic,
                    "compilation_issue_count": len((validation or {}).get("compilation_issues") or []) if isinstance(validation, dict) else 0,
                })
        except (UnicodeDecodeError, json.JSONDecodeError) as error:
            semantic_status.update({"reason": "json_parse_error", "detail": str(error)})
    elif not semantic_generated and len(semantic_names) > 1:
        # Kept as a defensive invariant; package-authority resolution should have
        # converted this situation to an explicit ambiguity diagnostic above.
        semantic_status = {
            "available": True, "valid": False, "reason": "ambiguous_authoritative_runtime_plans",
            "paths": semantic_names, "authority": "ambiguous",
        }

    # KF-019 / alpha.20.0.32: a package that contains a Runtime Semantic Plan
    # must never silently downgrade to YAML execution when that plan is stale or invalid.
    # A present-but-invalid compiled semantic artifact means the package is internally
    # inconsistent, so fail at load time before any model call can be made.
    if semantic_status.get("available") and not semantic_status.get("valid"):
        reason = semantic_status.get("reason") or "invalid"
        detail = ""
        if reason == "source_sha256_mismatch":
            detail = f" declared={semantic_status.get('declared_source_sha256')} actual={semantic_status.get('actual_source_sha256')}"
        path = semantic_status.get("path")
        authority = semantic_status.get("authority")
        diagnostic = f" path={path} authority={authority}" if path or authority else ""
        raise ValueError(f"Runtime Semantic Plan is present but unusable: {reason}.{detail}{diagnostic}")

    # Legacy V6 plan remains supported as a separate compatibility path.
    plan_names = [name for name in names if name == "compiled/llm_execution_plan.json" or name.endswith("/compiled/llm_execution_plan.json")]
    if not plan_names:
        plan_names = [name for name in names if Path(name).name == "llm_execution_plan.json"]
    if len(plan_names) == 1:
        plan_name = plan_names[0]
        compiled_status = {"available": True, "valid": False, "reason": "invalid", "path": plan_name}
        try:
            candidate_plan = json.loads(archive.read(plan_name).decode("utf-8-sig"))
            if not isinstance(candidate_plan, dict):
                compiled_status["reason"] = "root_must_be_object"
            elif not isinstance(candidate_plan.get("elements"), dict):
                compiled_status["reason"] = "elements_must_be_object"
            elif candidate_plan.get("format") != "ordo.llm_execution_plan":
                compiled_status["reason"] = "unsupported_format"
                compiled_status["format"] = candidate_plan.get("format")
            elif str(candidate_plan.get("format_version") or "") != "2.0":
                compiled_status["reason"] = "unsupported_format_version"
                compiled_status["format_version"] = candidate_plan.get("format_version")
            elif str(candidate_plan.get("validation_status") or "") not in {"valid", "partial"}:
                compiled_status["reason"] = "compiled_plan_not_executable"
                compiled_status["validation_status"] = candidate_plan.get("validation_status")
            else:
                runtime_contract = candidate_plan.get("runtime_contract")
                expected_contract = {
                    "phase_names": ["enter", "respond"], "prompt_field": "prompt", "state_field": "required_state",
                    "resource_field": "required_resources", "route_field": "allowed_routes", "output_contract_field": "output_contract",
                    "fallback_policy": "yaml_fallback",
                }
                contract_errors = []
                if not isinstance(runtime_contract, dict): contract_errors.append("runtime_contract_missing")
                else:
                    for key, expected in expected_contract.items():
                        if runtime_contract.get(key) != expected: contract_errors.append(f"runtime_contract.{key}_mismatch")
                if contract_errors:
                    compiled_status.update({"reason": "runtime_contract_mismatch", "detail": contract_errors})
                else:
                    actual_sha = hashlib.sha256(source_raw or b"").hexdigest(); declared_sha = str(candidate_plan.get("source_sha256") or "").strip().lower()
                    if not declared_sha: compiled_status.update({"reason": "source_sha256_missing", "actual_source_sha256": actual_sha})
                    elif declared_sha != actual_sha: compiled_status.update({"reason": "source_sha256_mismatch", "declared_source_sha256": declared_sha, "actual_source_sha256": actual_sha})
                    else:
                        compiled_plan = candidate_plan
                        compiled_status.update({"valid": True, "reason": "ok", "source_sha256": actual_sha, "format": candidate_plan.get("format"), "format_version": candidate_plan.get("format_version"), "validation_status": candidate_plan.get("validation_status")})
        except (UnicodeDecodeError, json.JSONDecodeError) as error:
            compiled_status.update({"reason": "json_parse_error", "detail": str(error)})
    elif len(plan_names) > 1:
        compiled_status = {"available": True, "valid": False, "reason": "multiple_execution_plans", "paths": plan_names}

    package_manifest_v2_status = _validate_r3_package_manifest_v2(archive, infos, source_name, source_raw or b"", semantic_names, semantic_plan)

    package_id = hashlib.sha256(original_raw).hexdigest()[:16]
    PLAYBOOK_PACKAGE.update({
        "id": package_id, "filename": original_filename, "source_name": source_name, "source": source,
        "resources": resources, "manifest": manifest, "compiled_plan": compiled_plan,
        "compiled_plan_status": compiled_status, "semantic_plan": semantic_plan, "semantic_plan_status": semantic_status,
        "package_manifest_v2_status": package_manifest_v2_status, "preparation_report": preparation_report, "input_kind": input_kind, "raw_zip": raw,
        "runtime_plan_authority": copy.deepcopy(runtime_plan_resolution),
        "ignored_non_authoritative_runtime_plan_named_resources": ignored_runtime_plan_named_resources,
    })
    PLAYBOOK_PACKAGES[package_id] = copy.deepcopy(PLAYBOOK_PACKAGE)
    return {
        "id": package_id,
        "filename": original_filename,
        "source_name": source_name,
        "entry_node": _entry_id(source),
        "file_count": len(infos),
        "text_resource_count": len(resources),
        "compiled_plan_status": compiled_status,
        "semantic_plan_status": semantic_status,
        "package_manifest_v2_status": package_manifest_v2_status,
        "preparation_report": preparation_report,
        "input_kind": input_kind,
        "runtime_plan_authority": copy.deepcopy(runtime_plan_resolution),
        "ignored_non_authoritative_runtime_plan_named_resources": ignored_runtime_plan_named_resources,
        "interaction_contract": (semantic_plan.get("interaction_contract") if isinstance(semantic_plan, dict) else None),
        "manifest": manifest,
        "source": source,
        "graph": graph_view(source, resources),
    }

def _plausible_resource_path(value: Any) -> bool:
    """Return True only for strings that are safe to treat as filesystem/resource paths.

    This is deliberately conservative. Runtime/package resource resolution must never
    reinterpret arbitrary prompts, prose or command lines as filenames merely because
    they contain a slash.
    """
    if not isinstance(value, str):
        return False
    candidate=value.strip().replace("\\","/")
    if not candidate or len(candidate)>1024:
        return False
    if any(ch in candidate for ch in ("\x00","\n","\r","\t")):
        return False
    if candidate.startswith(("{","[","```","`")):
        return False
    # Shell-like command lines and prose containing an option are not file references.
    if re.search(r"\s--?[A-Za-z0-9]",candidate):
        return False
    try:
        path=Path(candidate)
    except (TypeError,ValueError,OSError):
        return False
    # Prevent platform filename limits from being reached during stat()/resolve().
    if any(len(part.encode("utf-8",errors="ignore"))>240 for part in path.parts):
        return False
    suffix=path.suffix.lower()
    # A bare path segment is accepted only when it is actually path-shaped.
    return "/" in candidate or suffix in _TEXT_RESOURCE_EXTENSIONS


def _string_paths(value: Any) -> list[str]:
    found: list[str] = []
    if isinstance(value, dict):
        for child in value.values():
            found.extend(_string_paths(child))
    elif isinstance(value, list):
        for child in value:
            found.extend(_string_paths(child))
    elif isinstance(value, str):
        candidate=value.strip().replace("\\","/")
        if _plausible_resource_path(candidate):
            found.append(candidate)
    return found

def _context_keywords(record: dict[str, Any]) -> set[str]:
    """Return coarse keywords for selecting only resources relevant to one element."""
    words: set[str] = set()
    def add_text(value: Any) -> None:
        if isinstance(value, str):
            for part in re.split(r"[^A-Za-z0-9_\-]+", value.lower()):
                for token in re.split(r"[_\-]+", part):
                    if len(token) >= 4 and token not in {"node", "gate", "next", "true", "false", "null", "ordo"}:
                        words.add(token)
        elif isinstance(value, dict):
            for key, child in value.items():
                add_text(str(key)); add_text(child)
        elif isinstance(value, list):
            for child in value: add_text(child)
    add_text(record)
    return words


def _package_context_for_record(record: dict[str, Any]) -> dict[str, Any]:
    """Resolve only resources explicitly referenced by the current element.

    The package is a repository, not a prompt.  If the element does not contain
    an explicit path-like reference, the model receives no package resource by
    default.  In particular, there is no fuzzy filename/keyword fallback.
    """
    resources = _active_playbook_package().get("resources") if isinstance(_active_playbook_package().get("resources"), dict) else {}
    refs = _string_paths(record)
    chosen: list[tuple[str, str, str]] = []
    seen: set[str] = set()

    for ref in refs:
        normalized = ref.lstrip("./")
        normalized_name=""
        if _plausible_resource_path(normalized):
            try:
                normalized_name=Path(normalized).name
            except (OSError,ValueError,TypeError):
                normalized_name=""
        matches=[
            name for name in resources
            if name == normalized
            or name.endswith("/"+normalized)
            or (normalized_name and Path(name).name==normalized_name)
        ]
        for name in matches:
            if name not in seen:
                chosen.append((name, resources[name], "explicit"))
                seen.add(name)

        # Generated run-local resources are first-class context too. This lets a
        # later model node summarize a deterministic report without making the
        # playbook copy generated files back into the source package.
        if _plausible_resource_path(normalized):
            try:
                rel=Path(normalized)
                runtime_file=None
                if not rel.is_absolute() and ".." not in rel.parts:
                    runtime_root=_runtime_workspace().resolve()
                    candidate=(_runtime_workspace()/rel).resolve()
                    try:
                        candidate.relative_to(runtime_root)
                        runtime_file=candidate
                    except ValueError:
                        runtime_file=None
                if runtime_file is not None and runtime_file.is_file() and normalized not in seen:
                    try:
                        runtime_text=runtime_file.read_text(encoding="utf-8")
                    except (UnicodeDecodeError,OSError):
                        runtime_text=""
                    if runtime_text:
                        chosen.append((normalized,runtime_text,"runtime_artifact"))
                        seen.add(normalized)
            except (OSError,ValueError,TypeError):
                # Resource discovery is best-effort context enrichment. A malformed,
                # oversized or platform-invalid path must never abort playbook execution.
                pass

    rendered = []
    # Explicit resources may still be large, so keep a hard context budget.
    budget = 18000
    used = 0
    for name, text, reason in chosen:
        chunk = text[:7000]
        if used + len(chunk) > budget:
            remaining = budget - used
            if remaining < 500:
                break
            chunk = chunk[:remaining]
        rendered.append({"path": name, "reason": reason, "content": chunk})
        used += len(chunk)
        if used >= budget:
            break
    return {"resolved_resources": rendered}


def _execution_record_for_model(record: dict[str, Any]) -> dict[str, Any]:
    """Remove graph/editor bookkeeping that does not help execute this element."""
    compact = copy.deepcopy(record)
    compact.pop("incoming_from", None)
    compact.pop("allowed_from", None)
    return compact


def _project_execution_record(record: dict[str, Any], kind: str, phase: str, human_policy: dict[str, Any]) -> dict[str, Any]:
    """Project an Ordo element to only the fields useful for this exact LLM phase.

    Runtime-owned response validation/routing fields are deliberately omitted on enter.
    This keeps the model focused on the work it must perform rather than the mechanics
    the Python runtime already enforces.
    """
    projected = _execution_record_for_model(record)
    if kind == "node":
        if phase == "enter":
            projected.pop("on_unmatched_input", None)
            projected.pop("on_answer", None)
            # answer_type is a runtime validation concern.  The question/proposal
            # instructions already tell the model what to present to the analyst.
            projected.pop("answer_type", None)
            projected.pop("allowed_values", None)
            # If proposal rules already name every expected field, the explicit
            # expected_fields list is duplicate prompt material.
            proposal = projected.get("proposal_generation")
            expected = projected.get("expected_fields")
            if isinstance(proposal, dict) and isinstance(proposal.get("rules"), dict) and isinstance(expected, list):
                if set(str(x) for x in expected) == set(str(x) for x in proposal["rules"].keys()):
                    projected.pop("expected_fields", None)
            # A question/instruction already carries the semantic purpose; avoid
            # repeating a human-facing title unless it is the only description.
            if projected.get("question"):
                projected.pop("title", None)
        elif phase == "respond":
            # Proposal generation has already happened; response-time work concerns
            # interpreting the analyst answer and applying the declared response rules.
            for key in ("proposal_generation", "draft_generation", "generation", "ai_prepare", "pre_answer_analysis"):
                projected.pop(key, None)
    return projected


def _state_reference_roots(value: Any) -> set[str]:
    """Collect top-level runtime-state keys explicitly referenced by an element."""
    roots: set[str] = set()
    def visit(item: Any) -> None:
        if isinstance(item, dict):
            for child in item.values(): visit(child)
        elif isinstance(item, list):
            for child in item: visit(child)
        elif isinstance(item, str):
            for match in re.finditer(r"(?:\bstate|\$state)\.([A-Za-z0-9_]+)", item):
                roots.add(match.group(1))
    visit(value)
    return roots


def _state_reference_paths(value: Any, state: dict[str, Any] | None = None) -> set[str]:
    """Collect PRE-STATE dependencies without mistaking output targets for inputs.

    Explicit ``state.foo``/``$state.foo`` references are authoritative everywhere.
    Bare canonical state names are considered only inside input-bearing declarations
    (``from``, ``inputs``, ``reads``...) and gate ``condition`` text.  This mirrors the
    compiler dependency model and prevents retry cycles from turning an already-existing
    output collection into a false input merely because its name appears in
    ``normalize``/``update_state``.
    """
    paths: set[str] = set()
    input_keys = {
        "source", "source_priority", "input", "inputs", "reads", "required_inputs",
        "state_inputs", "transform", "preconditions", "entry_preconditions",
        "required_preconditions", "depends_on", "from", "using", "based_on",
        "context_from", "generated_from", "derive_from", "condition",
    }

    def add_explicit(text: str) -> None:
        for match in re.finditer(r"(?:\bstate|\$state)\.([A-Za-z0-9_.]+)", text):
            paths.add(match.group(1).rstrip(".,;:)"))

    def add_bare(text: str) -> None:
        if not state:
            return
        candidates=set(str(k) for k in state)
        candidates.update(str(k).split(".")[0] for k in state)
        for candidate in sorted(candidates, key=len, reverse=True):
            if candidate and re.search(r"(?<![A-Za-z0-9_])"+re.escape(candidate)+r"(?![A-Za-z0-9_])", text):
                paths.add(candidate)

    def collect_input_value(item: Any) -> None:
        if isinstance(item, str):
            clean=item.removeprefix("state.").removeprefix("$state.")
            if state and _state_path_exists(state, clean):
                paths.add(clean)
            add_explicit(item)
            add_bare(item)
        elif isinstance(item, list):
            for child in item:
                collect_input_value(child)
        elif isinstance(item, dict):
            for child in item.values():
                collect_input_value(child)

    def visit(item: Any) -> None:
        if isinstance(item, dict):
            for key, child in item.items():
                if str(key) in input_keys:
                    collect_input_value(child)
                else:
                    visit(child)
        elif isinstance(item, list):
            for child in item:
                visit(child)
        elif isinstance(item, str):
            add_explicit(item)

    visit(value)
    return {p for p in paths if p}


def _project_runtime_state(state: Any, projected_record: dict[str, Any], kind: str, phase: str) -> Any:
    """Send only state fields explicitly referenced by this exact element/phase."""
    if not isinstance(state, dict):
        return _bounded_json_value(state, 24000)
    paths=_state_reference_paths(projected_record,state)
    if not paths and kind == "gate":
        gate_id=str(projected_record.get("id") or "").lower()
        roots=sorted({str(key).split(".")[0] for key in state}, key=len, reverse=True)
        matches=[root for root in roots if root.lower() in gate_id]
        if matches:
            longest=len(matches[0])
            paths.update(root for root in matches if len(root)==longest)
    if not paths:
        return {}
    projected: dict[str, Any]={}
    for path in sorted(paths):
        value=_state_subtree(state,path)
        if value is not None:
            projected[path]=value
    return _bounded_json_value(projected,24000)


def _project_live_history(history: Any, projected_record: dict[str, Any]) -> list[dict[str, Any]]:
    """Conversation history is opt-in. Structured runtime state is the source of truth."""
    encoded = json.dumps(projected_record, ensure_ascii=False, default=str).lower()
    markers = ("recent_history", "conversation_history", "dialog_history", "previous_message", "prior_message")
    if any(marker in encoded for marker in markers):
        return _compact_live_history(history)
    return []


def _bounded_json_value(value: Any, max_chars: int) -> Any:
    """Keep runtime memory useful while preventing accidental prompt explosions."""
    try:
        encoded = json.dumps(value, ensure_ascii=False)
    except Exception:
        return value
    if len(encoded) <= max_chars:
        return value
    if isinstance(value, dict):
        result: dict[str, Any] = {}
        used = 2
        # Prefer recent insertion-order keys; Python dicts preserve runtime update order.
        for key, child in reversed(list(value.items())):
            child_text = json.dumps(child, ensure_ascii=False, default=str)
            cost = len(str(key)) + len(child_text) + 6
            if used + cost > max_chars: continue
            result[key] = child; used += cost
        return {key: result[key] for key in reversed(list(result))}
    if isinstance(value, list):
        out=[]; used=2
        for child in reversed(value):
            text=json.dumps(child, ensure_ascii=False, default=str)
            if used + len(text) + 1 > max_chars: break
            out.append(child); used += len(text)+1
        return list(reversed(out))
    return str(value)[:max_chars]


def _compact_live_history(history: Any) -> list[dict[str, Any]]:
    if not isinstance(history, list): return []
    compact: list[dict[str, Any]] = []
    for item in history[-8:]:
        if not isinstance(item, dict): continue
        text = str(item.get("text") or "")
        compact.append({
            "role": str(item.get("role") or ""),
            "node_id": item.get("node_id"),
            "text": text[:1800],
        })
    return compact


def _entry_id(source: dict[str, Any]) -> str | None:
    for container_key in ("graph_contract", "playbook"):
        container = source.get(container_key)
        if isinstance(container, dict) and isinstance(container.get("entry_node"), str):
            return container["entry_node"]
    for key in ("entry_node", "entry", "start_node"):
        if isinstance(source.get(key), str):
            return source[key]
    nodes = source.get("nodes")
    if isinstance(nodes, list):
        for node in nodes:
            if isinstance(node, dict) and isinstance(node.get("id"), str):
                return node["id"]
    return None


def _record_by_id(source: dict[str, Any], record_id: str) -> tuple[dict[str, Any] | None, str | None]:
    for collection, kind in (("nodes", "node"), ("gates", "gate")):
        records = source.get(collection)
        if isinstance(records, list):
            for record in records:
                if isinstance(record, dict) and record.get("id") == record_id:
                    return record, kind
    return None, None


def _live_routes(record: dict[str, Any], kind: str) -> list[dict[str, str]]:
    source = _active_playbook_package().get("source") if isinstance(_active_playbook_package().get("source"), dict) else {}
    known = {str(item.get("id")) for key in ("nodes", "gates") for item in (source.get(key) or []) if isinstance(item, dict) and item.get("id")}
    known.update(str(x) for x in ((source.get("graph_contract") or {}).get("external_terminal_targets") or []) if isinstance(x, str))
    routes: list[dict[str, str]] = []
    seen_targets: set[str] = set()
    for raw in _shared_declared_routes(record, is_gate=(kind == "gate")):
        item = {"key": str(raw["key"]), "target": str(raw["target"])}
        if item not in routes:
            routes.append(item); seen_targets.add(item["target"])
    if known:
        for raw in _shared_generic_routes(record, known):
            target = str(raw["target"])
            if target in seen_targets:
                continue
            item = {"key": str(raw["key"]), "target": target}
            if item not in routes:
                routes.append(item); seen_targets.add(target)
    return routes

def _extract_response_text(response: dict[str, Any]) -> str:
    text = response.get("output_text")
    if isinstance(text, str) and text.strip(): return text.strip()
    for output in response.get("output", []) if isinstance(response.get("output"), list) else []:
        if not isinstance(output, dict): continue
        for content in output.get("content", []) if isinstance(output.get("content"), list) else []:
            if isinstance(content, dict):
                value = content.get("text")
                if isinstance(value, str) and value.strip(): return value.strip()
    return ""


def _normalize_base_url(value: str | None, default: str) -> str:
    url = str(value or default).strip().rstrip("/")
    if not url.startswith(("http://", "https://")):
        raise ValueError("Provider base URL must start with http:// or https://")
    return url


def _provider_models(provider: str, base_url: str, api_key: str | None = None) -> list[str]:
    if provider == "openai" and base_url.rstrip("/") == "https://api.openai.com/v1":
        return list(OPENAI_MODELS)
    url = _normalize_base_url(base_url, DEFAULT_MLX_BASE_URL) + "/models"
    headers = {"Accept": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    request = urllib.request.Request(url, headers=headers, method="GET")
    try:
        with urllib.request.urlopen(request, timeout=8) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        detail = error.read().decode("utf-8", errors="replace")
        raise ValueError(f"Provider models API error {error.code}: {detail[:800]}") from error
    except urllib.error.URLError as error:
        label = "Local MLX-LM API" if provider == "mlx" else "OpenAI-compatible API"
        raise ValueError(f"{label} unavailable at {base_url}. {error.reason}") from error
    models = []
    for item in data.get("data", []) if isinstance(data, dict) else []:
        if isinstance(item, dict) and isinstance(item.get("id"), str) and item["id"].strip():
            models.append(item["id"].strip())
    if not models:
        raise ValueError(f"Provider at {base_url} returned no models from /models.")
    return models


def _live_credentials(payload: dict[str, Any]) -> dict[str, Any]:
    session_id=str(payload.get("session_id") or "").strip()
    session=LIVE_SESSIONS.get(session_id,{}) if session_id else {}
    provider=str(session.get("provider") or LIVE_RUNTIME.get("provider") or "openai").strip().lower()
    if provider not in PROVIDERS:
        raise ValueError(f"Unsupported LLM provider: {provider}")
    model=str(session.get("model") or LIVE_RUNTIME.get("model") or "").strip()
    structured_output_mode=str(session.get("structured_output_mode") or LIVE_RUNTIME.get("structured_output_mode") or "auto").strip().lower()
    if structured_output_mode not in STRUCTURED_OUTPUT_MODES:
        raise ValueError(f"Unsupported structured_output_mode: {structured_output_mode}")

    if provider=="openai":
        base_url="https://api.openai.com/v1"
        api_key=LIVE_RUNTIME.get("api_key") or session.get("api_key")
        if not api_key:
            raise ValueError("Live execution needs an OpenAI API key. Configure a personal key in Run or start the server with a shared key.")
        if not model:
            raise ValueError("Live execution needs an OpenAI model selected in Run.")
        api_style="responses"
    elif provider=="mlx":
        base_url=_normalize_base_url(session.get("base_url"),DEFAULT_MLX_BASE_URL)
        api_key=str(session.get("api_key") or "local")
        if not model:
            raise ValueError("Select a model reported by the Local MLX /models endpoint.")
        available=_provider_models(provider,base_url,api_key)
        if model not in available:
            raise ValueError(f"Local MLX model {model!r} is not currently available. Refresh models and choose one of: {', '.join(available)}")
        api_style="chat_completions"
    else:
        base_url=_normalize_base_url(session.get("base_url"),DEFAULT_CUSTOM_BASE_URL)
        api_key=str(session.get("api_key") or "")
        if not model:
            raise ValueError("Select a model reported by the custom provider /models endpoint.")
        api_style="chat_completions"

    capability_profile=copy.deepcopy(session.get("capability_profile"))
    if not isinstance(capability_profile,dict):
        capability_profile=_cached_provider_capability(provider,base_url,model,api_style)

    return {
        "provider":provider,
        "api_key":api_key,
        "model":model,
        "base_url":base_url,
        "api_style":api_style,
        "structured_output_mode":structured_output_mode,
        "capability_profile":capability_profile,
    }


def _text_from_provider_value(value: Any) -> str:
    """Best-effort normalization of OpenAI-compatible text payload shapes."""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        parts: list[str] = []
        for item in value:
            text = _text_from_provider_value(item)
            if text:
                parts.append(text)
        return "\n".join(parts).strip()
    if isinstance(value, dict):
        for key in ("text", "content", "value", "output_text", "answer", "message"):
            if key in value:
                text = _text_from_provider_value(value.get(key))
                if text:
                    return text
        return ""
    return ""


def _extract_chat_response_text(response: dict[str, Any]) -> str:
    if not isinstance(response, dict):
        return ""
    choices = response.get("choices")
    if isinstance(choices, list):
        for choice in choices:
            if not isinstance(choice, dict):
                continue
            message = choice.get("message")
            if isinstance(message, dict):
                text = _text_from_provider_value(message.get("content"))
                if text:
                    return text
                for key in ("text", "answer", "final", "output_text"):
                    text = _text_from_provider_value(message.get(key))
                    if text:
                        return text
            for key in ("text", "content", "answer", "output_text"):
                text = _text_from_provider_value(choice.get(key))
                if text:
                    return text
    for key in ("output_text", "text", "content", "answer", "message", "response"):
        text = _text_from_provider_value(response.get(key))
        if text:
            return text
    return ""


def _parse_model_json(raw_text: str) -> dict[str, Any]:
    text = raw_text.strip()

    candidates: list[str] = [text]

    # OpenAI-compatible local models commonly wrap structured output in Markdown fences.
    fenced = re.search(r"```(?:json)?\s*(.*?)\s*```", text, flags=re.I | re.S)
    if fenced:
        candidates.insert(0, fenced.group(1).strip())

    # Also accept prose around one JSON object, but still fail closed if no valid object exists.
    start = text.find("{")
    if start >= 0:
        depth = 0
        in_string = False
        escaped = False
        for idx in range(start, len(text)):
            ch = text[idx]
            if in_string:
                if escaped:
                    escaped = False
                elif ch == "\\":
                    escaped = True
                elif ch == '"':
                    in_string = False
                continue
            if ch == '"':
                in_string = True
            elif ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    candidates.append(text[start:idx + 1])
                    break

    last_error: Exception | None = None
    for candidate in candidates:
        try:
            result = json.loads(candidate)
        except json.JSONDecodeError as error:
            last_error = error
            continue
        if not isinstance(result, dict):
            raise ValueError("Model result must be a JSON object.")
        return result

    raise ValueError(f"Model did not return valid JSON: {raw_text[:800]}") from last_error


def _extract_structured_chat_answer(parsed: Any) -> str:
    """Extract the user-facing assistant text from common JSON chat envelopes."""
    if not isinstance(parsed, dict):
        return ""
    for key in ("answer_markdown", "answer", "message", "content", "text", "response", "output_text", "final_response"):
        text=_text_from_provider_value(parsed.get(key))
        if text:
            return text
    nested=parsed.get("final")
    if isinstance(nested, dict):
        for key in ("answer_markdown", "answer", "message", "content", "text", "output_text"):
            text=_text_from_provider_value(nested.get(key))
            if text:
                return text
    for key in ("conversation", "messages", "history"):
        items=parsed.get(key)
        if not isinstance(items, list):
            continue
        for item in reversed(items):
            if not isinstance(item, dict):
                continue
            role=str(item.get("role") or "").strip().lower()
            if role not in {"assistant", "model", "ai"}:
                continue
            for text_key in ("content", "text", "message", "answer", "output_text"):
                text=_text_from_provider_value(item.get(text_key))
                if text:
                    return text
        for item in reversed(items):
            if not isinstance(item, dict):
                continue
            for text_key in ("content", "text", "message", "answer", "output_text"):
                text=_text_from_provider_value(item.get(text_key))
                if text:
                    return text
    return ""


STRICT_SCHEMA_ALLOWED_KEYWORDS = {
    "type", "properties", "required", "additionalProperties", "items",
    "enum", "const", "anyOf", "description", "minItems", "maxItems",
    "minLength", "maxLength", "minimum", "maximum",
}
STRICT_PROVIDER_PROFILES = {
    ("openai", "responses"),
    ("openai", "chat_completions"),
}
STRUCTURED_OUTPUT_MODES = {"auto", "strict_json_schema", "json_object", "plain"}


def _runtime_debug_log(event: str, payload: Any) -> None:
    blocked={"api_key","authorization","password","secret","token","headers"}
    def scrub(value: Any) -> Any:
        if isinstance(value,dict):
            return {str(k):("***REDACTED***" if str(k).lower() in blocked else scrub(v)) for k,v in value.items()}
        if isinstance(value,(list,tuple)):
            return [scrub(v) for v in value]
        if isinstance(value,str):
            return value if len(value)<=12000 else value[:12000]+f"...<truncated {len(value)-12000} chars>"
        if value is None or isinstance(value,(int,float,bool)):
            return value
        return str(value)
    try:
        encoded=json.dumps(scrub(payload),ensure_ascii=False,sort_keys=True)
    except Exception as exc:
        encoded=json.dumps({"log_encoding_error":str(exc),"payload_repr":repr(payload)[:4000]},ensure_ascii=False)
    print(f"[ORDO_RUNTIME_DEBUG] {event} {encoded}",flush=True)


def _provider_structured_output_mode(credentials: dict[str, Any]) -> tuple[str, str]:
    """Resolve provider structured-output strategy from explicit mode and recorded capability evidence.

    In auto mode, a recorded provider capability probe is authoritative for the active
    provider/model session. This closes the gap where the probe result was stored in
    credentials but ignored during actual runtime calls.
    """
    requested=str(credentials.get("structured_output_mode") or "auto").strip().lower()
    if requested not in STRUCTURED_OUTPUT_MODES:
        raise ValueError(f"Unsupported structured_output_mode: {requested}")
    if requested != "auto":
        return requested, "explicit_capability_profile"

    profile=credentials.get("capability_profile")
    if isinstance(profile,dict) and isinstance(profile.get("supports_json_schema"),bool):
        # Bind capability evidence to the same endpoint/model when those fields are recorded.
        same_model=not profile.get("model") or str(profile.get("model"))==str(credentials.get("model") or "")
        same_base=not profile.get("base_url") or str(profile.get("base_url")).rstrip("/")==str(credentials.get("base_url") or "").rstrip("/")
        same_style=not profile.get("api_style") or str(profile.get("api_style"))==str(credentials.get("api_style") or "")
        if same_model and same_base and same_style:
            if profile.get("supports_json_schema"):
                return "strict_json_schema", "recorded_capability_probe"
            return "json_object", "recorded_capability_probe"

    provider=str(credentials.get("provider") or "").strip().lower()
    api_style=str(credentials.get("api_style") or "").strip().lower()
    if (provider,api_style) in STRICT_PROVIDER_PROFILES:
        return "strict_json_schema", "built_in_provider_profile"
    return "json_object", "compatibility_default_until_probe"

def _runtime_strict_schema_compatible(schema: Any, credentials: dict[str, Any]) -> tuple[bool, str]:
    """Fail-closed strict JSON-schema compatibility for the active provider profile.

    Unknown providers never get a speculative strict request. This avoids unnecessary
    provider 400s and preserves runtime validation as the final enforcement layer.
    """
    api_style = str(credentials.get("api_style") or "").strip().lower()
    provider = str(credentials.get("provider") or "").strip().lower()
    mode,mode_source=_provider_structured_output_mode(credentials)
    if mode != "strict_json_schema":
        return False, f"strict schema disabled by structured-output mode {mode} ({mode_source}) for {provider}/{api_style}"
    max_schema_chars = credentials.get("max_json_schema_chars")
    if max_schema_chars not in (None, ""):
        try:
            max_schema_chars = int(max_schema_chars)
        except (TypeError, ValueError) as exc:
            raise ValueError(f"Invalid max_json_schema_chars capability: {max_schema_chars}") from exc
        if max_schema_chars <= 0:
            raise ValueError(f"Invalid max_json_schema_chars capability: {max_schema_chars}")
        schema_chars = len(json.dumps(schema, ensure_ascii=False, sort_keys=True, separators=(",", ":")))
        if schema_chars > max_schema_chars:
            return False, f"schema exceeds provider capability limit: {schema_chars}>{max_schema_chars} chars for {provider}/{api_style}"

    def walk(node: Any) -> bool:
        if not isinstance(node, dict):
            return True
        if any(k not in STRICT_SCHEMA_ALLOWED_KEYWORDS for k in node):
            return False
        if node.get("type") == "object" and isinstance(node.get("properties"), dict):
            props=set(node.get("properties") or {})
            required=set(node.get("required") or [])
            if node.get("additionalProperties") is not False or props != required:
                return False
        props=node.get("properties")
        if isinstance(props,dict) and not all(walk(v) for v in props.values()):
            return False
        items=node.get("items")
        if isinstance(items,dict) and not walk(items):
            return False
        any_of=node.get("anyOf")
        if isinstance(any_of,list) and not all(walk(v) for v in any_of):
            return False
        return True

    ok=walk(schema)
    return ok, ("compatible" if ok else f"schema incompatible with strict profile {provider}/{api_style}")

def _probe_provider_json_schema_capability(credentials: dict[str, Any]) -> dict[str, Any]:
    """Perform one minimal strict-schema request without silent downgrade.

    The result is evidence, not inference. Unsupported means the provider explicitly
    rejected the strict request (HTTP 400/404/415/422); transport failures remain errors.
    """
    schema={
        "type":"object","additionalProperties":False,"required":["probe"],
        "properties":{"probe":{"type":"string"}},
    }
    headers={"Content-Type":"application/json"}
    if credentials.get("api_key"): headers["Authorization"]=f"Bearer {credentials['api_key']}"
    if credentials.get("api_style")=="responses":
        endpoint="/responses"
        body={"model":credentials["model"],"input":[{"role":"system","content":"Return the JSON object required by the schema."},{"role":"user","content":"Capability probe. Set probe to ok."}],"text":{"format":{"type":"json_schema","name":"ordo_capability_probe","schema":schema,"strict":True}}}
    else:
        endpoint="/chat/completions"
        body={"model":credentials["model"],"messages":[{"role":"system","content":"Return the JSON object required by the schema."},{"role":"user","content":"Capability probe. Set probe to ok."}],"temperature":0,"max_tokens":64,"stream":False,"response_format":{"type":"json_schema","json_schema":{"name":"ordo_capability_probe","strict":True,"schema":schema}}}
    url=credentials["base_url"].rstrip("/")+endpoint
    request=urllib.request.Request(url,data=json.dumps(body).encode("utf-8"),headers=headers,method="POST")
    try:
        with urllib.request.urlopen(request,timeout=30) as response:
            data=json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        detail=error.read().decode("utf-8",errors="replace")
        if error.code in {400,404,415,422}:
            return {"status":"recorded","supports_json_schema":False,"provider":credentials.get("provider"),"api_style":credentials.get("api_style"),"model":credentials.get("model"),"base_url":credentials.get("base_url"),"http_status":error.code,"evidence":"provider_explicit_rejection","detail":detail[:500]}
        raise ValueError(f"Provider capability probe failed with HTTP {error.code}: {detail[:800]}") from error
    except urllib.error.URLError as error:
        raise ValueError(f"Provider capability probe connection error: {error.reason}") from error
    raw=_extract_response_text(data) if credentials.get("api_style")=="responses" else _extract_chat_response_text(data)
    parsed=_parse_model_json(raw)
    if not isinstance(parsed,dict) or not isinstance(parsed.get("probe"),str):
        raise ValueError("Provider accepted strict schema request but returned a non-conforming probe result")
    return {"status":"recorded","supports_json_schema":True,"provider":credentials.get("provider"),"api_style":credentials.get("api_style"),"model":credentials.get("model"),"base_url":credentials.get("base_url"),"http_status":200,"evidence":"strict_schema_request_accepted"}


def _context_limit_retry_tokens(detail: str, requested_tokens: int) -> int | None:
    """Parse common OpenAI-compatible context-window errors and return a safe retry budget."""
    text=str(detail or "")
    patterns=[
        r"maximum context length is\s*(\d+)\s*tokens.*?requested\s*(\d+)\s*output tokens.*?contains at least\s*(\d+)\s*input tokens",
        r"maximum context length is\s*(\d+)\s*tokens.*?prompt contains at least\s*(\d+)\s*input tokens",
    ]
    max_ctx=None
    input_tokens=None
    for index,pattern in enumerate(patterns):
        match=re.search(pattern,text,re.I|re.S)
        if not match:
            continue
        max_ctx=int(match.group(1))
        input_tokens=int(match.group(3) if index==0 else match.group(2))
        break
    if not max_ctx or input_tokens is None:
        return None

    # Keep a meaningful guard because local compatible servers often report
    # approximate prompt token counts rather than exact tokenizer results.
    available=max_ctx-input_tokens-256
    if available < 256:
        return None
    return max(256,min(int(requested_tokens or 0),available))


def _provider_api_call(credentials: dict[str, Any], system_text: str, context: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any], str, dict[str, int]]:
    request_context = copy.deepcopy(context)
    response_schema = request_context.pop("__response_json_schema", None)
    structured_mode,structured_mode_source=_provider_structured_output_mode(credentials)
    strict_compatible, strict_reason = _runtime_strict_schema_compatible(response_schema, credentials) if isinstance(response_schema, dict) else (False, "no schema")
    structured_debug={
        "requested_mode":str(credentials.get("structured_output_mode") or "auto"),
        "resolved_mode":structured_mode,
        "mode_source":structured_mode_source,
        "strict_compatible":strict_compatible,
        "strict_reason":strict_reason,
    }
    if credentials["api_style"] == "responses":
        text_format = {"type": "json_object"}
        if structured_mode == "plain":
            text_format = None
        elif isinstance(response_schema, dict) and strict_compatible:
            text_format = {"type": "json_schema", "name": "ordo_result", "schema": response_schema, "strict": True}
        req_body = {
            "model": credentials["model"],
            "input": [
                {"role": "system", "content": system_text},
                {"role": "user", "content": json.dumps(request_context, ensure_ascii=False)},
            ],
        }
        if text_format is not None:
            req_body["text"] = {"format": text_format}
        endpoint = "/responses"
    else:
        req_body = {
            "model": credentials["model"],
            "messages": [
                {"role": "system", "content": system_text},
                {"role": "user", "content": json.dumps(request_context, ensure_ascii=False)},
            ],
            "temperature": 0.2,
            "max_tokens": int(request_context.pop("__max_output_tokens", 12000 if str(request_context.get("element_kind") or "") == "model_gate" else 8000)),
            "stream": False,
        }
        if isinstance(response_schema, dict) and strict_compatible:
            req_body["response_format"] = {"type": "json_schema", "json_schema": {"name": "ordo_result", "strict": True, "schema": response_schema}}
        elif structured_mode == "json_object" or structured_mode == "strict_json_schema":
            req_body["response_format"] = {"type": "json_object"}
        # plain mode intentionally omits provider-side structured-output controls.
        endpoint = "/chat/completions"
    headers = {"Content-Type": "application/json"}
    if credentials.get("api_key"):
        headers["Authorization"] = f"Bearer {credentials['api_key']}"
    url = credentials["base_url"].rstrip("/") + endpoint
    _runtime_debug_log("provider.request",{
        "provider":credentials.get("provider"),"model":credentials.get("model"),
        "api_style":credentials.get("api_style"),"endpoint":endpoint,
        "structured_output":structured_debug,"request_body":req_body,
    })
    request = urllib.request.Request(url, data=json.dumps(req_body).encode("utf-8"), headers=headers, method="POST")
    try:
        with urllib.request.urlopen(request, timeout=180) as response:
            api_response = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        detail = error.read().decode("utf-8", errors="replace")
        if error.code == 400 and credentials["api_style"] != "responses":
            current_tokens=int(req_body.get("max_tokens") or 0)
            retry_tokens=_context_limit_retry_tokens(detail,current_tokens)
            if retry_tokens is not None and retry_tokens < current_tokens:
                retry_body=copy.deepcopy(req_body)
                retry_body["max_tokens"]=retry_tokens
                retry_request=urllib.request.Request(url,data=json.dumps(retry_body).encode("utf-8"),headers=headers,method="POST")
                try:
                    with urllib.request.urlopen(retry_request,timeout=180) as response:
                        api_response=json.loads(response.read().decode("utf-8"))
                    raw_text=_extract_chat_response_text(api_response)
                    usage=api_response.get("usage") if isinstance(api_response.get("usage"),dict) else {}
                    usage=dict(usage)
                    usage["ordo_requested_output_tokens"]=current_tokens
                    usage["ordo_retry_output_tokens"]=retry_tokens
                    usage["ordo_context_budget_retry"]=1
                    return retry_body,api_response,raw_text,usage
                except urllib.error.HTTPError as retry_error:
                    detail=retry_error.read().decode("utf-8",errors="replace")
                    error=retry_error
        if error.code == 400 and isinstance(response_schema, dict) and strict_compatible:
            # Some OpenAI-compatible local servers implement json_object but not json_schema.
            fallback_body = copy.deepcopy(req_body)
            if credentials["api_style"] == "responses":
                fallback_body["text"] = {"format": {"type": "json_object"}}
            else:
                fallback_body["response_format"] = {"type": "json_object"}
            fallback_request = urllib.request.Request(url, data=json.dumps(fallback_body).encode("utf-8"), headers=headers, method="POST")
            try:
                with urllib.request.urlopen(fallback_request, timeout=180) as response:
                    api_response = json.loads(response.read().decode("utf-8"))
                req_body = fallback_body
                if isinstance(api_response, dict):
                    api_response["_ordo_debug"] = {
                        "response_format_downgraded": True,
                        "reason": detail[:500],
                        "fallback_request_payload": fallback_body,
                    }
            except Exception as fallback_error:
                raise ValueError(f"{credentials['provider']} API rejected strict JSON schema and fallback failed: {detail[:800]}") from fallback_error
        else:
            raise ValueError(f"{credentials['provider']} API error {error.code}: {detail[:1200]}") from error
    except urllib.error.URLError as error:
        if credentials["provider"] == "mlx":
            raise ValueError(f"Local MLX-LM API unavailable at {credentials['base_url']}. Check that mlx_lm.server is running. ({error.reason})") from error
        raise ValueError(f"{credentials['provider']} API connection error: {error.reason}") from error
    if isinstance(api_response, dict):
        dbg=api_response.setdefault("_ordo_debug", {})
        if isinstance(dbg, dict):
            dbg.setdefault("strict_compatibility_profile", f"{credentials.get('provider')}/{credentials.get('api_style')}")
            dbg.setdefault("structured_output_mode", structured_mode)
            dbg.setdefault("structured_output_mode_source", structured_mode_source)
            dbg.setdefault("strict_compatible", bool(strict_compatible))
            dbg.setdefault("strict_fallback_reason", None if strict_compatible else strict_reason)
    _runtime_debug_log("provider.response",{
        "provider":credentials.get("provider"),"model":credentials.get("model"),
        "structured_output":structured_debug,"api_response":api_response,
    })
    if isinstance(api_response,dict):
        dbg=api_response.get("_ordo_debug") if isinstance(api_response.get("_ordo_debug"),dict) else {}
        dbg["structured_output"]=structured_debug
        api_response["_ordo_debug"]=dbg
    raw_text = _extract_response_text(api_response) if credentials["api_style"] == "responses" else _extract_chat_response_text(api_response)
    raw_usage = api_response.get("usage") if isinstance(api_response, dict) and isinstance(api_response.get("usage"), dict) else {}
    if credentials["api_style"] == "responses":
        input_details = raw_usage.get("input_tokens_details") if isinstance(raw_usage.get("input_tokens_details"), dict) else {}
        output_details = raw_usage.get("output_tokens_details") if isinstance(raw_usage.get("output_tokens_details"), dict) else {}
        usage = {
            "input_tokens": int(raw_usage.get("input_tokens") or 0),
            "output_tokens": int(raw_usage.get("output_tokens") or 0),
            "total_tokens": int(raw_usage.get("total_tokens") or 0),
            "cached_tokens": int(input_details.get("cached_tokens") or 0),
            "reasoning_tokens": int(output_details.get("reasoning_tokens") or 0),
        }
    else:
        prompt_details = raw_usage.get("prompt_tokens_details") if isinstance(raw_usage.get("prompt_tokens_details"), dict) else {}
        completion_details = raw_usage.get("completion_tokens_details") if isinstance(raw_usage.get("completion_tokens_details"), dict) else {}
        usage = {
            "input_tokens": int(raw_usage.get("prompt_tokens") or raw_usage.get("input_tokens") or 0),
            "output_tokens": int(raw_usage.get("completion_tokens") or raw_usage.get("output_tokens") or 0),
            "total_tokens": int(raw_usage.get("total_tokens") or 0),
            "cached_tokens": int(prompt_details.get("cached_tokens") or 0),
            "reasoning_tokens": int(completion_details.get("reasoning_tokens") or 0),
        }
    if not usage["total_tokens"]:
        usage["total_tokens"] = usage["input_tokens"] + usage["output_tokens"]
    return req_body, api_response, raw_text, usage




def _effective_file_ref_answer(record: dict[str, Any], kind: str, phase: str, analyst_input: str, attachment_debug: list[dict[str, Any]]) -> str:
    if phase != "respond" or kind != "node" or str(record.get("answer_type") or "").strip().lower() != "file_ref" or analyst_input.strip():
        return analyst_input
    if len(attachment_debug) == 1 and attachment_debug[0].get("stored_path"):
        return str(attachment_debug[0]["stored_path"])
    if attachment_debug:
        raise ValueError("file_ref analyst response requires exactly one attachment when no explicit path is supplied.")
    return analyst_input



def _assistant_response_synthesis_record(record: dict[str, Any]) -> bool:
    """Return True for nodes that instruct the assistant to synthesize/present a result.

    These records often look superficially interactive because they use
    ``question`` + ``answer_type`` + ``on_answer`` as the model output contract.
    The key distinction is that the node consumes existing state/resources and
    addresses the analyst as the recipient rather than asking the analyst for new
    evidence or a decision.
    """
    if not isinstance(record, dict):
        return False
    question = str(record.get("question") or "").strip()
    if not question or question.endswith("?"):
        return False
    answer_type = str(record.get("answer_type") or "").strip().lower()
    if answer_type not in {"structured_record", "structured", "json", "object", "text"}:
        return False
    on_answer = record.get("on_answer")
    if not isinstance(on_answer, dict) or not (on_answer.get("next") or on_answer.get("update_state")):
        return False
    if isinstance(record.get("allowed_values"), list) and record.get("allowed_values"):
        return False
    node_context = record.get("node_context") if isinstance(record.get("node_context"), dict) else {}
    refs = node_context.get("knowledge_refs") if isinstance(node_context.get("knowledge_refs"), list) else []
    if not refs:
        return False
    # Recipient-oriented synthesis/presentation verbs in the common analyst-facing
    # languages. This is intentionally narrow; ambiguous records remain human.
    normalized = " ".join(question.lower().split())
    synthesis_patterns = (
        r"\b(?:summari[sz]e|present|report|explain|respond|reply|give|provide)\b.*\b(?:analyst|user)\b",
        r"\b(?:дай|надай|покажи|поясни|сформуй|повідом)\b.*\b(?:аналітик\w*|користувач\w*)\b",
        r"\b(?:дай|предоставь|покажи|объясни|сформируй|сообщи)\b.*\b(?:аналитик\w*|пользовател\w*)\b",
    )
    return any(re.search(pattern, normalized, re.IGNORECASE) for pattern in synthesis_patterns)


def _runtime_artifact_metadata(relative_path: str) -> dict[str, Any] | None:
    relative = str(relative_path or "").strip().replace("\\", "/").lstrip("/")
    if not relative:
        return None
    rel = Path(relative)
    if rel.is_absolute() or ".." in rel.parts:
        return None
    workspace = _runtime_workspace().resolve()
    path = (workspace / rel).resolve()
    try:
        path.relative_to(workspace)
    except ValueError:
        return None
    if not path.is_file():
        return None
    data = path.read_bytes()
    return {
        "path": relative,
        "filename": path.name,
        "size": len(data),
        "sha256": hashlib.sha256(data).hexdigest(),
        "source": "runtime_workspace",
    }


def _runtime_artifacts_for_record(record: dict[str, Any]) -> list[dict[str, Any]]:
    """Discover run-local files referenced by a node after deterministic producers ran."""
    candidates: set[str] = set()
    def walk(value: Any) -> None:
        if isinstance(value, dict):
            for child in value.values():
                walk(child)
        elif isinstance(value, list):
            for child in value:
                walk(child)
        elif isinstance(value, str):
            for match in re.finditer(r"(?<![A-Za-z0-9_.-])((?:[A-Za-z0-9_.-]+/)+[A-Za-z0-9_.-]+\.[A-Za-z0-9]{1,12})(?![A-Za-z0-9_.-])", value):
                candidates.add(match.group(1))
    walk(record)
    out = []
    for candidate in sorted(candidates):
        meta = _runtime_artifact_metadata(candidate)
        if meta:
            out.append(meta)
    return out


def _human_interaction_policy(record: dict[str, Any], kind: str) -> dict[str, Any]:
    """Classify declared analyst interaction independently from model preference."""
    if kind == "gate":
        method = str(record.get("method") or "").strip().lower()
        trust_class = str(record.get("trust_class") or "").strip().lower()
        if method == "human" or trust_class == "human_decision":
            title = str(record.get("title") or record.get("id") or "Human decision").strip()
            condition = str(record.get("condition") or "").strip()
            question = str(record.get("question") or "").strip()
            if not question:
                question = f"{title}. Підтверджуєте? Відповідайте так або ні."
                if condition:
                    question += f"\n\nКритерій: {condition}"
            return {
                "requires_human": True,
                "direct_enter": True,
                "direct_respond": True,
                "answer_type": "human_decision",
                "question": question,
                "allowed_values": ["on_pass", "on_fail"],
            }
        return {"requires_human": False, "direct_enter": False, "direct_respond": False}
    if kind != "node":
        return {"requires_human": False, "direct_enter": False, "direct_respond": False}
    answer_type = str(record.get("answer_type") or "").strip().lower()
    question = str(record.get("question") or "").strip()
    on_answer = record.get("on_answer")
    allowed_values = record.get("allowed_values")
    encoded = json.dumps(record, ensure_ascii=False, default=str)
    requires_human = bool(answer_type or (question and isinstance(on_answer, dict)) or "$answer" in encoded or isinstance(allowed_values, list))
    if requires_human and _assistant_response_synthesis_record(record):
        requires_human = False
    if not requires_human:
        return {"requires_human": False, "direct_enter": False, "direct_respond": False}

    # Enter needs an LLM only when the playbook explicitly asks AI to prepare a
    # proposal/draft before the analyst answers. Interpretation/normalization are
    # response-time concerns and must not make the model answer the question itself.
    enter_ai_keys = {"draft_generation", "proposal_generation", "proposal_rule", "generation", "ai_prepare", "pre_answer_analysis"}
    has_enter_ai = any(key in record for key in enter_ai_keys)
    if question.lstrip().lower().startswith(("ai ", "ai:", "ai аналіз", "ai формує", "ai готує")):
        has_enter_ai = True

    response_ai = any(key in record for key in ("interpretation_policy", "response_processing"))
    if isinstance(on_answer, dict):
        if any(key in on_answer for key in ("normalize", "analysis", "transform", "interpret")):
            response_ai = True
        update_state = on_answer.get("update_state")
        if isinstance(update_state, dict):
            # Only literal $answer assignment is safe without interpretation.
            # Structured selectors such as $answer.name_ua require model parsing.
            if any(isinstance(value, str) and value.startswith("$answer.") for value in update_state.values()):
                response_ai = True

    return {
        "requires_human": True,
        "direct_enter": not has_enter_ai,
        "direct_respond": not response_ai,
        "answer_type": answer_type,
        "question": question,
        "allowed_values": allowed_values if isinstance(allowed_values, list) else [],
    }



def _inherited_human_gate_route(
    source: dict[str, Any],
    gate_record: dict[str, Any],
    current_id: str,
    routes: list[dict[str, str]],
    history: list[Any],
) -> tuple[dict[str, str] | None, dict[str, Any] | None]:
    """Safely inherit an immediately preceding analyst enum decision into a human gate.

    This is intentionally conservative. It only applies when all of the following
    are true:
    - the current element is a declared human/human_decision gate;
    - the gate condition explicitly says that an analyst statement/selection is
      itself part of the completion criterion;
    - the immediately preceding analyst turn belongs to a human enum node;
    - that exact enum branch routes directly to this gate; and
    - the gate exposes an on_pass route.

    The rule prevents redundant double-confirmation such as:
      "Do you need another source?" -> analyst: "ні"
      -> human gate whose condition includes "analyst has stated that no
         additional source is required".

    It does NOT infer a pass merely because a previous node happened to route to
    a human gate; ordinary confirmation gates still wait for an explicit answer.
    """
    method = str(gate_record.get("method") or "").strip().lower()
    trust_class = str(gate_record.get("trust_class") or "").strip().lower()
    # Human-decision gates and deterministic gates may both inherit an immediately
    # preceding enum decision when the gate condition explicitly declares that the
    # analyst's statement is itself part of the criterion. Deterministic gates are
    # handled conservatively below: only decision-like missing required_inputs may
    # be satisfied by the inherited branch.
    if method != "human" and trust_class not in {"human_decision", "deterministic"}:
        return None, None

    condition = " ".join(str(gate_record.get("condition") or "").strip().lower().split())
    analyst_statement_markers = (
        "analyst has stated",
        "analyst stated",
        "analyst has confirmed",
        "analyst confirmed",
        "analyst has selected",
        "analyst selected",
        "аналітик зазначив",
        "аналітик підтвердив",
        "аналітик обрав",
    )
    if not condition or not any(marker in condition for marker in analyst_statement_markers):
        return None, None

    if not isinstance(history, list):
        return None, None
    previous_turn = next(
        (
            item for item in reversed(history)
            if isinstance(item, dict)
            and str(item.get("role") or "").lower() == "analyst"
            and str(item.get("node_id") or "").strip()
        ),
        None,
    )
    if not previous_turn:
        return None, None

    previous_id = str(previous_turn.get("node_id") or "").strip()
    if not previous_id or previous_id == current_id:
        return None, None
    previous_record, previous_kind = _record_by_id(source, previous_id)
    if previous_kind != "node" or not isinstance(previous_record, dict):
        return None, None

    answer_type = str(previous_record.get("answer_type") or "").strip().lower()
    allowed_values = previous_record.get("allowed_values")
    if answer_type != "enum" and not isinstance(allowed_values, list):
        return None, None

    analyst_input = str(previous_turn.get("text") or "").strip()
    if not analyst_input:
        return None, None
    previous_routes = _live_routes(previous_record, "node")
    selected_previous = _select_direct_answer_route(previous_record, previous_routes, analyst_input)
    if not selected_previous or str(selected_previous.get("target") or "") != current_id:
        return None, None

    # Generic branch-evidence rule: do not synthesize state by guessing from a
    # required-input field name. If the canonical gate condition itself says that
    # the analyst's immediately preceding statement/selection is part of the pass
    # criterion, and that exact branch routes directly into this gate, the branch
    # is already formal evidence for the gate. This works for both human and
    # deterministic gate implementations without playbook-specific field heuristics.
    inherited_state_updates: dict[str, Any] = {}

    on_pass = next((route for route in routes if route.get("key") == "on_pass"), None)
    if not on_pass:
        return None, None

    evidence = {
        "previous_element_id": previous_id,
        "analyst_input": analyst_input,
        "previous_route_key": selected_previous.get("key"),
        "previous_route_target": selected_previous.get("target"),
        "gate_condition": gate_record.get("condition"),
        "inherited_state_updates": inherited_state_updates,
    }
    return on_pass, evidence



def _is_affirmative_confirmation(value: Any) -> bool:
    normalized = _normalize_answer_key(value)
    return normalized in {
        "yes", "y", "так", "підтверджую", "підтверджено", "confirm",
        "confirmed", "approve", "approved", "ok", "okay", "гаразд", "погоджую", "згоден", "згодна",
    }


def _previous_assistant_proposal(history: Any, current_id: str) -> str:
    """Return the latest analyst-facing assistant proposal for this element.

    A bare confirmation such as `так` is semantically incomplete unless runtime
    also supplies the proposal being confirmed. We intentionally look only at
    the same node to avoid leaking unrelated conversation turns.
    """
    if not isinstance(history, list):
        return ""
    for item in reversed(history):
        if not isinstance(item, dict):
            continue
        if str(item.get("role") or "").lower() != "assistant":
            continue
        if str(item.get("node_id") or "") != str(current_id):
            continue
        text = str(item.get("text") or "").strip()
        if text:
            return text[:24000]
    return ""


def _split_markdown_row(line: str) -> list[str]:
    line = str(line or "").strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    # Ordo tables currently do not rely on escaped pipes in cell values. Keep
    # this parser deliberately small/fail-closed instead of pretending to be a
    # full Markdown implementation.
    return [cell.strip().strip("`") for cell in line.split("|")]


def _markdown_tables(text: str) -> list[tuple[list[str], list[list[str]]]]:
    lines = str(text or "").splitlines()
    tables: list[tuple[list[str], list[list[str]]]] = []
    i = 0
    while i + 1 < len(lines):
        header_line = lines[i].strip()
        sep_line = lines[i + 1].strip()
        if "|" not in header_line or "|" not in sep_line:
            i += 1
            continue
        headers = _split_markdown_row(header_line)
        sep = _split_markdown_row(sep_line)
        if len(headers) < 2 or len(sep) != len(headers):
            i += 1
            continue
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in sep):
            i += 1
            continue
        rows: list[list[str]] = []
        j = i + 2
        while j < len(lines) and "|" in lines[j]:
            row = _split_markdown_row(lines[j])
            if len(row) != len(headers):
                break
            rows.append(row)
            j += 1
        if rows:
            tables.append((headers, rows))
        i = max(j, i + 1)
    return tables


def _normalized_header(value: str) -> str:
    value = re.sub(r"[^a-zA-Z0-9_]+", "_", str(value or "").strip().lower()).strip("_")
    aliases = {
        "datatype": "data_type",
        "fieldpath": "field_path",
        "fieldtype": "field_type",
        "fieldrole": "field_role",
        "calculationrelevance": "calculation_relevance",
        "sourceevidence": "source_evidence",
        "analyststatus": "analyst_status",
    }
    return aliases.get(value, value)


def _confirmed_attribute_rows_from_proposal(proposal: str, state: dict[str, Any]) -> list[dict[str, Any]]:
    """Parse the candidate table that the analyst explicitly confirmed.

    This preserves the proposal's full schema instead of asking an LLM to
    reconstruct a table from the one-word answer `так`. Compatibility aliases
    (`field`, `purpose`, `confirmed`) are added without removing canonical
    columns used by the validation gate.
    """
    required = {"module", "data_type", "field_path", "field_type", "field_role", "calculation_relevance", "source_evidence"}
    candidates = []
    for headers, rows in _markdown_tables(proposal):
        normalized = [_normalized_header(h) for h in headers]
        if not required.issubset(set(normalized)):
            continue
        mapped_rows: list[dict[str, Any]] = []
        for values in rows:
            row = {normalized[idx]: values[idx] for idx in range(len(normalized)) if normalized[idx]}
            if not str(row.get("field_path") or "").strip():
                continue
            row["analyst_status"] = "Confirmed"
            row["confirmed"] = True
            row["field"] = row["field_path"]
            role = _normalize_answer_key(row.get("field_role"))
            if "trigger" in role or "розрах" in role:
                row["purpose"] = "trigger_calculation"
            elif "client" in role or "клієн" in role:
                row["purpose"] = "client_output"
            # Resolve source_number conservatively from the already collected
            # source definition, matching module + product data type.
            source_rows = _state_subtree(state, "source_data_definition.rows")
            if isinstance(source_rows, list):
                match = next((src for src in source_rows if isinstance(src, dict)
                    and str(src.get("module") or "") == str(row.get("module") or "")
                    and str(src.get("data_type") or "") == str(row.get("data_type") or "")), None)
                if isinstance(match, dict) and match.get("source_number") is not None:
                    row["source_number"] = match.get("source_number")
            mapped_rows.append(row)
        if mapped_rows:
            candidates = mapped_rows
    return candidates



def _has_prior_analyst_interaction(history: Any, current_id: str) -> bool:
    return isinstance(history, list) and any(
        isinstance(item, dict) and item.get("role") == "analyst" and str(item.get("node_id") or "") == current_id
        for item in history
    )


def _markdown_table_from_rows(rows: list[dict[str, Any]], columns: list[str]) -> str:
    def cell(value: Any) -> str:
        if value is None:
            return "null"
        if isinstance(value, (dict, list)):
            value = json.dumps(value, ensure_ascii=False, separators=(",", ":"))
        return str(value).replace("|", "\\|").replace("\n", " ")
    header = "| " + " | ".join(columns) + " |"
    sep = "| " + " | ".join("---" for _ in columns) + " |"
    body = ["| " + " | ".join(cell(row.get(col)) for col in columns) + " |" for row in rows]
    return "\n".join([header, sep, *body])


def _contract_column_applicability(column: str, row: dict[str, Any]) -> str:
    """Return required / not_applicable for a canonical table cell.

    Applicability is derived from row semantics rather than from playbook IDs.
    In particular, pseudo-sources such as static/system/algorithm describe values
    produced by the runtime itself; an upstream-data fallback is not meaningful
    for those rows and must not be mislabeled as unresolved evidence.
    """
    source_kind = _normalize_answer_key(row.get("source_field"))
    if column == "basic_fallback" and source_kind in {"static", "system", "algorithm", "runtime", "derived"}:
        return "not_applicable"
    return "required"


def _migrate_rows_to_contract(rows: list[dict[str, Any]], columns: list[str], state: dict[str, Any]) -> tuple[list[dict[str, Any]], list[str]]:
    """Deterministically migrate older/narrower rows into the current table contract.

    Migration may only reuse row aliases and already-present runtime/source context.
    It never invents evidence.  Missing, unresolved and semantically-not-applicable
    values are distinct states.  Only applicable fields that cannot be grounded are
    returned as unresolved for constrained schema repair.
    """
    migrated: list[dict[str, Any]] = []
    unresolved: set[str] = set()
    for row in rows:
        item = copy.deepcopy(row)
        for column in columns:
            if item.get(column) in (None, "", "UNRESOLVED (no fallback evidence in runtime_state)"):
                if _contract_column_applicability(column, item) == "not_applicable":
                    item[column] = "NOT_APPLICABLE"
                else:
                    value = _coerce_proposal_column(column, item, state, "")
                    if value not in (None, "", "UNRESOLVED (no fallback evidence in runtime_state)"):
                        item[column] = value
        if item.get("field_path") and not item.get("field"):
            item["field"] = item["field_path"]
        if item.get("field_type") and not item.get("type"):
            item["type"] = item["field_type"]
        if item.get("analyst_status") == "Confirmed":
            item["confirmed"] = True
        for column in columns:
            if item.get(column) in (None, "", "UNRESOLVED (no fallback evidence in runtime_state)") and _contract_column_applicability(column, item) != "not_applicable":
                unresolved.add(column)
        migrated.append(item)
    return migrated, sorted(unresolved)


def _retry_existing_table_message(record: dict[str, Any], current_id: str, history: Any, state: dict[str, Any]) -> tuple[str, dict[str, Any]] | None:
    """Render an owned existing table on retry without regenerating a fresh draft.

    If an older structured value is narrower than the current canonical contract,
    first perform a deterministic contract-aware migration.  If some columns remain
    unresolved, return repair metadata instead of rendering misleading null cells;
    the normal LLM path can then perform schema repair under a no-invention contract.
    """
    if not _has_prior_analyst_interaction(history, current_id):
        return None
    existing = _existing_structured_outputs_for_record(record, state)
    if not existing:
        return None
    contract = _canonical_proposal_contract(record)
    contract_columns = contract.get("columns") if isinstance(contract.get("columns"), list) else []
    for target, value in existing.items():
        if not isinstance(value, list) or not value or not all(isinstance(x, dict) for x in value):
            continue
        columns = contract_columns or list(dict.fromkeys(k for row in value for k in row.keys()))
        if not columns:
            continue
        migrated = copy.deepcopy(value)
        unresolved: list[str] = []
        if contract_columns:
            migrated, unresolved = _migrate_rows_to_contract(migrated, contract_columns, state)
        audit = {
            "target": target, "row_count": len(migrated), "columns": columns,
            "migration_mode": "contract-aware" if contract_columns else "as-is",
            "unresolved_columns": unresolved,
            "migrated_value": migrated,
        }
        if unresolved:
            audit["requires_schema_repair"] = True
            return "", audit
        table = _markdown_table_from_rows(migrated, columns)
        question = str(record.get("question") or "Підтвердьте або скоригуйте поточне значення.")
        msg = (
            "Поточний структурований результат цього кроку вже збережений у state. "
            "Після повернення з validation він не генерується заново.\n\n" + table +
            "\n\n" + question
        )
        return msg, audit
    return None


def _generic_table_confirmation_result(record: dict[str, Any], current_id: str, analyst_input: str, history: Any, state: dict[str, Any], routes: list[dict[str, str]]) -> tuple[dict[str, Any], dict[str, str] | None, str] | None:
    """Commit the exact displayed canonical table for any affirmative table confirmation."""
    if not _is_affirmative_confirmation(analyst_input):
        return None
    if "table" not in str(record.get("answer_type") or "").lower():
        return None
    if "NORMALIZE_CONFIRMED_RELEVANT_FIELD_ROWS" in json.dumps(record, ensure_ascii=False, default=str):
        return None
    targets = _declared_update_targets(record)
    if len(targets) != 1:
        return None
    target = targets[0]
    proposal = _previous_assistant_proposal(history, current_id)
    if not proposal:
        return None
    table = _proposal_table_schema(proposal)
    if table is None:
        return None
    headers, rows = table
    contract = _canonical_proposal_contract(record)
    columns = contract.get("columns") if isinstance(contract.get("columns"), list) else []
    if columns and not set(columns).issubset(set(headers)):
        return None
    selected = routes[0] if len(routes) == 1 else next((r for r in routes if r.get("key") == "next"), None)
    if selected is None:
        return None
    return {target: rows}, selected, proposal

def _proposal_confirmation_result(record: dict[str, Any], current_id: str, analyst_input: str, history: Any, state: dict[str, Any], routes: list[dict[str, str]]) -> tuple[list[dict[str, Any]], dict[str, str] | None, str] | None:
    """Deterministically commit a previously displayed structured proposal.

    Currently activated by the canonical NORMALIZE_CONFIRMED_RELEVANT_FIELD_ROWS
    operation. This is operation-driven rather than ID-driven so equivalent
    playbooks get the same behavior.
    """
    if not _is_affirmative_confirmation(analyst_input):
        return None
    encoded = json.dumps(record, ensure_ascii=False, default=str)
    if "NORMALIZE_CONFIRMED_RELEVANT_FIELD_ROWS" not in encoded:
        return None
    proposal = _previous_assistant_proposal(history, current_id)
    if not proposal:
        return None
    rows = _confirmed_attribute_rows_from_proposal(proposal, state)
    if not rows:
        return None
    selected = routes[0] if len(routes) == 1 else next((r for r in routes if r.get("key") == "next"), None)
    if selected is None:
        return None
    return rows, selected, proposal



def _looks_like_proposal_confirmation_record(record: dict[str, Any]) -> bool:
    """Whether respond semantics depend on a proposal/draft shown on enter.

    This is deliberately semantic/shape based rather than element-ID based.  It
    catches table confirmations, confirmation-or-correction nodes and canonical
    normalizers whose job is to commit a previously proposed structure.
    """
    answer_type = str(record.get("answer_type") or "").lower()
    encoded = json.dumps(record, ensure_ascii=False, default=str).lower()
    markers = ("confirm", "confirmation", "correction", "proposal", "draft", "candidate")
    if any(marker in answer_type for marker in markers):
        return True
    return any(marker in encoded for marker in (
        "normalize_confirmed", "confirmed_relevant", "confirm the proposal",
        "confirm or correct", "підтверд", "виправ",
    ))


def _proposal_table_schema(proposal: str) -> tuple[list[str], list[dict[str, Any]]] | None:
    """Return the richest Markdown table from an analyst-facing proposal."""
    best = None
    for headers, rows in _markdown_tables(proposal):
        normalized = [_normalized_header(h) for h in headers]
        if not rows or not normalized:
            continue
        mapped = [{normalized[i]: values[i] for i in range(min(len(normalized), len(values))) if normalized[i]} for values in rows]
        candidate = (normalized, mapped)
        if best is None or len(normalized) > len(best[0]):
            best = candidate
    return best


def _source_context_for_row(state: dict[str, Any], row: dict[str, Any]) -> dict[str, Any]:
    sources = _state_subtree(state, "source_data_definition.rows")
    if not isinstance(sources, list):
        return {}
    source_number = row.get("source_number")
    if source_number is not None:
        match = next((x for x in sources if isinstance(x, dict) and x.get("source_number") == source_number), None)
        if isinstance(match, dict):
            return match
    if len(sources) == 1 and isinstance(sources[0], dict):
        return sources[0]
    module = row.get("module")
    data_type = row.get("data_type")
    match = next((x for x in sources if isinstance(x, dict)
                  and (module is None or x.get("module") == module)
                  and (data_type is None or x.get("data_type") == data_type)), None)
    return match if isinstance(match, dict) else {}


def _coerce_proposal_column(column: str, row: dict[str, Any], state: dict[str, Any], analyst_input: str) -> Any:
    """Map compatibility aliases into a proposal column without inventing domain facts."""
    aliases = {
        "field_path": ("field_path", "field", "path", "attribute", "field_name"),
        "field_type": ("field_type", "type", "value_type"),
        "field_role": ("field_role", "role"),
        "calculation_relevance": ("calculation_relevance", "relevance", "logic", "rationale"),
        "source_evidence": ("source_evidence", "evidence"),
        "analyst_status": ("analyst_status", "status"),
        "module": ("module",),
        "data_type": ("data_type",),
    }
    for key in aliases.get(column, (column,)):
        if key in row and row.get(key) not in (None, ""):
            return copy.deepcopy(row.get(key))
    source = _source_context_for_row(state, row)
    if column in {"module", "data_type"} and source.get(column) not in (None, ""):
        return copy.deepcopy(source.get(column))
    if column == "field_role" and row.get("purpose"):
        purpose = _normalize_answer_key(row.get("purpose"))
        if "trigger" in purpose or "calcul" in purpose or "розрах" in purpose:
            return "Trigger"
        if "client" in purpose or "output" in purpose or "клієн" in purpose:
            return "Client Data"
        return str(row.get("purpose"))
    if column == "calculation_relevance" and row.get("purpose"):
        return str(row.get("purpose"))
    if column == "source_evidence":
        context = source.get("context") or source.get("description")
        if context:
            return str(context)
        if analyst_input.strip():
            return "Confirmed/corrected by analyst in this interaction"
    if column == "analyst_status" and (row.get("confirmed") is True or "підтвер" in analyst_input.lower() or "confirm" in analyst_input.lower()):
        return "Confirmed"
    return None


def _deep_merge_preserving_existing(existing: Any, incoming: Any) -> Any:
    """Deep merge correction data over an existing structured value.

    Explicit incoming values (including null/false/empty collections) win; keys the
    analyst/model did not mention are preserved.  This matches confirmation-or-
    correction semantics and avoids destructive partial replacement.
    """
    if not isinstance(existing, dict) or not isinstance(incoming, dict):
        return copy.deepcopy(incoming)
    merged = copy.deepcopy(existing)
    for key, value in incoming.items():
        if key in merged and isinstance(merged[key], dict) and isinstance(value, dict):
            merged[key] = _deep_merge_preserving_existing(merged[key], value)
        else:
            merged[key] = copy.deepcopy(value)
    return merged


def _canonical_proposal_contract(record: dict[str, Any]) -> dict[str, Any]:
    draft = record.get("draft_generation")
    if not isinstance(draft, dict):
        return {}
    contract: dict[str, Any] = {}
    if isinstance(draft.get("columns"), list):
        contract["columns"] = [_normalized_header(x) for x in draft["columns"] if str(x).strip()]
    if isinstance(draft.get("required_coverage"), list):
        contract["required_coverage"] = [str(x) for x in draft["required_coverage"]]
    if isinstance(draft.get("scope_constraints"), dict):
        contract["scope_constraints"] = copy.deepcopy(draft["scope_constraints"])
    if draft.get("rule") is not None:
        contract["rule"] = copy.deepcopy(draft.get("rule"))
    return contract




def _declared_update_targets(record: dict[str, Any]) -> list[str]:
    on_answer = record.get("on_answer")
    if not isinstance(on_answer, dict):
        return []
    update_state = on_answer.get("update_state")
    if not isinstance(update_state, dict):
        return []
    return [str(k) for k in update_state.keys() if str(k).strip()]


def _existing_structured_outputs_for_record(record: dict[str, Any], state: dict[str, Any]) -> dict[str, Any]:
    """Return already materialized outputs owned by a proposal/correction node.

    This is intentionally generic: retry/review entry should refine the state that
    this node previously produced instead of regenerating a fresh proposal from its
    upstream draft sources.
    """
    if not _looks_like_proposal_confirmation_record(record):
        return {}
    existing: dict[str, Any] = {}
    for target in _declared_update_targets(record):
        if _state_path_exists(state, target):
            value = _state_subtree(state, target)
            if value not in (None, "", [], {}):
                existing[target] = value
    return existing


def _confirmed_source_field_variants(state: dict[str, Any]) -> dict[str, str]:
    rows = _state_subtree(state, "source_attribute_mapping.rows")
    if not isinstance(rows, list):
        return {}
    index: dict[str, str] = {}
    for row in rows:
        if not isinstance(row, dict):
            continue
        status = str(row.get("analyst_status") or "").strip().lower()
        confirmed = row.get("confirmed") is True or status == "confirmed"
        if not confirmed:
            continue
        field = str(row.get("field_path") or row.get("field") or "").strip()
        if not field:
            continue
        module = str(row.get("module") or "").strip()
        data_type = str(row.get("data_type") or "").strip()
        variants = {field}
        if data_type:
            variants.add(f"{data_type}.{field}")
        if module:
            variants.add(f"{module}.{field}")
        if module and data_type:
            variants.add(f"{module}.{data_type}.{field}")
        for variant in variants:
            index[variant] = field
    return index


def _canonicalize_confirmed_source_references(state: dict[str, Any], updates: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any] | None]:
    """Canonicalize source_field references against confirmed source mappings.

    Namespace-qualified variants such as `module.record.field` are reduced to
    the canonical confirmed `field_path` only when the mapping is unambiguous.
    Unmatched values (including static/system/algorithm pseudo-sources) are preserved.
    """
    variants = _confirmed_source_field_variants(state)
    if not variants or not isinstance(updates, dict):
        return updates, None
    out = copy.deepcopy(updates)
    changes = []
    for target, value in list(out.items()):
        if not isinstance(value, list):
            continue
        new_rows = []
        touched = False
        for row in value:
            if not isinstance(row, dict):
                new_rows.append(row)
                continue
            item = copy.deepcopy(row)
            source_field = item.get("source_field")
            if isinstance(source_field, str) and source_field in variants and variants[source_field] != source_field:
                item["source_field"] = variants[source_field]
                changes.append({"target": str(target), "from": source_field, "to": variants[source_field]})
                touched = True
            new_rows.append(item)
        if touched:
            out[target] = new_rows
    return out, ({"mode": "confirmed-source-field-canonicalization", "changes": changes} if changes else None)

def _reconcile_structured_proposal_updates(record: dict[str, Any], current_id: str, analyst_input: str, history: Any, state: dict[str, Any], updates: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any] | None]:
    """Preserve proposal schema across enter→respond confirmation/correction.

    Models may correctly understand analyst corrections but emit a compact runtime
    representation (for example `field/purpose/type`) that drops the columns the
    analyst actually reviewed (`module/data_type/field_path/...`).  For nodes whose
    canonical semantics are proposal confirmation/correction, reconcile list rows
    back to the displayed proposal schema using only model data, canonical source
    state and analyst provenance.  No element IDs are referenced here.
    """
    if not _looks_like_proposal_confirmation_record(record):
        return updates, None
    proposal = _previous_assistant_proposal(history, current_id)
    if not proposal:
        return updates, None
    table = _proposal_table_schema(proposal)
    contract = _canonical_proposal_contract(record)
    contract_columns = contract.get("columns") if isinstance(contract.get("columns"), list) else []
    if table is None and not contract_columns:
        # Non-tabular confirmation/correction still gets non-destructive object merge.
        reconciled = copy.deepcopy(updates)
        merged_targets = []
        for target, value in list(reconciled.items()):
            existing = _state_subtree(state, str(target))
            if isinstance(existing, dict) and isinstance(value, dict):
                reconciled[target] = _deep_merge_preserving_existing(existing, value)
                merged_targets.append(str(target))
        return reconciled, {"mode": "proposal-context-and-object-merge" if merged_targets else "proposal-context-only", "proposal_chars": len(proposal), "merged_targets": merged_targets, "proposal_contract": contract}
    table_columns, proposal_rows = table if table is not None else ([], [])
    columns = contract_columns or table_columns
    if len(columns) < 2:
        return updates, {"mode": "proposal-context-only", "proposal_chars": len(proposal), "proposal_contract": contract}
    reconciled = copy.deepcopy(updates)
    # First preserve existing object fields for any object-valued correction target.
    merged_object_targets = []
    for target, value in list(reconciled.items()):
        existing = _state_subtree(state, str(target))
        if isinstance(existing, dict) and isinstance(value, dict):
            reconciled[target] = _deep_merge_preserving_existing(existing, value)
            merged_object_targets.append(str(target))
    changed_targets = []
    unresolved = {}
    for target, value in list(reconciled.items()):
        if not isinstance(value, list) or not value or not all(isinstance(x, dict) for x in value):
            continue
        # Only reconcile when the response rows are visibly schema-poorer than the
        # proposal. This avoids rewriting unrelated list outputs.
        response_keys = set().union(*(set(x.keys()) for x in value))
        missing_columns = [c for c in columns if c not in response_keys]
        if not missing_columns:
            continue
        enriched_rows = []
        target_unresolved = set()
        for row in value:
            enriched = copy.deepcopy(row)
            for column in columns:
                if enriched.get(column) in (None, ""):
                    coerced = _coerce_proposal_column(column, enriched, state, analyst_input)
                    if coerced not in (None, ""):
                        enriched[column] = coerced
            # Preserve useful compatibility aliases; canonical proposal columns are
            # additive and never destroy the compact representation.
            if enriched.get("field_path") and not enriched.get("field"):
                enriched["field"] = enriched["field_path"]
            if enriched.get("field_type") and not enriched.get("type"):
                enriched["type"] = enriched["field_type"]
            if enriched.get("analyst_status") == "Confirmed":
                enriched["confirmed"] = True
            for column in columns:
                if enriched.get(column) in (None, ""):
                    target_unresolved.add(column)
            enriched_rows.append(enriched)
        reconciled[target] = enriched_rows
        changed_targets.append(target)
        if target_unresolved:
            unresolved[target] = sorted(target_unresolved)
    if not changed_targets:
        return updates, {"mode": "proposal-context-only", "proposal_chars": len(proposal), "table_columns": columns, "proposal_contract": contract}
    return reconciled, {
        "mode": "schema-preserving-reconciliation",
        "proposal_chars": len(proposal),
        "proposal_columns": columns,
        "proposal_row_count": len(proposal_rows),
        "proposal_contract": contract,
        "targets": changed_targets,
        "merged_object_targets": merged_object_targets,
        "unresolved_columns": unresolved,
    }

def _selected_branch_requires_ai(record: dict[str, Any], analyst_input: str) -> bool:
    on_answer = record.get("on_answer")
    if not isinstance(on_answer, dict) or not isinstance(on_answer.get("branch"), dict):
        return False
    normalized = _normalize_answer_key(analyst_input)
    for key, branch_value in on_answer["branch"].items():
        if _normalize_answer_key(key) != normalized:
            continue
        return isinstance(branch_value, dict) and any(k in branch_value for k in ("analysis", "normalize", "transform", "interpret"))
    return False

def _normalize_answer_key(value: Any) -> str:
    return " ".join(str(value or "").strip().lower().split())


def _select_direct_answer_route(record: dict[str, Any], routes: list[dict[str, str]], analyst_input: str) -> dict[str, str] | None:
    on_answer = record.get("on_answer")
    normalized = _normalize_answer_key(analyst_input)

    # Human-decision gates are runtime-controlled: the analyst selects pass/fail,
    # never the model. Accept canonical route keys plus common yes/no forms.
    method = str(record.get("method") or "").strip().lower()
    trust_class = str(record.get("trust_class") or "").strip().lower()
    if method == "human" or trust_class == "human_decision":
        direct = next((r for r in routes if _normalize_answer_key(r.get("key")) == normalized), None)
        if direct:
            return direct
        yes = {
            "yes", "y", "так", "підтверджую", "підтверджено", "confirm", "confirmed",
            "approve", "approved", "pass", "on_pass",
            "погодити", "критерій виконано", "погодити — критерій виконано", "погодити - критерій виконано",
        }
        no = {
            "no", "n", "ні", "не підтверджую", "reject", "rejected", "fail", "on_fail",
            "потрібне виправлення", "критерій не виконано",
            "потрібне виправлення — критерій не виконано", "потрібне виправлення - критерій не виконано",
        }
        if normalized in yes:
            return next((r for r in routes if r.get("key") == "on_pass"), None)
        if normalized in no:
            return next((r for r in routes if r.get("key") == "on_fail"), None)
        return None
    if isinstance(on_answer, dict) and isinstance(on_answer.get("branch"), dict):
        branch = on_answer["branch"]
        for key in branch:
            if _normalize_answer_key(key) == normalized:
                return next((r for r in routes if _normalize_answer_key(r.get("key")) == _normalize_answer_key(key)), None)
        return None
    if len(routes) == 1:
        return routes[0]
    return next((r for r in routes if _normalize_answer_key(r.get("key")) == normalized), None)


def _set_dotted_state(target: dict[str, Any], dotted_key: str, value: Any) -> None:
    parts = [part for part in str(dotted_key).split(".") if part]
    if not parts:
        return
    cursor = target
    for part in parts[:-1]:
        child = cursor.get(part)
        if not isinstance(child, dict):
            child = {}
            cursor[part] = child
        cursor = child
    cursor[parts[-1]] = value


def _resolve_direct_human_value(value: Any, analyst_input: str) -> Any:
    if value == "$answer":
        return analyst_input
    if isinstance(value, str) and value in {"$runtime.timestamp", "$runtime.now", "$runtime.datetime"}:
        return datetime.now(timezone.utc).isoformat()
    if isinstance(value, list):
        return [_resolve_direct_human_value(v, analyst_input) for v in value]
    if isinstance(value, dict):
        return {k: _resolve_direct_human_value(v, analyst_input) for k, v in value.items()}
    return copy.deepcopy(value)

def _apply_direct_answer_updates(record: dict[str, Any], state: dict[str, Any], analyst_input: str) -> tuple[dict[str, Any], dict[str, Any]]:
    """Apply state writes owned by a deterministic human answer.

    Supports both a shared ``on_answer.update_state`` mapping and state writes
    declared inside the selected branch/choice. Graph selection and state
    mutation must be one atomic runtime decision; choosing a branch without
    committing that branch's declared updates leaves downstream gates with an
    impossible state.
    """
    new_state = copy.deepcopy(state) if isinstance(state, dict) else {}
    updates: dict[str, Any] = {}
    on_answer = record.get("on_answer")
    if not isinstance(on_answer, dict):
        return new_state, updates

    normalized = _normalize_answer_key(analyst_input)

    def apply_mapping(mapping: Any) -> None:
        if not isinstance(mapping, dict):
            return
        for key, value in mapping.items():
            if value == "$increment":
                previous = _state_subtree(new_state, str(key))
                resolved = (previous if isinstance(previous, int) else 0) + 1
            else:
                resolved = _resolve_direct_human_value(value, analyst_input)
            _set_dotted_state(new_state, str(key), resolved)
            updates[str(key)] = resolved

    # Shared writes apply to every accepted answer.
    apply_mapping(on_answer.get("update_state"))

    # Branch can be explicitly nested under ``branch`` ...
    selected_branch = None
    branch_map = on_answer.get("branch")
    if isinstance(branch_map, dict):
        selected_branch = next((v for k, v in branch_map.items() if _normalize_answer_key(k) == normalized), None)

    # ... or represented directly as named choices under ``on_answer``.
    if selected_branch is None:
        selected_branch = next((v for k, v in on_answer.items()
                                if k not in {"update_state", "normalize", "branch"}
                                and _normalize_answer_key(k) == normalized), None)

    if isinstance(selected_branch, dict):
        apply_mapping(selected_branch.get("update_state"))

    return new_state, updates




def _apply_declared_human_runtime_expressions(
    record: dict[str, Any], state: dict[str, Any], phase: str, updates: dict[str, Any], state_patch: dict[str, Any] | None
) -> tuple[dict[str, Any], dict[str, Any] | None, dict[str, Any]]:
    """Apply deterministic runtime-owned expressions declared by a human answer contract.

    The playbook owns *which* path is incremented; runtime only implements the
    generic `$increment` operator.  This prevents a model from having to invent
    counter arithmetic during a structured analyst response.
    """
    if phase != "respond":
        return updates, state_patch, {}
    on_answer = record.get("on_answer") if isinstance(record, dict) else None
    mapping = on_answer.get("update_state") if isinstance(on_answer, dict) else None
    if not isinstance(mapping, dict):
        return updates, state_patch, {}
    forced: dict[str, Any] = {}
    for path, spec in mapping.items():
        if spec == "$increment":
            previous = _state_subtree(state if isinstance(state, dict) else {}, str(path))
            forced[str(path)] = (previous if isinstance(previous, int) else 0) + 1
    if not forced:
        return updates, state_patch, {}
    merged = copy.deepcopy(updates) if isinstance(updates, dict) else {}
    merged.update(copy.deepcopy(forced))
    patch = copy.deepcopy(state_patch) if isinstance(state_patch, dict) else None
    if patch is not None:
        ops = patch.get("operations") if isinstance(patch.get("operations"), list) else []
        forced_paths = set(forced)
        ops = [op for op in ops if not (isinstance(op, dict) and str(op.get("path")) in forced_paths)]
        for path, value in forced.items():
            ops.append({"op": "set", "path": path, "value": copy.deepcopy(value), "basis": "runtime_declared_increment"})
        patch["operations"] = ops
    return merged, patch, forced



def _canonical_respond_contract(record: dict[str, Any], routes: list[dict[str, str]]) -> dict[str, Any]:
    """Return the runtime-owned fixed respond contract, when canonical YAML defines one.

    This deliberately uses canonical on_answer semantics rather than model prose.
    A fixed `on_answer.next` means that, once the answer has been successfully
    normalized into all declared update targets, graph orchestration belongs to
    the runtime.
    """
    on_answer = record.get("on_answer")
    if not isinstance(on_answer, dict):
        return {"next_route": None, "update_targets": []}
    next_target = on_answer.get("next")
    next_route = None
    if isinstance(next_target, str) and next_target and not next_target.startswith("$"):
        next_route = next((r for r in routes if r.get("key") == "next" and r.get("target") == next_target), None)
    mapping = on_answer.get("update_state")
    targets = sorted(str(k) for k in mapping.keys()) if isinstance(mapping, dict) else []
    return {"next_route": next_route, "update_targets": targets}


def _updates_cover_canonical_targets(updates: dict[str, Any], targets: list[str]) -> bool:
    if not targets:
        return False
    return all(_state_path_exists(updates, path) for path in targets)


def _resolve_respond_orchestration(
    record: dict[str, Any], kind: str, phase: str, routes: list[dict[str, str]], result: dict[str, Any]
) -> tuple[bool, dict[str, str] | None, str | None]:
    """Resolve model/runtime ownership of `await_analyst` and routing.

    General invariants:
    * a model-selected *allowed* route and `await_analyst=true` are contradictory;
      the explicit route wins because runtime owns graph traversal;
    * for a node with canonical `on_answer.next`, if the model produced every
      canonical update target, the respond phase is complete even if the model
      forgot the route or tried to ask another question;
    * incomplete/unmatched answers remain allowed to wait for clarification.
    """
    await_analyst = bool(result.get("await_analyst")) if kind != "gate" else False
    requested_key = result.get("route_key")
    selected = next((r for r in routes if r.get("key") == requested_key), None)
    reason = None
    if kind == "node" and phase == "respond":
        if selected is not None and await_analyst:
            await_analyst = False
            reason = "allowed-route-overrides-model-await"
        elif selected is None and await_analyst:
            contract = _canonical_respond_contract(record, routes)
            updates = result.get("state_updates") if isinstance(result.get("state_updates"), dict) else {}
            fixed = contract.get("next_route")
            targets = contract.get("update_targets") or []
            if isinstance(fixed, dict) and _updates_cover_canonical_targets(updates, targets):
                selected = fixed
                await_analyst = False
                reason = "canonical-on-answer-complete"
    return await_analyst, selected, reason


def _runtime_only_live_result(*, credentials: dict[str, Any], record: dict[str, Any], kind: str, current_id: str, phase: str, state: dict[str, Any], routes: list[dict[str, str]], assistant_message: str, await_analyst: bool, selected: dict[str, str] | None = None, updates: dict[str, Any] | None = None, new_state: dict[str, Any] | None = None, reason: str = "runtime-human-input", extra_runtime: dict[str, Any] | None = None, current_revision: int = 0) -> dict[str, Any]:
    updates = updates or {}
    new_state = new_state if isinstance(new_state, dict) else (copy.deepcopy(state) if isinstance(state, dict) else {})
    usage = {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0, "cached_tokens": 0, "reasoning_tokens": 0}
    context_breakdown = {"system_chars": 0, "element_chars": 0, "state_chars": 0, "history_chars": 0, "resources_chars": 0}
    revision_after = int(current_revision) + (1 if canonicalize_runtime_state(new_state) != canonicalize_runtime_state(state) else 0)
    debug = {
        "provider": credentials.get("provider"), "base_url": credentials.get("base_url"), "api_style": "runtime_only",
        "model": credentials.get("model"), "current_id": current_id, "element_kind": kind, "phase": phase,
        "usage": usage, "context_breakdown": context_breakdown,
        "input": {"system_text": "", "context": {"runtime_only": True, "reason": reason}, "request_payload": None},
        "output": {"raw_text": "", "parsed_result": None, "api_response": None},
        "runtime": {"state_before": state, "state_updates": updates, "state_after": new_state, "revision_before": int(current_revision), "revision_after": revision_after, "allowed_routes": routes,
                    "requested_route_key": None, "selected_route_key": selected["key"] if selected else None,
                    "next_id": selected["target"] if selected else None, "await_analyst": await_analyst,
                    "run_status": "running", "completion_reason": None, "llm_call_skipped": True, "reason": reason},
    }
    if isinstance(extra_runtime, dict):
        debug["runtime"].update(copy.deepcopy(extra_runtime))
        debug["input"]["context"].update(copy.deepcopy(extra_runtime))
    if "normalized_execution_result" not in debug["runtime"]:
        debug["runtime"]["normalized_execution_result"] = {
            "element_id": current_id, "phase": phase,
            "status": "awaiting_analyst" if await_analyst else "completed",
            "state_updates": copy.deepcopy(updates),
            "route_key": selected["key"] if selected else None,
            "next_id": selected["target"] if selected else None,
        }
    return {"assistant_message": assistant_message, "state_revision": revision_after, "rationale_short": reason, "await_analyst": await_analyst, "phase": phase,
            "route_key": selected["key"] if selected else None, "next_id": selected["target"] if selected else None,
            "state": new_state, "routes": routes, "usage": usage, "context_breakdown": context_breakdown, "debug": debug,
            "provider": credentials.get("provider"), "model": credentials.get("model"), "terminal": False,
            "run_status": "running", "completion_reason": None, "llm_call_skipped": True}

def _system_contract_for_call(kind: str, phase: str, human_policy: dict[str, Any]) -> str:
    """Return only the runtime invariants relevant to this exact API call."""
    schema = (
        "Return ONLY JSON: assistant_message:string, route_key:string|null, "
        "state_updates:object, rationale_short:string, await_analyst:boolean. "
    )
    base = "Execute exactly one Ordo element. Runtime owns graph traversal; never invent or skip elements. Do the actual work, not a paraphrase. "
    tail = "Do not claim that a graph transition occurred; runtime reports transitions."
    if kind == "gate":
        return base + schema + "Evaluate the gate from supplied context, never ask the analyst, and choose only a supplied allowed route. " + tail
    if phase == "respond":
        return base + schema + "Process the supplied analyst answer and update state. Ask one necessary follow-up or choose only a supplied allowed route. " + tail
    if human_policy.get("requires_human"):
        return base + schema + "Prepare the requested result/proposal, show it to the analyst, ask one concrete question, set await_analyst=true and route_key=null. " + tail
    return base + schema + "Perform all possible work; if input is truly required ask one concrete question, otherwise choose only a supplied allowed route. " + tail



def _package_resource_text(ref: str) -> tuple[str, str] | tuple[None, None]:
    resources = _active_playbook_package().get("resources") if isinstance(_active_playbook_package().get("resources"), dict) else {}
    normalized = str(ref or "").lstrip("./")
    matches = [(name, text) for name, text in resources.items()
               if name == normalized or name.endswith("/" + normalized)]
    if len(matches) == 1:
        return matches[0]
    exact_name = [(name, text) for name, text in resources.items() if Path(name).name == Path(normalized).name]
    if len(exact_name) == 1:
        return exact_name[0]
    return (None, None)


def _active_run_key() -> tuple[str,str,str]:
    ctx=_ACTIVE_RUN_CONTEXT.get() or {}
    return (str(ctx.get("package_id") or ""),str(ctx.get("session_id") or ""),str(ctx.get("run_id") or ""))

def _paths_overlap(a: str, b: str) -> bool:
    return a==b or a.startswith(b+'.') or b.startswith(a+'.')

def _update_artifact_registry(*, state_lineage: list[dict[str,Any]] | None = None, artifact_lineage: dict[str,Any] | None = None) -> None:
    key=_active_run_key()
    if not all(key): return
    registry=RUN_ARTIFACT_REGISTRY.setdefault(key,{})
    if artifact_lineage and artifact_lineage.get("path"):
        row=copy.deepcopy(artifact_lineage)
        row["freshness_status"]="fresh" if row.get("depends_on_paths") else "unknown_dependencies"
        row["stale_by_state_lineage"]=[]
        registry[str(row["path"])]=row
    for change in state_lineage or []:
        cp=str(change.get("path") or "")
        rev=int(change.get("revision") or 0)
        for row in registry.values():
            if rev <= int(row.get("materialized_from_revision") or -1): continue
            deps=[str(x) for x in (row.get("depends_on_paths") or [])]
            if deps and any(_paths_overlap(cp,d) for d in deps):
                row["freshness_status"]="stale"
                row.setdefault("stale_by_state_lineage",[]).append({"path":cp,"revision":rev,"producer_element_id":change.get("producer_element_id")})

def _run_artifact_status(path: str) -> dict[str,Any] | None:
    key=_active_run_key(); registry=RUN_ARTIFACT_REGISTRY.get(key) or {}
    return copy.deepcopy(registry.get(str(path))) if str(path) in registry else None

def _artifact_dependency_paths_for_element(element_id: str) -> list[str]:
    """Return compiler-declared state dependencies for an artifact producer.

    The package parser stores the Runtime Semantic Plan under ``semantic_plan``.
    Artifact freshness must use that canonical package key rather than a second
    alias, otherwise dependencies silently become unknown.
    """
    package=_active_playbook_package()
    plan=package.get("semantic_plan") if isinstance(package,dict) else None
    element=((plan or {}).get("elements") or {}).get(str(element_id or "")) if isinstance(plan,dict) else None
    sc=(element or {}).get("state_contract") if isinstance(element,dict) else {}
    deps=set(str(x) for x in (sc.get("reads_hint") or []) if str(x))
    deps.update(str(x) for x in (((sc.get("declared_inputs_by_class") or {}).get("state")) or []) if str(x))
    return sorted(deps)

def _runtime_workspace(*, package_id: str | None = None, session_id: str | None = None, run_id: str | None = None) -> Path:
    context = _ACTIVE_RUN_CONTEXT.get() or {}
    package = str(package_id or context.get("package_id") or _active_playbook_package().get("id") or "unloaded")
    session = str(session_id or context.get("session_id") or "legacy-session")
    run = str(run_id or context.get("run_id") or "legacy-run")
    safe = lambda value: re.sub(r"[^A-Za-z0-9._-]+", "_", value)[:160] or "unknown"
    root = Path(tempfile.gettempdir()) / "ordo_tree_editor_runs" / safe(package) / safe(session) / safe(run)
    root.mkdir(parents=True, exist_ok=True)
    return root




def _deep_merge_dicts(dst: dict[str, Any], src: dict[str, Any], *, overwrite: bool = True) -> dict[str, Any]:
    for key, value in src.items():
        if isinstance(value, dict) and isinstance(dst.get(key), dict):
            _deep_merge_dicts(dst[key], value, overwrite=overwrite)
        elif overwrite or key not in dst:
            dst[key] = copy.deepcopy(value)
    return dst


def _canonicalize_runtime_state(state: Any) -> dict[str, Any]:
    """Return one canonical nested state tree from mixed nested + dotted-key state.

    Existing nested values are authoritative on input; dotted keys fill missing leaves.
    New updates are applied separately with overwrite semantics by
    _apply_state_updates_canonical().
    """
    if not isinstance(state, dict):
        return {}
    canonical: dict[str, Any] = {}
    dotted: list[tuple[str, Any]] = []
    for key, value in state.items():
        if isinstance(key, str) and "." in key:
            dotted.append((key, value))
        else:
            canonical[key] = copy.deepcopy(value)
    for path, value in dotted:
        parts = [part for part in path.split(".") if part]
        if not parts:
            continue
        cursor = canonical
        blocked = False
        for part in parts[:-1]:
            existing = cursor.get(part)
            if existing is None:
                existing = {}
                cursor[part] = existing
            elif not isinstance(existing, dict):
                blocked = True
                break
            cursor = existing
        if blocked:
            continue
        leaf = parts[-1]
        if leaf not in cursor:
            cursor[leaf] = copy.deepcopy(value)
    return canonical


def _canonical_value_hash(value: Any) -> str:
    def norm(v: Any) -> Any:
        if isinstance(v, dict): return {str(k): norm(v[k]) for k in sorted(v, key=lambda x: str(x))}
        if isinstance(v, list): return [norm(x) for x in v]
        return v
    raw=json.dumps(norm(value),ensure_ascii=False,sort_keys=True,separators=(",",":"),default=str).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def _state_lineage_entries(*, patch: dict[str, Any], new_state: dict[str, Any], revision: int, producer_element_id: str, source_run_id: str, runtime_updates: dict[str, Any] | None = None) -> list[dict[str, Any]]:
    entries=[]
    for op in (patch.get("operations") or []) if isinstance(patch,dict) else []:
        if not isinstance(op,dict) or not op.get("path"): continue
        path=str(op["path"]); value=_state_subtree(new_state,path)
        entries.append({"path":path,"revision":int(revision),"producer_element_id":producer_element_id,"operation":str(op.get("op") or ""),"basis":op.get("basis"),"reason":op.get("reason"),"source_run_id":source_run_id,"value_hash":_canonical_value_hash(value)})
    for path,value in (runtime_updates or {}).items():
        if any(e["path"]==str(path) for e in entries): continue
        entries.append({"path":str(path),"revision":int(revision),"producer_element_id":producer_element_id,"operation":"runtime_set","basis":"derived","reason":"runtime-owned gate branch update","source_run_id":source_run_id,"value_hash":_canonical_value_hash(_state_subtree(new_state,str(path)))})
    return entries


def _apply_state_updates_canonical(state: Any, updates: Any) -> dict[str, Any]:
    """Apply state updates into the canonical nested tree. Dotted update paths overwrite leaves."""
    new_state = _canonicalize_runtime_state(state)
    if not isinstance(updates, dict):
        return new_state
    for key, value in updates.items():
        if isinstance(key, str) and "." in key:
            _assign_nested_path(new_state, key, value)
        elif isinstance(value, dict) and isinstance(new_state.get(key), dict):
            _deep_merge_dicts(new_state[key], value, overwrite=True)
        else:
            new_state[key] = copy.deepcopy(value)
    return new_state

def _state_subtree(state: dict[str, Any], path: str) -> Any:
    """Resolve a state path while supporting both nested and flat dotted runtime keys."""
    clean = path.removeprefix("state.").removeprefix("$state.")
    descendants = {k[len(clean)+1:]: v for k, v in state.items() if isinstance(k, str) and k.startswith(clean + ".")}
    if descendants:
        root: dict[str, Any] = {}
        for suffix, value in descendants.items():
            cursor = root
            parts = suffix.split(".")
            for part in parts[:-1]:
                cursor = cursor.setdefault(part, {})
            cursor[parts[-1]] = copy.deepcopy(value)
        base = state.get(clean)
        if isinstance(base, dict):
            merged = copy.deepcopy(base)
            def merge(dst: dict[str, Any], src: dict[str, Any]) -> None:
                for k, v in src.items():
                    if isinstance(v, dict) and isinstance(dst.get(k), dict): merge(dst[k], v)
                    else: dst[k] = v
            merge(merged, root)
            return merged
        return root
    if clean in state:
        return copy.deepcopy(state[clean])
    cur: Any = state
    for part in clean.split("."):
        if not isinstance(cur, dict) or part not in cur:
            return None
        cur = cur[part]
    return copy.deepcopy(cur)


def _render_markdown_value(value: Any) -> str:
    if value is None or value == "":
        return "не надано"
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (str, int, float)):
        return str(value)
    if isinstance(value, list):
        if not value: return "— немає —"
        if all(not isinstance(x, (dict, list)) for x in value):
            return "\n".join(f"- {x}" for x in value)
        chunks=[]
        for item in value:
            if isinstance(item, dict):
                chunks.append("\n".join(f"- **{k}**: {_render_markdown_value(v)}" for k,v in item.items()))
            else: chunks.append(f"- {_render_markdown_value(item)}")
        return "\n\n".join(chunks)
    if isinstance(value, dict):
        rows=[]
        for k,v in value.items():
            rendered=_render_markdown_value(v)
            rows.append(f"**{k}**: {rendered}" if "\n" not in rendered else f"**{k}**:\n{rendered}")
        return "\n\n".join(rows) if rows else "— немає —"
    return str(value)


def _binding_value(spec: Any, state: dict[str, Any]) -> Any:
    if isinstance(spec, str) and (spec.startswith("state.") or spec.startswith("$state.")):
        return _state_subtree(state, spec)
    if isinstance(spec, dict) and "source" in spec:
        value = _binding_value(spec.get("source"), state)
        if value is None or value == "" or value == [] or value == {}:
            return copy.deepcopy(spec.get("default"))
        return value
    return copy.deepcopy(spec)


def _lookup_bound_path(bound: dict[str, Any], expr: str) -> Any:
    parts=expr.split(".")
    cur: Any=bound.get(parts[0])
    for part in parts[1:]:
        if not isinstance(cur, dict) or part not in cur: return None
        cur=cur[part]
    return cur


def _document_binding_compatibility(bound: dict[str, Any], state: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, str]]]:
    """Return declared document bindings unchanged.

    Generic runtime must not know domain-specific aliases between a playbook's
    state schema and its templates.  Compatibility aliases belong in the
    playbook bindings/resources or in an explicitly declared migration contract.
    """
    return copy.deepcopy(bound), []

def _apply_declared_state_updates(record: dict[str, Any], state: dict[str, Any], derived: dict[str, Any] | None = None) -> tuple[dict[str, Any], dict[str, Any]]:
    updates: dict[str, Any] = {}
    for key, spec in (record.get("update_state") or {}).items() if isinstance(record.get("update_state"), dict) else []:
        value=spec
        if isinstance(spec, str) and spec.startswith("$derived."):
            value=_state_subtree(derived or {}, spec.removeprefix("$derived."))
        elif spec == "$increment":
            previous=_state_subtree(state, key)
            value=(previous if isinstance(previous, int) else 0)+1
        updates[str(key)]=copy.deepcopy(value)
    new_state=_apply_state_updates_canonical(state, updates)
    return new_state, updates


def _package_table_contracts() -> list[dict[str, Any]]:
    """Discover table contracts from package YAML resources without filename conventions."""
    resources = _active_playbook_package().get("resources") if isinstance(_active_playbook_package().get("resources"), dict) else {}
    found: list[dict[str, Any]] = []
    for name, text in resources.items():
        if not str(name).lower().endswith((".yaml", ".yml")):
            continue
        try:
            doc = yaml.safe_load(text)
        except Exception:
            continue
        tables = doc.get("tables") if isinstance(doc, dict) else None
        if not isinstance(tables, dict):
            continue
        for table_id, spec in tables.items():
            if not isinstance(spec, dict) or not isinstance(spec.get("collection"), str) or not isinstance(spec.get("columns"), list):
                continue
            columns=[]
            for col in spec["columns"]:
                if isinstance(col, dict) and col.get("key"):
                    columns.append({"key":str(col["key"]),"label":str(col.get("label") or col["key"]),"required":bool(col.get("required",False))})
            if columns:
                found.append({"resource":name,"table_id":str(table_id),"collection":spec["collection"],"columns":columns})
    return found


def _table_contract_for_placeholder(expr: str) -> dict[str, Any] | None:
    root = expr.split(".")[0]
    matches=[]
    for contract in _package_table_contracts():
        clean=str(contract["collection"]).removeprefix("state.").removeprefix("$state.")
        if clean == expr or clean.startswith(root + ".") or clean.split(".")[0] == root:
            matches.append(contract)
    return matches[0] if len(matches)==1 else None


def _canonical_table_cell(row: dict[str, Any], key: str, state: dict[str, Any], contract: dict[str, Any] | None = None, row_index: int = 0) -> Any:
    """Render only the column value declared by the table contract.

    The runtime does not infer Risk-Factor-specific aliases, test-id prefixes or
    semantic fallback columns.  A playbook that needs compatibility mappings must
    declare them in its table/binding contract.
    """
    if row.get(key) not in (None, ""):
        return row.get(key)
    if key in {"manual_check", "automation", "e2e_check"}:
        return "NOT_APPLICABLE"
    if key == "notes":
        return row.get("rationale") or "NOT_APPLICABLE"
    return "NOT_APPLICABLE"

def _render_contract_table(value: Any, contract: dict[str, Any], state: dict[str, Any]) -> str:
    rows = value.get("rows") if isinstance(value, dict) else value
    if not isinstance(rows, list):
        rows=[]
    columns=contract["columns"]
    labels=[c["label"] for c in columns]
    def cell(v: Any) -> str:
        if isinstance(v,(dict,list)):
            v=json.dumps(v,ensure_ascii=False,separators=(",",":"))
        if v in (None,""):
            v="NOT_APPLICABLE"
        return str(v).replace("|","\\|").replace("\n"," ")
    lines=["| " + " | ".join(labels) + " |", "| " + " | ".join("---" for _ in labels) + " |"]
    for row_index,row in enumerate(rows):
        if not isinstance(row,dict):
            row={"value":row}
        vals=[]
        for c in columns:
            v=_canonical_table_cell(row,c["key"],state,contract,row_index)
            vals.append(cell(v))
        lines.append("| " + " | ".join(vals) + " |")
    if len(lines)==2:
        lines.append("| " + " | ".join("NOT_APPLICABLE" for _ in labels) + " |")
    return "\n".join(lines)


def _align_rendered_document_to_validator_contract(rendered: str, record: dict[str, Any] | None = None) -> tuple[str, list[dict[str,str]]]:
    """Normalize structural labels only against validators downstream of this document.

    Validator discovery follows canonical graph routes for a bounded number of hops,
    so unrelated validators in the package cannot rewrite another artifact.
    """
    resources=_active_playbook_package().get("resources") if isinstance(_active_playbook_package().get("resources"),dict) else {}
    source=_active_playbook_package().get("source") if isinstance(_active_playbook_package().get("source"),dict) else {}
    allowed_validator_refs:set[str]=set()
    if isinstance(record,dict) and record.get("id"):
        frontier=[str(record.get("id"))]; seen=set()
        for _ in range(4):
            nxt=[]
            for rid in frontier:
                if rid in seen: continue
                seen.add(rid)
                rec,kind=_record_by_id(source,rid)
                if not isinstance(rec,dict): continue
                if rec.get("validator"):
                    allowed_validator_refs.add(str(rec.get("validator")))
                for route in _live_routes(rec,kind or "node"):
                    target=route.get("target")
                    if target and target not in seen: nxt.append(str(target))
            frontier=nxt
    candidates=[]
    for name,text in resources.items():
        if not str(name).lower().endswith(".py") or "REQUIRED_HEADINGS" not in text:
            continue
        if allowed_validator_refs:
            normalized=str(name).lstrip("./")
            if not any(normalized==ref.lstrip("./") or normalized.endswith("/"+ref.lstrip("./")) for ref in allowed_validator_refs):
                continue
        try:
            import ast
            tree=ast.parse(text); vals={}
            for node in tree.body:
                if isinstance(node,ast.Assign) and len(node.targets)==1 and isinstance(node.targets[0],ast.Name):
                    if node.targets[0].id in {"REQUIRED_HEADINGS","REQUIRED_IDENTITY_ROWS"}:
                        vals[node.targets[0].id]=ast.literal_eval(node.value)
            headings=vals.get("REQUIRED_HEADINGS") or []
            score=0
            for h in headings:
                hm=re.search(r"^(#{2,3})\s+(\d+(?:\.\d+)*\.)",h)
                if hm and re.search(r"^"+re.escape(hm.group(1))+r"\s+"+re.escape(hm.group(2)),rendered,re.M):
                    score+=1
            if headings and score >= max(1,len(headings)//2):
                candidates.append((score,name,vals))
        except Exception:
            continue
    if not candidates:
        return rendered,[]
    candidates.sort(reverse=True,key=lambda x:x[0])
    _,name,vals=candidates[0]
    audit=[]; out=rendered
    for required in vals.get("REQUIRED_HEADINGS") or []:
        m=re.match(r"^(#{2,3})\s+(\d+(?:\.\d+)*\.)\s+",required)
        if not m: continue
        pat=re.compile(r"^"+re.escape(m.group(1))+r"\s+"+re.escape(m.group(2))+r"\s+.*$",re.M)
        found=pat.search(out)
        if found and found.group(0)!=required:
            audit.append({"from":found.group(0),"to":required,"validator":name})
            out=out[:found.start()]+required+out[found.end():]
    for label in vals.get("REQUIRED_IDENTITY_ROWS") or []:
        exact=re.compile(r"^\|\s*"+re.escape(str(label))+r"\s*\|",re.M)
        if exact.search(out): continue
        cand=re.compile(r"^\|\s*("+re.escape(str(label))+r"[^|]*)\|",re.M)
        matches=list(cand.finditer(out))
        if len(matches)==1:
            old=matches[0].group(1).strip()
            audit.append({"from":old,"to":str(label),"validator":name})
            start,end=matches[0].span(1)
            replacement=str(label) + (" " if matches[0].group(1).endswith(" ") else "")
            out=out[:start]+replacement+out[end:]
    return out,audit


def _derive_document_contract(record: dict[str, Any], state: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str,str]]]:
    """Resolve only explicitly machine-readable derivation contracts.

    Supported generic forms:
    - a literal artifact/path-like value;
    - ``state.<path>`` or ``$state.<path>`` copy from canonical state;
    - ``{source: state.<path>, default: ...}``.

    Natural-language derivation instructions are intentionally not interpreted by
    the Editor.  They require a model node, a declared expression resolver, or a
    playbook-owned deterministic validator/transform.
    """
    contract=record.get("derivation_contract")
    if not isinstance(contract,dict):
        return {},[]
    derived: dict[str,Any]={}; audit=[]
    for target,instruction in contract.items():
        if target == "revision_policy":
            continue
        value=None; mode=None
        if isinstance(instruction, str):
            raw=instruction.strip()
            if raw.startswith("state.") or raw.startswith("$state."):
                source=raw.removeprefix("$state.").removeprefix("state.")
                value=copy.deepcopy(_state_subtree(state,source)); mode="declared-state-copy"
            elif re.fullmatch(r"[A-Za-z0-9_./-]+\.(?:md|json|yaml|yml|txt|zip)", raw):
                value=raw; mode="declared-literal"
        elif isinstance(instruction, dict) and "source" in instruction:
            src=str(instruction.get("source") or "")
            if src.startswith("state.") or src.startswith("$state."):
                source=src.removeprefix("$state.").removeprefix("state.")
                value=copy.deepcopy(_state_subtree(state,source))
                if value in (None,"",[],{}): value=copy.deepcopy(instruction.get("default"))
                mode="declared-state-copy"
        if value not in (None,""):
            _set_dotted_state(derived,str(target),value)
            audit.append({"target":str(target),"mode":str(mode)})
    return derived,audit

def _execute_document_generate(credentials: dict[str, Any], record: dict[str, Any], current_id: str, kind: str, phase: str, state: dict[str, Any], routes: list[dict[str,str]]) -> dict[str, Any]:
    template_ref=str(record.get("template") or "")
    bindings_ref=str(record.get("bindings") or "")
    output_ref=str(record.get("output") or "")
    t_name, template=_package_resource_text(template_ref)
    b_name, bindings_text=_package_resource_text(bindings_ref)
    if not template or not bindings_text or not output_ref:
        raise ValueError(f"DOCUMENT.GENERATE requires resolvable template, bindings and output ({template_ref}, {bindings_ref}, {output_ref}).")
    bindings_doc=yaml.safe_load(bindings_text)
    if not isinstance(bindings_doc, dict) or not isinstance(bindings_doc.get("bindings"), dict):
        raise ValueError("DOCUMENT.GENERATE bindings file must contain a bindings mapping.")
    derived: dict[str, Any]={}
    derive=record.get("derive_before_generate")
    if isinstance(derive, dict):
        for key,spec in derive.items():
            if isinstance(spec, str) and spec.startswith("state."):
                derived[key]=_state_subtree(state,spec)
            elif isinstance(spec, str) and (spec.startswith("AI.") or re.match(r"^[A-Z][A-Z0-9_.]+\(", spec)):
                raise ValueError(f"DOCUMENT.GENERATE has undeclared derivation resolver: {spec}. Move this rule into a playbook-owned transform/model node or declare a generic resolver contract.")
            else:
                derived[key]=copy.deepcopy(spec)
    contract_derived, derivation_audit = _derive_document_contract(record, state)
    if contract_derived:
        derived = _deep_merge_preserving_existing(derived, contract_derived)
    render_state=copy.deepcopy(state)
    render_state=_deep_merge_preserving_existing(render_state, derived)
    bound: dict[str, Any]={}
    for name,spec in bindings_doc["bindings"].items():
        _set_dotted_state(bound,str(name),_binding_value(spec, render_state))
    bound, compatibility_aliases=_document_binding_compatibility(bound, render_state)
    unresolved=[]
    missing_leaf_warnings=[]
    rendering_rules=bindings_doc.get("rendering_rules") if isinstance(bindings_doc.get("rendering_rules"),dict) else {}
    def repl(match: re.Match[str]) -> str:
        expr=match.group(1).strip()
        value=_lookup_bound_path(bound,expr)
        if value is None:
            # A missing leaf is explicit in the generated draft rather than being
            # silently invented. The warning remains machine-readable in debug.
            missing_leaf_warnings.append(expr)
            return "не надано"
        table_contract=_table_contract_for_placeholder(expr)
        if table_contract is not None and (isinstance(value,list) or (isinstance(value,dict) and isinstance(value.get("rows"),list))):
            return _render_contract_table(value,table_contract,render_state)
        return _render_markdown_value(value)
    rendered=re.sub(r"\{\{\s*([^{}]+?)\s*\}\}", repl, template)
    rendered, validator_alignment_audit = _align_rendered_document_to_validator_contract(rendered, record)
    residual=sorted(set(re.findall(r"\{\{\s*([^{}]+?)\s*\}\}", rendered)))
    if residual:
        result=_runtime_only_live_result(credentials=credentials,record=record,kind=kind,current_id=current_id,phase=phase,state=state,routes=routes,assistant_message="",await_analyst=False,selected=None,updates={},new_state=dict(state),reason="deterministic-document-generate-failed")
        result["rationale_short"]="Unresolved template placeholders: " + ", ".join(residual[:20])
        result["debug"]["runtime"]["materialization"]={"status":"failed","reason":"unresolved_placeholders","placeholders":residual,"template_resource":t_name,"bindings_resource":b_name,"output":output_ref,"compatibility_aliases":compatibility_aliases}
        return result
    workspace=_runtime_workspace(); out=workspace / output_ref
    out.parent.mkdir(parents=True,exist_ok=True); out.write_text(rendered,encoding="utf-8")
    new_state,updates=_apply_declared_state_updates(record,state,derived)
    selected=next((r for r in routes if r.get("key")=="next"), routes[0] if len(routes)==1 else None)
    message=f"Матеріалізовано документ: {output_ref} ({out.stat().st_size} bytes)."
    result=_runtime_only_live_result(credentials=credentials,record=record,kind=kind,current_id=current_id,phase=phase,state=state,routes=routes,assistant_message=message,await_analyst=False,selected=selected,updates=updates,new_state=new_state,reason="deterministic-document-generate")
    result["debug"]["runtime"]["artifact"]={"path":output_ref,"workspace_path":str(out),"size":out.stat().st_size,"sha256":hashlib.sha256(out.read_bytes()).hexdigest(),"template_resource":t_name,"bindings_resource":b_name,"compatibility_aliases":compatibility_aliases,"missing_leaf_warnings":sorted(set(missing_leaf_warnings)),"rendering_rules":rendering_rules,"derivation_audit":derivation_audit,"validator_alignment_audit":validator_alignment_audit}
    if missing_leaf_warnings:
        result["debug"]["runtime"]["materialization"]={"status":"generated_with_warnings","missing_leaf_warnings":sorted(set(missing_leaf_warnings)),"compatibility_aliases":compatibility_aliases}
    else:
        result["debug"]["runtime"]["materialization"]={"status":"generated","compatibility_aliases":compatibility_aliases}
    return result



def _execute_delivery_package_build(credentials: dict[str, Any], record: dict[str, Any], current_id: str, kind: str, phase: str, state: dict[str, Any], routes: list[dict[str,str]]) -> dict[str, Any]:
    """Build a ZIP strictly from the playbook-declared ``package`` contract.

    Packaging is generic mechanical orchestration.  The Editor does not add
    Risk-Factor/Jira/Passport artifacts that were not declared by the package.
    """
    workspace = _runtime_workspace()
    package_spec = record.get("package") if isinstance(record.get("package"), dict) else {}
    output_ref = str(package_spec.get("path") or "generated_outputs/ORDO_DELIVERY_PACKAGE.zip")
    out = workspace / output_ref
    out.parent.mkdir(parents=True, exist_ok=True)

    # Materialize runtime-owned package metadata without mutating semantic state.
    generated = workspace / "generated_outputs"
    generated.mkdir(parents=True, exist_ok=True)
    snapshot = generated / "RUNTIME_STATE_SNAPSHOT.json"
    snapshot.write_text(json.dumps(state, ensure_ascii=False, indent=2, default=str) + "\n", encoding="utf-8")
    metadata = generated / "MATERIALIZATION_METADATA.json"

    requested = package_spec.get("include") if isinstance(package_spec.get("include"), list) else []
    concrete: list[Path] = []
    missing_required: list[str] = []
    stale_required: list[dict[str,Any]] = []
    for item in requested:
        item_s = str(item or "").strip()
        if not item_s:
            continue
        if item_s == "validation reports":
            concrete.extend(sorted(generated.glob("*VALIDATION*.json")))
            concrete.extend(sorted(generated.glob("*POST_MATERIALIZATION*.json")))
            continue
        if item_s == "materialization metadata":
            concrete.append(metadata)
            continue
        if item_s == "runtime state snapshot":
            concrete.append(snapshot)
            continue
        path = workspace / item_s
        if path.is_file():
            status=_run_artifact_status(item_s)
            if status and status.get("freshness_status")=="stale":
                stale_required.append(status)
            concrete.append(path)
        else:
            missing_required.append(item_s)

    if stale_required:
        result = _runtime_only_live_result(
            credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
            state=state, routes=routes,
            assistant_message="Не вдалося сформувати пакет: один або більше оголошених артефактів застаріли після зміни залежного стану.",
            await_analyst=False, selected=None, updates={}, new_state=copy.deepcopy(state),
            reason="deterministic-delivery-package-failed",
            extra_runtime={"runtime_executor":"delivery_package_builder","stale_artifacts":stale_required},
        )
        result["run_status"]="halted"; result["completion_reason"]="delivery_package_stale_artifacts"
        result["debug"]["runtime"]["run_status"]="halted"; result["debug"]["runtime"]["completion_reason"]="delivery_package_stale_artifacts"
        return result

    if missing_required:
        result = _runtime_only_live_result(
            credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
            state=state, routes=routes,
            assistant_message="Не вдалося сформувати пакет: відсутні оголошені обов'язкові артефакти.",
            await_analyst=False, selected=None, updates={}, new_state=copy.deepcopy(state),
            reason="deterministic-delivery-package-failed",
            extra_runtime={"runtime_executor":"delivery_package_builder","missing_artifacts":sorted(set(missing_required))},
        )
        result["run_status"] = "halted"
        result["completion_reason"] = "delivery_package_missing_artifacts"
        result["debug"]["runtime"]["run_status"] = "halted"
        result["debug"]["runtime"]["completion_reason"] = "delivery_package_missing_artifacts"
        return result

    # De-duplicate while preserving stable archive names.
    unique: dict[str, Path] = {}
    for path in concrete:
        if path.is_file():
            try:
                arc = path.relative_to(workspace).as_posix()
            except ValueError:
                arc = path.name
            unique[arc] = path

    metadata_doc = {
        "format": "ordo.delivery_package",
        "format_version": "1",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "node_id": current_id,
        "source_action": str(record.get("action") or ""),
        "included_files": sorted(unique.keys()),
        "semantic_state_mutated": False,
    }
    metadata.write_text(json.dumps(metadata_doc, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    unique[metadata.relative_to(workspace).as_posix()] = metadata
    unique[snapshot.relative_to(workspace).as_posix()] = snapshot

    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for arc, path in sorted(unique.items()):
            archive.write(path, arcname=arc)
    digest = hashlib.sha256(out.read_bytes()).hexdigest()
    selected = next((r for r in routes if r.get("key") == "next"), routes[0] if len(routes) == 1 else None)
    result = _runtime_only_live_result(
        credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
        state=state, routes=routes,
        assistant_message=f"Сформовано пакет: {output_ref} ({out.stat().st_size} bytes).",
        await_analyst=False, selected=selected, updates={}, new_state=copy.deepcopy(state),
        reason="deterministic-delivery-package-build",
        extra_runtime={"runtime_executor":"delivery_package_builder","package_entries":sorted(unique.keys())},
    )
    result["debug"]["runtime"]["artifact"] = {
        "path": output_ref,
        "filename": out.name,
        "workspace_path": str(out),
        "size": out.stat().st_size,
        "sha256": digest,
        "warning_count": 0,
    }
    return result

def _declared_test_coverage_requirements(record: dict[str, Any]) -> list[str]:
    """Return coverage IDs declared by the playbook, without domain inference."""
    explicit = record.get("coverage_requirements")
    if isinstance(explicit, list):
        return [str(x).strip() for x in explicit if str(x).strip()]
    registry = record.get("coverage_registry")
    if isinstance(registry, dict):
        return [str(k).strip() for k,v in registry.items() if str(k).strip() and (not isinstance(v,dict) or v.get("required",True))]
    return []

def _generic_coverage_rows(state: dict[str, Any], record: dict[str, Any]) -> list[dict[str, Any]]:
    """Collect rows only from playbook-declared catalogue paths.

    No Risk-Factor catalogue names or outcome enums are embedded in the editor.
    """
    paths = record.get("coverage_catalog_paths")
    if not isinstance(paths, list):
        # Generic fallback: inspect state values that expose a rows collection and whose
        # rows explicitly carry a `covers` field. This is structural, not domain semantic.
        candidates=[]
        for key,value in state.items():
            if isinstance(value,dict) and isinstance(value.get("rows"),list):
                if any(isinstance(r,dict) and "covers" in r for r in value["rows"]):
                    candidates.append(str(key)+".rows")
        paths=candidates
    rows=[]
    for path in paths:
        value=_state_subtree(state,str(path))
        if isinstance(value,list):
            rows.extend(r for r in value if isinstance(r,dict))
    return rows

def _evaluate_test_coverage_gate(record: dict[str, Any], state: dict[str, Any]) -> tuple[str,str,dict[str,Any]] | None:
    requirements=_declared_test_coverage_requirements(record)
    if not requirements:
        return None
    rows=_generic_coverage_rows(state,record)
    ids=[str(r.get("tc_id") or "").strip() for r in rows]
    dup=sorted({x for x in ids if x and ids.count(x)>1})
    valid=[]
    invalid=[]
    overcredited=[]
    state_conflicts=[]
    allowed=set(requirements)
    unknown=set()
    policy=record.get("coverage_policy") if isinstance(record.get("coverage_policy"),dict) else {}
    max_ids=int(policy.get("max_coverage_ids_per_row") or record.get("coverage_max_ids_per_row") or 2)
    registry=record.get("coverage_registry") if isinstance(record.get("coverage_registry"),dict) else {}
    for row in rows:
        tc=str(row.get("tc_id") or "").strip()
        covers=row.get("covers")
        required_text_fields=("scenario","short_input","expected_result")
        missing_text=[k for k in required_text_fields if not str(row.get(k) or "").strip()]
        if not tc or not isinstance(covers,list) or not covers or missing_text:
            invalid.append({"tc_id":tc or "<missing-tc-id>","missing_fields":missing_text + ([] if isinstance(covers,list) and covers else ["covers"])})
            continue
        normalized=list(dict.fromkeys(str(x).strip() for x in covers if str(x).strip()))
        unknown.update(x for x in normalized if x not in allowed)
        if len(normalized)>max_ids:
            overcredited.append({"tc_id":tc,"coverage_ids":normalized,"max_allowed":max_ids})
        expected_state=str(row.get("expected_state") or "").strip()
        for req in normalized:
            spec=registry.get(req) if isinstance(registry.get(req),dict) else {}
            allowed_states=spec.get("expected_states") if isinstance(spec,dict) else None
            if isinstance(allowed_states,list) and expected_state and expected_state not in {str(x) for x in allowed_states}:
                state_conflicts.append({"tc_id":tc,"coverage_id":req,"expected_state":expected_state,"allowed_expected_states":[str(x) for x in allowed_states]})
        valid.append((tc,normalized))
    credited={req:[] for req in requirements}
    if not dup and not overcredited and not state_conflicts:
        for tc,covers in valid:
            for req in covers:
                if req in credited:
                    credited[req].append(tc)
    missing=[req for req in requirements if not credited[req]]
    extra={
        "coverage_requirements":{req:bool(credited[req]) for req in requirements},
        "coverage_details":{req:{"row_ids":credited[req],"source":"explicit_covers" if credited[req] else "none"} for req in requirements},
        "missing_coverage":missing,
        "duplicate_test_ids":dup,
        "invalid_coverage_rows":invalid,
        "overcredited_rows":overcredited,
        "expected_state_conflicts":state_conflicts,
        "unknown_coverage_ids":sorted(unknown),
        "max_coverage_ids_per_row":max_ids,
        "coverage_source":"playbook_declared_explicit_covers_with_evidence_shape",
        "acceptance_eligible":not missing and not dup and not invalid and not unknown and not overcredited and not state_conflicts,
    }
    if dup:
        return "fail","duplicate test IDs prevent structural coverage credit",extra
    if invalid:
        return "fail","coverage rows require tc_id, covers, scenario, short_input and expected_result",extra
    if unknown:
        return "fail","test rows declare coverage IDs outside the playbook registry: "+", ".join(sorted(unknown)),extra
    if overcredited:
        return "fail",f"one test row may credit at most {max_ids} independent coverage IDs",extra
    if state_conflicts:
        return "fail","coverage registry expected-state contract conflicts with test evidence",extra
    if missing:
        return "fail","missing declared structural test coverage: "+", ".join(missing),extra
    return "pass","all playbook-declared structural coverage IDs have explicit valid test evidence",extra



def _gate_condition_literal(value: str) -> Any:
    """Parse a small, deterministic literal subset used by mechanical gate conditions."""
    text=str(value or "").strip()
    low=text.lower()
    if low in {"true","yes"}: return True
    if low in {"false","no"}: return False
    if low in {"null","none"}: return None
    if (len(text) >= 2 and text[0] == text[-1] and text[0] in {"'", '"'}):
        return text[1:-1]
    try:
        if re.fullmatch(r"[-+]?\d+", text): return int(text)
        if re.fullmatch(r"[-+]?(?:\d+\.\d*|\d*\.\d+)", text): return float(text)
    except Exception:
        pass
    return text


def _evaluate_mechanical_condition(condition: str, state: dict[str, Any]) -> tuple[str, str, dict[str, Any]]:
    """Evaluate the generic, fail-closed condition subset used by source gates.

    Supported clauses (combined with top-level ``and`` / ``or``):
      - ``state.path is one of A, B``
      - ``state.path is not one of A, B``
      - ``state.path is not empty`` / ``is empty``
      - ``state.path == VALUE`` / ``!= VALUE`` / ``equals VALUE``
      - ``state.path is true`` / ``is false``

    No Python ``eval`` is used.  Unknown syntax is UNRESOLVED, never guessed.
    """
    raw=" ".join(str(condition or "").strip().split())
    if not raw:
        return "unresolved", "mechanical gate condition is empty", {"condition": raw, "unsupported_condition": True}

    # The current source grammar uses simple Boolean composition without quoted
    # boolean operators.  Parse OR groups first, then AND clauses.
    or_groups=[g.strip() for g in re.split(r"\s+or\s+", raw, flags=re.I) if g.strip()]
    group_results=[]
    all_refs=[]
    missing_refs=[]
    unsupported=[]
    clause_debug=[]

    path_pat=r"state\.([A-Za-z_][A-Za-z0-9_.]*)"

    def eval_clause(clause: str):
        c=clause.strip()
        m=re.fullmatch(path_pat+r"\s+is\s+(not\s+)?one\s+of\s+(.+)", c, flags=re.I)
        if m:
            path=m.group(1); negate=bool(m.group(2)); values_raw=m.group(3).strip()
            # Canonical source uses comma-separated enum values.  Also accept a
            # pipe separator as a harmless authoring convenience.
            values=[_gate_condition_literal(x.strip()) for x in re.split(r"\s*[,|]\s*", values_raw) if x.strip()]
            if not values:
                return None,{"clause":c,"path":path,"error":"empty one-of set"}
            all_refs.append(path)
            if not _state_path_exists(state,path):
                missing_refs.append(path); return None,{"clause":c,"path":path,"missing":True}
            actual=_state_subtree(state,path)
            result=actual in values
            if negate: result=not result
            return result,{"clause":c,"path":path,"actual":copy.deepcopy(actual),"operator":"not_one_of" if negate else "one_of","expected":copy.deepcopy(values),"result":result}

        m=re.fullmatch(path_pat+r"\s+is\s+(not\s+)?empty", c, flags=re.I)
        if m:
            path=m.group(1); negate=bool(m.group(2)); all_refs.append(path)
            if not _state_path_exists(state,path):
                missing_refs.append(path); return None,{"clause":c,"path":path,"missing":True}
            actual=_state_subtree(state,path)
            empty=actual in (None,"",[],{})
            result=(not empty) if negate else empty
            return result,{"clause":c,"path":path,"actual":copy.deepcopy(actual),"operator":"not_empty" if negate else "empty","result":result}

        m=re.fullmatch(path_pat+r"\s*(==|!=|equals)\s*(.+)", c, flags=re.I)
        if m:
            path=m.group(1); op=m.group(2).lower(); expected=_gate_condition_literal(m.group(3)); all_refs.append(path)
            if not _state_path_exists(state,path):
                missing_refs.append(path); return None,{"clause":c,"path":path,"missing":True}
            actual=_state_subtree(state,path)
            result=(actual != expected) if op=="!=" else (actual == expected)
            return result,{"clause":c,"path":path,"actual":copy.deepcopy(actual),"operator":op,"expected":copy.deepcopy(expected),"result":result}

        m=re.fullmatch(path_pat+r"\s+is\s+(true|false)", c, flags=re.I)
        if m:
            path=m.group(1); expected=m.group(2).lower()=="true"; all_refs.append(path)
            if not _state_path_exists(state,path):
                missing_refs.append(path); return None,{"clause":c,"path":path,"missing":True}
            actual=_state_subtree(state,path); result=actual is expected
            return result,{"clause":c,"path":path,"actual":copy.deepcopy(actual),"operator":"is_bool","expected":expected,"result":result}

        unsupported.append(c)
        return None,{"clause":c,"unsupported":True}

    for group in or_groups:
        clauses=[c.strip() for c in re.split(r"\s+and\s+", group, flags=re.I) if c.strip()]
        values=[]
        for clause in clauses:
            val,dbg=eval_clause(clause); values.append(val); clause_debug.append(dbg)
        if any(v is None for v in values):
            group_results.append(None)
        else:
            group_results.append(all(bool(v) for v in values))

    extra={
        "condition": raw,
        "referenced_state_paths": sorted(set(all_refs)),
        "condition_clauses": clause_debug,
    }
    if unsupported:
        extra["unsupported_clauses"]=unsupported
        return "unresolved", "unsupported mechanical gate condition syntax: "+"; ".join(unsupported), extra
    if missing_refs:
        extra["missing_required_inputs"]=sorted(set(missing_refs))
        return "unresolved", "missing required inputs for gate condition: "+", ".join(sorted(set(missing_refs))), extra
    if any(v is None for v in group_results):
        return "unresolved", "mechanical gate condition could not be fully evaluated", extra
    passed=any(bool(v) for v in group_results)
    extra["condition_result"]=passed
    return ("pass" if passed else "fail"), ("condition evaluated true" if passed else "condition evaluated false"), extra



def _compiled_artifact_validation_adapter(semantic_element: dict[str, Any] | None) -> dict[str, Any] | None:
    if not isinstance(semantic_element, dict):
        return None
    adapter = semantic_element.get("execution_adapter")
    if not isinstance(adapter, dict) or str(adapter.get("runtime_executor") or "") != "artifact_validation":
        return None
    spec = adapter.get("artifact_validation")
    return spec if isinstance(spec, dict) else None


def _ensure_runtime_package_workspace() -> Path:
    package=_active_playbook_package()
    raw_zip=package.get("raw_zip") if isinstance(package,dict) else None
    if not isinstance(raw_zip,(bytes,bytearray)):
        raise ValueError("Deterministic package validation requires the loaded source ZIP bytes.")
    run_ws=_runtime_workspace()
    package_root=run_ws/"package_tool_workspace"
    marker=package_root/".ordo_extracted"
    if not marker.exists():
        if package_root.exists():
            shutil.rmtree(package_root)
        package_root.mkdir(parents=True,exist_ok=True)
        with zipfile.ZipFile(io.BytesIO(bytes(raw_zip))) as archive:
            infos=[i for i in archive.infolist() if not i.is_dir()]
            if any(not _safe_zip_name(i.filename) for i in infos):
                raise ValueError("Package validation source archive contains unsafe paths.")
            archive.extractall(package_root)
        marker.write_text("ok\n",encoding="utf-8")
    return package_root


def _artifact_validation_decision(spec: dict[str, Any], state: dict[str, Any]) -> tuple[str, str, dict[str, Any]]:
    """Execute only compiled deterministic artifact/archive checks; never infer semantics from prose."""
    checks: list[dict[str, Any]]=[]
    def add(cid: str, ok: bool, evidence: Any=None, detail: str | None=None):
        row={"check_id":cid,"status":"pass" if ok else "fail"}
        if evidence not in (None, "", [], {}): row["evidence"]=copy.deepcopy(evidence)
        if detail: row["detail"]=detail
        checks.append(row)
        return ok

    state_path=str(spec.get("state_path") or "")
    expected=str(spec.get("expected_path") or "").replace("\\", "/")
    actual=_state_subtree(state,state_path) if state_path else None
    if not add("artifact_state_path_populated", actual not in (None,"",[],{}), {"state_path":state_path,"value":actual}):
        return "fail", f"artifact state path is not populated: {state_path}", {"checks":checks,"failed_checks":[checks[-1]],"artifact_validation":True}
    actual_norm=str(actual).replace("\\", "/")
    if not add("artifact_expected_path_matches_state", actual_norm==expected, {"expected":expected,"actual":actual_norm}):
        return "fail", "materialized artifact path does not match declared expected_path", {"checks":checks,"failed_checks":[checks[-1]],"artifact_validation":True}

    run_ws=_runtime_workspace().resolve()
    rel=Path(expected)
    if rel.is_absolute() or ".." in rel.parts:
        return "unresolved","compiled artifact validation expected_path is unsafe",{"checks":checks,"execution_error":True,"artifact_validation":True}
    target=(run_ws/rel).resolve()
    if run_ws not in (target,*target.parents):
        return "unresolved","compiled artifact validation escaped runtime workspace",{"checks":checks,"execution_error":True,"artifact_validation":True}
    exists=target.exists(); add("artifact_exists",exists,expected)
    regular=target.is_file() if exists else False; add("artifact_regular_file",regular,expected)
    non_empty=(target.stat().st_size>0) if regular else False; add("artifact_non_empty",non_empty,{"path":expected,"size":target.stat().st_size if regular else 0})
    if not (exists and regular and non_empty):
        failed=[c for c in checks if c["status"]=="fail"]
        return "fail","artifact filesystem validation failed",{"checks":checks,"failed_checks":failed,"artifact_validation":True,"artifact_path":expected}

    is_archive=str(spec.get("output_type") or "").lower() in {"archive","package","zip"} or expected.lower().endswith(".zip")
    if is_archive:
        readable=zipfile.is_zipfile(target); add("archive_readable",readable,expected)
        if not readable:
            return "fail","archive is not a readable ZIP",{"checks":checks,"failed_checks":[checks[-1]],"artifact_validation":True,"artifact_path":expected}
        try:
            with zipfile.ZipFile(target) as zf:
                names=[i.filename for i in zf.infolist() if not i.is_dir()]
                safe=all(_safe_zip_name(n) for n in names); add("archive_member_paths_safe",safe,names)
                corrupt=zf.testzip(); add("archive_crc_integrity",corrupt is None, {"first_bad_member":corrupt})
                required=[str(x) for x in (spec.get("required_members") or [])]
                if required:
                    missing=sorted(set(required)-set(names)); add("archive_required_members",not missing,{"required":required,"missing":missing})
                forbidden=[str(x) for x in (spec.get("forbidden_members") or [])]
                if forbidden:
                    present=sorted(set(forbidden)&set(names)); add("archive_forbidden_members",not present,{"forbidden":forbidden,"present":present})
                hashes=spec.get("member_hashes") if isinstance(spec.get("member_hashes"),dict) else {}
                for member,expected_hash in hashes.items():
                    ok=False; actual_hash=None
                    if member in names:
                        actual_hash=hashlib.sha256(zf.read(member)).hexdigest(); ok=actual_hash==str(expected_hash)
                    add(f"archive_member_sha256:{member}",ok,{"expected":str(expected_hash),"actual":actual_hash})
        except Exception as exc:
            return "unresolved",f"archive validation execution failed: {exc}",{"checks":checks,"execution_error":True,"artifact_validation":True}
        archive_hash=spec.get("archive_sha256")
        if archive_hash:
            actual_hash=hashlib.sha256(target.read_bytes()).hexdigest(); add("archive_sha256",actual_hash==str(archive_hash),{"expected":str(archive_hash),"actual":actual_hash})

    validators=[str(x) for x in (spec.get("validators") or []) if str(x).strip()]
    validator_runs=[]
    if validators:
        try:
            package_root=_ensure_runtime_package_workspace()
            for ref in validators:
                if not ref.lower().endswith('.py'):
                    return "unresolved",f"unsupported deterministic validator type: {ref}",{"checks":checks,"execution_error":True,"artifact_validation":True}
                base, validator_file=_package_tool_base_and_path(package_root,ref)
                cp=subprocess.run([sys.executable,str(validator_file)],cwd=str(base),capture_output=True,text=True,timeout=60)
                payload=None
                try:
                    parsed=json.loads((cp.stdout or '').strip())
                    if isinstance(parsed,dict): payload=parsed
                except Exception:
                    pass
                declared_status = str(payload.get('status') or '').upper() if isinstance(payload,dict) else ''
                ok=cp.returncode==0 and (not isinstance(payload,dict) or declared_status in {'','PASS','PASSED','OK','VALID'})
                validator_runs.append({"validator":ref,"exit_code":cp.returncode,"status":(payload or {}).get('status') if isinstance(payload,dict) else None,"stdout":(cp.stdout or '')[-1200:],"stderr":(cp.stderr or '')[-1200:]})
                add(f"validator:{ref}",ok,validator_runs[-1])
                if not ok:
                    # A validator may report a normal deterministic FAIL as structured evidence.
                    # A non-zero process with no machine-readable FAIL envelope is a validator
                    # execution error (crash/import/runtime failure), not a business/gate failure.
                    if cp.returncode != 0 and declared_status not in {'FAIL','FAILED','INVALID','BLOCKED'}:
                        return "unresolved",f"package-local validator execution failed: {ref}",{
                            "checks":checks,"validator_runs":validator_runs,"artifact_validation":True,"execution_error":True,
                            "validator_execution_error":{"validator":ref,"exit_code":cp.returncode,"stderr":(cp.stderr or '')[-1200:]},
                        }
                    return "fail",f"package-local validator failed: {ref}",{"checks":checks,"failed_checks":[checks[-1]],"validator_runs":validator_runs,"artifact_validation":True}
        except subprocess.TimeoutExpired as exc:
            return "unresolved",f"validator execution timed out: {exc}",{"checks":checks,"execution_error":True,"artifact_validation":True}
        except Exception as exc:
            return "unresolved",f"validator execution failed: {exc}",{"checks":checks,"execution_error":True,"artifact_validation":True}

    failed=[c for c in checks if c["status"]=="fail"]
    if failed:
        return "fail","deterministic artifact/archive validation failed",{"checks":checks,"failed_checks":failed,"validator_runs":validator_runs,"artifact_validation":True}
    contract_gaps=[str(x) for x in (spec.get("validation_contract_gaps") or []) if str(x).strip()]
    if contract_gaps:
        return "unresolved","generated-playbook artifact/archive validation contract is incomplete",{
            "checks":checks,"validator_runs":validator_runs,"artifact_validation":True,"artifact_path":expected,
            "execution_class":"deterministic_artifact_validation","profile_contract_gap":True,
            "validation_contract_gaps":contract_gaps,
        }
    return "pass","deterministic artifact/archive validation passed",{"checks":checks,"validator_runs":validator_runs,"artifact_validation":True,"artifact_path":expected,"execution_class":"deterministic_artifact_validation"}

def _deterministic_gate_decision(record: dict[str, Any], state: dict[str, Any], semantic_element: dict[str, Any] | None = None) -> tuple[str, str, dict[str, Any]]:
    """Return PASS / FAIL / UNRESOLVED for a deterministic gate.

    Missing required input is not the same thing as a false condition.  A gate may
    only FAIL when it had enough data to evaluate and the criterion was false.
    Missing producer data is UNRESOLVED and must halt/fallback rather than silently
    taking on_fail, which otherwise creates artificial retry loops.
    """
    gid=str(record.get("id") or "")
    workspace=_runtime_workspace()
    artifact_adapter=_compiled_artifact_validation_adapter(semantic_element)
    if artifact_adapter is not None:
        return _artifact_validation_decision(artifact_adapter,state)
    gate_contract = semantic_element.get("gate_contract") if isinstance(semantic_element, dict) and isinstance(semantic_element.get("gate_contract"), dict) else {}
    assertion = str(gate_contract.get("assert") or record.get("assert") or "").strip().upper()
    assertion_source = str(gate_contract.get("source") or record.get("source") or "").strip()
    if assertion == "FIELD_PRESENT" and assertion_source:
        path = assertion_source[6:] if assertion_source.startswith("state.") else assertion_source
        exists = bool(path) and _state_path_exists(state, path)
        value = _state_subtree(state, path) if exists else None
        present = exists and value not in (None, "", [], {})
        checks=[{"check_id":"FIELD_PRESENT","status":"pass" if present else "fail","evidence":{"source":assertion_source,"value":copy.deepcopy(value)}}]
        return ("pass" if present else "fail"), ("canonical FIELD_PRESENT assertion passed" if present else "canonical FIELD_PRESENT assertion failed"), {"checks":checks,"failed_checks":[] if present else checks,"canonical_assertion":True}
    required=record.get("required_inputs")
    if isinstance(required,list):
        missing=[str(x) for x in required if not _state_path_exists(state,str(x))]
        if missing:
            return "unresolved", "missing required inputs: "+", ".join(missing), {"missing_required_inputs": missing}
        empty=[str(x) for x in required if _state_subtree(state,str(x)) in (None,"",[],{})]
        if empty:
            return "fail", "required inputs are empty: "+", ".join(empty), {"empty_required_inputs": empty}
        return "pass", "required inputs present", {}
    coverage_result=_evaluate_test_coverage_gate(record,state)
    if coverage_result is not None:
        return coverage_result
    if record.get("validator"):
        validator_ref=str(record.get("validator")); v_name,v_text=_package_resource_text(validator_ref)
        if not v_text: return "unresolved",f"validator resource not found: {validator_ref}",{}
        vpath=workspace/validator_ref; vpath.parent.mkdir(parents=True,exist_ok=True); vpath.write_text(v_text,encoding="utf-8")
        command_value = record.get("command")
        if command_value:
            argv = shlex.split(str(command_value))
            # KF-013: validator commands are package contracts, but the Python
            # interpreter is runtime-owned.  Do not depend on a PATH alias named
            # `python`/`python3`; use the interpreter selected for the editor.
            if argv and argv[0] in {"python", "python3"}:
                argv[0] = os.environ.get("ORDO_PYTHON") or sys.executable
        else:
            # Generic validator-gate contract: when no explicit command is declared,
            # infer the artifact from the sole incoming DOCUMENT.GENERATE producer.
            # A validator without its document argument is not an executable gate.
            source = _active_playbook_package().get("source") if isinstance(_active_playbook_package().get("source"), dict) else {}
            incoming = record.get("incoming_from") if isinstance(record.get("incoming_from"), list) else []
            candidate_outputs: list[str] = []
            for producer_id in incoming:
                producer, producer_kind = _record_by_id(source, str(producer_id))
                if producer_kind == "node" and isinstance(producer, dict) and str(producer.get("action") or "").upper() == "DOCUMENT.GENERATE":
                    output_ref = str(producer.get("output") or (producer.get("artifact") or {}).get("expected_path") or "").strip()
                    if output_ref:
                        candidate_outputs.append(output_ref)
            candidate_outputs = sorted(set(candidate_outputs))
            if len(candidate_outputs) != 1:
                return "unresolved", (
                    "validator command is absent and artifact inference is ambiguous: "
                    + (", ".join(candidate_outputs) if candidate_outputs else "no incoming DOCUMENT.GENERATE output")
                ), {"candidate_validator_artifacts": candidate_outputs}
            argv = [sys.executable, validator_ref, candidate_outputs[0]]
            report_candidates: list[str] = []
            for update_map_name in ("on_pass_update_state", "on_fail_update_state"):
                update_map = record.get(update_map_name)
                if not isinstance(update_map, dict):
                    continue
                for update_key, update_value in update_map.items():
                    if "validation_report_path" in str(update_key) and isinstance(update_value, str) and update_value.strip() and not update_value.startswith("$"):
                        report_candidates.append(update_value.strip())
            report_candidates = sorted(set(report_candidates))
            if len(report_candidates) == 1:
                argv.extend(["--report", report_candidates[0]])
        # R2 validator-state contract: every deterministic validator gets a
        # read-only canonical runtime-state snapshot in the workspace. Playbooks
        # may opt into consuming it via e.g. `--state runtime_state.json`.
        # This is generic validator plumbing; validators remain responsible for
        # declaring which state invariants they enforce.
        runtime_state_snapshot = workspace / "runtime_state.json"
        runtime_state_snapshot.write_text(
            json.dumps(state, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        try:
            cp=subprocess.run(argv,cwd=workspace,text=True,capture_output=True,timeout=20)
        except Exception as exc:
            return "unresolved",f"validator execution failed: {exc}",{}
        extra: dict[str, Any] = {}
        # KF-021: deterministic validators may publish structured recovery metadata.
        # Prefer the explicit --report artifact; fall back to stdout when it is JSON.
        validator_payload = None
        report_path = None
        if "--report" in argv:
            try:
                report_path = workspace / argv[argv.index("--report") + 1]
            except Exception:
                report_path = None
        candidates = []
        if report_path is not None and report_path.is_file():
            candidates.append(report_path.read_text(encoding="utf-8"))
        if cp.stdout.strip():
            candidates.append(cp.stdout)
        for raw_candidate in candidates:
            try:
                parsed_candidate = json.loads(raw_candidate)
            except Exception:
                continue
            if isinstance(parsed_candidate, dict):
                validator_payload = parsed_candidate
                break
        if isinstance(validator_payload, dict):
            checks = validator_payload.get("checks") if isinstance(validator_payload.get("checks"), list) else []
            failed_checks = [
                {"check_id": str(item.get("id") or item.get("check_id") or "VALIDATOR"),
                 "summary": str(item.get("message") or item.get("summary") or "validator check failed"),
                 "severity": "error"}
                for item in checks if isinstance(item, dict) and str(item.get("status") or "").upper() == "FAIL"
            ]
            if failed_checks:
                extra["validator_failed_checks"] = failed_checks
            affected = validator_payload.get("affected_state")
            if isinstance(affected, list):
                extra["affected_state"] = [str(x) for x in affected if str(x).strip()]
            target = str(validator_payload.get("recommended_recovery_target") or "").strip()
            if target:
                extra["recommended_recovery_target"] = target
            extra["validator_report_status"] = str(validator_payload.get("status") or "")
        return ("pass" if cp.returncode==0 else "fail"), f"validator exit={cp.returncode}; stdout={cp.stdout[-1200:]}; stderr={cp.stderr[-1200:]}",extra
    condition=str(record.get("condition") or "").strip()
    if condition:
        return _evaluate_mechanical_condition(condition,state)
    return "unresolved","deterministic gate has no executable runtime rule",{}


def _coverage_recovery_progress(history: Any, gate_id: str, current_missing: list[str]) -> dict[str, Any]:
    """Compare deterministic coverage failures by missing-set monotonicity.

    Progress is strict set inclusion. Equal sets are a stall; a set containing any
    previously-resolved requirement is a regression. The limit lives in runtime,
    not in the external R1.15 runner.
    """
    prior=[]
    if isinstance(history,list):
        for item in history:
            if not isinstance(item,dict) or str(item.get("node_id") or "") != gate_id:
                continue
            debug=item.get("debug") if isinstance(item.get("debug"),dict) else {}
            runtime=debug.get("runtime") if isinstance(debug.get("runtime"),dict) else {}
            det=runtime.get("deterministic_gate") if isinstance(runtime.get("deterministic_gate"),dict) else {}
            missing=det.get("missing_coverage")
            if isinstance(missing,list) and det.get("result") == "fail":
                prior.append(sorted({str(x) for x in missing}))
    current=sorted({str(x) for x in current_missing})
    out={"failure_index":len(prior)+1,"current_missing":current,"previous_missing":prior[-1] if prior else None,"max_recovery_rounds":3}
    if not prior:
        out.update({"classification":"first_failure","progress_detected":None,"stop":False})
        return out
    before=set(prior[-1]); after=set(current)
    if after == before:
        out.update({"classification":"stall","progress_detected":False,"stop":True,"stop_reason":"no_progress_recovery_loop"})
        return out
    if not after < before:
        out.update({"classification":"regression","progress_detected":False,"stop":True,"stop_reason":"coverage_recovery_regression"})
        return out
    if len(prior) >= 3:
        out.update({"classification":"max_rounds_exceeded","progress_detected":True,"stop":True,"stop_reason":"coverage_recovery_round_limit"})
        return out
    out.update({"classification":"progress","progress_detected":True,"resolved_ids":sorted(before-after),"stop":False})
    return out


def _execute_deterministic_gate(credentials: dict[str, Any], record: dict[str, Any], current_id: str, state: dict[str, Any], routes: list[dict[str,str]], history: Any=None) -> dict[str, Any]:
    semantic_element=_semantic_plan_element(current_id)
    gate_result,reason,extra=_deterministic_gate_decision(record,state,semantic_element)
    selected=None
    updates={}
    new_state=copy.deepcopy(state)
    if gate_result in {"pass","fail"}:
        passed=gate_result=="pass"
        key="on_pass" if passed else "on_fail"
        selected=next((r for r in routes if r.get("key")==key),None)
        updates_spec=record.get("on_pass_update_state" if passed else "on_fail_update_state")
        updates=copy.deepcopy(updates_spec) if isinstance(updates_spec,dict) else {}
        for update_key, update_value in list(updates.items()):
            if update_value == "$gate.failure_reason": updates[update_key]=None if passed else reason
            elif update_value == "$gate.failed_checks": updates[update_key]=[] if passed else [reason]
            elif update_value == "$gate.missing_coverage": updates[update_key]=copy.deepcopy(extra.get("missing_coverage") or [])
            elif update_value == "$gate.coverage_requirements": updates[update_key]=copy.deepcopy(extra.get("coverage_requirements") or {})
        # KF-016: coverage bookkeeping is runtime-owned derived state.  The gate's
        # mechanical catalogue evaluation is authoritative, so stale booleans must
        # never survive a later successful recovery round.
        if isinstance(extra, dict) and isinstance(extra.get("coverage_requirements"), dict):
            updates["test_coverage_requirements"] = copy.deepcopy(extra.get("coverage_requirements") or {})
            updates["missing_test_coverage"] = copy.deepcopy(extra.get("missing_coverage") or [])
        new_state=_apply_state_updates_canonical(new_state, updates)
    coverage_progress=None
    if gate_result == "fail" and isinstance(extra.get("missing_coverage") if isinstance(extra,dict) else None,list):
        coverage_progress=_coverage_recovery_progress(history,current_id,extra.get("missing_coverage") or [])
        if coverage_progress.get("stop"):
            selected=None
    result=_runtime_only_live_result(credentials=credentials,record=record,kind="gate",current_id=current_id,phase="enter",state=state,routes=routes,assistant_message="",await_analyst=False,selected=selected,updates=updates,new_state=new_state,reason="deterministic-gate")
    result["rationale_short"]=reason
    result["debug"]["runtime"]["deterministic_gate"]={"result":gate_result,"passed": True if gate_result=="pass" else False if gate_result=="fail" else None,"reason":reason,**extra}
    result["debug"]["runtime"]["mechanical_model_calls"] = 0
    if coverage_progress is not None:
        result["debug"]["runtime"]["coverage_recovery_progress"]=copy.deepcopy(coverage_progress)
        if coverage_progress.get("stop"):
            result["run_status"]="halted"
            result["completion_reason"]=str(coverage_progress.get("stop_reason") or "no_progress_recovery_loop")
            result["failure_class"]="technical_stop"
            result["debug"]["runtime"]["run_status"]="halted"
            result["debug"]["runtime"]["completion_reason"]=result["completion_reason"]
    result["debug"]["runtime"]["normalized_execution_result"]={
        "element_id": current_id, "phase": "enter", "status": "completed" if gate_result in {"pass","fail"} else "unresolved",
        "gate_result": gate_result, "state_updates": copy.deepcopy(updates),
        "route_key": selected.get("key") if selected else None, "next_id": selected.get("target") if selected else None,
    }
    alpha20_failure = None
    if gate_result == "fail":
        missing = extra.get("missing_required_inputs") if isinstance(extra, dict) else None
        missing_coverage = extra.get("missing_coverage") if isinstance(extra, dict) else None
        structured_failed = extra.get("validator_failed_checks") if isinstance(extra, dict) else None
        structured_affected = extra.get("affected_state") if isinstance(extra, dict) else None
        alpha20_failure = normalize_gate_failure(
            current_id,
            failed_checks=structured_failed if isinstance(structured_failed, list) and structured_failed else [{"check_id": current_id, "summary": reason, "severity": "error"}],
            missing_information=[{"path": str(path), "needed": "required gate input", "why_needed": reason} for path in (missing if isinstance(missing, list) else [])],
            missing_coverage=missing_coverage if isinstance(missing_coverage, list) else [],
            affected_state=(structured_affected if isinstance(structured_affected, list) and structured_affected else (["functional_test_catalog", "unit_test_catalog", "edge_case_catalog"] if isinstance(missing_coverage, list) and missing_coverage else (missing if isinstance(missing, list) else []))),
            evidence=[{"source": "deterministic_gate", "fact": reason}],
            suggested_recovery_scope=("local" if (isinstance(structured_affected, list) and structured_affected) or (isinstance(missing_coverage, list) and missing_coverage) else ("single_node" if isinstance(missing, list) and len(missing) == 1 else "unknown")),
        )
        if isinstance(extra, dict) and extra.get("recommended_recovery_target"):
            alpha20_failure["recommended_target"] = str(extra.get("recommended_recovery_target"))
        if coverage_progress is not None and coverage_progress.get("stop"):
            alpha20_failure["failure_class"]="technical_stop"
            alpha20_failure["technical_stop_reason"]=str(coverage_progress.get("stop_reason") or "no_progress_recovery_loop")
    result["debug"]["alpha20"] = {
        "contract_version": "alpha.20.0.34",
        "state_patch": None,
        "state_patch_validation": "not_applicable_deterministic_gate",
        "state_patch_commit": "not_applicable_deterministic_gate",
        "gate_failure": alpha20_failure,
    }
    if gate_result=="unresolved":
        result["run_status"]="halted"
        result["completion_reason"]="unresolved_required_input"
        result["debug"]["runtime"]["run_status"]="halted"
        result["debug"]["runtime"]["completion_reason"]="unresolved_required_input"
    return result


def _compiled_element_for(current_id: str, phase: str) -> tuple[dict[str, Any] | None, dict[str, Any] | None]:
    """Return canonical V3-schema (element, phase_spec) for enter/respond only."""
    plan = _active_playbook_package().get("compiled_plan")
    status = _active_playbook_package().get("compiled_plan_status")
    if not isinstance(plan, dict) or not isinstance(status, dict) or not status.get("valid"):
        return None, None
    elements = plan.get("elements")
    if not isinstance(elements, dict):
        return None, None
    element = elements.get(current_id)
    if not isinstance(element, dict) or str(element.get("id") or "") != current_id:
        return None, None
    phases = element.get("phases")
    if not isinstance(phases, dict):
        return element, None
    phase_spec = phases.get(phase)
    return element, phase_spec if isinstance(phase_spec, dict) else None


def _compiled_prompt_text(element: dict[str, Any], phase_spec: dict[str, Any]) -> str:
    value = phase_spec.get("prompt") if isinstance(phase_spec, dict) else None
    return value.strip() if isinstance(value, str) and value.strip() else ""


def _compiled_list(element: dict[str, Any], phase_spec: dict[str, Any], key: str) -> list[str]:
    value = phase_spec.get(key) if isinstance(phase_spec, dict) else None
    if not isinstance(value, list):
        return []
    return [str(item).removeprefix("state.").removeprefix("$state.") for item in value if isinstance(item, str) and item.strip()]


def _compiled_route_map(phase_spec: dict[str, Any]) -> dict[str, str]:
    value = phase_spec.get("allowed_routes") if isinstance(phase_spec, dict) else None
    if not isinstance(value, dict):
        return {}
    return {str(k): str(v) for k, v in value.items() if isinstance(k, str) and k and isinstance(v, str) and v}


def _compiled_route_allowlist(phase_spec: dict[str, Any]) -> list[str]:
    contract = phase_spec.get("output_contract") if isinstance(phase_spec, dict) else None
    value = contract.get("route_key_allowlist") if isinstance(contract, dict) else None
    return [str(item) for item in value if isinstance(item, str) and item] if isinstance(value, list) else []


def _compiled_state_update_allowlist(element: dict[str, Any], phase_spec: dict[str, Any]) -> list[str]:
    contract = phase_spec.get("output_contract") if isinstance(phase_spec, dict) else None
    value = contract.get("state_update_allowlist") if isinstance(contract, dict) else None
    if not isinstance(value, list):
        value = element.get("state_update_allowlist") if isinstance(element, dict) else None
    return [str(item) for item in value if isinstance(item, str) and item] if isinstance(value, list) else []


def _alpha20_write_allowlist(record: dict[str, Any], compiled_element: dict[str, Any] | None = None, compiled_phase: dict[str, Any] | None = None, semantic_element: dict[str, Any] | None = None) -> list[str]:
    """Return authoritative write surface without widening compiled contracts.

    V7 semantic plan wins when present. V6 phase contract wins when active.
    YAML declarations are consulted only for YAML fallback/direct runtime execution.
    """
    if isinstance(semantic_element, dict):
        writes = ((semantic_element.get("state_contract") or {}).get("writes") if isinstance(semantic_element.get("state_contract"), dict) else None)
        return sorted({str(x).removeprefix("state.").removeprefix("$state.") for x in (writes or []) if isinstance(x, str) and x.strip()})
    if isinstance(compiled_element, dict) and isinstance(compiled_phase, dict):
        return sorted(set(_compiled_state_update_allowlist(compiled_element, compiled_phase)))
    return sorted(_shared_declared_writes(record).keys())

def _compiled_state_defaults(phase_spec: dict[str, Any]) -> dict[str, Any]:
    """Return compiler-declared deterministic state defaults (Compiler Kit V6.3+)."""
    value = phase_spec.get("state_defaults") if isinstance(phase_spec, dict) else None
    if not isinstance(value, dict):
        return {}
    return {
        str(path).removeprefix("state.").removeprefix("$state."): copy.deepcopy(default)
        for path, default in value.items()
        if isinstance(path, str) and path.strip()
    }


def _apply_compiled_state_defaults(projected: Any, defaults: dict[str, Any]) -> Any:
    """Materialize deterministic defaults into the compiled runtime_state projection.

    Defaults are only used when the projected path is absent. They never overwrite
    a real runtime value. This keeps V6.3 default-satisfiable dependencies usable
    without turning them back into hard runtime requirements.
    """
    result = _canonicalize_runtime_state(projected)
    for path, value in defaults.items():
        if not _state_path_exists(result, path):
            _assign_nested_path(result, path, value)
    return _bounded_json_value(result, 120000)


def _fallback_runtime_state(state: Any) -> Any:
    """Provide YAML fallback with the full canonical state, not a dependency guess.

    Compiled rejection is a safety event. Falling back must not compound a compiler
    problem by hiding valid runtime evidence from the YAML execution path.
    """
    if not isinstance(state, dict):
        return _bounded_json_value(state, 120000)
    return _bounded_json_value(_canonicalize_runtime_state(state), 120000)


def _compiled_element_validation(
    *,
    record: dict[str, Any],
    model_record: dict[str, Any],
    kind: str,
    phase: str,
    state: dict[str, Any],
    routes: list[dict[str, str]],
    element: dict[str, Any],
    phase_spec: dict[str, Any],
) -> tuple[bool, list[str]]:
    """Strict compiled-plan compatibility check; reject to YAML fallback on any mismatch.

    Route validation is phase-aware: a human-assisted node enter phase may deliberately
    expose no route and wait for analyst input, while respond owns the actual YAML route.
    """
    reasons: list[str] = []

    expected_kind = "model_gate" if kind == "gate" else "llm_node"
    if str(element.get("kind") or "") != expected_kind:
        reasons.append(f"compiled kind mismatch: compiled={element.get('kind')!r}, runtime={expected_kind!r}")

    if not isinstance(phase_spec.get("prompt"), str) or not phase_spec.get("prompt", "").strip():
        reasons.append("compiled phase.prompt is missing or empty")

    for required_field in ("required_state", "required_resources", "allowed_routes", "output_contract"):
        if required_field not in phase_spec:
            reasons.append(f"compiled phase.{required_field} is missing")

    declared_state = _compiled_list(element, phase_spec, "required_state")
    yaml_projection = _project_runtime_state(state, model_record, kind, phase)
    expected_paths = list(yaml_projection.keys()) if isinstance(yaml_projection, dict) else []

    def covered(expected: str) -> bool:
        return any(expected == declared or expected.startswith(declared + ".") for declared in declared_state)

    missing = sorted(path for path in expected_paths if not covered(path))
    if missing:
        reasons.append("compiled required_state misses YAML/runtime dependencies: " + ", ".join(missing))

    compiled_route_map = _compiled_route_map(phase_spec)
    runtime_route_map = {str(route.get("key")): str(route.get("target")) for route in routes if route.get("key") and route.get("target")}
    route_allowlist = _compiled_route_allowlist(phase_spec)

    # Phase-aware routing contract.  For an enter phase that is explicitly defined to
    # await the analyst and return route_key=null, graph traversal belongs to respond.
    # In that case an empty compiled route map is correct even when the YAML node has
    # a later `next` route. Gates and non-awaiting phases still require exact routes.
    output_contract = phase_spec.get("output_contract") if isinstance(phase_spec, dict) else None
    properties = output_contract.get("properties") if isinstance(output_contract, dict) else None
    await_spec = properties.get("await_analyst") if isinstance(properties, dict) else None
    route_spec = properties.get("route_key") if isinstance(properties, dict) else None
    enter_waits_for_analyst = (
        kind == "node"
        and phase == "enter"
        and isinstance(await_spec, dict)
        and await_spec.get("const") is True
        and isinstance(route_spec, dict)
        and (route_spec.get("type") == ["null"] or route_spec.get("type") == "null")
    )
    expected_route_map = {} if enter_waits_for_analyst else runtime_route_map

    if compiled_route_map != expected_route_map:
        reasons.append(
            f"compiled allowed_routes do not match phase/runtime routes: compiled={compiled_route_map}, expected={expected_route_map}, phase={phase}"
        )
    if set(route_allowlist) != set(expected_route_map):
        reasons.append(
            f"compiled route_key_allowlist does not match phase/runtime routes: compiled={route_allowlist}, expected={list(expected_route_map)}, phase={phase}"
        )

    # A declared required_state path must actually be available at runtime.  This catches
    # semantically wrong compiler dependencies before they reach the model. Null is a valid
    # value; absence is determined structurally rather than by value.
    unavailable = sorted(path for path in declared_state if not _state_path_exists(state, path))
    if unavailable:
        reasons.append("compiled required_state is unavailable in runtime state: " + ", ".join(unavailable))

    semantic = element.get("semantic_validation")
    if not isinstance(semantic, dict):
        reasons.append("compiled semantic_validation is missing")
    else:
        for flag in ("state_dependencies_match", "route_keys_match", "route_targets_match", "state_update_allowlist_match"):
            if semantic.get(flag) is not True:
                reasons.append(f"compiled semantic_validation.{flag} is not true")
        if semantic.get("errors") not in ([], None):
            reasons.append("compiled semantic_validation.errors is not empty")

    if not isinstance(output_contract, dict) or output_contract.get("additionalProperties") is not False:
        reasons.append("compiled output_contract must be strict with additionalProperties=false")

    return not reasons, reasons

def _state_path_exists(state: Any, path: str) -> bool:
    if not isinstance(state, dict):
        return False
    clean = path.removeprefix("state.").removeprefix("$state.")
    if clean in state:
        return True
    if any(isinstance(k, str) and k.startswith(clean + ".") for k in state):
        return True
    cur: Any = state
    for part in clean.split("."):
        if not isinstance(cur, dict) or part not in cur:
            return False
        cur = cur[part]
    return True


def _assign_nested_path(target: dict[str, Any], path: str, value: Any) -> None:
    parts = [part for part in path.split(".") if part]
    if not parts:
        return
    cursor = target
    for part in parts[:-1]:
        existing = cursor.get(part)
        if not isinstance(existing, dict):
            existing = {}
            cursor[part] = existing
        cursor = existing
    cursor[parts[-1]] = copy.deepcopy(value)


def _project_state_by_paths(state: Any, paths: list[str]) -> Any:
    if not isinstance(state, dict) or not paths:
        return {}
    projected: dict[str, Any] = {}
    for path in paths:
        if not _state_path_exists(state, path):
            continue
        value = _state_subtree(state, path)
        # Compiled prompts address runtime_state using canonical dotted paths. Emit a
        # nested JSON object so runtime_state.a.b means what the prompt says it means.
        _assign_nested_path(projected, path, value)
    return _bounded_json_value(projected, 24000)


def _compiled_resources(element: dict[str, Any], phase_spec: dict[str, Any]) -> list[dict[str, Any]]:
    refs = _compiled_list(element, phase_spec, "required_resources")
    resolved: list[dict[str, Any]] = []
    for ref in refs:
        name, text = _package_resource_text(ref)
        if name is not None and text is not None:
            resolved.append({"path": name, "content": text})
    return resolved

def _safe_attachment_name(name: str) -> str:
    base = Path(str(name or "attachment")).name
    safe = re.sub(r"[^A-Za-z0-9._()\- ]+", "_", base).strip(" .")
    return safe[:180] or "attachment"


def _extract_attachment_text(name: str, mime: str, raw: bytes) -> tuple[str, str]:
    suffix = Path(name).suffix.lower()
    text_suffixes = {
        ".txt", ".md", ".markdown", ".json", ".yaml", ".yml", ".csv", ".tsv", ".xml", ".html", ".htm",
        ".py", ".js", ".ts", ".tsx", ".jsx", ".css", ".scss", ".sql", ".sh", ".bash", ".zsh", ".java", ".kt",
        ".go", ".rs", ".c", ".h", ".cpp", ".hpp", ".cs", ".php", ".rb", ".toml", ".ini", ".cfg", ".conf", ".log",
    }
    if suffix in text_suffixes or mime.startswith("text/") or mime in {"application/json", "application/xml", "application/yaml", "application/x-yaml"}:
        for encoding in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
            try:
                return raw.decode(encoding), "text"
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace"), "text"
    if suffix == ".pdf" or mime == "application/pdf":
        try:
            from pypdf import PdfReader  # type: ignore
            reader = PdfReader(io.BytesIO(raw))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
            return text, "pdf_text"
        except Exception as error:
            return "", f"pdf_unreadable:{type(error).__name__}"
    if suffix == ".docx" or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            from docx import Document  # type: ignore
            doc = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + "\t".join(cell.text for cell in row.cells)
            return text, "docx_text"
        except Exception as error:
            return "", f"docx_unreadable:{type(error).__name__}"
    return "", "binary_metadata_only"


def _prepare_analyst_attachments(payload: Any, session_id: str) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    if not payload:
        return [], []
    if not isinstance(payload, list):
        raise ValueError("attachments must be a list.")
    if len(payload) > 8:
        raise ValueError("At most 8 attachments are allowed per analyst message.")
    prepared: list[dict[str, Any]] = []
    debug_meta: list[dict[str, Any]] = []
    total_bytes = 0
    total_text_chars = 0
    workspace = _runtime_workspace() / "analyst_attachments" / re.sub(r"[^A-Za-z0-9._-]+", "_", session_id or "session")
    workspace.mkdir(parents=True, exist_ok=True)
    for index, item in enumerate(payload):
        if not isinstance(item, dict):
            raise ValueError("Each attachment must be an object.")
        name = _safe_attachment_name(str(item.get("name") or f"attachment-{index+1}"))
        mime = str(item.get("type") or "application/octet-stream")[:200]
        encoded = item.get("data_base64")
        if not isinstance(encoded, str) or not encoded:
            raise ValueError(f"Attachment {name!r} has no file data.")
        try:
            raw = base64.b64decode(encoded, validate=True)
        except Exception as error:
            raise ValueError(f"Attachment {name!r} contains invalid base64 data.") from error
        if len(raw) > 10 * 1024 * 1024:
            raise ValueError(f"Attachment {name!r} exceeds the 10 MB per-file limit.")
        total_bytes += len(raw)
        if total_bytes > 24 * 1024 * 1024:
            raise ValueError("Total attachment size exceeds the 24 MB per-message limit.")
        digest = hashlib.sha256(raw).hexdigest()
        stored_name = f"{digest[:12]}-{name}"
        stored = workspace / stored_name
        stored.write_bytes(raw)
        text, extraction = _extract_attachment_text(name, mime, raw)
        remaining = max(0, 80000 - total_text_chars)
        if text and remaining:
            text = text[: min(30000, remaining)]
            total_text_chars += len(text)
        else:
            text = ""
        model_item = {
            "name": name,
            "mime_type": mime,
            "size": len(raw),
            "sha256": digest,
            "extraction": extraction,
        }
        if text:
            model_item["content"] = text
        else:
            model_item["note"] = "File is attached, but no text representation is available to this chat-completions runtime."
        prepared.append(model_item)
        debug_meta.append({
            "name": name, "mime_type": mime, "size": len(raw), "sha256": digest,
            "extraction": extraction, "extracted_chars": len(text),
            "stored_path": str(stored.relative_to(_runtime_workspace())),
        })
    return prepared, debug_meta


def _interaction_contract() -> dict[str, Any]:
    plan = _active_playbook_package().get("semantic_plan")
    contract = plan.get("interaction_contract") if isinstance(plan, dict) else None
    if isinstance(contract, dict):
        locale = str(contract.get("locale") or "uk-UA")
        language = str(contract.get("model_output_language") or ("uk" if locale.lower().startswith("uk") else "en"))
        return {"locale": locale, "model_output_language": language}
    source = _active_playbook_package().get("source")
    im = source.get("interaction_model") if isinstance(source, dict) and isinstance(source.get("interaction_model"), dict) else {}
    locale = str(im.get("locale") or im.get("interaction_locale") or "uk-UA")
    language = str(im.get("model_output_language") or im.get("language") or ("uk" if locale.lower().startswith("uk") else "en"))
    return {"locale": locale, "model_output_language": language}


def _analyst_language_instruction() -> str:
    c = _interaction_contract()
    return (f"Analyst-facing language contract: locale={c['locale']}, language={c['model_output_language']}. "
            "Write every assistant_message, analyst_explanation, question, validation explanation and recovery explanation in that language. "
            "Keep technical IDs, route keys, state paths, field names and machine-readable enum values unchanged. Never switch analyst-facing prose to English merely because internal contracts are English.")


def _semantic_plan_element(element_id: str) -> dict[str, Any] | None:
    plan = _active_playbook_package().get("semantic_plan")
    elements = plan.get("elements") if isinstance(plan, dict) else None
    value = elements.get(element_id) if isinstance(elements, dict) else None
    return value if isinstance(value, dict) else None


def _semantic_resource_context(element: dict[str, Any], budget_chars: int = 18000) -> list[dict[str, Any]]:
    plan = _active_playbook_package().get("semantic_plan")
    catalog = plan.get("resources") if isinstance(plan, dict) else None
    out = []
    used = 0
    for ref in element.get("resources") or []:
        meta = catalog.get(ref) if isinstance(catalog, dict) else None
        text = (_active_playbook_package().get("resources") or {}).get(ref)
        if text is None:
            matches = [v for k, v in (_active_playbook_package().get("resources") or {}).items() if k.endswith("/" + ref) or k == ref]
            text = matches[0] if len(matches) == 1 else None
        content = text if isinstance(text, str) else (json.dumps((meta or {}).get("structured_content"), ensure_ascii=False) if isinstance(meta, dict) else None)
        if not isinstance(content, str):
            continue
        remaining = budget_chars - used
        if remaining <= 0:
            out.append({"path": ref, "content": "", "__truncated__": True, "reason": "resource_budget_exhausted"})
            continue
        clipped = content[:remaining]
        item = {"path": ref, "content": clipped}
        if len(clipped) < len(content):
            item["__truncated__"] = True
            item["original_chars"] = len(content)
        out.append(item)
        used += len(clipped)
    return out


def _semantic_instruction_view(value: Any) -> Any:
    """Remove runtime-owned graph mechanics while preserving semantic instructions."""
    runtime_keys = {"id", "incoming_from", "allowed_from", "next", "pass_to", "fail_to", "on_pass", "on_fail", "on_yes", "on_no", "on_approved", "on_rejected", "on_unknown", "routes", "branches", "declared_dynamic_routes", "declared_exception_routes", "navigation_contract", "transitions", "terminal"}
    if isinstance(value, dict):
        out = {}
        for key, child in value.items():
            if key in runtime_keys:
                continue
            if key in {"update_state", "on_pass_update_state", "on_fail_update_state"}:
                continue
            out[key] = _semantic_instruction_view(child)
        return out
    if isinstance(value, list):
        return [_semantic_instruction_view(x) for x in value]
    return copy.deepcopy(value)


def _authority_value_missing(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return not value.strip()
    if isinstance(value, (list, dict, tuple, set)):
        return len(value) == 0
    return False


def _authority_contract(element: dict[str, Any]) -> dict[str, Any]:
    source = element.get("semantic_source") if isinstance(element.get("semantic_source"), dict) else {}
    contract = source.get("authority_contract")
    return contract if isinstance(contract, dict) else {}


def _authority_missing_clarifications(state: dict[str, Any], contract: dict[str, Any]) -> list[str]:
    fields = [str(x) for x in (contract.get("clarification_only_fields") or []) if isinstance(x, str)]
    missing: list[str] = []
    for path in fields:
        value = _state_subtree(state, path) if _state_path_exists(state, path) else None
        if _authority_value_missing(value):
            missing.append(path)
    return missing


def _authority_selector_values(state: Any, selector: str) -> list[Any]:
    parts=[p for p in str(selector or "").split(".") if p]
    current=[state]
    for part in parts:
        nxt=[]
        for obj in current:
            if part == "*":
                if isinstance(obj, dict):
                    nxt.extend(obj.values())
                elif isinstance(obj, list):
                    nxt.extend(obj)
                continue
            if isinstance(obj, dict) and part in obj:
                nxt.append(obj[part])
            elif isinstance(obj, list):
                try:
                    idx=int(part)
                except (TypeError, ValueError):
                    continue
                if 0 <= idx < len(obj):
                    nxt.append(obj[idx])
        current=nxt
        if not current:
            break
    out=[]
    def flatten(value: Any):
        if isinstance(value, list):
            for item in value: flatten(item)
        elif isinstance(value, (str, int, float, bool)) and not _authority_value_missing(value):
            out.append(value)
    for value in current:
        flatten(value)
    return out


def _authority_search_text(value: Any) -> str:
    if isinstance(value, str):
        return value
    try:
        return json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
    except Exception:
        return str(value)


def _authority_literal_errors(target: str, value: Any, spec: dict[str, Any], state: dict[str, Any]) -> list[str]:
    errors=[]
    text=_authority_search_text(value)
    for selector in [str(x) for x in (spec.get("must_include_from") or []) if isinstance(x, str)]:
        for literal in _authority_selector_values(state, selector):
            needle=str(literal)
            if needle not in text:
                errors.append(f"authority-derived path {target} omits canonical literal from {selector}: {needle!r}")
    return errors


def _apply_authority_contract_to_candidate(
    candidate: dict[str, Any], *, semantic_element: dict[str, Any], state: dict[str, Any], phase: str
) -> tuple[dict[str, Any], list[dict[str, Any]], list[str]]:
    """Enforce generic canonical-state authority declared by a playbook element.

    The runtime does not know domain values. It only enforces the source-declared
    relationship: fields marked authority-derived must be produced from already
    available canonical inputs, while clarification-only fields may not be invented
    by the derivation node. Open-question state is runtime-owned from missingness.
    """
    contract = _authority_contract(semantic_element)
    if not contract:
        return candidate, [], []
    patch = candidate.get("state_patch")
    if not isinstance(patch, dict) or not isinstance(patch.get("operations"), list):
        return candidate, [], ["authority_contract requires a structured state_patch"]
    ops = patch["operations"]
    adaptations: list[dict[str, Any]] = []
    errors: list[str] = []
    by_path: dict[str, list[dict[str, Any]]] = {}
    for op in ops:
        if isinstance(op, dict) and isinstance(op.get("path"), str):
            by_path.setdefault(str(op.get("path")), []).append(op)

    derived = contract.get("derived_targets") if isinstance(contract.get("derived_targets"), dict) else {}
    for target, spec in derived.items():
        target = str(target)
        spec = spec if isinstance(spec, dict) else {}
        sources = [str(x) for x in (spec.get("sources") or []) if isinstance(x, str)]
        source_values = []
        sources_available = bool(sources)
        for source in sources:
            value = _state_subtree(state, source) if _state_path_exists(state, source) else None
            source_values.append(value)
            if _authority_value_missing(value):
                sources_available = False
        current = _state_subtree(state, target) if _state_path_exists(state, target) else None
        current_present = not _authority_value_missing(current)
        current_literal_errors = _authority_literal_errors(target, current, spec, state) if current_present else []
        current_valid = current_present and not current_literal_errors
        target_ops = by_path.get(target, [])
        if current_valid:
            # Existing target is reusable only when it still carries every canonical literal
            # explicitly required by the authority contract.
            for op in target_ops:
                if op.get("value") != current and not (phase == "respond" and op.get("basis") == "analyst_input"):
                    errors.append(f"authority-derived path {target} already has a canonical-complete value and cannot be replaced by model derivation")
        elif sources_available:
            if not target_ops:
                errors.append(f"authority-derived path {target} must be populated from declared canonical sources before clarification")
            else:
                for op in target_ops:
                    if op.get("basis") not in {"confirmed_state", "derived"}:
                        errors.append(f"authority-derived path {target} requires basis confirmed_state|derived")
                    errors.extend(_authority_literal_errors(target, op.get("value"), spec, state))
        # If declared sources are incomplete, the node may leave the target unresolved;
        # the playbook can route through its normal sufficiency/clarification contour.

    clarification = [str(x) for x in (contract.get("clarification_only_fields") or []) if isinstance(x, str)]
    for path in clarification:
        for op in by_path.get(path, []):
            if not (phase == "respond" and op.get("basis") == "analyst_input"):
                errors.append(f"clarification-only path {path} cannot be invented by model derivation")

    open_path = str(contract.get("open_questions_path") or "").strip()
    if open_path:
        missing = _authority_missing_clarifications(state, contract)
        existing = by_path.get(open_path, [])
        canonical_op = {
            "op": "set", "path": open_path, "value": missing, "basis": "derived",
            "reason": "runtime-owned authority contract: missing clarification-only fields",
            "row_key": None, "row_match": None,
        }
        if existing:
            # Replace every model attempt to author open-question state with one runtime-owned value.
            ops[:] = [op for op in ops if not (isinstance(op, dict) and op.get("path") == open_path)]
        ops.append(canonical_op)
        candidate["state_patch"] = patch
        adaptations.append({
            "kind": "authority_open_questions_runtime_owned",
            "path": open_path, "missing_fields": missing,
        })
    return candidate, adaptations, errors


def _assemble_runtime_semantic_call(element: dict[str, Any], current_id: str, phase: str, state: dict[str, Any], history: Any, analyst_input: str) -> tuple[str, dict[str, Any], dict[str, Any]]:
    traits = element.get("execution_traits") if isinstance(element.get("execution_traits"), dict) else {}
    source = element.get("semantic_source") if isinstance(element.get("semantic_source"), dict) else {}
    task = _semantic_instruction_view(source)
    routes = element.get("routes") if isinstance(element.get("routes"), list) else []
    resources = _semantic_resource_context(element)
    state_contract = element.get("state_contract") if isinstance(element.get("state_contract"), dict) else {}
    semantic_objects = [str(x) for x in state_contract.get("semantic_objects") or [] if isinstance(x, str)]
    reads = [str(x) for x in state_contract.get("reads_hint") or [] if isinstance(x, str)]
    preload = sorted(set(semantic_objects + [x.split(".")[0] for x in reads]))
    full_projected_state: dict[str, Any] = {}
    projection_defaults = (((state_contract.get("projection_defaults") or {}).get(phase)) if isinstance(state_contract.get("projection_defaults"), dict) else {}) or {}
    default_materialized: list[str] = []
    missing_produced_state: list[str] = []
    for path in preload:
        if _state_path_exists(state, path):
            _assign_nested_path(full_projected_state, path, _state_subtree(state, path))
        elif path in projection_defaults:
            _assign_nested_path(full_projected_state, path, copy.deepcopy(projection_defaults[path]))
            default_materialized.append(path)
        else:
            missing_produced_state.append(path)
    if not preload:
        full_projected_state = _fallback_runtime_state(state)
    state_budget = 180000 if str(element.get("kind") or "").endswith("_gate") else 100000
    compact_state = _bounded_json_value(full_projected_state, state_budget)
    projected_keys = set(compact_state.keys()) if isinstance(compact_state, dict) else set()
    full_keys = set(full_projected_state.keys()) if isinstance(full_projected_state, dict) else set()
    truncated_by_budget = sorted(full_keys - projected_keys)
    unavailable_runtime_context: list[str] = []
    context_complete = not missing_produced_state and not truncated_by_budget and not unavailable_runtime_context
    context_status = {
        "context_complete": context_complete,
        "default_materialized": sorted(default_materialized),
        "missing_produced_state": sorted(missing_produced_state),
        "truncated_by_budget": sorted(truncated_by_budget),
        "unavailable_runtime_context": sorted(unavailable_runtime_context),
    }
    if isinstance(compact_state, dict):
        compact_state = dict(compact_state)
        compact_state["__context_status__"] = copy.deepcopy(context_status)
    recent_history = _compact_live_history(history if isinstance(history, list) else [])
    system = (
        _system_contract_for_call("gate" if str(element.get("kind", "")).endswith("_gate") else "node", phase, {"requires_human": bool(traits.get("requires_analyst"))})
        + "\n\nAlpha.20 Runtime Semantic Plan instruction. The semantic_task below is the task content, not graph authority. "
          "Do not invent routes or state paths. Runtime validates route selection and every StatePatch. "
          "Return JSON only. For model nodes include assistant_message, state_patch, route_key (or null), needs_analyst, next_intent and rationale_short. "
          "StatePatch base_revision is runtime-owned: omit it or treat any supplied value as non-authoritative. For every StatePatch operation, basis is provenance classification only: analyst_input, confirmed_state, derived, generated, recovery, or legacy_unknown; put rule/action identifiers and derivation explanations in reason, never in basis. IMPORTANT: basis and reason are operation-envelope fields and MUST be siblings of op/path/value; never place basis or reason inside the value object. "
          "For collection StatePatch operations: append or merge_row MUST send exactly one row object in value, never an array; set or replace MUST send the complete array value. For functional_test_catalog.rows, unit_test_catalog.rows, and edge_case_catalog.rows, use append only for a NEW tc_id. On recovery/revisit, if tc_id already exists, correct that row with merge_row using row_key='tc_id' and row_match=<existing tc_id>; never append a duplicate tc_id. Never set, replace, or remove whole test catalogues. "
          "For model gates also include status and check_results. For every declared check return exactly one check_results item with check_id, status=pass|fail|not_run, evidence (max 2 concise items), remediation, and not_run_reason. If no checks are declared return check_results=[]. failed_checks is compatibility-only and runtime derives it from check_results. Also include invalid_state, missing_information, missing_coverage, affected_state and evidence."
        + "\n\n" + _analyst_language_instruction()
    )
    context = {
        "execution_phase": phase,
        "element_id": current_id,
        "element_kind": element.get("kind"),
        "semantic_task": task,
        "output_contract": element.get("output_contract"),
        "state_patch_template": state_contract.get("patch_template") or [],
        "allowed_routes": routes,
        "runtime_state": compact_state,
        "recent_history": recent_history,
        "context_complete": context_complete,
        "context_truncated_objects": truncated_by_budget,
        "context_status": copy.deepcopy(context_status),
    }
    authority = _authority_contract(element)
    if authority:
        context["authority_contract"] = copy.deepcopy(authority)
        missing_clarifications = _authority_missing_clarifications(state, authority)
        context["authority_missing_clarifications"] = missing_clarifications
        system += (
            "\n\nCanonical-state authority rule: this element declares an authority_contract. "
            "Existing non-empty canonical target values are authoritative and must be preserved. "
            "For every authority_contract.derived_targets entry, when all declared source paths are present, "
            "derive the target only from those canonical sources; do not substitute generic defaults, common conventions, "
            "assumed class names, assumed module architecture, or project patterns that are not present in runtime_state. "
            "Do not ask the analyst for authority-derived fields merely because the target field is currently empty. "
            "Ask only for fields listed in clarification_only_fields that are actually missing. "
            "The runtime owns open_questions_path and will normalize it from missing clarification-only fields."
        )
    if analyst_input:
        context["analyst_input"] = analyst_input
    if resources:
        context["resources"] = resources
    output_contract = element.get("output_contract") if isinstance(element.get("output_contract"), dict) else {}
    response_schema = copy.deepcopy(output_contract.get("json_schema"))
    strict_api_compatible = True  # provider-specific compatibility is resolved at send time
    if isinstance(response_schema, dict):
        try:
            op_schema = response_schema["properties"]["state_patch"]["properties"]["operations"]["items"]
            if isinstance(op_schema, dict):
                basis_schema = (op_schema.get("properties") or {}).get("basis")
                if isinstance(basis_schema, dict):
                    basis_schema.clear()
                    basis_schema.update({"type": ["string", "null"], "enum": ["analyst_input", "confirmed_state", "derived", "generated", "recovery", "legacy_unknown", None]})
        except (KeyError, TypeError):
            pass
    revisit = _revisit_context(history, current_id, state, element)
    if revisit:
        context["revisit_context"] = revisit
        system += "\n\nRevisit rule: previous answers are first-class evidence. Preserve the prior answer when still valid; ask only for the delta when state dependencies changed. If revisit_context.previous_answer_status is needs_extension, you MUST address revisit_context.required_extension before proceeding. A routed gate failure is authoritative evidence that the prior answer is not sufficient even when canonical state dependencies did not otherwise change. If revisit_context.can_confirm_without_changes is true, allow the analyst to confirm the prior answer without regeneration."
    if isinstance(response_schema, dict):
        context["__response_json_schema"] = response_schema
    context["response_schema_mode"] = "provider_profile_resolved_at_runtime"
    return system, context, {"resolved_resources": resources, "preload_paths": preload, "runtime_state": compact_state, "recent_history": recent_history, "context_complete": context_complete, "context_truncated_objects": truncated_by_budget, "context_status": copy.deepcopy(context_status), "default_materialized": sorted(default_materialized), "missing_produced_state": sorted(missing_produced_state), "truncated_by_budget": sorted(truncated_by_budget), "unavailable_runtime_context": sorted(unavailable_runtime_context)}


def _latest_routed_gate_failure(history: Any, current_id: str) -> dict[str, Any] | None:
    """Return the latest structured gate failure only when that failure routed to current_id.

    This makes recovery evidence route-aware: an old failure elsewhere in the run must not
    invalidate an unrelated revisit, while the exact recovery target receives the failed
    gate's missing coverage/information as first-class evidence.
    """
    if not isinstance(history, list):
        return None
    for item in reversed(history):
        if not isinstance(item, dict):
            continue
        debug = item.get("debug") if isinstance(item.get("debug"), dict) else {}
        alpha20 = debug.get("alpha20") if isinstance(debug.get("alpha20"), dict) else {}
        failure = alpha20.get("gate_failure") if isinstance(alpha20.get("gate_failure"), dict) else None
        runtime = debug.get("runtime") if isinstance(debug.get("runtime"), dict) else {}
        if failure and str(runtime.get("next_id") or "") == str(current_id):
            return {
                "gate_id": str(failure.get("gate_id") or debug.get("current_id") or ""),
                "failed_checks": copy.deepcopy(failure.get("failed_checks") or []),
                "invalid_state": copy.deepcopy(failure.get("invalid_state") or []),
                "missing_information": copy.deepcopy(failure.get("missing_information") or []),
                "missing_coverage": copy.deepcopy(failure.get("missing_coverage") or []),
                "affected_state": copy.deepcopy(failure.get("affected_state") or []),
                "evidence": copy.deepcopy(failure.get("evidence") or []),
                "suggested_recovery_scope": failure.get("suggested_recovery_scope"),
            }
    return None


def _revisit_context(history: Any, current_id: str, state: dict[str, Any], semantic_element: dict[str, Any]) -> dict[str, Any] | None:
    if not isinstance(history, list): return None
    previous = [item for item in history if isinstance(item, dict) and item.get("node_id") == current_id]
    if not previous: return None
    prior_answer = next((str(item.get("text") or "") for item in reversed(previous) if item.get("role") == "assistant" and item.get("text")), "")
    prior_state = None
    for item in reversed(previous):
        debug = item.get("debug") if isinstance(item.get("debug"), dict) else {}
        runtime = debug.get("runtime") if isinstance(debug.get("runtime"), dict) else {}
        if isinstance(runtime.get("state_after"), dict):
            prior_state = canonicalize_runtime_state(runtime.get("state_after")); break
    contract = semantic_element.get("state_contract") if isinstance(semantic_element, dict) else {}
    paths = list(dict.fromkeys((contract.get("reads_hint") or []) + (contract.get("writes") or [])))
    changed = []
    if prior_state is not None:
        current = canonicalize_runtime_state(state)
        for path in paths:
            old_val = _state_subtree(prior_state, path) if _state_path_exists(prior_state, path) else None
            new_val = _state_subtree(current, path) if _state_path_exists(current, path) else None
            if old_val != new_val: changed.append(path)
    recovery_failure = _latest_routed_gate_failure(history, current_id)
    status = "still_valid" if prior_state is not None and not changed and recovery_failure is None else "needs_extension"
    result = {
        "previous_answer_status": status,
        "previous_answer": prior_answer[:4000],
        "changed_state_paths": changed,
        "can_confirm_without_changes": status == "still_valid" and bool(prior_answer),
        "policy": "preserve_previous_answer_and_ask_only_delta",
    }
    if recovery_failure is not None:
        result["recovery_gate_failure"] = recovery_failure
        result["required_extension"] = {
            "missing_information": copy.deepcopy(recovery_failure.get("missing_information") or []),
            "missing_coverage": copy.deepcopy(recovery_failure.get("missing_coverage") or []),
            "failed_checks": copy.deepcopy(recovery_failure.get("failed_checks") or []),
        }
    return result


def _normalize_semantic_model_result(result: dict[str, Any]) -> dict[str, Any]:
    out = dict(result)
    if "await_analyst" not in out and "needs_analyst" in out:
        out["await_analyst"] = bool(out.get("needs_analyst"))
    if "route_key" not in out:
        route = out.get("route") or out.get("requested_route")
        if isinstance(route, str): out["route_key"] = route
    return out


def _adapt_runtime_owned_node_envelope(
    candidate: dict[str, Any],
    *,
    semantic_element: dict[str, Any] | None,
    semantic_traits: dict[str, Any] | None,
    phase: str,
) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    """Normalize only runtime-owned/mechanical NodeExecutionResult envelope fields.

    This intentionally does *not* invent semantic state values, routes, gate verdicts, or
    business content. It exists for providers that guarantee JSON but do not enforce our
    JSON Schema. The adaptation is explicit in debug evidence so a live PASS never hides
    how much mechanical help the runtime supplied.
    """
    out = dict(candidate)
    adaptations: list[dict[str, Any]] = []
    contract = semantic_element.get("output_contract") if isinstance(semantic_element, dict) and isinstance(semantic_element.get("output_contract"), dict) else {}
    if str(contract.get("contract") or "") != "NodeExecutionResult":
        return out, adaptations

    # A recurring Gemma error is to put the StatePatch envelope into legacy
    # state_updates. Reinterpret it only when the object is unambiguously a patch
    # envelope; arbitrary state_updates still go through the normal legacy adapter.
    if not isinstance(out.get("state_patch"), dict):
        legacy = out.get("state_updates")
        if isinstance(legacy, dict) and isinstance(legacy.get("operations"), list) and set(legacy).issubset({"base_revision", "operations", "semantic_summary"}):
            out["state_patch"] = copy.deepcopy(legacy)
            out.pop("state_updates", None)
            adaptations.append({"kind": "state_patch_envelope_from_state_updates"})
        elif isinstance(out.get("operations"), list):
            patch = {"base_revision": int(out.get("base_revision") or 0), "operations": copy.deepcopy(out.get("operations") or [])}
            if out.get("semantic_summary") is not None:
                patch["semantic_summary"] = out.get("semantic_summary")
            out["state_patch"] = patch
            for key in ("base_revision", "operations", "semantic_summary"):
                out.pop(key, None)
            adaptations.append({"kind": "state_patch_envelope_from_top_level"})

    # row_key / row_match are transport-shape fields in the strict operation schema.
    # They carry semantic meaning only for merge_row.  Older/live JSON-only providers
    # frequently omit them for set/append/merge operations.  Materialize explicit nulls
    # as runtime-owned envelope metadata, but never invent merge_row identity.
    patch_obj = out.get("state_patch") if isinstance(out.get("state_patch"), dict) else None
    patch_ops = patch_obj.get("operations") if isinstance(patch_obj, dict) and isinstance(patch_obj.get("operations"), list) else []
    for index, operation in enumerate(patch_ops):
        if not isinstance(operation, dict) or str(operation.get("op") or "") == "merge_row":
            continue
        added: list[str] = []
        for field in ("row_key", "row_match"):
            if field not in operation:
                operation[field] = None
                added.append(field)
        if added:
            adaptations.append({"kind": "derive_non_merge_row_metadata", "operation_index": index, "fields": added})

    traits = semantic_traits if isinstance(semantic_traits, dict) else {}
    requires_analyst = bool(traits.get("requires_analyst"))
    route_key = out.get("route_key")

    # needs_analyst is orchestration metadata, not business semantics. Runtime already
    # owns the final analyst/route decision, so derive a conservative value when omitted.
    if "needs_analyst" not in out:
        inferred = bool(requires_analyst and (phase == "enter" or route_key in (None, "")))
        out["needs_analyst"] = inferred
        adaptations.append({"kind": "derive_needs_analyst", "value": inferred})
    if "await_analyst" not in out:
        out["await_analyst"] = bool(out.get("needs_analyst"))

    # next_intent is descriptive orchestration metadata. Do not derive a route; merely
    # describe the already supplied route/await state when the model omits the string.
    if "next_intent" not in out:
        if bool(out.get("needs_analyst")):
            next_intent = "await_analyst_input"
        elif isinstance(route_key, str) and route_key:
            next_intent = f"route:{route_key}"
        else:
            next_intent = "continue_default_path"
        out["next_intent"] = next_intent
        adaptations.append({"kind": "derive_next_intent", "value": next_intent})

    # action is required by the compiled schema in some nodes but is not used as graph
    # authority. Fill it only when absent, from the already resolved orchestration shape.
    required = contract.get("required") if isinstance(contract.get("required"), list) else []
    if "action" in required and "action" not in out:
        action = "AWAIT_ANALYST" if bool(out.get("needs_analyst")) else "CONTINUE"
        out["action"] = action
        adaptations.append({"kind": "derive_action", "value": action})
    return out, adaptations



def _alpha20_value_schemas(semantic_element: dict[str, Any] | None) -> dict[str, Any]:
    if not isinstance(semantic_element, dict): return {}
    return (((semantic_element.get("output_contract") or {}).get("state_patch") or {}).get("value_schema_by_path") or {})

def _alpha20_operation_variants(semantic_element: dict[str, Any] | None) -> list[dict[str, Any]]:
    if not isinstance(semantic_element, dict):
        return []
    variants = (((semantic_element.get("output_contract") or {}).get("state_patch") or {}).get("operation_variants") or [])
    return [copy.deepcopy(v) for v in variants if isinstance(v, dict)]

def _gate_declared_check_ids(semantic_element: dict[str, Any] | None) -> list[str]:
    if not isinstance(semantic_element, dict):
        return []
    oc = semantic_element.get("output_contract") if isinstance(semantic_element.get("output_contract"), dict) else {}
    explicit = oc.get("declared_check_ids") if isinstance(oc.get("declared_check_ids"), list) else None
    if explicit is not None:
        return [str(x) for x in explicit if str(x or "").strip()]
    gc = semantic_element.get("gate_contract") if isinstance(semantic_element.get("gate_contract"), dict) else {}
    checks = gc.get("checks_inline") if isinstance(gc.get("checks_inline"), list) else []
    ids=[]
    for item in checks:
        if not isinstance(item, dict):
            continue
        cid=item.get("id") or item.get("check_id")
        if cid is not None and str(cid).strip():
            ids.append(str(cid))
    if ids:
        return ids
    ext = gc.get("external_specification") if isinstance(gc.get("external_specification"), dict) else {}
    for item in ext.get("checks") or ext.get("assertions") or []:
        if isinstance(item, dict):
            cid=item.get("id") or item.get("check_id")
            if cid is not None and str(cid).strip():
                ids.append(str(cid))
    return ids

def _validate_gate_check_results(candidate: dict[str, Any], semantic_element: dict[str, Any]) -> tuple[list[str], dict[str, Any]]:
    errors: list[str] = []
    declared = _gate_declared_check_ids(semantic_element)
    raw = candidate.get("check_results")
    if not isinstance(raw, list):
        return ["gate result requires check_results array"], {"declared_check_ids":declared,"execution_status":"not_evaluated","executed_check_ids":[],"verdicts":{},"not_run_reasons":{},"per_check_accounting":"applicable" if declared else "not_applicable"}
    if not declared and raw:
        errors.append("gate with no declared checks requires check_results=[]")
    seen: dict[str, dict[str, Any]] = {}
    for idx,item in enumerate(raw):
        if not isinstance(item, dict):
            errors.append(f"check_results[{idx}] must be an object"); continue
        required={"check_id","status","evidence","remediation","not_run_reason"}
        missing=sorted(required-set(item))
        if missing: errors.append(f"check_results[{idx}] missing required fields: {missing}")
        extra=sorted(set(item)-required)
        if extra: errors.append(f"check_results[{idx}] has unknown fields: {extra}")
        cid=str(item.get("check_id") or "")
        if not cid: errors.append(f"check_results[{idx}].check_id is empty"); continue
        if cid in seen: errors.append(f"duplicate check_id: {cid}")
        seen[cid]=item
        if declared and cid not in declared: errors.append(f"unknown check_id: {cid}")
        status=str(item.get("status") or "")
        if status not in {"pass","fail","not_run"}: errors.append(f"check_results[{idx}].status is invalid: {status!r}")
        evidence=item.get("evidence")
        if not isinstance(evidence,list): errors.append(f"check_results[{idx}].evidence must be an array")
        else:
            if len(evidence)>2: errors.append(f"check_results[{idx}].evidence exceeds maxItems=2")
            for j,val in enumerate(evidence):
                if not isinstance(val,str): errors.append(f"check_results[{idx}].evidence[{j}] must be string")
                elif len(val)>320: errors.append(f"check_results[{idx}].evidence[{j}] exceeds maxLength=320")
        for fld in ("remediation","not_run_reason"):
            val=item.get(fld)
            if val is not None and not isinstance(val,str): errors.append(f"check_results[{idx}].{fld} must be string|null")
            elif isinstance(val,str) and len(val)>320: errors.append(f"check_results[{idx}].{fld} exceeds maxLength=320")
        if status=="not_run" and not str(item.get("not_run_reason") or "").strip(): errors.append(f"check_results[{idx}].not_run_reason required for status=not_run")
    if declared:
        missing_ids=[cid for cid in declared if cid not in seen]
        if missing_ids: errors.append(f"missing check_results for declared checks: {missing_ids}")
    status=str(candidate.get("status") or "").lower()
    if status in {"passed","pass"} and declared:
        nonpass=[cid for cid in declared if str((seen.get(cid) or {}).get("status") or "")!="pass"]
        if nonpass: errors.append(f"gate PASS requires every declared check PASS; non-pass={nonpass}")
    executed=[cid for cid,item in seen.items() if str(item.get("status")) in {"pass","fail"}]
    not_run={cid:str(item.get("not_run_reason") or "") for cid,item in seen.items() if str(item.get("status"))=="not_run"}
    if not declared:
        exec_status="evaluated" if status in {"passed","pass","failed","fail"} else "not_evaluated"
        accounting="not_applicable"
    elif len(executed)==len(declared):
        exec_status="evaluated"
        accounting="applicable"
    elif executed or not_run:
        exec_status="partially_evaluated"
        accounting="applicable"
    else:
        exec_status="not_evaluated"
        accounting="applicable"
    accounting_data={"declared_check_ids":declared,"execution_status":exec_status,"executed_check_ids":executed,"verdicts":seen,"not_run_reasons":not_run,"per_check_accounting":accounting}
    return errors, accounting_data

def _derive_failed_checks_from_results(candidate: dict[str, Any]) -> list[dict[str, Any]]:
    out=[]
    for item in candidate.get("check_results") if isinstance(candidate.get("check_results"),list) else []:
        if not isinstance(item,dict) or str(item.get("status"))!="fail": continue
        evidence=item.get("evidence") if isinstance(item.get("evidence"),list) else []
        summary=str(item.get("remediation") or (evidence[0] if evidence else "check failed"))
        out.append({"check_id":str(item.get("check_id") or ""),"summary":summary,"severity":"error"})
    return out

def _incomplete_gate_result(current_id: str, semantic_element: dict[str, Any], semantic_call_meta: dict[str, Any]) -> dict[str, Any]:
    routes = [r for r in (semantic_element.get("routes") or []) if isinstance(r, dict)]
    failure_route = next((r for r in routes if str(r.get("key") or "").lower() in {"on_fail","fail","failed","revise"}), None)
    status=semantic_call_meta.get("context_status") if isinstance(semantic_call_meta.get("context_status"),dict) else {}
    missing=[]
    for obj in status.get("missing_produced_state") or []:
        missing.append({"path":obj,"needed":"required producer state before gate evaluation","why_needed":"missing_produced_state"})
    for obj in status.get("truncated_by_budget") or []:
        missing.append({"path":obj,"needed":"runtime context within budget","why_needed":"truncated_by_budget"})
    for obj in status.get("unavailable_runtime_context") or []:
        missing.append({"path":obj,"needed":"available runtime context","why_needed":"unavailable_runtime_context"})
    declared = _gate_declared_check_ids(semantic_element)
    check_results=[{"check_id":cid,"status":"not_run","evidence":[],"remediation":None,"not_run_reason":"runtime context incomplete before gate evaluation"} for cid in declared]
    return {"status":"failed","gate_id":current_id,"assistant_message":"Gate evaluation stopped because required runtime context was incomplete.","route_key": failure_route.get("key") if failure_route else None,"check_results":check_results,"failed_checks":[],"invalid_state":[],"missing_information":missing or [{"path":"runtime_context","needed":"complete state projection"}],"missing_coverage":[],"affected_state":list(semantic_call_meta.get("preload_paths") or []),"evidence":[{"kind":"runtime_context","context_complete":False,"context_status":status}],"suggested_recovery_scope":"context","failure_class":"technical_stop"}

def _adapt_misnested_patch_metadata(
    patch: dict[str, Any], *, allowed_paths: list[str], current_revision: int,
    value_schemas: dict[str, Any], operation_variants: list[dict[str, Any]],
) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    """R3 generic repair for misplaced StatePatch envelope metadata.

    Some models occasionally place ``basis``/``reason`` inside a row ``value``.
    We only hoist those keys when the original patch is invalid AND the fully
    adapted patch validates against the exact declared state/operation contract.
    This cannot broaden write authority or relax the row schema.
    """
    if not isinstance(patch, dict) or not isinstance(patch.get("operations"), list):
        return patch, []
    original_validation = validate_state_patch(
        patch, allowed_paths=allowed_paths, current_revision=current_revision,
        value_schemas=value_schemas, operation_variants=operation_variants,
    )
    if original_validation.get("valid"):
        return patch, []
    candidate = copy.deepcopy(patch)
    adaptations: list[dict[str, Any]] = []
    changed = False
    for idx, op in enumerate(candidate.get("operations") or []):
        if not isinstance(op, dict) or not isinstance(op.get("value"), dict):
            continue
        value = op["value"]
        moved = []
        if "basis" in value and "basis" not in op and value.get("basis") in {"analyst_input","confirmed_state","derived","generated","recovery","legacy_unknown"}:
            op["basis"] = value.pop("basis")
            moved.append("basis")
        if "reason" in value and "reason" not in op and isinstance(value.get("reason"), str):
            op["reason"] = value.pop("reason")
            moved.append("reason")
        if moved:
            changed = True
            adaptations.append({"kind":"hoist_statepatch_operation_metadata","operation_index":idx,"fields":moved})
    if not changed:
        return patch, []
    adapted_validation = validate_state_patch(
        candidate, allowed_paths=allowed_paths, current_revision=current_revision,
        value_schemas=value_schemas, operation_variants=operation_variants,
    )
    if adapted_validation.get("valid"):
        return candidate, adaptations
    return patch, []



_PACKAGE_TOOL_COMMAND_RE = re.compile(r"`([^`]*\b(?:python|python3)\s+[^`]+)`", re.IGNORECASE | re.DOTALL)
_PACKAGE_TOOL_STATE_PLACEHOLDER_RE = re.compile(r"<[^>]*?state\.([A-Za-z_][A-Za-z0-9_.]*)[^>]*>", re.IGNORECASE)


def _compiled_package_tool_adapter(semantic_element: dict[str, Any] | None) -> dict[str, Any] | None:
    if not isinstance(semantic_element, dict):
        return None
    adapter = semantic_element.get("execution_adapter")
    if not isinstance(adapter, dict) or str(adapter.get("runtime_executor") or "") != "package_tool":
        return None
    spec = adapter.get("package_tool")
    return spec if isinstance(spec, dict) else None


def _package_tool_allowed_paths(record: dict[str, Any], semantic_element: dict[str, Any] | None = None) -> list[str]:
    adapter_spec = _compiled_package_tool_adapter(semantic_element)
    if adapter_spec is not None:
        tool_ref = adapter_spec.get("tool_ref")
        raw = [tool_ref] if isinstance(tool_ref, str) and tool_ref.strip() else []
    else:
        ctx = record.get("node_context") if isinstance(record.get("node_context"), dict) else {}
        raw = ctx.get("allowed_tools") if isinstance(ctx.get("allowed_tools"), list) else []
    out=[]
    for value in raw:
        if not isinstance(value, str):
            continue
        path=value.strip().replace("\\\\", "/")
        if path and not path.startswith(("/", "http://", "https://")) and ".." not in Path(path).parts:
            out.append(path)
    return list(dict.fromkeys(out))


def _package_tool_base_and_path(package_root: Path, declared_tool: str) -> tuple[Path, Path]:
    direct=(package_root/declared_tool).resolve()
    root=package_root.resolve()
    if direct.is_file() and root in direct.parents:
        return root, direct
    matches=[p.resolve() for p in package_root.rglob(Path(declared_tool).name) if p.is_file() and p.as_posix().endswith('/'+declared_tool)]
    if len(matches) != 1:
        raise ValueError(f"Package tool could not be resolved uniquely: {declared_tool}")
    tool=matches[0]
    base=tool
    for _ in Path(declared_tool).parts:
        base=base.parent
    if package_root.resolve() not in (base, *base.parents):
        raise ValueError("Resolved package tool escaped package root.")
    return base, tool


def _package_tool_command(record: dict[str, Any], state: dict[str, Any], allowed_tool: str, semantic_element: dict[str, Any] | None = None) -> tuple[list[str], str | None]:
    adapter_spec = _compiled_package_tool_adapter(semantic_element)
    if adapter_spec is not None:
        args = adapter_spec.get("args") if isinstance(adapter_spec.get("args"), list) else []
        argv = ["python", allowed_tool] + [str(x) for x in args]
        output_ref = None
        if "--output" in argv:
            idx = argv.index("--output")
            if idx + 1 >= len(argv):
                raise ValueError("Package tool --output argument is missing a value.")
            output_ref = argv[idx + 1]
        return argv, output_ref
    text="\n".join(str(record.get(k) or "") for k in ("question","purpose","description","instruction"))
    match=_PACKAGE_TOOL_COMMAND_RE.search(text)
    if not match:
        raise ValueError(f"Package tool node declares {allowed_tool!r} but no explicit backticked python invocation was found.")
    command=" ".join(match.group(1).split())
    def repl(m: re.Match[str]) -> str:
        path=m.group(1)
        value=_state_subtree(state, path)
        if value in (None, "", [], {}):
            raise ValueError(f"Package tool required state value is unresolved: {path}")
        return shlex.quote(str(value))
    command=_PACKAGE_TOOL_STATE_PLACEHOLDER_RE.sub(repl, command)
    argv=shlex.split(command)
    if len(argv) < 2 or Path(argv[1]).as_posix().lstrip('./') != Path(allowed_tool).as_posix().lstrip('./'):
        raise ValueError("Declared package tool and command invocation do not match.")
    output_ref=None
    if "--output" in argv:
        idx=argv.index("--output")
        if idx+1 >= len(argv): raise ValueError("Package tool --output argument is missing a value.")
        output_ref=argv[idx+1]
    return argv, output_ref


def _resolve_machine_answer_value(value: Any, answer: dict[str, Any]) -> Any:
    if isinstance(value, str) and value.startswith("$answer."):
        cur: Any=answer
        for part in value[len("$answer."):].split('.'):
            if not isinstance(cur, dict) or part not in cur:
                return None
            cur=cur[part]
        return copy.deepcopy(cur)
    if value == "$answer": return copy.deepcopy(answer)
    if isinstance(value, list): return [_resolve_machine_answer_value(v,answer) for v in value]
    if isinstance(value, dict): return {k:_resolve_machine_answer_value(v,answer) for k,v in value.items()}
    return copy.deepcopy(value)


def _execute_package_tool(
    credentials: dict[str, Any], record: dict[str, Any], current_id: str, kind: str, phase: str,
    state: dict[str, Any], routes: list[dict[str, Any]], current_revision: int = 0, semantic_element: dict[str, Any] | None = None,
) -> dict[str, Any]:
    if phase != "enter":
        raise ValueError(f"package_tool executor only supports enter phase for {current_id}")
    allowed=_package_tool_allowed_paths(record, semantic_element)
    if len(allowed) != 1:
        raise ValueError(f"package_tool executor requires exactly one allowed tool for {current_id}; found {len(allowed)}")
    declared_tool=allowed[0]
    package=_active_playbook_package()
    raw_zip=package.get("raw_zip") if isinstance(package,dict) else None
    if not isinstance(raw_zip,(bytes,bytearray)):
        raise ValueError("Package tool execution requires the loaded source ZIP bytes.")
    run_ws=_runtime_workspace()
    package_root=run_ws/"package_tool_workspace"
    marker=package_root/".ordo_extracted"
    if not marker.exists():
        if package_root.exists():
            shutil.rmtree(package_root)
        package_root.mkdir(parents=True,exist_ok=True)
        with zipfile.ZipFile(io.BytesIO(bytes(raw_zip))) as archive:
            infos=[i for i in archive.infolist() if not i.is_dir()]
            if any(not _safe_zip_name(i.filename) for i in infos):
                raise ValueError("Package tool source archive contains unsafe paths.")
            archive.extractall(package_root)
        marker.write_text("ok\n",encoding="utf-8")
    base, tool_file=_package_tool_base_and_path(package_root,declared_tool)
    adapter_spec = _compiled_package_tool_adapter(semantic_element)
    if adapter_spec is not None:
        state_input = str(adapter_spec.get("state_input") or "runtime/state.yaml")
        rel = Path(state_input)
        if rel.is_absolute() or ".." in rel.parts:
            raise ValueError("Compiled package-tool state_input must be package-relative.")
        state_file = (base / rel).resolve()
        if base.resolve() not in (state_file, *state_file.parents):
            raise ValueError("Compiled package-tool state_input escaped package root.")
        state_file.parent.mkdir(parents=True, exist_ok=True)
        state_file.write_text(yaml.safe_dump(state, allow_unicode=True, sort_keys=False), encoding="utf-8")
    argv, output_ref=_package_tool_command(record,state,declared_tool,semantic_element)
    argv[0]=sys.executable
    argv[1]=str(tool_file)
    runtime_root=run_ws.resolve()
    package_base=base.resolve()
    declared_output_ref=output_ref
    resolved_args=[]
    expect_output=False
    runtime_output_dest: Path | None = None
    for idx,arg in enumerate(argv):
        if idx < 2:
            resolved_args.append(arg); continue
        if expect_output:
            rel=Path(arg)
            if rel.is_absolute() or ".." in rel.parts: raise ValueError("Package tool output path must be package-relative.")
            dest=(run_ws/"package_tool_outputs"/rel).resolve()
            if runtime_root not in dest.parents: raise ValueError("Package tool output escaped runtime workspace.")
            dest.parent.mkdir(parents=True,exist_ok=True)
            runtime_output_dest = dest
            resolved_args.append(str(dest)); expect_output=False; continue
        if arg == "--output":
            resolved_args.append(arg); expect_output=True; continue
        if arg.startswith("-"):
            resolved_args.append(arg); continue
        candidate=Path(arg)
        if candidate.is_absolute():
            resolved=candidate.resolve()
            if runtime_root not in (resolved,*resolved.parents) and package_base not in (resolved,*resolved.parents):
                raise ValueError("Package tool input path is outside allowed runtime/package roots.")
            resolved_args.append(str(resolved)); continue
        pkg=(package_base/candidate).resolve()
        run=(run_ws/candidate).resolve()
        if pkg.exists() and package_base in (pkg,*pkg.parents): resolved_args.append(str(pkg))
        elif run.exists() and runtime_root in (run,*run.parents): resolved_args.append(str(run))
        else: resolved_args.append(arg)
    completed=subprocess.run(resolved_args,cwd=str(package_base),capture_output=True,text=True,timeout=120)
    machine: dict[str,Any]
    try:
        machine=json.loads((completed.stdout or "").strip())
        if not isinstance(machine,dict): raise ValueError("root")
    except Exception:
        machine={"status":"ERROR","exit_code":completed.returncode,"stderr":(completed.stderr or "")[-2000:]}
    runtime_artifact = None
    runtime_artifacts: list[dict[str, Any]] = []
    if adapter_spec is not None:
        for declared in adapter_spec.get("declared_outputs") or []:
            if not isinstance(declared, str) or not declared.strip():
                continue
            rel = Path(declared)
            if rel.is_absolute() or ".." in rel.parts:
                raise ValueError("Compiled package-tool declared output must be package-relative.")
            source_artifact = (base / rel).resolve()
            if base.resolve() not in (source_artifact, *source_artifact.parents):
                raise ValueError("Compiled package-tool declared output escaped package root.")
            if not source_artifact.is_file():
                continue
            canonical = (run_ws / rel).resolve()
            if runtime_root not in (canonical, *canonical.parents):
                raise ValueError("Compiled package-tool declared output escaped runtime workspace.")
            canonical.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source_artifact, canonical)
            meta = _runtime_artifact_metadata(rel.as_posix())
            if meta:
                runtime_artifacts.append(meta)
        if runtime_artifacts:
            runtime_artifact = runtime_artifacts[0]
    if declared_output_ref:
        machine.setdefault("report_ref",declared_output_ref)
        # Expose tool output under its declared run-relative path. Package tools
        # execute in an isolated output directory, but downstream nodes and the
        # artifact endpoint must see the canonical report_ref from the playbook.
        if runtime_output_dest is not None and runtime_output_dest.is_file():
            canonical = (run_ws / declared_output_ref).resolve()
            if runtime_root not in (canonical, *canonical.parents):
                raise ValueError("Package tool declared output escaped runtime workspace.")
            canonical.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(runtime_output_dest, canonical)
            runtime_artifact = _runtime_artifact_metadata(declared_output_ref)
            if runtime_artifact and all(x.get("path") != runtime_artifact.get("path") for x in runtime_artifacts):
                runtime_artifacts.append(runtime_artifact)
    machine.setdefault("exit_code",completed.returncode)
    on_answer=record.get("on_answer") if isinstance(record.get("on_answer"),dict) else {}
    if adapter_spec is not None:
        raw_updates = machine.get("state_updates")
        if raw_updates is None:
            updates = {}
        elif not isinstance(raw_updates, dict):
            raise ValueError("Compiled package-tool result state_updates must be an object.")
        else:
            writes = ((semantic_element or {}).get("state_contract") or {}).get("writes") if isinstance((semantic_element or {}).get("state_contract"), dict) else []
            allowed_writes = {str(x) for x in (writes or [])}
            updates = {str(k): copy.deepcopy(v) for k, v in raw_updates.items()}
            unauthorized = sorted(k for k in updates if k not in allowed_writes)
            if unauthorized:
                raise ValueError(f"Compiled package-tool returned unauthorized state_updates: {unauthorized}")
    else:
        mapping=on_answer.get("update_state") if isinstance(on_answer.get("update_state"),dict) else {}
        updates={str(k):_resolve_machine_answer_value(v,machine) for k,v in mapping.items()}
    new_state=_apply_state_updates_canonical(state,updates)
    requested_route_key = str(machine.get("route_key") or "").strip()
    selected = None
    closure_reason = None
    if requested_route_key:
        selected = next((r for r in routes if str(r.get("key") or "") == requested_route_key), None)
        if selected is None:
            closure_reason = "unknown_route_key"
    else:
        # A deterministic tool may omit route_key when the source declares one
        # unambiguous continuation. Prefer the canonical `next` route, then a
        # sole route, then an explicit on_answer.next target. Never guess among
        # multiple unrelated branches.
        selected = next((r for r in routes if r.get("key") == "next"), routes[0] if len(routes) == 1 else None)
        if selected is None:
            target = str(on_answer.get("next") or "")
            selected = next((r for r in routes if r.get("target") == target), None)
        if selected is None:
            closure_reason = "route_not_selected"

    if selected is not None:
        source = package.get("source") if isinstance(package, dict) and isinstance(package.get("source"), dict) else {}
        executable_ids = {
            str(item.get("id"))
            for collection in (source.get("nodes", []), source.get("gates", []))
            if isinstance(collection, list)
            for item in collection
            if isinstance(item, dict) and item.get("id")
        }
        terminal_ids = _declared_terminal_ids(source) if isinstance(source, dict) else set()
        if (executable_ids or terminal_ids) and selected.get("target") not in executable_ids and selected.get("target") not in terminal_ids:
            closure_reason = "unresolved_target"
            selected = None

    closure = {
        "code": "DETERMINISTIC_EXECUTION_ROUTE_CLOSURE",
        "status": "FAIL" if closure_reason else "PASS",
        "reason": closure_reason,
        "requested_route_key": requested_route_key or None,
        "selected_route_key": selected.get("key") if selected else None,
        "next_id": selected.get("target") if selected else None,
        "allowed_routes": copy.deepcopy(routes),
    }
    message=f"Deterministic package tool completed: {machine.get('status','completed')}."
    result = _runtime_only_live_result(
        credentials=credentials,record=record,kind=kind,current_id=current_id,phase=phase,state=state,routes=routes,
        assistant_message=message,await_analyst=False,selected=selected,updates=updates,new_state=new_state,
        reason="alpha20-package-tool",extra_runtime={"runtime_executor":"package_tool","tool":declared_tool,"exit_code":completed.returncode,"machine_result":machine,"command":[Path(x).name if i<2 else x for i,x in enumerate(resolved_args)],"artifact":runtime_artifact,"artifacts":runtime_artifacts,"deterministic_execution_route_closure":closure,"execution_adapter":copy.deepcopy((semantic_element or {}).get("profile_adapter") or {})},current_revision=int(current_revision),
    )
    if closure_reason:
        result["route_key"] = None
        result["next_id"] = None
        result["run_status"] = "halted"
        result["completion_reason"] = "deterministic_route_closure_failed"
        result["rationale_short"] = f"DETERMINISTIC_EXECUTION_ROUTE_CLOSURE failed: {closure_reason}"
        runtime_debug = result.setdefault("debug", {}).setdefault("runtime", {})
        runtime_debug["selected_route_key"] = None
        runtime_debug["next_id"] = None
        runtime_debug["run_status"] = "halted"
        runtime_debug["completion_reason"] = "deterministic_route_closure_failed"
        normalized = runtime_debug.get("normalized_execution_result")
        if isinstance(normalized, dict):
            normalized["route_key"] = None
            normalized["next_id"] = None
            normalized["status"] = "halted"
    return result


def _semantic_model_phase_enabled(semantic_traits: dict[str, Any], phase: str) -> bool:
    """Return whether the Runtime Semantic Plan authorizes semantic-model dispatch.

    `respond` is dynamically valid for a semantic_model executor because a model
    node can legitimately decide at runtime to await analyst clarification even
    when analyst interaction was not statically declared by the source element.
    The caller reaches respond only after an analyst reply for the current node.
    """
    phases = set(str(x) for x in (semantic_traits.get("model_executed_phases") or []) if isinstance(x, str))
    executor = str(semantic_traits.get("runtime_executor") or "")
    return bool(
        phase in phases
        or (not phases and semantic_traits.get("model_executed"))
        or (phase == "respond" and executor == "semantic_model")
    )


def _execute_alpha20_runtime_executor(
    *, credentials: dict[str, Any], record: dict[str, Any], kind: str, current_id: str,
    phase: str, state: dict[str, Any], routes: list[dict[str, Any]], semantic_element: dict[str, Any], executor: str, current_revision: int = 0,
) -> dict[str, Any] | None:
    """Execute small deterministic alpha.20 primitives without an LLM fallback."""
    if phase != "enter":
        return None
    if executor == "terminal":
        result = _runtime_only_live_result(
            credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
            state=state, routes=routes, assistant_message="Процес завершено.", await_analyst=False,
            selected=None, updates={}, new_state=copy.deepcopy(state), reason="alpha20-terminal",
            extra_runtime={"runtime_executor": "terminal"},
        )
        result["terminal"] = True
        result["run_status"] = "completed"
        result["completion_reason"] = "terminal"
        result["debug"]["runtime"]["run_status"] = "completed"
        result["debug"]["runtime"]["completion_reason"] = "terminal"
        return result
    if executor == "delivery_package_builder":
        return _execute_delivery_package_build(credentials, record, current_id, kind, phase, state, routes)
    if executor == "package_tool":
        return _execute_package_tool(credentials, record, current_id, kind, phase, state, routes, current_revision=current_revision, semantic_element=semantic_element)
    if executor not in {"state_patch_template", "artifact_presenter"}:
        return None
    state_contract = semantic_element.get("state_contract") if isinstance(semantic_element.get("state_contract"), dict) else {}
    allow = _alpha20_write_allowlist(record, semantic_element=semantic_element)
    operations = []
    now = datetime.now(timezone.utc).isoformat()
    for spec in state_contract.get("patch_template") or []:
        if not isinstance(spec, dict) or not spec.get("path"):
            continue
        path = str(spec["path"])
        source_class = str(spec.get("source_class") or "")
        if source_class == "constant":
            value = copy.deepcopy(spec.get("value"))
        elif source_class == "runtime_value":
            source = str(spec.get("source") or "")
            value = now if source in {"$runtime.timestamp", "$runtime.now", "$runtime.datetime"} else now
        elif source_class == "confirmed_state":
            source = str(spec.get("source") or "").removeprefix("$state.")
            value = copy.deepcopy(_state_subtree(state, source))
        elif source_class == "analyst_answer":
            raise ValueError(f"runtime executor {executor} cannot materialize analyst_answer source for {path}; semantic respond execution is required")
        else:
            raise ValueError(f"runtime executor {executor} has unsupported patch source_class={source_class!r} for {path}")
        operations.append({"op": "set", "path": path, "value": value, "basis": "derived" if source_class != "constant" else "generated", "reason": f"runtime executor {executor}"})
    patch = {"base_revision": int(current_revision), "operations": operations}
    new_state, commit = apply_state_patch_atomic(canonicalize_runtime_state(state), patch, allowed_paths=allow, current_revision=int(current_revision), value_schemas=_alpha20_value_schemas(semantic_element))
    if not commit.get("committed"):
        raise ValueError("alpha.20 deterministic executor StatePatch rejected: " + "; ".join(commit.get("errors") or []))
    selected = next((r for r in routes if r.get("key") == "next"), routes[0] if len(routes) == 1 else None)
    message = ""
    if executor == "artifact_presenter":
        artifact = record.get("artifact") if isinstance(record.get("artifact"), dict) else {}
        path = str(artifact.get("path") or artifact.get("expected_path") or "")
        message = f"Артефакт готовий до перегляду: {path}" if path else "Артефакт готовий до перегляду."
    updates = {op["path"]: copy.deepcopy(op.get("value")) for op in operations}
    return _runtime_only_live_result(
        credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
        state=state, routes=routes, assistant_message=message, await_analyst=False,
        selected=selected, updates=updates, new_state=new_state,
        reason=f"alpha20-{executor}",
        extra_runtime={"runtime_executor": executor, "state_patch": patch, "state_patch_commit": commit},
        current_revision=int(current_revision),
    )


def _semantic_fallback_policy(payload: dict[str, Any]) -> str:
    raw = str(payload.get("semantic_fallback_policy") or "automatic_safe").strip().lower().replace("-", "_")
    aliases = {
        "automatic": "automatic_safe",
        "auto": "automatic_safe",
        "automatic_safe_fallback": "automatic_safe",
        "ask_before_fallback": "ask",
        "disabled": "disabled",
        "off": "disabled",
    }
    raw = aliases.get(raw, raw)
    return raw if raw in {"disabled", "ask", "automatic_safe"} else "automatic_safe"


def _semantic_recovery_schema() -> dict[str, Any]:
    return {
        "type": "object",
        "additionalProperties": False,
        "required": ["status", "assistant_message", "reason", "state_patch", "next_id"],
        "properties": {
            "status": {"type": "string", "enum": ["resolved", "needs_analyst", "unsupported"]},
            "assistant_message": {"type": "string"},
            "reason": {"type": "string"},
            "next_id": {"type": ["string", "null"]},
            "state_patch": {
                "type": "object",
                "additionalProperties": False,
                "required": ["base_revision", "operations"],
                "properties": {
                    "base_revision": {"type": "integer", "minimum": 0},
                    "operations": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": ["op", "path", "value", "basis", "reason"],
                            "properties": {
                                "op": {"type": "string", "enum": ["set", "replace", "append", "merge", "remove"]},
                                "path": {"type": "string"},
                                "value": {},
                                "basis": {"type": "string", "enum": ["derived", "generated", "confirmed", "analyst"]},
                                "reason": {"type": "string"},
                            },
                        },
                    },
                },
            },
        },
    }




def _runtime_model_candidate_guard(
    candidate: Any, *,
    current_revision: int,
    state: dict[str,Any],
    allowed_paths: list[str] | set[str],
    allowed_route_keys: list[str] | set[str] | None = None,
    value_schemas: dict[str,Any] | None = None,
    operation_variants: dict[str,Any] | None = None,
    require_patch_when_writable: bool = False,
) -> tuple[dict[str,Any],list[str],dict[str,Any]]:
    """Generic VALIDATE -> NORMALIZE -> dry-COMMIT guard for model-produced runtime data."""
    errors=[]
    adaptations=[]
    if not isinstance(candidate,dict):
        return {},["model response must be a JSON object"],{"adaptations":[]}

    normalized=copy.deepcopy(candidate)
    allowed=set(str(x) for x in (allowed_paths or []) if str(x))
    route_allow=set(str(x) for x in (allowed_route_keys or []) if str(x))
    schemas=value_schemas or {}
    variants=operation_variants or {}

    patch=normalized.get("state_patch")
    legacy=normalized.get("state_updates") if isinstance(normalized.get("state_updates"),dict) else None
    if not isinstance(patch,dict):
        if legacy is not None:
            patch=legacy_updates_to_state_patch(legacy,base_revision=current_revision)
            normalized["state_patch"]=patch
            adaptations.append({"kind":"legacy_state_updates_to_state_patch"})
        elif require_patch_when_writable and allowed:
            errors.append("state_patch is required when this element can write state")
            patch={"base_revision":current_revision,"operations":[]}
            normalized["state_patch"]=patch
        else:
            patch={"base_revision":current_revision,"operations":[]}
            normalized["state_patch"]=patch
            adaptations.append({"kind":"insert_empty_state_patch"})
    else:
        patch=copy.deepcopy(patch)
        normalized["state_patch"]=patch

    patch["base_revision"]=int(current_revision)
    patch,metadata_adaptations=_adapt_misnested_patch_metadata(
        patch,allowed_paths=allowed,current_revision=current_revision,
        value_schemas=schemas,operation_variants=variants,
    )
    if metadata_adaptations:
        adaptations.extend(metadata_adaptations)
        normalized["state_patch"]=patch

    validation=validate_state_patch(
        patch,allowed_paths=allowed,current_revision=current_revision,
        value_schemas=schemas,operation_variants=variants,
    )
    errors.extend(str(e) for e in (validation.get("errors") or []))

    route_key=normalized.get("route_key")
    if route_key not in (None,"") and route_allow and str(route_key) not in route_allow:
        errors.append(f"route_key {route_key!r} is not allowed; allowed={sorted(route_allow)}")

    dry_state,dry_commit=apply_state_patch_atomic(
        canonicalize_runtime_state(state),patch,
        allowed_paths=allowed,current_revision=current_revision,
        value_schemas=schemas,operation_variants=variants,
    )
    if not dry_commit.get("committed"):
        for err in (dry_commit.get("errors") or ["StatePatch candidate is not commit-able"]):
            if str(err) not in errors:
                errors.append(str(err))

    return normalized,errors,{
        "adaptations":adaptations,
        "state_patch_validation":validation,
        "dry_run_commit":dry_commit,
    }


def _runtime_model_call_with_guard(
    *, credentials: dict[str,Any], system_text: str, context: dict[str,Any],
    current_revision: int, state: dict[str,Any],
    allowed_paths: list[str] | set[str],
    allowed_route_keys: list[str] | set[str] | None = None,
    value_schemas: dict[str,Any] | None = None,
    operation_variants: dict[str,Any] | None = None,
    require_patch_when_writable: bool = False,
    max_attempts: int = 3,
) -> tuple[dict[str,Any],dict[str,Any],dict[str,Any],str,Any,list[dict[str,Any]]]:
    """Bounded model-call pipeline: parse -> normalize -> validate -> dry commit -> repair retry."""
    attempts=[]
    last_errors=[]
    last_req=last_api=last_raw=last_usage=None
    for attempt in range(1,max(1,int(max_attempts))+1):
        repair_suffix=""
        call_context=context
        if attempt>1:
            repair_suffix=(
                "\n\nThe previous structured response was rejected by the runtime contract. "
                "Repair ONLY the structured response. Preserve the task semantics and do not widen state writes or routing."
            )
            call_context=copy.deepcopy(context)
            call_context["runtime_validation_errors"]=last_errors
            call_context["repair_attempt"]=attempt
        req,api,raw,usage=_provider_api_call(credentials,system_text+repair_suffix,call_context)
        last_req,last_api,last_raw,last_usage=req,api,raw,usage
        try:
            parsed=_parse_model_json(raw)
            normalized,errors,debug=_runtime_model_candidate_guard(
                parsed,current_revision=current_revision,state=state,
                allowed_paths=allowed_paths,allowed_route_keys=allowed_route_keys,
                value_schemas=value_schemas,operation_variants=operation_variants,
                require_patch_when_writable=require_patch_when_writable,
            )
        except Exception as exc:
            normalized={}
            errors=[f"parse/normalization error: {exc}"]
            debug={"adaptations":[]}
        _runtime_debug_log("runtime_model_guard.attempt",{
            "attempt":attempt,"raw_text":raw,"normalized_candidate":normalized,
            "errors":errors,"guard":debug,
        })
        attempts.append({
            "attempt":attempt,
            "errors":copy.deepcopy(errors),
            "guard":debug,
            "usage":copy.deepcopy(usage or {}),
        })
        if not errors:
            return normalized,req,api,raw,usage,attempts
        last_errors=errors

    raise ValueError(
        "model response could not satisfy runtime contract after "
        f"{max(1,int(max_attempts))} attempts: " + "; ".join(last_errors or ["unknown validation error"])
    )



def _normalize_semantic_recovery_state_patch(
    patch: Any, *,
    current_revision: int,
    allowed_paths: set[str],
) -> tuple[dict[str,Any],list[dict[str,Any]]]:
    """Normalize only structurally-equivalent StatePatch spellings; never coerce values."""
    adaptations=[]
    if not isinstance(patch,dict):
        return {"base_revision":current_revision,"operations":[]},adaptations

    # If the model explicitly supplied `operations`, preserve an invalid container
    # type so the normal StatePatch validator can reject it and trigger repair.
    if "operations" in patch and not isinstance(patch.get("operations"),list):
        canonical=copy.deepcopy(patch)
        canonical["base_revision"]=int(current_revision)
        return canonical,adaptations

    if isinstance(patch.get("operations"),list):
        canonical=copy.deepcopy(patch)
        canonical["base_revision"]=int(current_revision)
        ops=[]
        for index,item in enumerate(canonical.get("operations") or []):
            if not isinstance(item,dict):
                ops.append(item)
                continue
            op=copy.deepcopy(item)
            # Generic syntactic shorthand: path+value with omitted op means assignment.
            if not str(op.get("op") or "").strip() and "path" in op and "value" in op:
                op["op"]="set"
                adaptations.append({
                    "kind":"default_missing_op_to_set",
                    "operation_index":index,
                    "path":op.get("path"),
                })
            ops.append(op)
        canonical["operations"]=ops
        return canonical,adaptations

    # No `operations` key at all: treat remaining keys as a common flat path/value patch.
    reserved={"base_revision","operations"}
    flat_keys=[str(k) for k in patch.keys() if str(k) not in reserved]
    if flat_keys:
        operations=[
            {"op":"set","path":key,"value":copy.deepcopy(patch.get(key))}
            for key in flat_keys
        ]
        adaptations.append({"kind":"flat_state_patch_to_set_operations","paths":flat_keys})
        return {"base_revision":int(current_revision),"operations":operations},adaptations

    return {"base_revision":int(current_revision),"operations":[]},adaptations



def _normalize_semantic_recovery_envelope(candidate: Any) -> tuple[dict[str,Any],list[dict[str,Any]]]:
    """Normalize generic alternative recovery envelopes into the canonical contract.

    Accepted structural equivalents:
      {"resolved": {...}}
      {"needs_analyst": {...}}
      {"unsupported": {...}}

    No domain values are invented or coerced.
    """
    if not isinstance(candidate,dict):
        return {},[]
    status=str(candidate.get("status") or "").strip().lower()
    if status in {"resolved","needs_analyst","unsupported"}:
        return copy.deepcopy(candidate),[]

    wrappers=[key for key in ("resolved","needs_analyst","unsupported") if key in candidate]
    if len(wrappers)!=1:
        return copy.deepcopy(candidate),[]

    wrapper=wrappers[0]
    inner=candidate.get(wrapper)
    if not isinstance(inner,dict):
        return copy.deepcopy(candidate),[]

    normalized=copy.deepcopy(inner)
    normalized["status"]=wrapper

    if "state_patch" not in normalized and ("operations" in normalized or "base_revision" in normalized):
        normalized["state_patch"]={
            key:copy.deepcopy(normalized.get(key))
            for key in ("base_revision","operations")
            if key in normalized
        }
    if normalized.get("next_id") in (None,""):
        for alias in ("next_node","next_target"):
            if normalized.get(alias) not in (None,""):
                normalized["next_id"]=normalized.get(alias)
                break
    if not normalized.get("assistant_message") and normalized.get("message") not in (None,""):
        normalized["assistant_message"]=normalized.get("message")

    return normalized,[{"kind":"wrapped_recovery_outcome_to_canonical","wrapper":wrapper}]


def _semantic_recovery_candidate_from_raw(raw_text: str, current_revision: int, allowed_paths: set[str] | None = None) -> dict[str, Any]:
    parsed=_parse_model_json(raw_text)
    if not isinstance(parsed,dict):
        raise ValueError("semantic recovery model result must be a JSON object")
    candidate,envelope_adaptations=_normalize_semantic_recovery_envelope(parsed)
    status=str(candidate.get("status") or "").strip().lower()
    if status not in {"resolved","needs_analyst","unsupported"}:
        raise ValueError("semantic recovery status must be resolved|needs_analyst|unsupported")
    patch,adaptations=_normalize_semantic_recovery_state_patch(
        candidate.get("state_patch"),
        current_revision=current_revision,
        allowed_paths=set(allowed_paths or set()),
    )
    candidate=copy.deepcopy(candidate)
    candidate["status"]=status
    candidate["state_patch"]=patch
    all_adaptations=list(envelope_adaptations)+list(adaptations)
    if all_adaptations:
        candidate["_ordo_normalization"]=all_adaptations
    if candidate.get("next_id") in (None,""):
        for alias in ("next","next_node","next_target"):
            if candidate.get(alias) not in (None,""):
                candidate["next_id"]=candidate.get(alias)
                if alias != "next_id":
                    candidate.pop(alias, None)
                candidate.setdefault("_ordo_normalization",[]).append({"kind":f"{alias}_to_next_id"})
                break
    return candidate


def _semantic_recovery_validate_candidate(
    candidate: dict[str,Any], *,
    allowed_paths: set[str], current_revision: int,
    value_schemas: dict[str,Any], operation_variants: dict[str,Any],
    allowed_targets: list[str],
) -> tuple[bool,list[str]]:
    errors=[]
    patch=candidate.get("state_patch") if isinstance(candidate.get("state_patch"),dict) else {"base_revision":current_revision,"operations":[]}
    validation=validate_state_patch(
        patch,allowed_paths=allowed_paths,current_revision=current_revision,
        value_schemas=value_schemas,operation_variants=operation_variants,
    )
    if not validation.get("valid"):
        errors.extend(validation.get("errors") or [])
    next_id=candidate.get("next_id")
    if next_id not in ("",None) and str(next_id) not in allowed_targets:
        errors.append(f"next_id {next_id!r} is not allowed; allowed={allowed_targets}")
    status=str(candidate.get("status") or "")
    if status=="needs_analyst" and isinstance(patch.get("operations"),list) and patch.get("operations"):
        errors.append("needs_analyst must not commit state before analyst input")
    return (not errors),errors



def _operation_contract_summary(operation_variants: list[dict[str,Any]]) -> dict[str,Any]:
    """Return compact generic requirements derived from declared operation schemas."""
    required=set()
    op_values=set()
    basis_values=set()
    for variant in operation_variants or []:
        if not isinstance(variant,dict):
            continue
        for field in variant.get("required") or []:
            if isinstance(field,str):
                required.add(field)
        props=variant.get("properties") if isinstance(variant.get("properties"),dict) else {}
        op_schema=props.get("op") if isinstance(props.get("op"),dict) else {}
        basis_schema=props.get("basis") if isinstance(props.get("basis"),dict) else {}
        for value in op_schema.get("enum") or []:
            if isinstance(value,str): op_values.add(value)
        for value in basis_schema.get("enum") or []:
            if isinstance(value,str): basis_values.add(value)
    return {
        "required_fields":sorted(required),
        "allowed_op_values":sorted(op_values),
        "allowed_basis_values":sorted(basis_values),
        "operation_variants":copy.deepcopy(operation_variants or []),
    }


def _safe_semantic_model_recovery(
    *, credentials: dict[str, Any], record: dict[str, Any], kind: str, current_id: str,
    phase: str, state: dict[str, Any], routes: list[dict[str, Any]], semantic_element: dict[str, Any],
    current_revision: int, failure_class: str, failure_detail: dict[str, Any] | str,
) -> dict[str, Any] | None:
    """Bounded model recovery for unsupported/unresolved semantics only.

    This is intentionally NOT used for deterministic validator FAIL, package-tool
    execution errors, missing package resources, authority violations, or
    security/permission failures.
    """
    allowed_paths = _alpha20_write_allowlist(record, semantic_element=semantic_element)
    value_schemas = _alpha20_value_schemas(semantic_element)
    operation_variants = _alpha20_operation_variants(semantic_element)
    allowed_targets = sorted({str(r.get("target")) for r in routes if isinstance(r, dict) and r.get("target")})
    semantic_source = semantic_element.get("semantic_source") if isinstance(semantic_element, dict) else {}
    resources = _package_context_for_record(record).get("resolved_resources") or []

    system = (
        "You are Ordo's SAFE semantic recovery executor. The deterministic runtime could not interpret "
        "a declarative semantic construct. Resolve ONLY the supplied execution gap. Do not reinterpret "
        "validated deterministic evidence, do not invent missing domain facts, do not run or simulate package tools, "
        "and do not bypass authority or state contracts. Return JSON only.\n\n"
        "Allowed outcomes:\n"
        "- resolved: only when the supplied state and element semantics are sufficient to choose an allowed next target "
        "and/or produce a StatePatch using ONLY allowed write paths.\n"
        "- needs_analyst: when genuine information is missing. Do not fabricate it.\n"
        "- unsupported: when the construct cannot be safely resolved.\n\n"
        "The runtime will reject any unauthorized write or next target.\n\n"
        "Canonical StatePatch shape: "
        "{\"base_revision\": <runtime-owned>, \"operations\": [{\"op\": \"set\", \"path\": \"allowed.path\", "
        "\"value\": <schema-valid value>, \"basis\": \"<allowed basis>\", \"reason\": \"<short provenance reason>\"}]}. "
        "Use the supplied operation_contract and allowed_value_schemas exactly."
    )
    context = {
        "element_id": current_id,
        "element_kind": kind,
        "phase": phase,
        "failure_class": failure_class,
        "failure_detail": failure_detail,
        "semantic_source": semantic_source,
        "runtime_state": copy.deepcopy(state),
        "allowed_write_paths": sorted(allowed_paths),
        "allowed_next_targets": allowed_targets,
        "allowed_value_schemas": copy.deepcopy(value_schemas),
        "operation_contract": _operation_contract_summary(operation_variants),
        "routes": copy.deepcopy(routes),
        "resources": resources,
        "__response_json_schema": _semantic_recovery_schema(),
    }
    req_body, api_response, raw_text, usage = _provider_api_call(credentials, system, context)
    _runtime_debug_log("semantic_recovery.initial_raw",{
        "element_id":current_id,"failure_class":failure_class,"raw_text":raw_text,
        "usage":usage,"allowed_write_paths":sorted(allowed_paths),"allowed_next_targets":allowed_targets,
    })

    # Bounded contract recovery: exactly one schema-repair retry.
    repair_debug=None
    candidate=None
    validation_errors=[]
    try:
        candidate=_semantic_recovery_candidate_from_raw(raw_text,current_revision,allowed_paths)
        valid,validation_errors=_semantic_recovery_validate_candidate(
            candidate,
            allowed_paths=allowed_paths,current_revision=current_revision,
            value_schemas=value_schemas,operation_variants=operation_variants,
            allowed_targets=allowed_targets,
        )
    except Exception as exc:
        valid=False
        validation_errors=[str(exc)]

    _runtime_debug_log("semantic_recovery.initial_validation",{
        "element_id":current_id,"candidate":candidate,"valid":valid,"validation_errors":validation_errors,
    })
    if not valid:
        repair_system=(
            system
            + "\n\nYour previous response was structurally invalid. Repair ONLY the JSON response shape/contract. "
              "Do not change the underlying semantic conclusion unless required by the validation errors. "
              "Do not add writes outside allowed_write_paths and do not add a next target outside allowed_next_targets."
        )
        repair_context=copy.deepcopy(context)
        repair_context["previous_raw_response"]=raw_text
        repair_context["previous_candidate"]=candidate
        repair_context["validation_errors"]=validation_errors
        repair_context["allowed_value_schemas"]=copy.deepcopy(value_schemas)
        repair_context["operation_contract"]=_operation_contract_summary(operation_variants)
        repair_context["repair_instruction"]=(
            "Return one corrected JSON object matching __response_json_schema. "
            "Each state_patch.operations item MUST satisfy one declared operation variant, including every required metadata field. "
            "Respect allowed_value_schemas and operation_contract exactly; do not coerce or invent domain values."
        )
        repair_req,repair_api,repair_raw,repair_usage=_provider_api_call(credentials,repair_system,repair_context)
        _runtime_debug_log("semantic_recovery.repair_raw",{
            "element_id":current_id,"raw_text":repair_raw,"usage":repair_usage,
            "validation_errors_from_initial":validation_errors,
        })
        repaired=None
        repaired_errors=[]
        try:
            repaired=_semantic_recovery_candidate_from_raw(repair_raw,current_revision,allowed_paths)
            repaired_valid,repaired_errors=_semantic_recovery_validate_candidate(
                repaired,
                allowed_paths=allowed_paths,current_revision=current_revision,
                value_schemas=value_schemas,operation_variants=operation_variants,
                allowed_targets=allowed_targets,
            )
        except Exception as exc:
            repaired_valid=False
            repaired_errors=[str(exc)]
        _runtime_debug_log("semantic_recovery.repair_validation",{
            "element_id":current_id,"candidate":repaired,"valid":repaired_valid,"validation_errors":repaired_errors,
        })
        repair_debug={
            "initial_errors":validation_errors,
            "request_body":repair_req,
            "api_response":repair_api,
            "usage":repair_usage,
            "repaired_candidate":repaired,
            "repaired_errors":repaired_errors,
        }
        if not repaired_valid or repaired is None:
            _runtime_debug_log("semantic_recovery.halt",{
                "element_id":current_id,"failure_class":"contract_unsatisfiable_by_model",
                "initial_errors":validation_errors,"repair_errors":repaired_errors,
            })
            halted=_runtime_only_live_result(
                credentials=credentials,record=record,kind=kind,current_id=current_id,phase=phase,
                state=state,routes=routes,
                assistant_message="Semantic recovery could not produce a valid structured result after one repair attempt.",
                await_analyst=False,selected=None,updates={},new_state=copy.deepcopy(state),
                reason="semantic-recovery-contract-unsatisfied",
                extra_runtime={
                    "runtime_executor":"semantic_model_recovery",
                    "failure_class":"contract_unsatisfiable_by_model",
                    "semantic_fallback":{
                        "status":"unsupported",
                        "failure_class":failure_class,
                        "initial_errors":validation_errors,
                        "schema_repair":repair_debug,
                    },
                },
                current_revision=current_revision,
            )
            halted["run_status"]="halted"
            halted["completion_reason"]="contract_unsatisfiable_by_model"
            halted["failure_class"]="contract_unsatisfiable_by_model"
            return halted
        candidate=repaired
        usage={"initial":usage,"repair":repair_usage}

    status=str(candidate.get("status") or "").strip().lower()
    reason=str(candidate.get("reason") or "").strip()
    assistant_message=str(candidate.get("assistant_message") or "").strip()
    patch=candidate.get("state_patch") if isinstance(candidate.get("state_patch"),dict) else {"base_revision":current_revision,"operations":[]}
    patch["base_revision"]=int(current_revision)
    next_id = candidate.get("next_id")
    if next_id in ("", None):
        next_id = None
    else:
        next_id = str(next_id)
        if next_id not in allowed_targets:
            raise ValueError(f"semantic recovery next_id {next_id!r} is not allowed; allowed={allowed_targets}")

    if status == "unsupported":
        return None
    if status == "needs_analyst":
        if patch.get("operations"):
            raise ValueError("semantic recovery needs_analyst must not commit state before analyst input")
        return _runtime_only_live_result(
            credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
            state=state, routes=routes, assistant_message=assistant_message or "Additional analyst input is required.",
            await_analyst=True, selected=None, updates={}, new_state=copy.deepcopy(state),
            reason="semantic-model-recovery-needs-analyst",
            extra_runtime={
                "runtime_executor": "semantic_model_recovery",
                "semantic_fallback": {"status": status, "failure_class": failure_class, "reason": reason, "usage": usage},
            },
            current_revision=current_revision,
        )

    new_state, commit = apply_state_patch_atomic(
        canonicalize_runtime_state(state), patch, allowed_paths=allowed_paths,
        current_revision=current_revision, value_schemas=value_schemas, operation_variants=operation_variants,
    )
    if not commit.get("committed"):
        raise ValueError("semantic recovery StatePatch commit rejected: " + "; ".join(commit.get("errors") or []))
    selected = next((r for r in routes if str(r.get("target") or "") == str(next_id or "")), None)
    if next_id is None and len(routes) == 1:
        selected = routes[0]
    updates = {str(op.get("path")): copy.deepcopy(op.get("value")) for op in patch.get("operations") or [] if isinstance(op, dict) and op.get("path")}
    return _runtime_only_live_result(
        credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
        state=state, routes=routes, assistant_message=assistant_message,
        await_analyst=False, selected=selected, updates=updates, new_state=new_state,
        reason="semantic-model-recovery",
        extra_runtime={
            "runtime_executor": "semantic_model_recovery",
            "semantic_fallback": {
                "status": status, "failure_class": failure_class, "reason": reason,
                "usage": usage, "request_body": req_body, "api_response": api_response,
                "state_patch": patch, "state_patch_commit": commit,
                "schema_repair": repair_debug,
            },
        },
        current_revision=current_revision,
    )


def _validate_live_transition_provenance(source: dict[str, Any], target_node_id: str, previous_node_id: str | None, entry_mode: str) -> dict[str, Any]:
    graph = source.get("graph_contract") if isinstance(source.get("graph_contract"), dict) else {}
    contract = graph.get("transition_provenance") if isinstance(graph.get("transition_provenance"), dict) else {}
    if not contract.get("enabled", False):
        return {"status": "passed", "mode": "transition_provenance_not_enabled", "issues": []}
    record, kind = _record_by_id(source, target_node_id)
    if not record or not kind:
        return {"status": "blocked", "mode": "transition_provenance", "issues": [{"severity":"error","code":"RUNTIME_TARGET_NODE_MISSING","message":f"Target graph vertex not found: {target_node_id}","source_node":previous_node_id,"target_node":target_node_id,"direction":"entry"}]}
    if previous_node_id in (None, ""):
        modes = set(str(x) for x in (record.get("entry_modes") or []) if isinstance(x, str))
        if entry_mode in modes:
            return {"status":"passed","mode":"transition_provenance","entry_mode":entry_mode,"issues":[]}
        return {"status":"blocked","mode":"transition_provenance_recovery","issues":[{"severity":"error","code":"RUNTIME_ENTRY_PROVENANCE_MISSING","message":"Graph vertex entry requires previous_node_id or an explicitly allowed entry mode.","source_node":previous_node_id,"target_node":target_node_id,"direction":"entry","entry_mode":entry_mode}]}
    allowed = set(str(x) for x in (record.get("allowed_from") or record.get("incoming_from") or []) if isinstance(x, str))
    if str(previous_node_id) not in allowed:
        return {"status":"blocked","mode":"transition_provenance_recovery","issues":[{"severity":"error","code":"RUNTIME_PREDECESSOR_NOT_ALLOWED","message":f"Graph vertex {target_node_id} does not accept direct entry from {previous_node_id}.","source_node":previous_node_id,"target_node":target_node_id,"direction":"inbound","allowed_from":sorted(allowed)}]}
    return {"status":"passed","mode":"transition_provenance","source_node":str(previous_node_id),"target_node":target_node_id,"issues":[]}


def _transition_provenance_block_result(credentials: dict[str, Any], record: dict[str, Any], kind: str, current_id: str, phase: str, state: dict[str, Any], routes: list[dict[str, str]], current_revision: int, result: dict[str, Any]) -> dict[str, Any]:
    first = (result.get("issues") or [{}])[0]
    message = str(first.get("message") or "Transition provenance blocked execution.")
    out = _runtime_only_live_result(
        credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
        state=state, routes=routes, assistant_message=message, await_analyst=False, selected=None,
        updates={}, new_state=copy.deepcopy(state), reason="transition-provenance-blocked",
        extra_runtime={"transition_provenance": copy.deepcopy(result)}, current_revision=current_revision,
    )
    out["next_id"] = None
    out["route_key"] = None
    out["run_status"] = "halted"
    out["completion_reason"] = "transition_provenance"
    out["debug"]["runtime"]["run_status"] = "halted"
    out["debug"]["runtime"]["completion_reason"] = "transition_provenance"
    return out


def _call_openai_live(payload: dict[str, Any]) -> dict[str, Any]:
    package_id = str(payload.get("package_id") or "")
    package = PLAYBOOK_PACKAGES.get(package_id)
    # Backward-compatible fallback for a server started before the registry was populated.
    if package is None and package_id and package_id == str(PLAYBOOK_PACKAGE.get("id") or ""):
        package = PLAYBOOK_PACKAGE
    if package is None:
        raise ValueError("Live execution requires a playbook ZIP package loaded by this editor server.")
    package_token = _ACTIVE_PLAYBOOK_PACKAGE.set(package)
    run_context = {
        "package_id": package_id,
        "session_id": str(payload.get("session_id") or "").strip(),
        "run_id": str(payload.get("run_id") or "").strip(),
    }
    run_token = _ACTIVE_RUN_CONTEXT.set(run_context)
    try:
        return _call_openai_live_impl(payload)
    finally:
        _ACTIVE_RUN_CONTEXT.reset(run_token)
        _ACTIVE_PLAYBOOK_PACKAGE.reset(package_token)

def _call_openai_live_impl(payload: dict[str, Any]) -> dict[str, Any]:
    # Historical function name retained for endpoint compatibility; it now dispatches to the selected provider.
    credentials = _live_credentials(payload)
    semantic_fallback_policy = _semantic_fallback_policy(payload)
    source = payload.get("source")
    state = _canonicalize_runtime_state(payload.get("state") or {})
    current_revision = int(payload.get("state_revision") or 0)
    if current_revision < 0: raise ValueError("state_revision must be non-negative")
    history = payload.get("history") or []
    current_id = str(payload.get("current_id") or "")
    analyst_input = str(payload.get("analyst_input") or "")
    analyst_override_context = str(payload.get("analyst_override_context") or "").strip()
    session_id = str(payload.get("session_id") or "")
    run_id = str(payload.get("run_id") or "").strip() or session_id
    analyst_attachments, analyst_attachment_debug = _prepare_analyst_attachments(payload.get("attachments") or [], session_id)
    if not isinstance(source, dict): raise ValueError("Live execution requires a loaded playbook source.")
    record, kind = _record_by_id(source, current_id)
    if not record or not kind: raise ValueError(f"Unknown current playbook element: {current_id}")
    routes = _live_routes(record, kind)
    route_keys = [r["key"] for r in routes]
    phase = str(payload.get("phase") or ("respond" if analyst_input else "enter"))
    if phase not in {"enter", "respond"}: raise ValueError("Live execution phase must be 'enter' or 'respond'.")
    if phase == "enter":
        previous_node_id = payload.get("previous_node_id")
        entry_mode = str(payload.get("entry_mode") or ("root" if current_id == str((source.get("graph_contract") or {}).get("entry_node") or "") and previous_node_id in (None, "") else "transition"))
        provenance = _validate_live_transition_provenance(source, current_id, None if previous_node_id in (None, "") else str(previous_node_id), entry_mode)
        if provenance.get("status") == "blocked":
            return _transition_provenance_block_result(credentials, record, kind, current_id, phase, state if isinstance(state, dict) else {}, routes, current_revision, provenance)
    if semantic_fallback_policy == "ask" and phase == "respond" and _is_affirmative_confirmation(analyst_input):
        semantic_fallback_policy = "automatic_safe"
    analyst_input = _effective_file_ref_answer(record, kind, phase, analyst_input, analyst_attachment_debug)

    human_policy = _human_interaction_policy(record, kind)
    semantic_pre = _semantic_plan_element(current_id)
    semantic_pre_traits = semantic_pre.get("execution_traits") if isinstance(semantic_pre, dict) and isinstance(semantic_pre.get("execution_traits"), dict) else {}
    semantic_executor = str(semantic_pre_traits.get("runtime_executor") or "")
    semantic_requires_analyst = bool(semantic_pre_traits.get("requires_analyst"))
    if semantic_executor and semantic_executor not in {"human_interaction", "human_gate"} and not semantic_requires_analyst:
        # Runtime Semantic Plan is authoritative: presentation metadata cannot
        # reclassify model/deterministic execution as analyst-owned interaction.
        human_policy = {"requires_human": False, "direct_enter": False, "direct_respond": False}
    if kind == "gate" and phase == "enter" and str(record.get("trust_class") or "").lower() == "deterministic":
        inherited_route, inherited_evidence = _inherited_human_gate_route(
            source, record, current_id, routes, history if isinstance(history, list) else []
        )
        if inherited_route is not None:
            current_state = state if isinstance(state, dict) else {}
            result = _runtime_only_live_result(
                credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
                state=current_state, routes=routes, assistant_message="", await_analyst=False,
                selected=inherited_route, updates={}, new_state=copy.deepcopy(current_state),
                reason="inherited-branch-evidence",
                extra_runtime={
                    "inherited_branch_evidence": inherited_evidence,
                    "normalized_execution_result": {
                        "element_id": current_id, "phase": phase, "status": "completed",
                        "gate_result": "pass", "state_updates": {},
                        "route_key": inherited_route.get("key"), "next_id": inherited_route.get("target"),
                    },
                },
            )
            result["rationale_short"] = "Gate passed from explicit incoming analyst branch evidence declared by the canonical gate condition."
            return result
        gate_result = _execute_deterministic_gate(credentials, record, current_id, state if isinstance(state, dict) else {}, routes, history)
        gate_debug = ((gate_result.get("debug") or {}).get("runtime") or {}).get("deterministic_gate") if isinstance(gate_result, dict) else None
        if isinstance(gate_debug, dict) and gate_debug.get("result") == "unresolved":
            # Canonical mechanical gates are owned by code/runtime/script. An unsupported
            # mechanical rule must fail closed as a deterministic capability/contract gap;
            # it must never be silently reclassified as an LLM semantic judgment.
            gate_result["run_status"] = "halted"
            gate_result["completion_reason"] = "deterministic_mechanical_validation_unresolved"
            gate_result["failure_class"] = "deterministic_validation_gap"
            gate_result.setdefault("debug",{}).setdefault("runtime",{})["semantic_fallback_available"] = False
            gate_result["debug"]["runtime"]["mechanical_model_calls"] = 0
            if gate_debug.get("profile_contract_gap"):
                gate_result["completion_reason"] = "generated_profile_validation_contract_incomplete"
                gate_result["failure_class"] = "profile_contract_gap"
            elif gate_debug.get("execution_error"):
                gate_result["completion_reason"] = "deterministic_validation_execution_error"
                gate_result["failure_class"] = "deterministic_execution_error"
        return gate_result

    if (
        kind == "node" and phase == "enter"
        and str(record.get("action") or "").upper() == "DOCUMENT.GENERATE"
        and (semantic_pre is None or semantic_executor == "document_generate")
    ):
        # Runtime Semantic Plan is authoritative. This legacy source-action
        # executor is used only when there is no semantic element, or when the
        # semantic plan itself explicitly selected document_generate. In
        # particular, a package_tool semantic executor must never be intercepted
        # by this compatibility branch.
        return _execute_document_generate(credentials, record, current_id, kind, phase, state if isinstance(state, dict) else {}, routes)

    retry_schema_repair: dict[str, Any] | None = None
    if kind == "node" and phase == "enter" and _looks_like_proposal_confirmation_record(record):
        retry_table = _retry_existing_table_message(record, current_id, history, state if isinstance(state, dict) else {})
        if retry_table is not None:
            message, retry_audit = retry_table
            if retry_audit.get("requires_schema_repair"):
                retry_schema_repair = retry_audit
            else:
                migrated_value = retry_audit.get("migrated_value")
                target = str(retry_audit.get("target") or "")
                migrated_state = state if isinstance(state, dict) else {}
                if target and migrated_value is not None:
                    migrated_state = _apply_state_updates_canonical(migrated_state, {target: migrated_value})
                return _runtime_only_live_result(
                    credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
                    state=state if isinstance(state, dict) else {}, routes=routes, assistant_message=message, await_analyst=True,
                    new_state=migrated_state, updates=({target: migrated_value} if target and migrated_value is not None else {}),
                    reason="retry-existing-structured-output", extra_runtime={"retry_existing_output": retry_audit}
                )

    if human_policy.get("requires_human") and phase == "enter" and human_policy.get("direct_enter"):
        inherited_route, inherited_evidence = _inherited_human_gate_route(
            source, record, current_id, routes, history if isinstance(history, list) else []
        )
        if inherited_route is not None:
            return _runtime_only_live_result(credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
                state=state if isinstance(state, dict) else {}, routes=routes, assistant_message="", await_analyst=False,
                selected=inherited_route, reason="inherited-human-decision", extra_runtime={"inherited_human_decision": inherited_evidence})
        semantic_for_question = _semantic_plan_element(current_id)
        semantic_ai = semantic_for_question.get("analyst_interaction") if isinstance(semantic_for_question, dict) and isinstance(semantic_for_question.get("analyst_interaction"), dict) else {}
        question = str(human_policy.get("question") or record.get("question") or semantic_ai.get("question") or "Потрібне рішення аналітика.")
        if kind == "gate" and not (human_policy.get("question") or record.get("question")):
            condition = str(record.get("condition") or (semantic_for_question.get("semantic_source") or {}).get("condition") if isinstance(semantic_for_question, dict) else "").strip()
            if condition:
                question = f"Потрібне рішення аналітика для {current_id}. Перевірте критерій: {condition}\n\nЯкщо критерій виконано — оберіть «Погодити». Якщо ні — оберіть «Потрібне виправлення»; система покаже маршрут відновлення та передасть контекст помилки."
        return _runtime_only_live_result(credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
            state=state if isinstance(state, dict) else {}, routes=routes, assistant_message=question, await_analyst=True,
            reason="declared-human-input")

    if human_policy.get("requires_human") and phase == "respond" and human_policy.get("direct_respond") and not _selected_branch_requires_ai(record, analyst_input):
        selected = _select_direct_answer_route(record, routes, analyst_input)
        if selected is None:
            allowed = human_policy.get("allowed_values") or [r.get("key") for r in routes]
            choices = ", ".join(str(v) for v in allowed if v)
            message = f"Невідома відповідь. Дозволені значення: {choices}." if choices else "Не вдалося визначити дозволений перехід для цієї відповіді."
            return _runtime_only_live_result(credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
                state=state if isinstance(state, dict) else {}, routes=routes, assistant_message=message, await_analyst=True,
                reason="unmatched-human-input")
        new_state, updates = _apply_direct_answer_updates(record, state if isinstance(state, dict) else {}, analyst_input)
        return _runtime_only_live_result(credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
            state=state if isinstance(state, dict) else {}, routes=routes, assistant_message="", await_analyst=False, selected=selected,
            updates=updates, new_state=new_state, reason="deterministic-human-route", extra_runtime={
                "normalized_execution_result": {
                    "element_id": current_id, "phase": phase, "status": "completed",
                    "decision": {"raw_answer": analyst_input, "route_key": selected.get("key"), "route_target": selected.get("target")},
                    "state_updates": copy.deepcopy(updates), "route_key": selected.get("key"), "next_id": selected.get("target"),
                }
            })

    # A one-word confirmation is not enough information for an LLM to recreate a
    # structured proposal. For canonical proposal-confirmation operations, commit
    # the exact table previously shown to the analyst instead of regenerating it.
    if kind == "node" and phase == "respond":
        generic_commit = _generic_table_confirmation_result(
            record, current_id, analyst_input, history, state if isinstance(state, dict) else {}, routes
        )
        if generic_commit is not None:
            generic_updates, selected, proposal_text = generic_commit
            generic_updates, source_ref_audit = _canonicalize_confirmed_source_references(
                state if isinstance(state, dict) else {}, generic_updates
            )
            new_state = _apply_state_updates_canonical(state if isinstance(state, dict) else {}, generic_updates)
            return _runtime_only_live_result(
                credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
                state=state if isinstance(state, dict) else {}, routes=routes,
                assistant_message="Дякую. Підтверджений структурований результат зафіксовано без повторної генерації.",
                await_analyst=False, selected=selected, updates=generic_updates, new_state=new_state,
                reason="generic-proposal-preserving-confirmation",
                extra_runtime={
                    "confirmed_proposal": {"source":"previous-assistant-message","proposal_chars":len(proposal_text)},
                    "source_reference_reconciliation": source_ref_audit,
                },
            )

        proposal_commit = _proposal_confirmation_result(
            record, current_id, analyst_input, history, state if isinstance(state, dict) else {}, routes
        )
        if proposal_commit is not None:
            confirmed_rows, selected, proposal_text = proposal_commit
            updates = {"source_attribute_mapping.rows": confirmed_rows}
            new_state = copy.deepcopy(state) if isinstance(state, dict) else {}
            new_state["source_attribute_mapping.rows"] = copy.deepcopy(confirmed_rows)
            _set_dotted_state(new_state, "source_attribute_mapping.rows", confirmed_rows)
            return _runtime_only_live_result(
                credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
                state=state if isinstance(state, dict) else {}, routes=routes,
                assistant_message="Дякую. Підтверджену таблицю атрибутів зафіксовано без втрати її структури.",
                await_analyst=False, selected=selected, updates=updates, new_state=new_state,
                reason="proposal-preserving-confirmation",
                extra_runtime={
                    "confirmed_proposal": {
                        "source": "previous-assistant-message",
                        "row_count": len(confirmed_rows),
                        "proposal_chars": len(proposal_text),
                    }
                },
            )

    # alpha.20 Runtime Semantic Plan is a separate execution path from legacy V6 compiled phases.
    semantic_element = _semantic_plan_element(current_id)
    semantic_traits = semantic_element.get("execution_traits") if isinstance(semantic_element, dict) and isinstance(semantic_element.get("execution_traits"), dict) else {}
    semantic_runtime_executor = str(semantic_traits.get("runtime_executor") or "")
    semantic_model_execution = bool(semantic_element and _semantic_model_phase_enabled(semantic_traits, phase))
    semantic_source = semantic_element.get("semantic_source") if isinstance(semantic_element, dict) and isinstance(semantic_element.get("semantic_source"), dict) else {}
    semantic_action = str(semantic_source.get("action") or record.get("action") or "")
    if semantic_action.startswith("PACKAGE.BUILD_") and isinstance(record.get("package"), dict):
        semantic_model_execution = False
        semantic_runtime_executor = "delivery_package_builder"
    if semantic_element is not None and not semantic_model_execution:
        deterministic_result = _execute_alpha20_runtime_executor(
            credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
            state=state if isinstance(state, dict) else {}, routes=routes,
            semantic_element=semantic_element, executor=semantic_runtime_executor, current_revision=current_revision,
        )
        if deterministic_result is not None:
            return deterministic_result

    # Prefer a validated compiled execution artifact. V7 semantic execution never silently falls back to YAML.
    compiled_element, compiled_phase = _compiled_element_for(current_id, phase)
    compiled_prompt = _compiled_prompt_text(compiled_element, compiled_phase) if compiled_element and compiled_phase else ""

    system_text = _system_contract_for_call(kind, phase, human_policy) + "\n\n" + _analyst_language_instruction()
    model_record = _project_execution_record(record, kind, phase, human_policy)
    compiled_element_rejection_reasons: list[str] = []
    semantic_call_meta: dict[str, Any] = {}
    if semantic_model_execution:
        execution_mechanism = "runtime_semantic_plan"
    elif semantic_element is not None:
        # Reaching the model dispatch for a non-model semantic element is a configuration
        # error. Dedicated deterministic/human executors must have intercepted it earlier.
        if semantic_fallback_policy == "automatic_safe":
            recovered = _safe_semantic_model_recovery(
                credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
                state=state if isinstance(state, dict) else {}, routes=routes, semantic_element=semantic_element,
                current_revision=current_revision, failure_class="unsupported_runtime_executor",
                failure_detail={"runtime_executor": semantic_runtime_executor or "unset", "semantic_kind": semantic_element.get("kind")},
            )
            if recovered is not None:
                return recovered
        if semantic_fallback_policy == "ask":
            ask_result = _runtime_only_live_result(
                credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
                state=state if isinstance(state, dict) else {}, routes=routes,
                assistant_message="This runtime semantic construct is unsupported deterministically. Model recovery is available if you approve it.",
                await_analyst=True, selected=None, updates={}, new_state=copy.deepcopy(state if isinstance(state,dict) else {}),
                reason="semantic-fallback-approval-required",
                extra_runtime={"semantic_fallback_available":True,"semantic_fallback_policy":"ask","unsupported_runtime_executor":semantic_runtime_executor or "unset"},
                current_revision=current_revision,
            )
            return ask_result
        raise ValueError(f"Runtime Semantic Plan has no executor for {current_id}/{phase}: kind={semantic_element.get('kind')}, runtime_executor={semantic_runtime_executor or 'unset'}; safe model recovery did not resolve the construct.")
    elif compiled_prompt and compiled_element and compiled_phase:
        compiled_ok, compiled_element_rejection_reasons = _compiled_element_validation(
            record=record, model_record=model_record, kind=kind, phase=phase,
            state=state if isinstance(state, dict) else {}, routes=routes,
            element=compiled_element, phase_spec=compiled_phase,
        )
        execution_mechanism = "compiled_llm_plan" if compiled_ok else "yaml_fallback"
    else:
        execution_mechanism = "yaml_fallback"
        if compiled_element and compiled_phase and not compiled_prompt:
            compiled_element_rejection_reasons = ["compiled element has no executable prompt for this phase"]

    if execution_mechanism == "runtime_semantic_plan":
        system_text, context, semantic_call_meta = _assemble_runtime_semantic_call(
            semantic_element, current_id, phase, state if isinstance(state, dict) else {}, history, analyst_input
        )
        package_context = {"resolved_resources": semantic_call_meta.get("resolved_resources") or []}
        compact_state = semantic_call_meta.get("runtime_state") or {}
        compact_history = semantic_call_meta.get("recent_history") or []
    elif execution_mechanism == "compiled_llm_plan":
        system_text += "\n\nCompiled Ordo instruction:\n" + compiled_prompt
        required_state = _compiled_list(compiled_element, compiled_phase, "required_state")
        compact_state = _project_state_by_paths(state, required_state)
        compact_state = _apply_compiled_state_defaults(compact_state, _compiled_state_defaults(compiled_phase))
        resolved_resources = _compiled_resources(compiled_element, compiled_phase)
        package_context = {"resolved_resources": resolved_resources}
        compact_history = []
        context = {
            "execution_phase": phase,
            "element_id": current_id,
            "element_kind": kind,
        }
        output_contract = compiled_phase.get("output_contract") if isinstance(compiled_phase, dict) else None
        if output_contract is None and isinstance(compiled_element, dict):
            output_contract = compiled_element.get("output_contract")
        if isinstance(output_contract, (dict, list)):
            context["output_contract"] = output_contract
    else:
        package_context = _package_context_for_record(model_record)
        # YAML fallback is the safety path. If a compiled phase was rejected, pass the
        # full canonical runtime state so a compiler/dependency defect cannot become a
        # false-negative gate result. For packages with no compiled phase at all, retain
        # the normal YAML dependency projection to keep prompts compact.
        if compiled_element_rejection_reasons:
            compact_state = _fallback_runtime_state(state)
            fallback_state_mode = "full_canonical_after_compiled_rejection"
        else:
            compact_state = _project_runtime_state(state, model_record, kind, phase)
            fallback_state_mode = "yaml_dependency_projection"
        compact_history = _project_live_history(history, model_record)
        context = {
            "execution_phase": phase,
            "current_element": model_record,
            "element_kind": kind,
            "runtime_state_mode": fallback_state_mode,
        }
    if kind == "node" and phase == "enter" and _looks_like_proposal_confirmation_record(record):
        existing_outputs = _existing_structured_outputs_for_record(record, state if isinstance(state, dict) else {})
        if retry_schema_repair:
            repair_target = str(retry_schema_repair.get("target") or "")
            repair_value = retry_schema_repair.get("migrated_value")
            context["existing_structured_output"] = ({repair_target: repair_value} if repair_target else existing_outputs)
            context["retry_mode"] = "schema_repair_existing_output"
            context["schema_repair_contract"] = {
                "target": repair_target,
                "required_columns": retry_schema_repair.get("columns") or [],
                "unresolved_columns": retry_schema_repair.get("unresolved_columns") or [],
                "rules": [
                    "preserve every existing row and confirmed value",
                    "fill a missing field only from supplied runtime_state, source context, or explicit row aliases",
                    "never invent source evidence, identifiers, field names, values, or provenance",
                    "if a field cannot be grounded, mark it UNRESOLVED and explain exactly what evidence is missing",
                    "return the repaired structure as the analyst-facing proposal; do not select a graph route",
                ],
            }
            system_text += (
                "\n\nRuntime contract-aware schema-repair rule: the existing structured output is older or narrower "
                "than the current canonical contract. Repair schema only. Preserve all existing rows and values. "
                "Use only supplied runtime_state/source evidence and explicit aliases to populate missing contract fields. "
                "Never invent evidence or domain facts. If a required field cannot be grounded, render it explicitly as "
                "UNRESOLVED and state what evidence is missing. Do not generate a fresh unrelated proposal and do not "
                "change graph routing; show the repaired structure to the analyst for confirmation/correction."
            )
        elif existing_outputs:
            context["existing_structured_output"] = existing_outputs
            context["retry_mode"] = "refine_existing_output"
            system_text += (
                "\n\nRuntime retry-continuity rule: this node already owns structured output in runtime state. "
                "Treat context.existing_structured_output as the authoritative starting point. "
                "Refine or explain only what still needs correction; do not regenerate an unrelated fresh draft, "
                "do not discard analyst-confirmed rows/fields, and preserve unchanged values exactly."
            )

    # Routes are useful only when the model is actually allowed to choose one.
    # Human-controlled enter is runtime-forced to wait, so do not expose a route there.
    if not (human_policy.get("requires_human") and phase == "enter"):
        context["allowed_routes"] = routes
    if phase == "enter" and analyst_override_context:
        context["analyst_override_context"] = analyst_override_context
        system_text += (
            "\n\nAnalyst recovery clarification: context.analyst_override_context contains a supplemental "
            "instruction supplied immediately before entering this element. Apply it to this element's task only. "
            "It may refine requested content or coverage, but it MUST NOT override canonical state paths, output schema, "
            "allowed routes, deterministic evidence, resource truth, or any playbook constraint. Never treat the clarification "
            "as evidence for facts it does not itself establish, and do not persist it beyond this call unless the element's "
            "canonical state updates explicitly encode grounded results."
        )
    if phase == "respond" and analyst_input:
        context["analyst_input"] = analyst_input
        prior_proposal = _previous_assistant_proposal(history, current_id)
        if prior_proposal and _looks_like_proposal_confirmation_record(record):
            context["prior_structured_proposal"] = prior_proposal
            proposal_contract = _canonical_proposal_contract(record)
            if proposal_contract:
                context["proposal_contract"] = proposal_contract
            system_text += (
                "\n\nRuntime proposal-continuity rule: this respond phase confirms or corrects "
                "the analyst-facing proposal shown on enter. Treat "
                "context.prior_structured_proposal as the exact structure under review. "
                "Apply analyst corrections to that structure, preserve unchanged rows/fields "
                "and schema coverage, and never replace a rich proposal with a narrower schema."
            )
            if _is_affirmative_confirmation(analyst_input):
                context["confirmed_proposal"] = prior_proposal
    if phase == "respond" and analyst_attachments:
        context["analyst_attachments"] = analyst_attachments
    if compact_state:
        context["runtime_state"] = compact_state
    if compact_history:
        context["recent_history"] = compact_history
    if package_context.get("resolved_resources"):
        context["resources"] = package_context["resolved_resources"]
    context_breakdown = {
        "system_chars": len(system_text),
        "element_chars": len(compiled_prompt) if execution_mechanism == "compiled_llm_plan" else len(json.dumps(model_record, ensure_ascii=False)),
        "state_chars": len(json.dumps(compact_state, ensure_ascii=False)),
        "history_chars": len(json.dumps(compact_history, ensure_ascii=False)),
        "resources_chars": sum(len(str(item.get("content") or "")) for item in package_context.get("resolved_resources", [])),
        "attachments_chars": sum(len(str(item.get("content") or "")) for item in analyst_attachments),
        "analyst_override_chars": len(analyst_override_context),
    }
    semantic_attempts: list[dict[str, Any]] = []
    if execution_mechanism == "runtime_semantic_plan":
        result = None
        last_error = None
        req_body = api_response = raw_text = usage = None
        semantic_contract = semantic_element.get("output_contract") if isinstance(semantic_element, dict) else {}
        semantic_contract_name = str((semantic_contract or {}).get("contract") or "")
        if semantic_contract_name == "GateFailureOrPass" and not semantic_call_meta.get("context_complete", True):
            result = _incomplete_gate_result(current_id, semantic_element, semantic_call_meta)
            semantic_attempts.append({"attempt":0,"skipped_model_call":True,"reason":"context_incomplete","context_truncated_objects":semantic_call_meta.get("context_truncated_objects") or []})
        for attempt in range(3) if result is None else range(0):
            repair_suffix = "" if not last_error else ("\n\nPrevious structured result was rejected by runtime validation. Correct it without changing the task. Error: " + last_error)
            recorded_result = payload.get("recorded_model_result") if attempt == 0 else None
            if isinstance(recorded_result, dict):
                replay_provenance = payload.get("recorded_model_provenance") if isinstance(payload.get("recorded_model_provenance"), dict) else {}
                req_body = {"replay": True, "source": "recorded_model_result", "provenance": copy.deepcopy(replay_provenance)}
                api_response = {"replay": True, "provenance": copy.deepcopy(replay_provenance)}
                raw_text = json.dumps(recorded_result, ensure_ascii=False)
                usage = {"input_tokens":0,"output_tokens":0,"total_tokens":0,"cached_tokens":0,"reasoning_tokens":0}
            else:
                req_body, api_response, raw_text, usage = _provider_api_call(credentials, system_text + repair_suffix, context)
            try:
                candidate = _normalize_semantic_model_result(_parse_model_json(raw_text))
                candidate, runtime_adaptations = _adapt_runtime_owned_node_envelope(
                    candidate, semantic_element=semantic_element, semantic_traits=semantic_traits, phase=phase
                )
                _oc=(semantic_element.get("output_contract") or {}) if isinstance(semantic_element,dict) else {}
                if str(_oc.get("contract") or "")=="GateFailureOrPass" and not _gate_declared_check_ids(semantic_element) and candidate.get("check_results") not in (None,[]):
                    candidate["check_results"]=[]
                    runtime_adaptations.append({"kind":"drop_undeclared_gate_check_results","reason":"zero_declared_checks"})
                allow = _alpha20_write_allowlist(record, semantic_element=semantic_element)
                patch = candidate.get("state_patch")
                legacy_updates = candidate.get("state_updates") if isinstance(candidate.get("state_updates"), dict) else None
                if not isinstance(patch, dict):
                    if legacy_updates is not None:
                        patch = legacy_updates_to_state_patch(legacy_updates, base_revision=current_revision)
                        candidate["state_patch"] = patch
                    elif allow and semantic_contract_name != "GateFailureOrPass":
                        raise ValueError("state_patch is required when this element can write state")
                    else:
                        patch = {"base_revision": current_revision, "operations": []}
                        candidate["state_patch"] = patch
                patch["base_revision"] = current_revision  # R2-D00: runtime-owned; ignore any model-supplied revision
                value_schemas = _alpha20_value_schemas(semantic_element)
                operation_variants = _alpha20_operation_variants(semantic_element)
                patch, metadata_adaptations = _adapt_misnested_patch_metadata(
                    patch, allowed_paths=allow, current_revision=current_revision,
                    value_schemas=value_schemas, operation_variants=operation_variants,
                )
                if metadata_adaptations:
                    candidate["state_patch"] = patch
                    runtime_adaptations.extend(metadata_adaptations)
                candidate, authority_adaptations, authority_errors = _apply_authority_contract_to_candidate(
                    candidate, semantic_element=semantic_element, state=state if isinstance(state, dict) else {}, phase=phase
                )
                if authority_adaptations:
                    runtime_adaptations.extend(authority_adaptations)
                patch = candidate.get("state_patch") if isinstance(candidate.get("state_patch"), dict) else patch
                validation = validate_state_patch(patch, allowed_paths=allow, current_revision=current_revision, value_schemas=value_schemas, operation_variants=operation_variants)
                result_errors = list(validation.get("errors") or []) + list(authority_errors or [])
                contract = semantic_element.get("output_contract") if isinstance(semantic_element, dict) else {}
                contract_name = str((contract or {}).get("contract") or "")
                if contract_name == "NodeExecutionResult":
                    for field in ("assistant_message", "state_patch", "needs_analyst", "next_intent"):
                        if field not in candidate:
                            result_errors.append(f"missing required NodeExecutionResult field: {field}")
                    if bool(candidate.get("needs_analyst")) and not bool(semantic_traits.get("requires_analyst")):
                        result_errors.append("needs_analyst=true is not allowed for an element without requires_analyst")
                allowed_route_keys = {str(r.get("key")) for r in (semantic_element.get("routes") or []) if isinstance(r, dict) and r.get("key") is not None}
                route_key = candidate.get("route_key")
                if route_key is not None and str(route_key) not in allowed_route_keys:
                    result_errors.append(f"route_key {route_key!r} is not allowed; allowed={sorted(allowed_route_keys)}")
                if contract_name == "GateFailureOrPass":
                    status = str(candidate.get("status") or "").lower()
                    if status not in {"passed", "pass", "failed", "fail"}:
                        result_errors.append("gate result requires status passed|failed")
                    if status in {"passed", "pass"} and not semantic_call_meta.get("context_complete", True):
                        result_errors.append("gate cannot PASS because semantic runtime context was truncated or incomplete")
                    check_errors, gate_accounting = _validate_gate_check_results(candidate, semantic_element)
                    result_errors.extend(check_errors)
                    if not check_errors:
                        candidate["failed_checks"] = _derive_failed_checks_from_results(candidate)
                        candidate["gate_accounting"] = gate_accounting
                # alpha.20.0.30: a model candidate must be not only schema-valid but also
                # atomically commit-able against the exact same runtime state/write contract.
                # Any commit rejection is a MODEL CONTRACT rejection and must stay inside the
                # bounded semantic retry loop; it must never escape later as a generic
                # Execution error. This specifically covers hallucinated paths such as
                # `riskfactorproposal` and operation/target mismatches such as replace-on-missing.
                _dry_state, dry_commit = apply_state_patch_atomic(
                    canonicalize_runtime_state(state),
                    patch,
                    allowed_paths=allow,
                    current_revision=current_revision,
                    value_schemas=_alpha20_value_schemas(semantic_element),
                    operation_variants=_alpha20_operation_variants(semantic_element),
                )
                if not dry_commit.get("committed"):
                    for err in (dry_commit.get("errors") or ["StatePatch candidate is not commit-able"]):
                        if str(err) not in result_errors:
                            result_errors.append(str(err))
                combined_validation = {
                    "valid": not result_errors,
                    "errors": result_errors,
                    "state_patch": validation,
                    "dry_run_commit": dry_commit,
                }
                semantic_attempt = {"attempt": attempt + 1, "validation": combined_validation, "usage": copy.deepcopy(usage or {})}
                if runtime_adaptations:
                    semantic_attempt["runtime_owned_envelope_adaptations"] = copy.deepcopy(runtime_adaptations)
                semantic_attempts.append(semantic_attempt)
                if not result_errors:
                    result = candidate
                    break
                last_error = "; ".join(result_errors)
            except ValueError as exc:
                last_error = str(exc)
                semantic_attempts.append({"attempt": attempt + 1, "parse_error": last_error, "usage": copy.deepcopy(usage or {})})
        if result is None:
            halted=_runtime_only_live_result(credentials=credentials,record=record,kind=kind,current_id=current_id,phase=phase,state=state if isinstance(state,dict) else {},routes=routes,assistant_message="Model output could not satisfy the required structured contract after 3 attempts.",await_analyst=False,selected=None,updates={},new_state=copy.deepcopy(state if isinstance(state,dict) else {}),reason="contract-unsatisfiable-by-model",extra_runtime={"failure_class":"contract_unsatisfiable_by_model","attempts":3,"last_validation_errors":[str(last_error or "unknown validation error")],"semantic_model_attempts":semantic_attempts})
            halted["run_status"]="halted"
            halted["completion_reason"]="contract_unsatisfiable_by_model"
            halted["failure_class"]="contract_unsatisfiable_by_model"
            halted["debug"]["runtime"]["run_status"]="halted"
            halted["debug"]["runtime"]["completion_reason"]="contract_unsatisfiable_by_model"
            return halted
    else:
        legacy_allow=_alpha20_write_allowlist(record,compiled_element,compiled_phase,semantic_element)
        legacy_route_keys=[str(r.get("key")) for r in routes if isinstance(r,dict) and r.get("key") is not None]
        try:
            result,req_body,api_response,raw_text,usage,legacy_model_attempts=_runtime_model_call_with_guard(
                credentials=credentials,system_text=system_text,context=context,
                current_revision=current_revision,state=state if isinstance(state,dict) else {},
                allowed_paths=legacy_allow,allowed_route_keys=legacy_route_keys,
                value_schemas=_alpha20_value_schemas(semantic_element),
                operation_variants=_alpha20_operation_variants(semantic_element),
                require_patch_when_writable=False,max_attempts=3,
            )
        except ValueError as exc:
            halted=_runtime_only_live_result(
                credentials=credentials,record=record,kind=kind,current_id=current_id,phase=phase,
                state=state if isinstance(state,dict) else {},routes=routes,
                assistant_message="Model output could not satisfy the runtime contract after 3 attempts.",
                await_analyst=False,selected=None,updates={},new_state=copy.deepcopy(state if isinstance(state,dict) else {}),
                reason="contract-unsatisfiable-by-model",
                extra_runtime={"failure_class":"contract_unsatisfiable_by_model","error":str(exc)},
            )
            halted["run_status"]="halted"
            halted["completion_reason"]="contract_unsatisfiable_by_model"
            halted["failure_class"]="contract_unsatisfiable_by_model"
            return halted
    structured_proposal_reconciliation = None
    if kind == "node" and phase == "respond" and isinstance(result.get("state_updates"), dict):
        reconciled_updates, structured_proposal_reconciliation = _reconcile_structured_proposal_updates(
            record, current_id, analyst_input, history, state if isinstance(state, dict) else {}, result.get("state_updates") or {}
        )
        result["state_updates"] = reconciled_updates
    source_reference_reconciliation = None
    if kind == "node" and phase == "respond" and isinstance(result.get("state_updates"), dict):
        canonical_updates, source_reference_reconciliation = _canonicalize_confirmed_source_references(
            state if isinstance(state, dict) else {}, result.get("state_updates") or {}
        )
        result["state_updates"] = canonical_updates
    await_analyst, selected, orchestration_override_reason = _resolve_respond_orchestration(record, kind, phase, routes, result)
    # Human-input declarations are runtime authority. On enter, a model may prepare
    # a proposal, but it may never answer the declared analyst question or advance.
    if human_policy.get("requires_human") and phase == "enter":
        await_analyst = True
        result["route_key"] = None
        selected = None
        orchestration_override_reason = None
    requested_key = result.get("route_key")
    if selected is None:
        selected = next((r for r in routes if r["key"] == requested_key), None)
    if selected is None and len(routes) == 1 and not await_analyst: selected = routes[0]
    if await_analyst and selected is not None: selected = None
    if selected is None and routes and requested_key is not None and not await_analyst:
        raise ValueError(f"Model proposed disallowed route {requested_key!r}; allowed routes: {route_keys}")
    if selected is None and len(routes) > 1 and not await_analyst:
        raise ValueError(f"Model did not select a route; allowed routes: {route_keys}")

    terminal_ids = _declared_terminal_ids(source)
    current_is_terminal = bool(record.get("terminal") is True or current_id in terminal_ids)
    executable_ids = {str(item.get("id")) for collection in (source.get("nodes", []), source.get("gates", [])) if isinstance(collection, list) for item in collection if isinstance(item, dict) and item.get("id")}
    run_status = "running"
    completion_reason = None
    if not await_analyst and selected is None and not routes:
        if current_is_terminal:
            run_status = "completed"
            completion_reason = "terminal"
        else:
            run_status = "halted"
            completion_reason = "dead_end"
    elif selected is not None and selected["target"] not in executable_ids and selected["target"] not in terminal_ids:
        run_status = "halted"
        completion_reason = "unresolved_target"
        selected = None

    updates = result.get("state_updates") if isinstance(result.get("state_updates"), dict) else {}
    # alpha.20.0.7: StatePatch is the actual commit path with per-path value schema enforcement.
    # Legacy model output is projected with explicit unknown provenance, validated
    # against the authoritative declared write surface, and committed atomically.
    alpha20_write_allowlist = _alpha20_write_allowlist(record, compiled_element, compiled_phase, semantic_element)
    alpha20_current_revision = current_revision
    alpha20_state_patch_raw = result.get("state_patch") if isinstance(result.get("state_patch"), dict) else None
    updates, alpha20_state_patch_raw, runtime_human_expression_updates = _apply_declared_human_runtime_expressions(
        record, state if isinstance(state, dict) else {}, phase, updates, alpha20_state_patch_raw
    )
    alpha20_state_patch = alpha20_state_patch_raw if isinstance(alpha20_state_patch_raw, dict) else legacy_updates_to_state_patch(updates, base_revision=alpha20_current_revision)
    alpha20_patch_validation = validate_state_patch(
        alpha20_state_patch,
        allowed_paths=alpha20_write_allowlist,
        current_revision=alpha20_current_revision,
        value_schemas=_alpha20_value_schemas(semantic_element),
        operation_variants=_alpha20_operation_variants(semantic_element),
    )
    new_state, alpha20_patch_commit = apply_state_patch_atomic(
        canonicalize_runtime_state(state),
        alpha20_state_patch,
        allowed_paths=alpha20_write_allowlist,
        current_revision=alpha20_current_revision,
        value_schemas=_alpha20_value_schemas(semantic_element),
        operation_variants=_alpha20_operation_variants(semantic_element),
    )
    if not alpha20_patch_commit.get("committed"):
        raise ValueError("alpha.20 StatePatch rejected: " + "; ".join(alpha20_patch_commit.get("errors") or ["unknown validation error"]))
    runtime_owned_gate_updates = {}
    semantic_contract_name_final = str(((semantic_element.get("output_contract") if isinstance(semantic_element, dict) else {}) or {}).get("contract") or "")
    if semantic_contract_name_final == "GateFailureOrPass" and selected is not None:
        runtime_owned_gate_updates = _runtime_owned_gate_branch_updates(record, selected, result)
        # Branch effects are compiled/source-declared exact state paths.  Keep the
        # editor service independent of alpha20_runtime private path helpers and
        # enforce the same declared write surface by exact canonical path here.
        allowset = {str(path).strip(".") for path in alpha20_write_allowlist}
        undeclared = [k for k in runtime_owned_gate_updates if str(k).strip(".") not in allowset]
        if undeclared:
            raise ValueError("runtime-owned gate branch update is outside write allowlist: " + ", ".join(sorted(undeclared)))
        if runtime_owned_gate_updates:
            new_state = _apply_state_updates_canonical(new_state, runtime_owned_gate_updates)
            updates = {**updates, **copy.deepcopy(runtime_owned_gate_updates)}
    alpha20_gate_failure = None
    if kind == "gate" and not await_analyst and selected is not None and selected.get("key") in {"on_fail", "fail", "failed", "revise"}:
        alpha20_gate_failure = normalize_gate_failure(
            current_id,
            failed_checks=result.get("failed_checks"),
            invalid_state=result.get("invalid_state"),
            missing_information=result.get("missing_information"),
            missing_coverage=result.get("missing_coverage"),
            affected_state=result.get("affected_state"),
            evidence=result.get("evidence"),
            suggested_recovery_scope=str(result.get("suggested_recovery_scope") or ("context" if result.get("failure_class") == "technical_stop" else "unknown")),
        )
        alpha20_gate_failure["failure_class"] = str(result.get("failure_class") or "business_stop")
    revision_after = alpha20_current_revision + (1 if canonicalize_runtime_state(new_state) != canonicalize_runtime_state(state) else 0)
    state_lineage = _state_lineage_entries(
        patch=alpha20_state_patch, new_state=new_state, revision=revision_after, producer_element_id=current_id,
        source_run_id=run_id, runtime_updates=runtime_owned_gate_updates,
    ) if revision_after != alpha20_current_revision else []

    # R3: declarative post-commit artifact rematerialization. A playbook may attach
    # a generic ``rematerialization`` contract to any model node. The runtime must
    # execute that contract against the POST-COMMIT state; otherwise lifecycle
    # changes can leave a previously materialized artifact stale even though the
    # semantic source explicitly requested rematerialization. No domain identifiers
    # or artifact names are interpreted here.
    post_commit_artifact = None
    post_commit_materialization = None
    rematerialization = record.get("rematerialization") if isinstance(record, dict) else None
    if kind == "node" and phase == "enter" and isinstance(rematerialization, dict):
        synthetic_generate = {
            "action": "DOCUMENT.GENERATE",
            "template": rematerialization.get("template"),
            "bindings": rematerialization.get("bindings"),
            "output": rematerialization.get("output"),
        }
        remat_result = _execute_document_generate(
            credentials, synthetic_generate, current_id, kind, phase,
            new_state if isinstance(new_state, dict) else {}, []
        )
        remat_runtime = ((remat_result.get("debug") or {}).get("runtime") or {}) if isinstance(remat_result, dict) else {}
        post_commit_artifact = copy.deepcopy(remat_runtime.get("artifact")) if isinstance(remat_runtime.get("artifact"), dict) else None
        post_commit_materialization = copy.deepcopy(remat_runtime.get("materialization")) if isinstance(remat_runtime.get("materialization"), dict) else None
        if post_commit_artifact is None:
            reason = str(remat_result.get("rationale_short") or "declared rematerialization failed") if isinstance(remat_result, dict) else "declared rematerialization failed"
            halted = _runtime_only_live_result(
                credentials=credentials, record=record, kind=kind, current_id=current_id, phase=phase,
                state=state if isinstance(state, dict) else {}, routes=routes,
                assistant_message="Declared artifact rematerialization failed.", await_analyst=False,
                selected=None, updates=updates, new_state=new_state,
                reason="declared-rematerialization-failed",
                extra_runtime={"failure_class":"artifact_error","rematerialization":copy.deepcopy(rematerialization),"detail":reason},
            )
            halted["state_revision"] = revision_after
            halted["run_status"] = "halted"
            halted["completion_reason"] = "declared_rematerialization_failed"
            halted["failure_class"] = "artifact_error"
            halted["debug"]["runtime"]["state_lineage"] = copy.deepcopy(state_lineage)
            halted["debug"]["runtime"]["revision_before"] = alpha20_current_revision
            halted["debug"]["runtime"]["revision_after"] = revision_after
            return halted
        post_commit_artifact["materialized_from_revision"] = revision_after
        post_commit_artifact["producer_node"] = current_id
        post_commit_artifact["rematerialized_post_commit"] = True

    runtime_referenced_artifacts = _runtime_artifacts_for_record(record)
    debug = {
        "provider": credentials["provider"], "base_url": credentials["base_url"], "api_style": credentials["api_style"],
        "model": credentials["model"], "current_id": current_id, "element_kind": kind, "phase": phase,
        "analyst_attachments": analyst_attachment_debug,
        "execution_mechanism": execution_mechanism,
        "semantic_plan_status": copy.deepcopy(_active_playbook_package().get("semantic_plan_status")) if execution_mechanism == "runtime_semantic_plan" else None,
        "semantic_model_attempts": semantic_attempts if execution_mechanism == "runtime_semantic_plan" else [],
        "compiled_plan_status": _active_playbook_package().get("compiled_plan_status"),
        "compiled_element_rejection_reasons": compiled_element_rejection_reasons,
        "usage": usage, "context_breakdown": context_breakdown,
        "input": {"system_text": system_text, "context": context, "request_payload": req_body},
        "output": {"raw_text": raw_text, "parsed_result": result, "api_response": api_response},
        "alpha20": {
            "contract_version": "alpha.20.0.34",
            "state_patch": alpha20_state_patch,
            "state_patch_validation": alpha20_patch_validation,
            "state_patch_commit": alpha20_patch_commit,
            "revision_before": alpha20_current_revision,
            "revision_after": revision_after,
            "state_lineage": copy.deepcopy(state_lineage),
            "write_allowlist": alpha20_write_allowlist,
            "revision_check": f"runtime_owned_revision_{alpha20_current_revision}",
            "gate_failure": alpha20_gate_failure,
            "runtime_owned_gate_state_updates": copy.deepcopy(runtime_owned_gate_updates),
        },
        "runtime": {"state_before": state, "state_updates": updates, "state_after": new_state, "revision_before": alpha20_current_revision, "revision_after": revision_after, "state_lineage": copy.deepcopy(state_lineage), "allowed_routes": routes, "requested_route_key": requested_key, "selected_route_key": selected["key"] if selected else None, "next_id": selected["target"] if selected else None, "await_analyst": await_analyst, "run_status": run_status, "completion_reason": completion_reason,
                    "orchestration_override_reason": orchestration_override_reason,
                    "structured_proposal_reconciliation": structured_proposal_reconciliation,
                    "source_reference_reconciliation": source_reference_reconciliation,
                    "gate_accounting": copy.deepcopy(result.get("gate_accounting")) if isinstance(result.get("gate_accounting"), dict) else None,
                    "context_status": copy.deepcopy(semantic_call_meta.get("context_status")) if execution_mechanism == "runtime_semantic_plan" else None,
                    "failure_class": str(result.get("failure_class") or ("business_stop" if alpha20_gate_failure else "")) or None,
                    "normalized_execution_result": {
                        "element_id": current_id, "phase": phase,
                        "status": "awaiting_analyst" if await_analyst else ("completed" if run_status == "running" else run_status),
                        "state_updates": copy.deepcopy(updates),
                        "route_key": selected["key"] if selected else None,
                        "next_id": selected["target"] if selected else None,
                    },
                    "artifact": post_commit_artifact or (runtime_referenced_artifacts[0] if runtime_referenced_artifacts else None),
                    "artifacts": ([post_commit_artifact] if post_commit_artifact else []) + runtime_referenced_artifacts,
                    "materialization": post_commit_materialization},
    }
    return {"assistant_message": str(result.get("assistant_message") or ""), "state_revision": revision_after, "rationale_short": str(result.get("rationale_short") or ""), "await_analyst": await_analyst, "phase": phase, "route_key": selected["key"] if selected else None, "next_id": selected["target"] if selected else None, "state": new_state, "routes": routes, "usage": usage, "context_breakdown": context_breakdown, "debug": debug, "provider": credentials["provider"], "model": credentials["model"], "terminal": current_is_terminal, "run_status": run_status, "completion_reason": completion_reason}


def _runtime_owned_gate_branch_updates(record: dict[str, Any], selected: dict[str, Any] | None, result: dict[str, Any]) -> dict[str, Any]:
    """Materialize declarative gate branch updates as runtime-owned effects.

    A model gate decides only the verdict/route.  `on_pass_update_state` and
    `on_fail_update_state` belong to the compiled/source graph contract and must not
    be delegated back to the model merely because the gate declares writes.
    """
    if not isinstance(selected, dict):
        return {}
    key=str(selected.get("key") or "")
    map_name = "on_pass_update_state" if key in {"on_pass","pass","passed"} else ("on_fail_update_state" if key in {"on_fail","fail","failed","revise"} else None)
    if not map_name:
        return {}
    spec=record.get(map_name)
    if not isinstance(spec, dict):
        return {}
    def resolve(value: Any) -> Any:
        if isinstance(value, str) and value.startswith("$gate."):
            cur: Any=result
            for part in value[len("$gate."):].split("."):
                if isinstance(cur, dict) and part in cur:
                    cur=cur[part]
                else:
                    return None
            return copy.deepcopy(cur)
        if isinstance(value, list):
            return [resolve(x) for x in value]
        if isinstance(value, dict):
            return {k:resolve(v) for k,v in value.items()}
        return copy.deepcopy(value)
    return {str(k):resolve(v) for k,v in spec.items()}


def _recovery_diagnosis(payload: dict[str, Any]) -> dict[str, Any]:
    """Explain a recorded validation failure without changing runtime state or choosing a route."""
    credentials = _live_credentials(payload)
    evidence = payload.get("evidence") if isinstance(payload.get("evidence"), dict) else {}
    choices = payload.get("choices") if isinstance(payload.get("choices"), list) else []
    allowed_targets = [str(x.get("target") or "") for x in choices if isinstance(x, dict) and x.get("target")]
    context = {
        "task": "recovery_diagnosis",
        "failed_gate": evidence.get("gate_id"),
        "failure_explanation": evidence.get("explanation"),
        "failure_rationale": evidence.get("rationale"),
        "failed_checks": evidence.get("failed_checks") or [],
        "affected_nodes": evidence.get("affected_nodes") or [],
        "available_recovery_targets": [
            {"target": str(x.get("target") or ""), "label": str(x.get("label") or x.get("target") or "")}
            for x in choices if isinstance(x, dict)
        ],
        "recent_gate_debug": evidence.get("debug") if isinstance(evidence.get("debug"), dict) else {},
        "interaction_contract": _interaction_contract(),
    }
    system_text = (
        "You are a diagnostic layer for an Ordo playbook executor. Explain why a recorded validation gate failed. "
        "Use ONLY supplied evidence. Never invent failed checks, state defects, or recovery targets. "
        "If evidence does not identify the concrete defect, return diagnosis_status='insufficient_evidence'. "
        "A recovery target may be recommended only when the evidence directly supports that target and it is in available_recovery_targets. "
        "Do not modify state and do not select a graph route. " + _analyst_language_instruction() + " "
        "Return ONLY JSON with keys: "
        "diagnosis_status ('identified' or 'insufficient_evidence'), summary, failed_checks (array), missing_evidence (array), "
        "likely_affected_state (array), recommended_recovery_target (string or null), recommendation_confidence ('high','medium','low','none'), analyst_explanation."
    )
    req_body, api_response, raw_text, usage = _provider_api_call(credentials, system_text, context)
    result = _parse_model_json(raw_text)
    status = str(result.get("diagnosis_status") or "insufficient_evidence")
    if status not in {"identified", "insufficient_evidence"}: status = "insufficient_evidence"
    target = result.get("recommended_recovery_target")
    target = str(target).strip() if target is not None else ""
    if target not in allowed_targets: target = ""
    confidence = str(result.get("recommendation_confidence") or "none").lower()
    if confidence not in {"high","medium","low","none"}: confidence = "none"
    if not target: confidence = "none"
    clean = {
        "diagnosis_status": status,
        "summary": str(result.get("summary") or "").strip(),
        "failed_checks": result.get("failed_checks") if isinstance(result.get("failed_checks"), list) else [],
        "missing_evidence": result.get("missing_evidence") if isinstance(result.get("missing_evidence"), list) else [],
        "likely_affected_state": result.get("likely_affected_state") if isinstance(result.get("likely_affected_state"), list) else [],
        "recommended_recovery_target": target or None,
        "recommendation_confidence": confidence,
        "analyst_explanation": str(result.get("analyst_explanation") or result.get("summary") or "").strip(),
    }
    return {"diagnosis": clean, "usage": usage, "debug": {"provider": credentials["provider"], "model": credentials["model"], "input": {"system_text": system_text, "context": context}, "output": {"raw_text": raw_text, "parsed_result": result}, "api_response": api_response}}


def _recovery_conversation(payload: dict[str, Any]) -> dict[str, Any]:
    """Free-form analyst conversation inside a validation-recovery node.

    The analyst remains on the recovery node. The model may explain the failure,
    propose a grounded repair target, request another gate evaluation, or submit
    a StatePatch limited to state roots explicitly named by the recorded failure.
    """
    credentials = _live_credentials(payload)
    evidence = payload.get("evidence") if isinstance(payload.get("evidence"), dict) else {}
    choices = payload.get("choices") if isinstance(payload.get("choices"), list) else []
    state = _canonicalize_runtime_state(payload.get("state") or {})
    current_revision = int(payload.get("state_revision") or 0)
    analyst_input = str(payload.get("analyst_input") or "").strip()
    history = payload.get("history") if isinstance(payload.get("history"), list) else []
    allowed_targets = [str(x.get("target") or "") for x in choices if isinstance(x, dict) and x.get("target")]
    affected = [str(x) for x in (evidence.get("affected_state") or []) if isinstance(x, str) and x.strip()]
    # Recovery chat is fail-closed for state writes: no affected-state evidence means no writes.
    allowed_paths = sorted(set(affected))
    context = {
        "task": "conversational_validation_recovery",
        "failed_gate": evidence.get("gate_id"),
        "failure_explanation": evidence.get("explanation"),
        "failed_checks": evidence.get("failed_checks") or [],
        "missing_information": evidence.get("missing_information") or [],
        "missing_coverage": evidence.get("missing_coverage") or [],
        "affected_state": affected,
        "available_recovery_targets": [
            {"target": str(x.get("target") or ""), "label": str(x.get("label") or x.get("target") or "")}
            for x in choices if isinstance(x, dict)
        ],
        "runtime_state": _bounded_json_value(state, 60000),
        "recent_recovery_dialog": _bounded_json_value(history[-12:], 18000),
        "analyst_input": analyst_input,
        "interaction_contract": _interaction_contract(),
    }
    system_text = (
        "You are the conversational recovery assistant inside an Ordo playbook. The analyst is intentionally staying on the current recovery node. "
        "Have a useful free-form diagnostic conversation while preserving graph authority. Explain what is known, what is not known, and what the analyst can provide. "
        "You MAY propose state corrections only when they are directly supported by the analyst's message and only under affected_state roots supplied in context. "
        "Never invent business facts. Never write outside affected_state. Never silently choose a graph route. "
        "If the analyst asks to re-run/recheck the failed gate, set suggested_action='retry_gate'. "
        "If the analyst explicitly asks to go to a repair node and evidence supports one of available_recovery_targets, set suggested_action='go_to_target' and recommended_recovery_target accordingly. "
        "Otherwise stay in recovery with suggested_action='stay'. "
        "StatePatch base_revision is runtime-owned; provide operations with op/path/value/basis/reason; basis must be analyst_input or recovery. "
        "StatePatch paths MUST use canonical dotted Ordo paths, never JSON Pointer slash paths. "
        "To correct an existing row in a collection, use merge_row on the collection path with row_key and row_match; do not replace by numeric array index. "
        "If the analyst requests a correction and then a recheck, use retry_gate only when the correction can be committed in this recovery turn. "
        "If a safe direct patch is not allowed, choose go_to_target with the most specific grounded repair target instead of retrying unchanged state. "
        + _analyst_language_instruction() + " "
        "Return ONLY JSON with keys: assistant_message, suggested_action ('stay','retry_gate','go_to_target'), recommended_recovery_target (string or null), "
        "state_patch ({base_revision,operations}), rationale_short."
    )
    try:
        result, req_body, api_response, raw_text, usage, recovery_model_attempts = _runtime_model_call_with_guard(
            credentials=credentials,system_text=system_text,context=context,
            current_revision=current_revision,state=state,
            allowed_paths=allowed_paths,allowed_route_keys=[],
            value_schemas={},operation_variants={},
            require_patch_when_writable=False,max_attempts=3,
        )
    except ValueError as exc:
        return {
            "assistant_message":"Не вдалося отримати структурно коректну відповідь для recovery після 3 спроб.",
            "state_revision":current_revision,
            "suggested_action":"stay",
            "recommended_recovery_target":None,
            "failed_gate":str(evidence.get("gate_id") or "") or None,
            "state":state,
            "state_patch":{"base_revision":current_revision,"operations":[]},
            "state_patch_validation":{"valid":False,"errors":[str(exc)]},
            "state_patch_commit":{"committed":False,"errors":[str(exc)]},
            "usage":{},
            "debug":{"failure_class":"contract_unsatisfiable_by_model","error":str(exc)},
        }
    action = str(result.get("suggested_action") or "stay").strip().lower()
    if action not in {"stay", "retry_gate", "go_to_target"}: action = "stay"
    target = str(result.get("recommended_recovery_target") or "").strip()
    if target not in allowed_targets:
        target = ""
        if action == "go_to_target": action = "stay"
    patch = result.get("state_patch") if isinstance(result.get("state_patch"), dict) else {"base_revision": current_revision, "operations": []}
    patch["base_revision"] = current_revision
    attempted_patch_operations = bool(isinstance(patch.get("operations"), list) and patch.get("operations"))
    validation = validate_state_patch(patch, allowed_paths=allowed_paths, current_revision=current_revision)
    if validation.get("valid"):
        new_state, commit = apply_state_patch_atomic(state, patch, allowed_paths=allowed_paths, current_revision=current_revision)
    else:
        new_state, commit = state, {"committed": False, **validation}
        patch = {"base_revision": current_revision, "operations": []}
    message = str(result.get("assistant_message") or "").strip()
    if not message:
        message = "Можу допомогти розібрати причину зупинки. Уточніть, що саме ви хочете перевірити або які дані готові надати."
    state_changed = canonicalize_runtime_state(new_state) != canonicalize_runtime_state(state)
    revision_after=current_revision + (1 if state_changed else 0)
    # KF-021 fail-closed guard: never re-run the same failed gate after an attempted
    # correction when no state change was actually committed. This prevents a false
    # conversational "fixed" message from becoming an immediate no-progress cycle.
    if action == "retry_gate" and attempted_patch_operations and not state_changed:
        action = "stay"
        message += "\n\nВиправлення не було застосовано до runtime state, тому повторна перевірка не запускається на незміненому стані."
    return {
        "assistant_message": message,
        "state_revision": revision_after,
        "suggested_action": action,
        "recommended_recovery_target": target or None,
        "failed_gate": str(evidence.get("gate_id") or "") or None,
        "state": new_state,
        "state_patch": patch,
        "state_patch_validation": validation,
        "state_patch_commit": commit,
        "usage": usage,
        "debug": {"provider": credentials["provider"], "model": credentials["model"], "input": {"system_text": system_text, "context": context}, "output": {"raw_text": raw_text, "parsed_result": result}, "api_response": api_response, "model_response_attempts": recovery_model_attempts},
    }


try:
    from .verification.runner import load_catalog as _verification_catalog, run_catalog as _run_verification_catalog
except ImportError:
    from verification.runner import load_catalog as _verification_catalog, run_catalog as _run_verification_catalog

VERIFICATION_RUNS: dict[str, dict[str, Any]] = {}
VERIFICATION_RUNS_LOCK = threading.Lock()

MODEL_CHAT_RUNS: dict[str, dict[str, Any]] = {}
MODEL_CHAT_RUNS_LOCK = threading.Lock()




_PLAYBOOK_SETTINGS_ROOTS = (
    "program_contract", "ordo", "interaction_model", "process_rail", "conversation_semantics",
    "hybrid_execution", "hybrid_execution_model", "execution_trace",
    "startup_package_profile", "runtime_capabilities", "conversation_scope_guard",
    "artifact_sync", "coverage_rules", "go_no_go", "prompt_registry", "flow_reuse",
)

_PLAYBOOK_SETTINGS_SCHEMA_MAP = {
    "program_contract": ("schemas/program_level_contract_schema.yaml", None),
    "interaction_model": ("schemas/interaction_model_schema.yaml", None),
    "process_rail": ("schemas/process_rail_schema.yaml", None),
    "conversation_semantics": ("schemas/conversation_semantics_schema.yaml", None),
    "startup_package_profile": ("schemas/startup_package_profile_schema.yaml", "startup_package_profile"),
    "artifact_sync": ("schemas/derived_artifact_sync_schema.yaml", "artifact_sync"),
}


def _settings_registry_dir() -> Path:
    editor_root = Path(__file__).resolve().parent
    bundled = editor_root / "verification" / "language" / "registry"
    if bundled.is_dir():
        return bundled
    return editor_root.parents[1] / "language" / "registry"


def _parse_markdown_value_registry(path: Path) -> dict[str, list[dict[str, str]]]:
    """Parse registry markdown sections with Value/Meaning style tables.

    The parser is deliberately generic: future language packages can add new
    headings/rows without requiring Editor UI changes.
    """
    if not path.exists():
        return {}
    lines=path.read_text(encoding="utf-8", errors="replace").splitlines()
    result: dict[str, list[dict[str, str]]] = {}
    section: str | None = None
    headers: list[str] | None = None
    for raw in lines:
        line=raw.strip()
        if line.startswith("## "):
            section=line[3:].strip().strip("`")
            headers=None
            continue
        if not section or not line.startswith("|"):
            continue
        cells=[c.strip().strip("`") for c in line.strip("|").split("|")]
        if not headers:
            lowered=[c.lower() for c in cells]
            if any(x in lowered for x in ("value","execution_mode","control_level")):
                headers=lowered
            continue
        if all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        if len(cells) < 2:
            continue
        value=cells[0]
        meaning=cells[1]
        guarantee=cells[2] if len(cells)>2 else ""
        if value and value.lower() not in {"value","kind","execution_mode","control_level"}:
            result.setdefault(section,[]).append({"value":value,"meaning":meaning,"guarantee":guarantee})
    return result


def _playbook_setting_value_catalog() -> dict[str, list[dict[str, str]]]:
    base=_settings_registry_dir()
    catalog: dict[str, list[dict[str,str]]] = {}
    for path in sorted(base.glob("*.md")):
        # Value registries are documentation-first and are the source for UI meanings.
        if not (path.name.endswith("_VALUES.md") or path.name in {"EXECUTION_MODES.md","CONTROL_LEVELS.md"}):
            continue
        for key, rows in _parse_markdown_value_registry(path).items():
            existing=catalog.setdefault(key,[])
            seen={r.get("value") for r in existing}
            for row in rows:
                if row.get("value") not in seen:
                    existing.append(row); seen.add(row.get("value"))
    # aliases from dedicated registries whose section heading is the document concept.
    exec_rows=[]
    control_rows=[]
    for path in [base/"EXECUTION_MODES.md", base/"CONTROL_LEVELS.md"]:
        if not path.exists(): continue
        lines=path.read_text(encoding="utf-8",errors="replace").splitlines()
        for line in lines:
            if not line.strip().startswith("|"): continue
            cells=[c.strip().strip("`") for c in line.strip().strip("|").split("|")]
            if len(cells)<2 or cells[0].lower() in {"execution_mode","control_level"} or set(cells[0]) <= {"-",":"}: continue
            row={"value":cells[0],"meaning":cells[1],"guarantee":cells[2] if len(cells)>2 else ""}
            (exec_rows if path.name=="EXECUTION_MODES.md" else control_rows).append(row)
    if exec_rows:
        by={r['value']:r for r in catalog.get('execution_mode',[])}
        for r in exec_rows: by.setdefault(r['value'],r)
        catalog['execution_mode']=list(by.values())
    if control_rows:
        by={r['value']:r for r in catalog.get('control_level',[])}
        for r in control_rows: by.setdefault(r['value'],r)
        catalog['control_level']=list(by.values())
    return catalog


def _flatten_setting_fields(value: Any, prefix: str = "") -> list[tuple[str, Any]]:
    out: list[tuple[str,Any]]=[]
    if isinstance(value, dict):
        for k,v in value.items():
            path=f"{prefix}.{k}" if prefix else str(k)
            if isinstance(v, dict): out.extend(_flatten_setting_fields(v,path))
            else: out.append((path,v))
    else:
        out.append((prefix,value))
    return out



def _settings_language_dir() -> Path:
    editor_root = Path(__file__).resolve().parent
    bundled = editor_root / "verification" / "language"
    if bundled.is_dir():
        return bundled
    return editor_root.parents[1] / "language"


def _setting_path_get(source: dict[str, Any], path: str) -> tuple[bool, Any]:
    cur: Any = source
    for part in str(path).split("."):
        if not isinstance(cur, dict) or part not in cur:
            return False, None
        cur = cur[part]
    return True, copy.deepcopy(cur)


def _schema_setting_rows(root: str, schema: dict[str, Any], prefix: str = "") -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    props = schema.get("properties") if isinstance(schema, dict) else None
    if not isinstance(props, dict):
        return rows
    for key, spec in props.items():
        if not isinstance(spec, dict):
            spec = {}
        local = f"{prefix}.{key}" if prefix else key
        typ = spec.get("type")
        if typ == "object" and isinstance(spec.get("properties"), dict):
            rows.extend(_schema_setting_rows(root, spec, local))
            continue
        rows.append({
            "path": f"{root}.{local}",
            "leaf": key,
            "label": key.replace("_", " ").title(),
            "description": str(spec.get("description") or "").strip(),
            "schema_enum": copy.deepcopy(spec.get("enum")) if isinstance(spec.get("enum"), list) else [],
            "schema_type": copy.deepcopy(typ),
        })
    return rows


_SETTING_REGISTRY_ALIASES = {
    "default_startup_mode": "Startup modes",
    "mode": "Startup modes",
    "audience": "Audiences",
    "role": "Entry file roles",
    "required": "Requiredness values",
    "visibility": "Visibility values",
    "profile": "Profiles",
    "severity": "Severity values",
    "decision": "Approval decisions",
    "status": "Status values",
    "capture_level": "Capture levels",
    "replay_mode": "Replay modes",
    "actor": "Actor types",
    "source": "Trace sources",
    "event_type": "Event types",
    "artifact_role": "Artifact roles",
    "derivation_method": "Derivation methods",
    "freshness_policy": "Freshness policies",
    "stale_action": "Stale actions",
}


def _all_registry_sections() -> dict[str, list[dict[str, str]]]:
    base=_settings_registry_dir()
    sections: dict[str, list[dict[str,str]]] = {}
    for path in sorted(base.glob("*.md")):
        parsed=_parse_markdown_value_registry(path)
        for key, rows in parsed.items():
            bucket=sections.setdefault(key,[])
            seen={str(r.get("value")) for r in bucket}
            for row in rows:
                if str(row.get("value")) not in seen:
                    bucket.append(row); seen.add(str(row.get("value")))
    return sections


def _options_for_setting(leaf: str, schema_enum: list[Any], catalog: dict[str, list[dict[str,str]]], sections: dict[str,list[dict[str,str]]]) -> list[dict[str,str]]:
    candidates=[]
    for key in [leaf, _SETTING_REGISTRY_ALIASES.get(leaf)]:
        if key and key in catalog:
            candidates.extend(copy.deepcopy(catalog[key]))
        elif key and key in sections:
            candidates.extend(copy.deepcopy(sections[key]))
    by_value: dict[str,dict[str,str]]={}
    for row in candidates:
        value=str(row.get("value") or "")
        if value:
            by_value.setdefault(value,row)
    if schema_enum:
        ordered=[]
        for value in schema_enum:
            sval=str(value)
            ordered.append(copy.deepcopy(by_value.get(sval,{"value":sval,"meaning":"Allowed by the Ordo language schema.","guarantee":""})))
        return ordered
    return list(by_value.values())


def _language_defined_settings_catalog(source: dict[str,Any]) -> list[dict[str,Any]]:
    lang=_settings_language_dir()
    value_catalog=_playbook_setting_value_catalog()
    sections=_all_registry_sections()
    groups: list[dict[str,Any]]=[]
    schema_paths_seen:set[str]=set()
    for root,(rel,nested) in _PLAYBOOK_SETTINGS_SCHEMA_MAP.items():
        p=lang/rel
        if not p.exists():
            continue
        try:
            schema=parse_yaml(p.read_text(encoding="utf-8"))
        except Exception:
            continue
        if nested and isinstance(schema,dict):
            schema=(schema.get("properties") or {}).get(nested) or {}
        rows=_schema_setting_rows(root,schema)
        fields=[]
        for row in rows:
            path=row["path"]; schema_paths_seen.add(path)
            specified,current=_setting_path_get(source,path)
            options=_options_for_setting(row["leaf"],row.get("schema_enum") or [],value_catalog,sections)
            fields.append({
                "path":path,
                "label":row["label"],
                "description":row.get("description") or "",
                "specified":specified,
                "current_value":current if specified else None,
                "options":options,
                "registry_backed":bool(options),
                "language_defined":True,
            })
        if fields:
            groups.append({"id":root,"title":root.replace("_"," ").title(),"fields":fields})
    # Preserve package-defined settings not represented by the standard schemas.
    extra_by_root={}
    for root in _PLAYBOOK_SETTINGS_ROOTS:
        if root not in source:
            continue
        for path,value in _flatten_setting_fields(source.get(root),root):
            if path in schema_paths_seen:
                continue
            leaf=path.split(".")[-1]
            options=_options_for_setting(leaf,[],value_catalog,sections)
            extra_by_root.setdefault(root,[]).append({
                "path":path,"label":leaf.replace("_"," ").title(),"description":"",
                "specified":True,"current_value":copy.deepcopy(value),"options":options,
                "registry_backed":bool(options),"language_defined":False,
            })
    for root,fields in extra_by_root.items():
        existing=next((g for g in groups if g["id"]==root),None)
        if existing: existing["fields"].extend(fields)
        else: groups.append({"id":root,"title":root.replace("_"," ").title(),"fields":fields})
    return groups


_LINEAGE_STATE_REF_RE = re.compile(r"(?:\$?state\.)([A-Za-z_][A-Za-z0-9_.-]*)")
_LINEAGE_TEMPLATE_REF_RE = re.compile(r"{{\s*([A-Za-z_][A-Za-z0-9_.-]*\.[A-Za-z0-9_.-]+)\s*}}")
_LINEAGE_ARTIFACT_PATH_RE = re.compile(
    r"(?<![A-Za-z0-9_.-])([A-Za-z0-9_.-]+(?:/[A-Za-z0-9_.-]+)+\.(?:md|markdown|json|ya?ml|txt|html?|pdf|docx?|xlsx?|csv|zip|tar|gz|tgz|xml))(?![A-Za-z0-9_.-])",
    re.IGNORECASE,
)

def _lineage_normalize_state_path(value: Any) -> str:
    text=str(value or "").strip().strip("`\"'")
    text=re.sub(r"^\$?state\.","",text)
    return text.strip(".")

def _lineage_state_refs(value: Any) -> list[str]:
    """Discover explicit state references and bare dotted Mustache placeholders."""
    found=[]
    def add(raw: Any) -> None:
        path=_lineage_normalize_state_path(raw)
        if path and path not in found: found.append(path)
    def walk(v: Any) -> None:
        if isinstance(v,dict):
            for child in v.values(): walk(child)
        elif isinstance(v,list):
            for child in v: walk(child)
        elif isinstance(v,str):
            for match in _LINEAGE_STATE_REF_RE.finditer(v): add(match.group(1))
            for match in _LINEAGE_TEMPLATE_REF_RE.finditer(v): add(match.group(1))
    walk(value)
    return found

def _lineage_artifact_refs(value: Any) -> list[str]:
    found=[]
    def walk(v: Any) -> None:
        if isinstance(v,dict):
            for child in v.values(): walk(child)
        elif isinstance(v,list):
            for child in v: walk(child)
        elif isinstance(v,str):
            for match in _LINEAGE_ARTIFACT_PATH_RE.finditer(v):
                path=match.group(1).strip()
                if path and path not in found: found.append(path)
    walk(value)
    return found

def _lineage_record_actor(record: dict[str,Any]) -> str:
    joined=" ".join(str(record.get(k) or "") for k in ("runtime_executor","executor","kind","action","method","interaction_kind")).lower()
    if any(x in joined for x in ("python","deterministic","mechanical","package_tool","tool_executor","script")): return "deterministic"
    if any(x in joined for x in ("model","llm","ai.","ai_","generate","respond")): return "model"
    if any(x in joined for x in ("human","analyst","interaction")): return "analyst"
    if record.get("question") is not None and isinstance(record.get("on_answer"),dict): return "analyst"
    return "internal"

def _lineage_state_writes(record: dict[str,Any]) -> list[dict[str,Any]]:
    rows=[]
    def add_mapping(mapping: Any, role: str) -> None:
        if not isinstance(mapping,dict): return
        for target,expr in mapping.items():
            path=_lineage_normalize_state_path(target)
            if path:
                rows.append({"path":path,"expression":copy.deepcopy(expr),"role":role})
    add_mapping(record.get("update_state"),"update_state")
    on_answer=record.get("on_answer") if isinstance(record.get("on_answer"),dict) else {}
    add_mapping(on_answer.get("update_state"),"on_answer.update_state")
    add_mapping(record.get("derive_before_generate"),"derive_before_generate")
    patch=record.get("state_patch") if isinstance(record.get("state_patch"),dict) else {}
    add_mapping(patch.get("updates"),"state_patch.updates")
    return rows

def _lineage_output_paths(record: dict[str,Any]) -> list[str]:
    out=[]
    def add(v: Any) -> None:
        if not isinstance(v,str): return
        text=v.strip().strip("`\"'")
        if _LINEAGE_ARTIFACT_PATH_RE.fullmatch(text) and text not in out: out.append(text)
    add(record.get("output")); add(record.get("output_path")); add(record.get("artifact_path")); add(record.get("report_path"))
    artifact=record.get("artifact") if isinstance(record.get("artifact"),dict) else {}
    for key in ("path","output","file","filename"): add(artifact.get(key))
    remat=record.get("rematerialization") if isinstance(record.get("rematerialization"),dict) else {}
    for key in ("output","output_path","artifact_path"): add(remat.get(key))
    package_block=record.get("package") if isinstance(record.get("package"),dict) else {}
    for key in ("path","output","output_path","artifact_path"): add(package_block.get(key))
    return out

def _lineage_artifact_kind(path: str) -> str:
    suffix=Path(path).suffix.lower()
    if suffix in {".zip",".tar",".gz",".tgz"}: return "archive"
    if suffix in {".md",".markdown",".html",".htm",".pdf",".doc",".docx",".json",".yaml",".yml",".txt",".csv",".xml",".xls",".xlsx"}: return "document"
    return "artifact"

def _lineage_runtime_value(runtime_state: dict[str,Any], path: str) -> tuple[bool,Any]:
    cur: Any=runtime_state
    for part in [p for p in path.split(".") if p]:
        if isinstance(cur,dict) and part in cur: cur=cur[part]
        else: return False,None
    return True,copy.deepcopy(cur)

def _lineage_declared_input_refs(record: dict[str,Any]) -> tuple[list[str],list[str]]:
    """Return declared state-like and artifact inputs that are not always written with state. prefixes."""
    states=[]; artifacts=[]
    def add_state(v: Any) -> None:
        if not isinstance(v,str): return
        text=v.strip().strip("`\"'")
        if not text or text.startswith("$") or " " in text or "/" in text: return
        if re.fullmatch(r"[A-Za-z_][A-Za-z0-9_.-]*",text):
            path=_lineage_normalize_state_path(text)
            if path and path not in states: states.append(path)
    def add_artifact(v: Any) -> None:
        if not isinstance(v,str): return
        text=v.strip().strip("`\"'")
        if _LINEAGE_ARTIFACT_PATH_RE.fullmatch(text) and text not in artifacts: artifacts.append(text)
    def walk_declared(v: Any) -> None:
        if isinstance(v,list):
            for x in v:
                if isinstance(x,str):
                    add_artifact(x); add_state(x)
                elif isinstance(x,(dict,list)): walk_declared(x)
        elif isinstance(v,dict):
            for x in v.values(): walk_declared(x)
        elif isinstance(v,str):
            add_artifact(v); add_state(v)
    for key in ("inputs","reads","input","source"):
        if key in record: walk_declared(record.get(key))
    for key in ("proposal_generation","draft_generation","interpretation_policy","derivation_contract","generation_contract"):
        block=record.get(key)
        if isinstance(block,dict):
            for subkey in ("source","sources","inputs","input"):
                if subkey in block: walk_declared(block.get(subkey))
    # Explicit state references remain authoritative and may appear anywhere in these blocks.
    for ref in _lineage_state_refs({k:record.get(k) for k in ("inputs","reads","proposal_generation","draft_generation","interpretation_policy","derivation_contract","generation_contract","preconditions") if k in record}):
        if ref not in states: states.append(ref)
    for art in _lineage_artifact_refs({k:record.get(k) for k in ("inputs","reads","package","preconditions") if k in record}):
        if art not in artifacts: artifacts.append(art)
    return states,artifacts


def _lineage_transform_profile(record: dict[str,Any], *, has_output: bool=False) -> dict[str,str]:
    action=str(record.get("action") or "")
    joined=" ".join(str(record.get(k) or "") for k in ("runtime_executor","executor","kind","action","method","interaction_kind")).lower()
    if isinstance(record.get("package"),dict) or action.upper().startswith("PACKAGE."):
        return {"kind":"transform_package","label":"Package / archive build","mechanism":action or "package"}
    if has_output and (record.get("template") or isinstance(record.get("rematerialization"),dict) or action.upper().startswith("DOCUMENT.")):
        return {"kind":"transform_template","label":"Document materialization","mechanism":action or "template rendering"}
    if any(x in joined for x in ("python","package_tool","tool_executor","script")) or record.get("validator") or isinstance(record.get("validation"),dict):
        return {"kind":"transform_tool","label":"Python / tool transformation","mechanism":action or str(record.get("runtime_executor") or "tool")}
    if record.get("proposal_generation") or record.get("draft_generation") or record.get("interpretation_policy") or action.upper().startswith("AI.") or any(x in joined for x in ("model","llm","ai.")):
        return {"kind":"transform_model","label":"AI transformation","mechanism":action or "model-assisted derivation"}
    actor=_lineage_record_actor(record)
    if actor=="analyst": return {"kind":"transform_analyst","label":"Analyst input / confirmation","mechanism":"analyst interaction"}
    # Literal state assignments with no model/tool markers are deterministic/default-like.
    return {"kind":"transform_deterministic","label":"Deterministic derivation","mechanism":action or "state derivation"}


def _lineage_registry_entries(resources: dict[str,Any]) -> list[dict[str,Any]]:
    """Read variable-registry style package resources without making the Editor depend on a domain-specific filename."""
    out=[]
    for resource_path,value in resources.items():
        if not isinstance(resource_path,str) or not re.search(r"\.ya?ml$",resource_path,re.I): continue
        text=None
        if isinstance(value,str): text=value
        elif isinstance(value,dict):
            for key in ("content_text","text","content"):
                if isinstance(value.get(key),str): text=value.get(key); break
        if not text or "variables:" not in text: continue
        try:
            parsed=yaml.safe_load(text)
        except Exception:
            continue
        variables=parsed.get("variables") if isinstance(parsed,dict) else None
        if not isinstance(variables,list): continue
        purpose=str(parsed.get("purpose") or "") if isinstance(parsed,dict) else ""
        # Accept registries that look like traceability metadata, regardless of filename.
        if not purpose and not any(isinstance(x,dict) and x.get("source_node") for x in variables): continue
        for item in variables:
            if not isinstance(item,dict) or not item.get("path"): continue
            row=copy.deepcopy(item); row["resource_path"]=resource_path; row["registry_purpose"]=purpose; out.append(row)
    return out


def _embedded_flow_resolve_resource(resources: dict[str,Any], bundle_path: str, ref: Any) -> tuple[str | None, str | None]:
    """Resolve one authoring-model resource relative to its bundle manifest.

    This is intentionally an Editor adapter contract, not Ordo syntax.  The
    manifest may live anywhere in the uploaded ZIP; canonical source references
    are resolved relative to that manifest first and then by a unique suffix.
    """
    if not isinstance(ref,str) or not ref.strip():
        return None,None
    raw=ref.strip().replace("\\","/").lstrip("/")
    rel=Path(bundle_path).parent.joinpath(raw).as_posix()
    candidates=[]
    for key,value in resources.items():
        if not isinstance(value,str): continue
        normalized=str(key).replace("\\","/")
        if normalized==rel or normalized==raw:
            return normalized,value
        if normalized.endswith("/"+rel) or normalized.endswith("/"+raw):
            candidates.append((normalized,value))
    if len(candidates)==1:
        return candidates[0]
    return None,None


def _embedded_flow_yaml(text: Any) -> dict[str,Any] | None:
    if not isinstance(text,str): return None
    try: value=yaml.safe_load(text)
    except yaml.YAMLError: return None
    return value if isinstance(value,dict) else None



def _canonical_data_layer_flow(package: dict[str,Any] | None) -> dict[str,Any] | None:
    """Adapt canonical Data Layer + Editor Projection into the read-only Data Flow UI graph.

    This is an Editor projection adapter only.  It does not extend canonical Ordo
    syntax and does not participate in runtime/compiler semantics.  The contract is
    intentionally generic: ``ordo.design.editor_projection.v1`` points at a
    ``ordo.authoring.canonical_data_layer.v1`` resource and provides node read/write
    projections.  The canonical Data Layer supplies state defaults, node metadata,
    gates, and declared outputs/artifacts.
    """
    resources=package.get("resources") if isinstance(package,dict) and isinstance(package.get("resources"),dict) else {}
    projections=[]
    for path,text in sorted(resources.items(),key=lambda item:str(item[0])):
        if Path(str(path)).suffix.lower() not in {".yaml",".yml"}: continue
        parsed=_embedded_flow_yaml(text)
        if not isinstance(parsed,dict): continue
        if str(parsed.get("schema") or "").strip()=="ordo.design.editor_projection.v1" and isinstance(parsed.get("data_layer"),str) and parsed.get("data_layer").strip():
            projections.append((str(path),parsed))
    if not projections:
        return None

    invalid=[]
    for projection_path,projection in projections:
        data_path,data_text=_embedded_flow_resolve_resource(resources,projection_path,projection.get("data_layer"))
        if not data_path or not data_text:
            invalid.append(f"{projection_path}: data_layer resource could not be resolved")
            continue
        data_layer=_embedded_flow_yaml(data_text)
        if not isinstance(data_layer,dict):
            invalid.append(f"{projection_path}: data_layer is not a YAML mapping")
            continue
        if str(data_layer.get("schema") or "").strip()!="ordo.authoring.canonical_data_layer.v1":
            invalid.append(f"{projection_path}: unsupported data_layer schema {data_layer.get('schema')!r}")
            continue
        program=data_layer.get("program_model") if isinstance(data_layer.get("program_model"),dict) else {}
        state=program.get("state") if isinstance(program.get("state"),dict) else {}
        state_schema=state.get("schema") if isinstance(state.get("schema"),dict) else {}
        # Canonical v1 data-class annotations live beside program_model so the
        # executable state schema remains plain Ordo state.  Older .181 fixtures
        # used program_model.state.variable_metadata; keep that as compatibility
        # fallback.  editor_projection.state_path_data_classes is a derived UI
        # projection and is used only when the canonical annotation is absent.
        state_path_annotations=data_layer.get("state_path_annotations") if isinstance(data_layer.get("state_path_annotations"),dict) else {}
        projected_data_classes=projection.get("state_path_data_classes") if isinstance(projection.get("state_path_data_classes"),dict) else {}
        legacy_variable_metadata=state.get("variable_metadata") if isinstance(state.get("variable_metadata"),dict) else {}
        allowed_data_classes={"business","technical","control","metadata"}
        canonical_nodes={str(x.get("id")):x for x in (program.get("nodes") or []) if isinstance(x,dict) and x.get("id")}
        canonical_gates={str(x.get("id")):x for x in (program.get("gates") or []) if isinstance(x,dict) and x.get("id")}
        canonical_outputs={str(x.get("id")):x for x in (program.get("outputs") or []) if isinstance(x,dict) and x.get("id")}
        projection_nodes=[x for x in (projection.get("nodes") or []) if isinstance(x,dict) and x.get("id")]
        projection_gates=[x for x in (projection.get("gates") or []) if isinstance(x,dict) and x.get("id")]
        projection_outputs=[str(x) for x in (projection.get("outputs") or []) if str(x).strip()]
        if not projection_nodes and not canonical_nodes:
            invalid.append(f"{projection_path}: no nodes are declared in projection or canonical Data Layer")
            continue

        graph_nodes=[]
        graph_edges=[]
        node_ids=set()
        variable_ids=set()

        def add_node(row: dict[str,Any]) -> None:
            identifier=str(row.get("id") or "")
            if identifier and identifier not in node_ids:
                graph_nodes.append(row); node_ids.add(identifier)

        def variable_node(path: str) -> str:
            path=str(path or "").strip()
            identifier=f"VAR::{path}"
            if identifier in variable_ids: return identifier
            default=copy.deepcopy(state_schema.get(path)) if path in state_schema else None
            inferred_type=("null" if default is None else type(default).__name__)
            canonical_meta=copy.deepcopy(state_path_annotations.get(path)) if isinstance(state_path_annotations.get(path),dict) else {}
            legacy_meta=copy.deepcopy(legacy_variable_metadata.get(path)) if isinstance(legacy_variable_metadata.get(path),dict) else {}
            projected_class=str(projected_data_classes.get(path) or "").strip().lower()
            canonical_class=str(canonical_meta.get("data_class") or "").strip().lower()
            legacy_class=str(legacy_meta.get("data_class") or "").strip().lower()
            declared_class=canonical_class or projected_class or legacy_class
            data_class=declared_class if declared_class in allowed_data_classes else "unclassified"
            provenance=(
                "canonical_data_layer.state_path_annotations" if canonical_class else
                "editor_projection.state_path_data_classes" if projected_class else
                "canonical_data_layer.program_model.state.variable_metadata" if legacy_class else
                "unclassified"
            )
            meta={"id":path,"path":path,"type":inferred_type,"default":default,"provenance":"canonical_data_layer.program_model.state.schema","data_class":data_class,"data_class_provenance":provenance}
            if declared_class and declared_class not in allowed_data_classes:
                meta["declared_data_class"]=declared_class
            for source_meta in (legacy_meta,canonical_meta):
                for key,value in source_meta.items():
                    if key not in {"data_class"}: meta[key]=value
            add_node({"id":identifier,"label":path,"type":"variable","variable_ref":path,"variable_metadata":meta,"section":"state"})
            variable_ids.add(identifier)
            return identifier

        # Operation/interaction nodes are retained in the graph because reads/writes
        # are relationships between canonical state variables and executable elements.
        projection_by_id={str(x.get("id")):x for x in projection_nodes}
        ordered_ids=[str(x.get("id")) for x in projection_nodes]
        for cid in canonical_nodes:
            if cid not in projection_by_id: ordered_ids.append(cid)
        for element_id in ordered_ids:
            projected=projection_by_id.get(element_id,{})
            canonical=canonical_nodes.get(element_id,{})
            contract=canonical.get("execution_contract") if isinstance(canonical.get("execution_contract"),dict) else {}
            owner=str(contract.get("owner") or "").lower()
            node_type="interaction" if owner=="human" or bool((canonical.get("node_context") or {}).get("user_question_allowed")) else "transformation"
            label=str(canonical.get("title") or canonical.get("purpose") or element_id)
            add_node({"id":element_id,"label":label,"type":node_type,"section":"operations","operation_metadata":{"id":element_id,"execution_owner":owner or None,"action":canonical.get("action"),"method":projected.get("method")}})
            reads=projected.get("reads") if isinstance(projected.get("reads"),list) else canonical.get("reads") if isinstance(canonical.get("reads"),list) else []
            writes=projected.get("writes") if isinstance(projected.get("writes"),list) else canonical.get("writes") if isinstance(canonical.get("writes"),list) else []
            for raw in reads:
                path=str(raw or "").strip()
                if path: graph_edges.append({"from":variable_node(path),"to":element_id,"type":"read","label":"read"})
            for raw in writes:
                path=str(raw or "").strip()
                if path: graph_edges.append({"from":element_id,"to":variable_node(path),"type":"write","label":"write"})

        # Gates remain explicit validation objects.  Their control links are shown as
        # validation relations, not runtime execution semantics reconstructed by UI.
        merged_gates={**canonical_gates,**{str(x.get('id')):{**canonical_gates.get(str(x.get('id')),{}) ,**x} for x in projection_gates}}
        for gate_id,gate in merged_gates.items():
            add_node({"id":gate_id,"label":str(gate.get("title") or gate_id),"type":"gate","section":"gates","gate_metadata":copy.deepcopy(gate)})
            for src in gate.get("allowed_from") or []:
                if str(src) in node_ids: graph_edges.append({"from":str(src),"to":gate_id,"type":"validation"})
            for key in ("on_pass","on_fail"):
                dst=str(gate.get(key) or "")
                if dst and dst in node_ids: graph_edges.append({"from":gate_id,"to":dst,"type":"validation","label":key})

        output_ids=projection_outputs or list(canonical_outputs)
        for output_id in output_ids:
            meta=copy.deepcopy(canonical_outputs.get(output_id) or {"id":output_id})
            artifact_id=f"ART::{output_id}"
            add_node({"id":artifact_id,"label":str(meta.get("label") or output_id),"type":"artifact","artifact_ref":output_id,"artifact_metadata":meta,"section":"artifacts"})
            producer=str(meta.get("producer") or "")
            if producer and producer in node_ids: graph_edges.append({"from":producer,"to":artifact_id,"type":"artifact_input"})

        # De-duplicate exact relations while preserving deterministic order.
        dedup=[]; seen=set()
        for edge in graph_edges:
            key=(str(edge.get("from")),str(edge.get("to")),str(edge.get("type")),str(edge.get("label") or ""))
            if key in seen: continue
            seen.add(key); dedup.append(edge)
        sections=[
            {"id":"state","label":"Canonical state variables","order":1},
            {"id":"operations","label":"Consumers / producers","order":2},
            {"id":"gates","label":"Validation gates","order":3},
            {"id":"artifacts","label":"Declared artifacts","order":4},
        ]
        type_counts={}
        for n in graph_nodes: type_counts[str(n.get("type") or "unknown")]=type_counts.get(str(n.get("type") or "unknown"),0)+1
        edge_counts={}
        for e in dedup: edge_counts[str(e.get("type") or "relation")]=edge_counts.get(str(e.get("type") or "relation"),0)+1
        data_class_counts={}
        for n in graph_nodes:
            if n.get("type")!="variable": continue
            cls=str((n.get("variable_metadata") or {}).get("data_class") or "unclassified")
            data_class_counts[cls]=data_class_counts.get(cls,0)+1
        return {
            "status":"passed","available":True,
            "bundle":{"path":projection_path,"schema_version":projection.get("schema"),"data_layer":data_path,"data_layer_schema":data_layer.get("schema"),"adapter":"canonical_data_layer_projection_v1"},
            "graph":{"model_id":str(projection.get("playbook") or data_layer.get("playbook") or "canonical_data_layer"),"revision":str(projection.get("version") or data_layer.get("version") or ""),"nodes":graph_nodes,"edges":dedup,"sections":sections},
            "catalogs":{"data_layer":{"path":data_path,"available":True},"editor_projection":{"path":projection_path,"available":True}},
            "summary":{"nodes":len(graph_nodes),"edges":len(dedup),"sections":len(sections),"gates":len(merged_gates),"node_types":type_counts,"edge_types":edge_counts,"data_classes":data_class_counts,"dangling_edges":0},
            "note":"Canonical Data Layer visualized through the package Editor Projection. This read-only adapter does not participate in Ordo execution semantics."
        }
    return {"status":"invalid","available":False,"error":"Canonical Data Layer projection candidate found, but it is invalid: "+"; ".join(invalid[:4]),"candidates":[p for p,_ in projections]}


def _discover_embedded_authoring_data_flow(package: dict[str,Any] | None) -> dict[str,Any]:
    """Discover an optional authoring data-flow bundle shipped beside an Ordo playbook.

    Discovery is contract-based: a YAML resource declares ``canonical_sources.graph``
    and the referenced graph contains ordinary model ``nodes`` and ``edges`` lists.
    No domain filename, node id, variable id, or Ordo-language extension is assumed.
    """
    resources=package.get("resources") if isinstance(package,dict) and isinstance(package.get("resources"),dict) else {}
    canonical_flow=_canonical_data_layer_flow(package)
    if canonical_flow is not None:
        return canonical_flow
    candidates=[]
    for path,text in sorted(resources.items(),key=lambda item:str(item[0])):
        if Path(str(path)).suffix.lower() not in {".yaml",".yml"}: continue
        parsed=_embedded_flow_yaml(text)
        canonical=parsed.get("canonical_sources") if isinstance(parsed,dict) else None
        if isinstance(canonical,dict) and isinstance(canonical.get("graph"),str) and canonical.get("graph").strip():
            candidates.append((str(path),parsed))
    if not candidates:
        return {"status":"not_present","available":False,"note":"No embedded authoring data-flow bundle was discovered."}

    invalid=[]
    for bundle_path,bundle in candidates:
        canonical=bundle.get("canonical_sources") or {}
        graph_path,graph_text=_embedded_flow_resolve_resource(resources,bundle_path,canonical.get("graph"))
        if not graph_path or not graph_text:
            invalid.append(f"{bundle_path}: canonical graph resource could not be resolved")
            continue
        graph=_embedded_flow_yaml(graph_text)
        if not isinstance(graph,dict) or not isinstance(graph.get("nodes"),list) or not isinstance(graph.get("edges"),list):
            invalid.append(f"{bundle_path}: canonical graph must contain nodes[] and edges[]")
            continue

        catalogs={}
        for key in ("variable_catalog","variable_group_catalog","artifact_catalog","playbook_projection","runtime_state_example"):
            if key not in canonical: continue
            resolved,text=_embedded_flow_resolve_resource(resources,bundle_path,canonical.get(key))
            parsed=_embedded_flow_yaml(text) if text else None
            catalogs[key]={"path":resolved,"data":parsed} if resolved and parsed is not None else {"path":resolved,"data":None}

        variables={}
        variable_data=(catalogs.get("variable_catalog") or {}).get("data")
        if isinstance(variable_data,dict):
            for row in variable_data.get("variables") or []:
                if isinstance(row,dict) and row.get("id"): variables[str(row["id"])]=copy.deepcopy(row)
        groups={}
        group_data=(catalogs.get("variable_group_catalog") or {}).get("data")
        if isinstance(group_data,dict):
            for row in group_data.get("groups") or []:
                if isinstance(row,dict) and row.get("id"): groups[str(row["id"])]=copy.deepcopy(row)
        artifacts={}
        artifact_data=(catalogs.get("artifact_catalog") or {}).get("data")
        if isinstance(artifact_data,dict):
            for row in artifact_data.get("artifacts") or []:
                if isinstance(row,dict) and row.get("id"): artifacts[str(row["id"])]=copy.deepcopy(row)

        logical_gates={str(g.get("id")):g for g in (graph.get("gates") or []) if isinstance(g,dict) and g.get("id")}
        fragment_to_gate={}
        for gid,g in logical_gates.items():
            for fragment in g.get("fragments") or []: fragment_to_gate[str(fragment)]=gid

        enriched=[]
        for raw_node in graph.get("nodes") or []:
            if not isinstance(raw_node,dict) or not raw_node.get("id"): continue
            node=copy.deepcopy(raw_node)
            if node.get("variable_ref") and str(node.get("variable_ref")) in variables:
                meta=copy.deepcopy(variables[str(node.get("variable_ref"))]); node["variable_metadata"]=meta
                group_id=str(meta.get("group_id") or "")
                if group_id and group_id in groups: node["group_metadata"]=copy.deepcopy(groups[group_id])
            if node.get("artifact_ref") and str(node.get("artifact_ref")) in artifacts:
                node["artifact_metadata"]=copy.deepcopy(artifacts[str(node.get("artifact_ref"))])
            gate_id=fragment_to_gate.get(str(node.get("id")))
            if gate_id and gate_id in logical_gates: node["gate_metadata"]=copy.deepcopy(logical_gates[gate_id])
            enriched.append(node)

        clean_edges=[]
        node_ids={str(n.get("id")) for n in enriched if n.get("id")}
        for raw_edge in graph.get("edges") or []:
            if not isinstance(raw_edge,dict): continue
            src=str(raw_edge.get("from") or ""); dst=str(raw_edge.get("to") or "")
            if not src or not dst: continue
            edge=copy.deepcopy(raw_edge); edge["from"]=src; edge["to"]=dst
            edge["dangling"]=src not in node_ids or dst not in node_ids
            clean_edges.append(edge)

        section_rows=[copy.deepcopy(x) for x in (graph.get("sections") or []) if isinstance(x,dict) and x.get("id")]
        section_rows.sort(key=lambda x:(x.get("order") if isinstance(x.get("order"),(int,float)) else 10**9,str(x.get("id"))))
        graph_out={k:copy.deepcopy(v) for k,v in graph.items() if k not in {"nodes","edges"}}
        graph_out.update({"nodes":enriched,"edges":clean_edges,"sections":section_rows})
        type_counts={}
        for n in enriched: type_counts[str(n.get("type") or "unknown")]=type_counts.get(str(n.get("type") or "unknown"),0)+1
        edge_counts={}
        for e in clean_edges: edge_counts[str(e.get("type") or "relation")]=edge_counts.get(str(e.get("type") or "relation"),0)+1
        return {
            "status":"passed","available":True,
            "bundle":{"path":bundle_path,"model_bundle_id":bundle.get("model_bundle_id"),"revision":bundle.get("revision"),"schema_version":bundle.get("schema_version"),"canonical_sources":copy.deepcopy(canonical)},
            "graph":graph_out,
            "catalogs":{k:{"path":v.get("path"),"available":isinstance(v.get("data"),dict)} for k,v in catalogs.items()},
            "summary":{"nodes":len(enriched),"edges":len(clean_edges),"sections":len(section_rows),"gates":len(logical_gates),"node_types":type_counts,"edge_types":edge_counts,"dangling_edges":sum(1 for e in clean_edges if e.get("dangling"))},
            "note":"Optional authoring data-flow model discovered from package resources. It is visualized by the Editor but does not participate in Ordo execution semantics."
        }
    return {"status":"invalid","available":False,"error":"Embedded authoring data-flow candidate found, but its canonical graph is invalid: "+"; ".join(invalid[:4]),"candidates":[p for p,_ in candidates]}


def _build_data_lineage(package: dict[str,Any] | None, source: dict[str,Any], runtime_state: dict[str,Any] | None=None) -> dict[str,Any]:
    runtime_state=runtime_state if isinstance(runtime_state,dict) else {}
    resources=package.get("resources") if isinstance(package,dict) and isinstance(package.get("resources"),dict) else {}
    records=[r for r in [*(source.get("nodes") or []),*(source.get("gates") or [])] if isinstance(r,dict) and r.get("id")]
    record_by_id={str(r.get("id")):r for r in records}
    nodes: dict[str,dict[str,Any]]={}
    edges: dict[tuple[str,str,str],dict[str,Any]]={}
    resource_state_mentions: dict[str,set[str]]={}
    registry_entries=_lineage_registry_entries(resources)
    registry_by_path={_lineage_normalize_state_path(x.get("path")):x for x in registry_entries if x.get("path")}

    for resource_path, resource_value in resources.items():
        text=None
        if isinstance(resource_value,str): text=resource_value
        elif isinstance(resource_value,dict):
            for key in ("content_text","text","content"):
                if isinstance(resource_value.get(key),str): text=resource_value.get(key); break
        if isinstance(text,str):
            for state_path in _lineage_state_refs(text): resource_state_mentions.setdefault(state_path,set()).add(str(resource_path))

    def ensure_state(path: str, *, origin: str="external", producer: str="") -> dict[str,Any]:
        path=_lineage_normalize_state_path(path)
        identifier=f"state:{path}"
        row=nodes.setdefault(identifier,{"id":identifier,"kind":"state","state_path":path,"label":path,"origin":origin,"producer_nodes":[],"consumer_nodes":[]})
        priority={"external":0,"internal":1,"deterministic":2,"model":3,"analyst_confirmed":3,"analyst":4}
        if priority.get(origin,0)>priority.get(str(row.get("origin") or "external"),0): row["origin"]=origin
        if producer and producer not in row["producer_nodes"]: row["producer_nodes"].append(producer)
        present,value=_lineage_runtime_value(runtime_state,path)
        row["current_value_available"]=present
        if present: row["current_value"]=value
        reg=registry_by_path.get(path)
        if reg:
            row["registry_metadata"]={k:copy.deepcopy(v) for k,v in reg.items() if k not in {"path"}}
            src=str(reg.get("source_node") or "")
            if src in record_by_id and src not in row["producer_nodes"]: row["producer_nodes"].append(src)
            rpath=str(reg.get("resource_path") or "")
            if rpath: resource_state_mentions.setdefault(path,set()).add(rpath)
        return row

    def ensure_artifact(path: str, *, producer: str="") -> dict[str,Any]:
        identifier=f"artifact:{path}"
        row=nodes.setdefault(identifier,{"id":identifier,"kind":_lineage_artifact_kind(path),"artifact_path":path,"label":Path(path).name or path,"producer_nodes":[],"consumer_nodes":[]})
        if producer and producer not in row["producer_nodes"]: row["producer_nodes"].append(producer)
        return row

    def ensure_transform(record: dict[str,Any], *, has_output: bool=False) -> dict[str,Any]:
        rid=str(record.get("id")); identifier=f"transform:{rid}"
        profile=_lineage_transform_profile(record,has_output=has_output)
        row=nodes.setdefault(identifier,{"id":identifier,"kind":profile["kind"],"label":profile["label"],"execution_node_id":rid,"execution_title":str(record.get("title") or rid),"mechanism":profile["mechanism"],"producer_nodes":[rid],"consumer_nodes":[]})
        return row

    def add_edge(source_id: str,target_id: str,relation: str,producer_node: str="",evidence: str="",hidden_projection: bool=False) -> None:
        if not source_id or not target_id or source_id==target_id: return
        key=(source_id,target_id,relation)
        row=edges.setdefault(key,{"source":source_id,"target":target_id,"relation":relation,"producer_nodes":[],"evidence":[]})
        if hidden_projection: row["hidden_projection"]=True
        if producer_node and producer_node not in row["producer_nodes"]: row["producer_nodes"].append(producer_node)
        if evidence and evidence not in row["evidence"]: row["evidence"].append(evidence)

    record_outputs: dict[str,list[str]]={}
    all_artifacts=set()
    try: declared=_projection_declared_outputs([r for r in (source.get("nodes") or []) if isinstance(r,dict)])
    except Exception: declared=[]
    for item in declared or []:
        if not isinstance(item,dict): continue
        path=str(item.get("path") or "").strip()
        if not path: continue
        all_artifacts.add(path); art=ensure_artifact(path)
        for producer in item.get("producers") or []:
            producer=str(producer)
            if producer and producer not in art["producer_nodes"]: art["producer_nodes"].append(producer)
            record_outputs.setdefault(producer,[]).append(path)
    for record in records:
        rid=str(record.get("id")); outputs=_lineage_output_paths(record)
        record_outputs.setdefault(rid,[])
        for path in outputs:
            if path not in record_outputs[rid]: record_outputs[rid].append(path)
            all_artifacts.add(path); ensure_artifact(path,producer=rid)

    # Build producer transformations for state changes. This preserves the mechanism between inputs and outputs instead of inventing direct state→state edges.
    for record in records:
        rid=str(record.get("id")); actor=_lineage_record_actor(record); writes=_lineage_state_writes(record); outputs=record_outputs.get(rid) or []
        declared_states,declared_artifacts=_lineage_declared_input_refs(record)
        source_refs=set(declared_states)
        for write in writes: source_refs.update(_lineage_state_refs(write.get("expression")))
        # Explicit transformation blocks often hold the real upstream dependency.
        for block_key in ("proposal_generation","draft_generation","interpretation_policy","derivation_contract","generation_contract","derive_before_generate"):
            if block_key in record: source_refs.update(_lineage_state_refs(record.get(block_key)))
        transform=None
        if writes or outputs:
            transform=ensure_transform(record,has_output=bool(outputs))
            for ref in sorted(source_refs):
                s=ensure_state(ref,origin="external")
                if rid not in s["consumer_nodes"]: s["consumer_nodes"].append(rid)
                add_edge(s["id"],transform["id"],"input_to",rid,"declared/read input")
            for dep_path in declared_artifacts:
                if dep_path in outputs: continue
                dep=ensure_artifact(dep_path)
                if rid not in dep["consumer_nodes"]: dep["consumer_nodes"].append(rid)
                add_edge(dep["id"],transform["id"],"input_to",rid,"declared artifact input")
        for write in writes:
            expr_refs=_lineage_state_refs(write.get("expression"))
            # A direct analyst answer with no upstream proposal/source is a true analyst-input root.
            from_answer=isinstance(write.get("expression"),str) and str(write.get("expression")).strip().startswith("$answer")
            has_upstream=bool(source_refs)
            origin=("analyst_confirmed" if actor=="analyst" and from_answer and has_upstream else actor)
            target=ensure_state(write["path"],origin=origin,producer=rid)
            if transform:
                add_edge(transform["id"],target["id"],"produces",rid,write.get("role") or "state write")
            # Preserve the previous machine-readable direct dependency contract for exact expression refs,
            # while the UI uses the richer transformation path.
            for ref in expr_refs:
                src=ensure_state(ref,origin="external")
                add_edge(src["id"],target["id"],"derived_from",rid,"direct expression dependency",hidden_projection=True)
            if actor=="analyst" and from_answer and not has_upstream:
                target["origin"]="analyst"
        # Every explicit read becomes a consumer relationship even when no write exists.
        for ref in set(_lineage_state_refs(record)) | set(declared_states):
            row=ensure_state(ref,origin="external")
            if rid not in row["consumer_nodes"]: row["consumer_nodes"].append(rid)

    # Enrich materialization/package transformations from templates, bindings, manifests and explicit package includes.
    for record in records:
        rid=str(record.get("id")); outputs=record_outputs.get(rid) or []
        if not outputs: continue
        transform=ensure_transform(record,has_output=True)
        state_refs=set(_lineage_state_refs(record)); declared_states,declared_artifacts=_lineage_declared_input_refs(record); state_refs.update(declared_states)
        artifact_refs=set(_lineage_artifact_refs(record)); artifact_refs.update(declared_artifacts)
        resource_evidence={}
        for ref in _generic_record_resource_references(record,resources):
            resolved,text=_resolve_package_resource(resources,str(ref.get("path") or ""))
            if isinstance(text,str):
                discovered_refs=_lineage_state_refs(text); state_refs.update(discovered_refs); artifact_refs.update(_lineage_artifact_refs(text))
                resolved_path=str((ref or {}).get("resolved_path") or (ref or {}).get("path") or "")
                for state_path in discovered_refs:
                    if resolved_path:
                        resource_state_mentions.setdefault(state_path,set()).add(resolved_path)
                        resource_evidence.setdefault(state_path,[]).append(resolved_path)
        for state_path in sorted(state_refs):
            s=ensure_state(state_path,origin="external")
            if rid not in s["consumer_nodes"]: s["consumer_nodes"].append(rid)
            add_edge(s["id"],transform["id"],"input_to",rid,", ".join(resource_evidence.get(state_path,[])) or "materialization input")
        for dep_path in sorted(artifact_refs):
            if dep_path in outputs: continue
            dep=ensure_artifact(dep_path)
            if rid not in dep["consumer_nodes"]: dep["consumer_nodes"].append(rid)
            add_edge(dep["id"],transform["id"],"input_to",rid,"artifact/package dependency")
        for output in outputs:
            out_node=ensure_artifact(output,producer=rid)
            relation="packages" if out_node["kind"]=="archive" else "materializes"
            add_edge(transform["id"],out_node["id"],relation,rid,"declared output")
            # Backward-compatible direct data→artifact edges remain machine-readable but hidden in the visual projection.
            for state_path in sorted(state_refs):
                state_node=ensure_state(state_path,origin="external")
                add_edge(state_node["id"],out_node["id"],"materializes",rid,"direct materialization dependency",hidden_projection=True)
            for dep_path in sorted(artifact_refs):
                if dep_path==output: continue
                dep=ensure_artifact(dep_path)
                direct_relation="packages" if out_node["kind"]=="archive" else "includes"
                add_edge(dep["id"],out_node["id"],direct_relation,rid,"direct artifact dependency",hidden_projection=True)

    # Registry source_node metadata supplements producer traceability; it never fabricates a consumer edge.
    for reg in registry_entries:
        path=_lineage_normalize_state_path(reg.get("path")); row=ensure_state(path,origin="external")
        src=str(reg.get("source_node") or "")
        if src in record_by_id and src not in row["producer_nodes"]: row["producer_nodes"].append(src)

    # Make read-only state references visible.
    for record in records:
        _,declared_artifacts=_lineage_declared_input_refs(record)
        for ref in _lineage_state_refs(record): ensure_state(ref,origin="external")
        for art in declared_artifacts: ensure_artifact(art)

    # Final user-facing state kinds. Analyst-confirmed values with upstream derivation stay derived, while direct answer roots remain analyst inputs.
    for row in nodes.values():
        if row.get("kind")=="state":
            origin=str(row.get("origin") or "external")
            row["kind"]="analyst_input" if origin=="analyst" else "derived_state"
            row["source_kind"]=origin

    edge_rows=list(edges.values())
    visible_edge_rows=[e for e in edge_rows if not e.get("hidden_projection")]
    for edge in visible_edge_rows:
        src=nodes.get(edge["source"]); dst=nodes.get(edge["target"])
        if src and dst and dst.get("producer_nodes"):
            for pid in dst.get("producer_nodes") or []:
                if pid not in src.setdefault("consumer_nodes",[]): src["consumer_nodes"].append(pid)

    # Semantic layers keep transformation mechanisms between data and artifacts.
    base={"analyst_input":0,"derived_state":2,"transform_analyst":1,"transform_model":1,"transform_deterministic":1,"transform_tool":1,"transform_template":3,"document":4,"artifact":4,"transform_package":5,"archive":6}
    indegree={nid:0 for nid in nodes}; outgoing={nid:[] for nid in nodes}
    for e in visible_edge_rows:
        if e["source"] in nodes and e["target"] in nodes:
            outgoing[e["source"]].append(e["target"]); indegree[e["target"]]+=1
    layers={nid:base.get(str(row.get("kind")),2) for nid,row in nodes.items()}
    queue=sorted([nid for nid,d in indegree.items() if d==0],key=lambda nid:(layers[nid],nid))
    while queue:
        nid=queue.pop(0)
        for target in outgoing.get(nid,[]):
            layers[target]=max(layers[target],layers[nid]+1)
            indegree[target]-=1
            if indegree[target]==0: queue.append(target)
    for nid,row in nodes.items():
        row["layer"]=int(layers.get(nid,base.get(str(row.get("kind")),2)))
        row["producer_nodes"]=sorted(set(row.get("producer_nodes") or [])); row["consumer_nodes"]=sorted(set(row.get("consumer_nodes") or []))
        row["incoming_count"]=sum(1 for e in visible_edge_rows if e["target"]==nid); row["outgoing_count"]=sum(1 for e in visible_edge_rows if e["source"]==nid)
        row["final_artifact"]=row.get("kind") in {"document","artifact","archive"} and row["outgoing_count"]==0
        if row.get("kind") in {"document","artifact","archive"}:
            artifact_path=str(row.get("artifact_path") or "")
            row["artifact_role"]="package_source_resource" if artifact_path in resources and not row.get("producer_nodes") else "generated_or_runtime_artifact"
        if row.get("kind") in {"analyst_input","derived_state"}:
            state_path=str(row.get("state_path") or ""); candidates=sorted(resource_state_mentions.get(state_path,set())); row["candidate_usage_resources"]=candidates
            if row["outgoing_count"]==0:
                if candidates: row["lineage_diagnostic"]={"code":"POTENTIALLY_UNRESOLVED_USAGE","label":"Potentially unresolved usage","message":"No downstream lineage edge was resolved, but package resources mention this state path. The extractor may not yet have resolved the complete binding/template relationship."}
                else: row["lineage_diagnostic"]={"code":"UNUSED_IN_DISCOVERED_LINEAGE","label":"Unused in discovered lineage","message":"No downstream use was discovered in source, resolved templates/bindings, registries, or known output contracts. This alone does not prove a playbook defect."}

    kind_counts={}
    for row in nodes.values(): kind_counts[row["kind"]]=kind_counts.get(row["kind"],0)+1
    return {"status":"passed","nodes":sorted(nodes.values(),key=lambda r:(r.get("layer",0),r.get("kind",""),r.get("label",""))),"edges":sorted(edge_rows,key=lambda e:(e["source"],e["target"],e["relation"])),"summary":{"entities":len(nodes),"relations":len(visible_edge_rows),"kinds":kind_counts,"final_artifacts":sum(1 for r in nodes.values() if r.get("final_artifact")),"transformations":sum(v for k,v in kind_counts.items() if str(k).startswith("transform_"))},"note":"Logical data lineage inferred from declared analyst/model/deterministic transformations, state reads/writes, output contracts, templates/bindings, package includes and package traceability registries. Transformation nodes show how data is produced; this is not executable control flow."}



def _provider_response_shape(value: Any, depth: int = 0) -> Any:
    """Return only keys/types for diagnostics; never copy provider message text."""
    if depth >= 4:
        return type(value).__name__
    if isinstance(value, dict):
        return {str(k): _provider_response_shape(v, depth + 1) for k, v in list(value.items())[:30] if k != "_ordo_debug"}
    if isinstance(value, list):
        return [_provider_response_shape(v, depth + 1) for v in value[:3]]
    return type(value).__name__


def _model_chat_archive_priority(path: str) -> tuple[int, int, str]:
    """Lower score means a file is more useful as initial workspace context."""
    value=str(path or "").replace("\\","/").lower()
    name=Path(value).name
    priority=50
    if any(token in name for token in ("start","startup","entrypoint","entry_point","bootstrap")):
        priority=0
    elif "prompt" in name:
        priority=2
    elif name.startswith("readme") or name in {"readme.md","readme.txt"}:
        priority=4
    elif any(token in name for token in ("instruction","guide","handoff","transfer")):
        priority=6
    elif "manifest" in name:
        priority=8
    elif Path(name).suffix in {".yaml",".yml"}:
        priority=20
    elif Path(name).suffix in {".md",".txt"}:
        priority=25
    return (priority,len(value),value)


def _compact_model_chat_zip(raw_zip: bytes, filename: str) -> dict[str, Any]:
    """Build a bounded workspace index plus a few likely entrypoint files."""
    result={"filename":filename,"media_type":"application/zip","archive_index":[],"entrypoint_context":[]}
    if not zipfile.is_zipfile(io.BytesIO(raw_zip)):
        result["note"]="ZIP attachment could not be read as an archive."
        return result

    candidates=[]
    total_files=0
    total_uncompressed=0
    with zipfile.ZipFile(io.BytesIO(raw_zip),"r") as z:
        for info in z.infolist():
            if info.is_dir():
                continue
            total_files+=1
            total_uncompressed+=int(info.file_size or 0)
            suffix=Path(info.filename).suffix.lower()
            text_candidate=suffix in (_TEXT_RESOURCE_EXTENSIONS|{".yaml",".yml",".md",".json",".py",".txt"})
            if len(result["archive_index"]) < 800:
                result["archive_index"].append({
                    "path":info.filename,
                    "bytes":int(info.file_size or 0),
                    "text_candidate":bool(text_candidate),
                })
            if text_candidate and info.file_size <= 350000:
                candidates.append(info)

        # Initial context is deliberately small. The archive index tells the model
        # what exists without spending most of the context window on file bodies.
        for info in sorted(candidates,key=lambda row:_model_chat_archive_priority(row.filename))[:6]:
            try:
                body=z.read(info.filename).decode("utf-8")
            except Exception:
                continue
            result["entrypoint_context"].append({
                "path":info.filename,
                "content_text":body[:24000],
                "truncated":len(body)>24000,
            })

    result["workspace_summary"]={
        "files":total_files,
        "uncompressed_bytes":total_uncompressed,
        "indexed_files":len(result["archive_index"]),
        "initial_text_files":len(result["entrypoint_context"]),
    }
    result["note"]="ZIP indexed as a local workspace. Only likely entrypoint/instruction files are embedded in the initial model request."
    return result



MODEL_CHAT_WORKSPACES = Path(tempfile.gettempdir()) / "ordo_model_chat_workspaces"

def _workspace_root(session_id: str) -> Path:
    safe=re.sub(r"[^A-Za-z0-9_.-]+","_",str(session_id or "default"))[:120] or "default"
    root=(MODEL_CHAT_WORKSPACES/safe).resolve()
    root.mkdir(parents=True,exist_ok=True)
    for sub in ("uploads","extracted","generated","tmp"):
        (root/sub).mkdir(exist_ok=True)
    return root

def _workspace_safe_path(root: Path, relative: str) -> Path:
    root=root.resolve()
    rel=str(relative or "").replace("\\","/").lstrip("/")
    target=(root/rel).resolve()
    if root!=target and root not in target.parents:
        raise ValueError("Workspace path escapes session root.")
    return target

def _workspace_index(root: Path, limit: int = 1200) -> list[dict[str, Any]]:
    out=[]
    for p in sorted(root.rglob("*")):
        if not p.is_file(): continue
        try: size=p.stat().st_size
        except OSError: continue
        out.append({"path":p.relative_to(root).as_posix(),"bytes":size,"suffix":p.suffix.lower()})
        if len(out)>=limit: break
    return out

def _workspace_head(root: Path) -> dict[str, Any]:
    """Small deterministic workspace summary safe to include in every model iteration."""
    files=[]
    total_bytes=0
    top_counts={}
    candidates=[]
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        try:
            size=p.stat().st_size
        except OSError:
            continue
        rel=p.relative_to(root).as_posix()
        files.append((rel,size))
        total_bytes+=size
        parts=Path(rel).parts
        if parts:
            top=parts[0] if len(parts)==1 else "/".join(parts[:2])
            top_counts[top]=top_counts.get(top,0)+1
        suffix=p.suffix.lower()
        if suffix in (_TEXT_RESOURCE_EXTENSIONS|{".yaml",".yml",".md",".json",".py",".txt",".ordo"}):
            candidates.append(rel)

    ranked=sorted(candidates,key=lambda value:_model_chat_archive_priority(value))
    entrypoints=[]
    for rel in ranked:
        score=_model_chat_archive_priority(rel)[0]
        # Only likely startup/instruction candidates. Do not turn arbitrary YAML/MD
        # files into a large first-turn list.
        if score>8:
            continue
        entrypoints.append(rel)
        if len(entrypoints)>=6:
            break

    top_dirs=[
        {"path":path,"files":count}
        for path,count in sorted(top_counts.items(),key=lambda item:(-item[1],item[0]))[:8]
    ]
    return {
        "files":len(files),
        "bytes":total_bytes,
        "top_dirs":top_dirs,
        "entrypoint_candidates":entrypoints,
        "note":"Use workspace tools to inspect files. File contents are not embedded in workspace_head.",
    }

def _workspace_store_attachments(root: Path, attachments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    stored=[]
    for item in attachments or []:
        if not isinstance(item,dict): continue
        filename=Path(str(item.get("filename") or "attachment")).name
        b64=item.get("content_base64")
        if not isinstance(b64,str): continue
        try: raw=base64.b64decode(b64,validate=True)
        except Exception: continue
        target=_workspace_safe_path(root,f"uploads/{filename}")
        target.write_bytes(raw)
        rec={
            "filename":filename,
            "media_type":str(item.get("media_type") or mimetypes.guess_type(filename)[0] or "application/octet-stream"),
            "size_bytes":len(raw),
            "workspace_path":target.relative_to(root).as_posix(),
            "bytes":len(raw),
        }
        if filename.lower().endswith(".zip") and zipfile.is_zipfile(io.BytesIO(raw)):
            dest=_workspace_safe_path(root,f"extracted/{Path(filename).stem}")
            if dest.exists(): shutil.rmtree(dest)
            dest.mkdir(parents=True,exist_ok=True)
            with zipfile.ZipFile(io.BytesIO(raw),"r") as z:
                for info in z.infolist():
                    if info.is_dir(): continue
                    out=_workspace_safe_path(dest,info.filename)
                    out.parent.mkdir(parents=True,exist_ok=True)
                    with z.open(info,"r") as src, out.open("wb") as dst:
                        shutil.copyfileobj(src,dst)
            rec["extracted_to"]=dest.relative_to(root).as_posix()
        stored.append(rec)
    return stored


def _workspace_file_snapshot(root: Path) -> dict[str, tuple[int,int]]:
    result={}
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        try:
            st=p.stat()
        except OSError:
            continue
        result[p.relative_to(root).as_posix()]=(int(st.st_mtime_ns),int(st.st_size))
    return result


def _workspace_changed_files(root: Path, before: dict[str, tuple[int,int]]) -> list[Path]:
    changed=[]
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        rel=p.relative_to(root).as_posix()
        if rel.startswith("uploads/") or rel.startswith("extracted/"):
            continue
        try:
            st=p.stat()
        except OSError:
            continue
        now=(int(st.st_mtime_ns),int(st.st_size))
        if before.get(rel)!=now:
            changed.append(p)
    return sorted(changed,key=lambda p:p.relative_to(root).as_posix())


def _workspace_tool_execute(root: Path, call: dict[str, Any]) -> dict[str, Any]:
    root=root.resolve()
    name=str(call.get("name") or "")
    args=call.get("arguments") if isinstance(call.get("arguments"),dict) else {}
    if name=="workspace.list":
        prefix=str(args.get("path") or "")
        base=_workspace_safe_path(root,prefix) if prefix else root
        if not base.exists(): return {"ok":False,"error":"path_not_found"}
        items=[]
        for p in sorted(base.iterdir()):
            items.append({"name":p.name,"path":p.relative_to(root).as_posix(),"type":"dir" if p.is_dir() else "file","bytes":p.stat().st_size if p.is_file() else None})
            if len(items)>=300: break
        return {"ok":True,"items":items}
    if name=="workspace.search":
        query=str(args.get("query") or "").lower()
        path=str(args.get("path") or "")
        base=_workspace_safe_path(root,path) if path else root
        results=[]
        for p in base.rglob("*"):
            if not p.is_file(): continue
            rel=p.relative_to(root).as_posix()
            if query and query in rel.lower():
                results.append({"path":rel,"match":"filename"})
            elif p.suffix.lower() in {".md",".txt",".yaml",".yml",".json",".py",".ordo"} and p.stat().st_size<=400000:
                try: text=p.read_text(encoding="utf-8",errors="ignore")
                except Exception: continue
                idx=text.lower().find(query) if query else -1
                if idx>=0:
                    results.append({"path":rel,"match":"content","excerpt":text[max(0,idx-120):idx+240]})
            if len(results)>=80: break
        return {"ok":True,"results":results}
    if name=="workspace.read":
        path=_workspace_safe_path(root,str(args.get("path") or ""))
        if not path.is_file(): return {"ok":False,"error":"file_not_found"}
        max_chars=max(1000,min(int(args.get("max_chars") or 30000),120000))
        if path.stat().st_size>2_000_000: return {"ok":False,"error":"file_too_large"}
        text=path.read_text(encoding="utf-8",errors="replace")
        return {"ok":True,"path":path.relative_to(root).as_posix(),"content":text[:max_chars],"truncated":len(text)>max_chars}
    if name=="workspace.stat":
        path=_workspace_safe_path(root,str(args.get("path") or ""))
        if not path.exists(): return {"ok":False,"error":"path_not_found"}
        return {"ok":True,"path":path.relative_to(root).as_posix(),"type":"dir" if path.is_dir() else "file","bytes":path.stat().st_size if path.is_file() else None}
    if name=="workspace.write":
        path=_workspace_safe_path(root,str(args.get("path") or "generated/output.txt"))
        content=str(args.get("content") or "")
        path.parent.mkdir(parents=True,exist_ok=True)
        path.write_text(content,encoding="utf-8")
        return {"ok":True,"path":path.relative_to(root).as_posix(),"bytes":len(content.encode("utf-8"))}
    if name=="workspace.archive":
        source=_workspace_safe_path(root,str(args.get("source") or ""))
        if not source.exists():
            return {"ok":False,"error":"source_not_found"}
        output_arg=str(args.get("output") or "").strip() or f"generated/{source.name or 'archive'}.zip"
        output=_workspace_safe_path(root,output_arg)
        if output.suffix.lower()!=".zip":
            output=output.with_suffix(".zip")
        output.parent.mkdir(parents=True,exist_ok=True)
        with zipfile.ZipFile(output,"w",zipfile.ZIP_DEFLATED) as archive:
            if source.is_file():
                archive.write(source,arcname=source.name)
            else:
                for child in sorted(source.rglob("*")):
                    if not child.is_file():
                        continue
                    if child.resolve()==output.resolve():
                        continue
                    archive.write(child,arcname=child.relative_to(source).as_posix())
        return {"ok":True,"path":output.relative_to(root).as_posix(),"bytes":output.stat().st_size,"source":source.relative_to(root).as_posix()}
    return {"ok":False,"error":"unsupported_tool","tool":name}

def _model_chat_tool_schema() -> list[dict[str, Any]]:
    return [
        {"name":"workspace.list","description":"List files/directories in the current local chat workspace.","arguments":{"path":"optional relative path"}},
        {"name":"workspace.search","description":"Search workspace filenames and supported text file contents.","arguments":{"query":"required string","path":"optional relative path"}},
        {"name":"workspace.read","description":"Read a text file from the workspace.","arguments":{"path":"required relative path","max_chars":"optional integer"}},
        {"name":"workspace.stat","description":"Get metadata for a workspace file/directory.","arguments":{"path":"required relative path"}},
        {"name":"workspace.write","description":"Create or replace a text file in generated workspace.","arguments":{"path":"relative path","content":"text"}},
        {"name":"workspace.archive","description":"Create a ZIP archive from a workspace file or directory. Use this when the user asks for a downloadable package/archive.","arguments":{"source":"required relative source file/directory","output":"optional relative .zip output path"}},
    ]

def _model_chat_parse_agent_turn(raw_text: str) -> dict[str, Any]:
    try:
        parsed=_parse_model_json(raw_text)
    except Exception:
        return {"type":"final","message":str(raw_text or "").strip()}
    if not isinstance(parsed,dict):
        return {"type":"final","message":str(raw_text or "").strip()}

    # Canonical form: {"type":"tool","tool":{"name":"...","arguments":{...}}}
    kind=str(parsed.get("type") or parsed.get("action") or "").lower()
    if kind in {"tool","tool_call","call_tool"}:
        tool=parsed.get("tool") if isinstance(parsed.get("tool"),dict) else parsed
        name=tool.get("name") or parsed.get("name")
        arguments=tool.get("arguments") if isinstance(tool.get("arguments"),dict) else parsed.get("arguments")
        return {"type":"tool","name":str(name or ""),"arguments":arguments if isinstance(arguments,dict) else {},"parsed":parsed}

    # Generic equivalent form used by some OpenAI-compatible models:
    # {"tool_call":{"name":"workspace.list","arguments":{...}}}
    direct=parsed.get("tool_call")
    if isinstance(direct,dict):
        name=direct.get("name")
        arguments=direct.get("arguments")
        return {"type":"tool","name":str(name or ""),"arguments":arguments if isinstance(arguments,dict) else {},"parsed":parsed}

    # Also accept a singular nested `call_tool` object as structural spelling.
    direct=parsed.get("call_tool")
    if isinstance(direct,dict):
        name=direct.get("name")
        arguments=direct.get("arguments")
        return {"type":"tool","name":str(name or ""),"arguments":arguments if isinstance(arguments,dict) else {},"parsed":parsed}

    msg=_extract_structured_chat_answer(parsed)
    return {"type":"final","message":msg or str(raw_text or "").strip(),"parsed":parsed}


def _model_chat_tool_activity_label(name: str, arguments: Any) -> str:
    labels={
        "workspace.list":"Inspecting workspace files",
        "workspace.search":"Searching workspace",
        "workspace.read":"Reading workspace file",
        "workspace.write":"Writing workspace file",
        "workspace.archive":"Creating ZIP archive",
        "workspace.extract":"Extracting archive",
        "workspace.run":"Running local command",
    }
    return labels.get(str(name or ""),"Using local workspace tool")

def _model_chat_tool_result_label(name: str, result: Any) -> str:
    ok=bool(result.get("ok")) if isinstance(result,dict) and "ok" in result else True
    if not ok:
        return "Workspace tool failed"
    labels={
        "workspace.list":"Workspace inspected",
        "workspace.search":"Workspace search completed",
        "workspace.read":"File read",
        "workspace.write":"File written",
        "workspace.archive":"ZIP archive created",
        "workspace.extract":"Archive extracted",
        "workspace.run":"Local command completed",
    }
    return labels.get(str(name or ""),"Workspace action completed")


class _ModelChatHistory(list):
    """List-compatible history carrying latest-turn attachment metadata internally."""
    def __init__(self, values: Any = (), current_attachments: list[dict[str, Any]] | None = None):
        super().__init__(values or [])
        self.current_attachments=copy.deepcopy(current_attachments or [])

def _model_chat_tool_scopes_current_attachment(
    name: str,
    arguments: Any,
    current_attachments: list[dict[str, Any]],
) -> bool:
    """Return True when a model workspace inspection is scoped to the latest attachment.

    Presence grounding by the host is useful evidence, but it is not proof that the
    provider actually inspected that evidence. For a turn containing a newly
    attached file, a final answer is only trustworthy after the provider has issued
    at least one read/list/search/stat call whose path is the attachment itself or a
    descendant of its extracted directory.
    """
    if str(name or "") not in {"workspace.list", "workspace.search", "workspace.read", "workspace.stat"}:
        return False
    args=arguments if isinstance(arguments,dict) else {}
    path=str(args.get("path") or "").strip().replace("\\","/").strip("/")
    if not path:
        return False
    for item in current_attachments or []:
        if not isinstance(item,dict):
            continue
        for key in ("workspace_path","extracted_to"):
            root=str(item.get(key) or "").strip().replace("\\","/").strip("/")
            if root and (path==root or path.startswith(root+"/")):
                return True
    return False

def _model_chat_agent_loop(
    credentials: dict[str, Any],
    root: Path,
    user_message: str,
    history: list[dict[str, Any]],
    max_iterations: int = 12,
    activity_callback: Any = None,
    cancel_check: Any = None,
) -> tuple[str,list[dict[str,Any]],dict[str,int],list[dict[str,Any]]]:
    trace=[]
    activities=[]
    usage_total={"prompt_tokens":0,"completion_tokens":0,"total_tokens":0}
    current_attachments=copy.deepcopy(getattr(history,"current_attachments",[]) or [])
    attachment_state=[]
    for item in current_attachments[:8]:
        if not isinstance(item,dict):
            continue
        attachment_state.append({
            "filename":str(item.get("filename") or ""),
            "workspace_path":str(item.get("workspace_path") or ""),
            "extracted_to":str(item.get("extracted_to") or ""),
            "media_type":str(item.get("media_type") or ""),
            "size_bytes":int(item.get("size_bytes") or item.get("bytes") or 0),
        })
    context={
        "task":"Model Chat agent turn",
        "workspace_head":_workspace_head(root),
        "current_attachments":current_attachments,
        "conversation":history[-12:],
        "user_message":user_message,
        "available_tools":_model_chat_tool_schema(),
        "agent_protocol":{
            "tool_call":{"type":"tool","tool":{"name":"workspace.read","arguments":{"path":"..."}}},
            "final":{"type":"final","message":"..."},
            "rule":"Return exactly one JSON object. Use tools when local workspace information is required. Do not invent file contents."
        },
    }
    if attachment_state:
        context["attachment_contract"]={
            "inspection_required_before_final":True,
            "model_inspection_satisfied":False,
            "final_allowed":False,
            "authoritative_paths":[
                {"filename":x.get("filename"),"workspace_path":x.get("workspace_path"),"extracted_to":x.get("extracted_to")}
                for x in attachment_state
            ],
            "required_tool_scope":"Before a final answer, call workspace.list/search/read/stat with path equal to or inside one authoritative workspace_path/extracted_to.",
        }
    attachment_notice=""
    if attachment_state:
        attachment_notice=(
            " LATEST TURN ATTACHMENTS ARE ALREADY PRESENT. This is authoritative runtime state, not a user claim. "
            + json.dumps(attachment_state,ensure_ascii=False)
            + " Do not ask the user to upload or re-attach these files. Use their exact workspace_path/extracted_to paths. "
        )
    system=("You are a tool-using assistant inside Ordo Editor Model Chat. You have a persistent local workspace. "
            "workspace_head is only a compact summary and may include a few likely entrypoint candidates; it is not a complete file listing. "
            "current_attachments is the authoritative list of files successfully attached to the latest user turn and stored in the workspace. "
            "If current_attachments is non-empty, never tell the user that no file was attached. A final answer is invalid until you have called workspace.list/search/read/stat with a path equal to or inside one of the current attachment workspace_path/extracted_to paths; the runtime will reject a premature final. "
            + attachment_notice +
            "When you need local files, use workspace.list/search/read. Use workspace.write to create text files and workspace.archive to create ZIP packages. "
            "Never claim that you cannot create a ZIP when workspace.archive is available. "
            "Never invent a path or claim to have read a file that was not returned by a tool. "
            "Do not rely on sandbox: markdown URLs for downloads; created files are surfaced by the host UI. "
            "Return exactly one JSON object: either a tool call or a final user-facing message.")
    attachment_grounding_done=False
    attachment_model_inspected=False
    attachment_final_rejections=0
    model_tool_calls=0
    for iteration in range(1,max_iterations+1):
        if callable(cancel_check) and cancel_check():
            return "",trace,usage_total,activities
        _,resp,raw,usage=_provider_api_call(credentials,system,context)
        for k in ("prompt_tokens","completion_tokens","total_tokens"):
            usage_total[k]+=int((usage or {}).get(k) or 0)
        turn=_model_chat_parse_agent_turn(raw)
        trace.append({"iteration":iteration,"model":turn})
        if turn.get("type")!="tool":
            # A latest-turn attachment is authoritative runtime state. Do not accept
            # a provider final until the provider has actually inspected that exact
            # attachment path. Host-side grounding proves presence, but the .149
            # debug showed that some local models can still ignore that evidence and
            # repeat a false "please upload" final. The runtime therefore enforces
            # one attachment-scoped model tool call before finalization.
            if current_attachments and not attachment_model_inspected:
                attachment_final_rejections+=1
                trace.append({
                    "iteration":iteration,
                    "runtime_guard":{
                        "kind":"attachment_inspection_required",
                        "rejected_final":str(turn.get("message") or ""),
                        "rejection_count":attachment_final_rejections,
                    },
                })
                grounding=[]
                if not attachment_grounding_done:
                    attachment_grounding_done=True
                    for attached in current_attachments[:8]:
                        if not isinstance(attached,dict):
                            continue
                        extracted=str(attached.get("extracted_to") or "").strip()
                        workspace_path=str(attached.get("workspace_path") or "").strip()
                        if extracted:
                            call={"name":"workspace.list","arguments":{"path":extracted}}
                        elif workspace_path:
                            call={"name":"workspace.stat","arguments":{"path":workspace_path}}
                        else:
                            continue
                        call_activity={"kind":"tool_call","name":call["name"],"label":"Grounding attached file"}
                        activities.append(call_activity)
                        if callable(activity_callback):
                            activity_callback(copy.deepcopy(call_activity))
                        result=_workspace_tool_execute(root,call)
                        grounding.append({
                            "attachment":str(attached.get("filename") or ""),
                            "workspace_path":workspace_path,
                            "extracted_to":extracted,
                            "tool":call,
                            "result":result,
                        })
                        result_activity={"kind":"tool_result","name":call["name"],"ok":bool(result.get("ok")) if isinstance(result,dict) else True,"label":"Attached file grounded"}
                        activities.append(result_activity)
                        if callable(activity_callback):
                            activity_callback(copy.deepcopy(result_activity))
                    context["attachment_grounding"]=grounding
                contract=context.get("attachment_contract") if isinstance(context.get("attachment_contract"),dict) else {}
                contract.update({
                    "inspection_required_before_final":True,
                    "model_inspection_satisfied":False,
                    "final_allowed":False,
                    "rejected_final_count":attachment_final_rejections,
                    "required_next_action":"Return a workspace.list/search/read/stat tool call scoped to one authoritative attachment path. Do not return a final yet.",
                })
                context["attachment_contract"]=contract
                context["conversation"]=context.get("conversation",[])+[
                    {"role":"runtime","content":{
                        "attachment_grounding":context.get("attachment_grounding",grounding),
                        "final_rejected":True,
                        "instruction":"The latest-turn attachments are confirmed present. A final answer is not allowed yet. Inspect one exact attachment workspace_path/extracted_to path with workspace.list/search/read/stat, then continue the user's request. Never ask for re-upload."
                    }}
                ]
                context["conversation"]=context["conversation"][-18:]
                continue
            return str(turn.get("message") or ""),trace,usage_total,activities
        model_tool_calls+=1
        tool_name=str(turn.get("name") or "")
        if current_attachments and _model_chat_tool_scopes_current_attachment(tool_name,turn.get("arguments"),current_attachments):
            attachment_model_inspected=True
            contract=context.get("attachment_contract") if isinstance(context.get("attachment_contract"),dict) else {}
            contract.update({"model_inspection_satisfied":True,"final_allowed":True,"satisfied_by":{"name":tool_name,"arguments":copy.deepcopy(turn.get("arguments") or {})}})
            context["attachment_contract"]=contract
        activity={"kind":"tool_call","name":tool_name,"label":_model_chat_tool_activity_label(tool_name,turn.get("arguments"))}
        activities.append(activity)
        if callable(activity_callback):
            activity_callback(copy.deepcopy(activity))
        if callable(cancel_check) and cancel_check():
            return "",trace,usage_total,activities
        result=_workspace_tool_execute(root,{"name":tool_name,"arguments":turn.get("arguments")})
        trace.append({"iteration":iteration,"tool":{"name":tool_name,"result":result}})
        result_activity={"kind":"tool_result","name":tool_name,"ok":bool(result.get("ok")) if isinstance(result,dict) else True,"label":_model_chat_tool_result_label(tool_name,result)}
        activities.append(result_activity)
        if callable(activity_callback):
            activity_callback(copy.deepcopy(result_activity))
        context["conversation"]=context.get("conversation",[])+[
            {"role":"assistant","content":{"tool_call":{"name":turn.get("name"),"arguments":turn.get("arguments")}}},
            {"role":"tool","content":result},
        ]
        # Bound tool-result growth.
        context["conversation"]=context["conversation"][-18:]
    if current_attachments and not attachment_model_inspected:
        names=[str(x.get("filename") or "") for x in current_attachments if isinstance(x,dict) and str(x.get("filename") or "")]
        shown=", ".join(names[:4]) or "the latest attachment"
        return (
            f"The attachment is present in the workspace ({shown}), but the configured model did not inspect it after repeated runtime retries. "
            "The file was not lost and does not need to be uploaded again. Please retry this turn or use another configured model.",
            trace,usage_total,activities
        )
    return "I reached the internal tool-iteration limit before producing a final answer.",trace,usage_total,activities



def _model_chat_attachment_metadata(items: Any) -> list[dict[str, Any]]:
    result=[]
    if not isinstance(items,list):
        return result
    for item in items:
        if not isinstance(item,dict):
            continue
        result.append({
            "filename":str(item.get("filename") or ""),
            "media_type":str(item.get("media_type") or ""),
            "size_bytes":int(item.get("size_bytes") or 0),
            "has_content_text":isinstance(item.get("content_text"),str),
            "has_content_base64":isinstance(item.get("content_base64"),str),
        })
    return result

def _model_chat_safe_messages(messages: Any) -> list[dict[str, Any]]:
    """Keep visible conversation plus attachment metadata, never attachment bodies."""
    safe=[]
    if not isinstance(messages,list):
        return safe
    for item in messages:
        if not isinstance(item,dict):
            continue
        row={
            "role":str(item.get("role") or ""),
            "content":str(item.get("content") or ""),
        }
        files=_model_chat_attachment_metadata(item.get("files"))
        if files:
            row["files"]=files
        activities=item.get("activities")
        if isinstance(activities,list):
            row["activities"]=copy.deepcopy(activities)
        safe.append(row)
    return safe

def _model_chat_export_attachment_metadata(messages: Any, pending: Any) -> list[dict[str, Any]]:
    result=[]
    if isinstance(messages,list):
        for index,item in enumerate(messages):
            if not isinstance(item,dict):
                continue
            for meta in _model_chat_attachment_metadata(item.get("files")):
                result.append({"message_index":index,"role":str(item.get("role") or ""),**meta})
    for meta in _model_chat_attachment_metadata(pending):
        result.append({"message_index":None,"role":"pending",**meta})
    return result

def _model_chat_export(payload: dict[str, Any]) -> dict[str, Any]:
    debug=bool(payload.get("debug"))
    session_id=str(payload.get("session_id") or "default").strip() or "default"
    messages=payload.get("messages") if isinstance(payload.get("messages"),list) else []
    safe_messages=_model_chat_safe_messages(messages)
    safe_session=re.sub(r"[^A-Za-z0-9_.-]+","_",session_id)[:100] or "default"

    if not debug:
        lines=["# Model Chat Export",""]
        for item in messages:
            if not isinstance(item,dict):
                continue
            role=str(item.get("role") or "unknown").strip().title()
            content=str(item.get("content") or "")
            lines.extend([f"## {role}","",content,""])
            files=item.get("files")
            if isinstance(files,list) and files:
                lines.append("Files:")
                for f in files:
                    if isinstance(f,dict):
                        name=str(f.get("filename") or f.get("name") or "file")
                        lines.append(f"- `{name}`")
                lines.append("")
        return {
            "status":"passed",
            "filename":f"MODEL_CHAT_{safe_session}.md",
            "media_type":"text/markdown",
            "content_text":"\n".join(lines),
        }

    root=_workspace_root(session_id)
    agent_trace=payload.get("agent_trace") if isinstance(payload.get("agent_trace"),list) else []
    usage_history=payload.get("usage_history") if isinstance(payload.get("usage_history"),list) else []
    errors=payload.get("errors") if isinstance(payload.get("errors"),list) else []
    generated_files=payload.get("generated_files") if isinstance(payload.get("generated_files"),list) else []
    attachments=_model_chat_export_attachment_metadata(messages,payload.get("attachments"))
    provider_info=payload.get("provider_info") if isinstance(payload.get("provider_info"),dict) else {}

    # Explicit allowlist: never serialize keys/secrets from browser/provider state.
    provider_safe={
        key:provider_info.get(key)
        for key in ("enabled","provider","base_url","model","structured_output_mode","semantic_fallback_policy")
        if key in provider_info
    }

    session_meta={
        "schema_version":"1.0",
        "export_kind":"model_chat_debug",
        "session_id":session_id,
        "editor_version":(UTILITY_ROOT/"VERSION").read_text(encoding="utf-8").strip() if (UTILITY_ROOT/"VERSION").exists() else "",
        "message_count":len(messages),
        "agent_trace_entries":len(agent_trace),
        "error_count":len(errors),
    }
    workspace_head=_workspace_head(root)
    workspace_index=_workspace_index(root,limit=10000)

    tool_calls=[]
    for row in agent_trace:
        if not isinstance(row,dict):
            continue
        tool=row.get("tool")
        if isinstance(tool,dict):
            tool_calls.append(tool)

    generated_meta=[]
    for item in generated_files:
        if not isinstance(item,dict):
            continue
        generated_meta.append({
            "filename":str(item.get("filename") or item.get("name") or ""),
            "media_type":str(item.get("media_type") or ""),
            "size_bytes":int(item.get("size_bytes") or 0),
            "source":str(item.get("source") or ""),
        })

    contents={
        "README.md":"""# Ordo Model Chat Debug Export

This archive contains diagnostic metadata for a Model Chat session.

Files:
- `conversation.json` — visible chat messages.
- `agent_trace.json` — internal model/tool iterations returned by the orchestrator.
- `tool_calls.json` — tool calls/results extracted from the trace.
- `usage_history.json` — model token/accounting data per user turn.
- `workspace_head.json` — compact workspace summary.
- `workspace_index.json` — full workspace path/size index for offline diagnosis.
- `attachments.json` — uploaded attachment metadata only; attachment bodies are not embedded.
- `generated_files.json` — generated-file metadata.
- `provider_info.json` — non-secret provider/model configuration.
- `errors.json` — UI/API errors recorded during the session.
- `session.json` — session/editor metadata.

Credentials and API keys are intentionally excluded.
""",
        "conversation.json":json.dumps(safe_messages,ensure_ascii=False,indent=2),
        "agent_trace.json":json.dumps(agent_trace,ensure_ascii=False,indent=2),
        "tool_calls.json":json.dumps(tool_calls,ensure_ascii=False,indent=2),
        "usage_history.json":json.dumps(usage_history,ensure_ascii=False,indent=2),
        "workspace_head.json":json.dumps(workspace_head,ensure_ascii=False,indent=2),
        "workspace_index.json":json.dumps(workspace_index,ensure_ascii=False,indent=2),
        "attachments.json":json.dumps(attachments,ensure_ascii=False,indent=2),
        "generated_files.json":json.dumps(generated_meta,ensure_ascii=False,indent=2),
        "provider_info.json":json.dumps(provider_safe,ensure_ascii=False,indent=2),
        "errors.json":json.dumps(errors,ensure_ascii=False,indent=2),
        "session.json":json.dumps(session_meta,ensure_ascii=False,indent=2),
    }

    buffer=io.BytesIO()
    with zipfile.ZipFile(buffer,"w",zipfile.ZIP_DEFLATED) as archive:
        for name,content in contents.items():
            archive.writestr(name,content)

    raw=buffer.getvalue()
    return {
        "status":"passed",
        "filename":f"MODEL_CHAT_DEBUG_{safe_session}.zip",
        "media_type":"application/zip",
        "size_bytes":len(raw),
        "content_base64":base64.b64encode(raw).decode("ascii"),
    }

def _model_chat(payload: dict[str, Any], activity_callback: Any = None, cancel_check: Any = None) -> dict[str, Any]:
    # Preserve the public/browser contract used by Model Chat and the configured
    # live provider session.
    credentials=_live_credentials(payload)
    messages=payload.get("messages") if isinstance(payload.get("messages"),list) else []
    payload_attachments=payload.get("attachments") if isinstance(payload.get("attachments"),list) else []
    session_id=str(payload.get("session_id") or "").strip() or "default"

    safe_history=[]
    latest_user_files=[]
    for raw_index,msg in enumerate(messages[-80:]):
        if not isinstance(msg,dict): continue
        role=str(msg.get("role") or "").lower().strip()
        if role not in {"user","assistant"}: continue
        content=str(msg.get("content") or "")
        row={"role":role,"content":content[:60000]}
        file_meta=_model_chat_attachment_metadata(msg.get("files"))
        if file_meta:
            row["files"]=file_meta
        safe_history.append(row)
        if role=="user":
            latest_user_files=msg.get("files") if isinstance(msg.get("files"),list) else []

    # The explicit request attachments are authoritative. Falling back to the
    # newest user message files makes the backend robust to a browser transport
    # that preserved the message file-card but omitted the parallel attachments field.
    attachments=payload_attachments or latest_user_files

    # The newest user message is the task for the agent loop. Do not duplicate it
    # in the historical context. Keep prior file metadata so free chat retains
    # awareness of previously supplied materials without embedding their bodies.
    user_message=""
    history=safe_history
    for idx in range(len(safe_history)-1,-1,-1):
        if safe_history[idx].get("role")=="user":
            user_message=str(safe_history[idx].get("content") or "")
            history=safe_history[:idx]
            break

    root=_workspace_root(session_id)
    stored=_workspace_store_attachments(root,attachments)
    workspace_before=_workspace_file_snapshot(root)

    agent_history=_ModelChatHistory(history,current_attachments=stored)
    answer,trace,usage,activities=_model_chat_agent_loop(
        credentials,
        root,
        user_message,
        agent_history,
        max_iterations=12,
        activity_callback=activity_callback,
        cancel_check=cancel_check,
    )
    if not answer:
        raise ValueError("Model Chat agent loop completed without a final message.")

    generated=[]
    for p in _workspace_changed_files(root,workspace_before):
        try:
            size=p.stat().st_size
        except OSError:
            continue
        if size>20_000_000:
            continue
        rel=p.relative_to(root).as_posix()
        item={
            "filename":rel,
            "display_name":p.name,
            "media_type":mimetypes.guess_type(p.name)[0] or "application/octet-stream",
            "size_bytes":size,
            "source":"workspace",
            "workspace_path":rel,
            "download_url":"/api/model-chat-workspace-file?session_id="
                + urllib.parse.quote(session_id,safe="")
                + "&path=" + urllib.parse.quote(rel,safe=""),
        }
        if size<=2_000_000 and p.suffix.lower() in _TEXT_RESOURCE_EXTENSIONS|{".yaml",".yml",".json",".py",".txt"}:
            item["content_text"]=p.read_text(encoding="utf-8",errors="replace")
        generated.append(item)
        if len(generated)>=80:
            break

    return {
        "status":"passed",
        "answer_markdown":answer,
        "files":generated,
        "workspace":{
            "session_id":session_id,
            "stored_attachments":stored,
            "head":_workspace_head(root),
        },
        "agent_trace":trace,"activity_events":activities,
        "usage":usage,
        "model":credentials.get("model"),
        "provider":credentials.get("provider"),
    }




def _model_chat_run_worker(run_id: str, payload: dict[str,Any]) -> None:
    def cancelled() -> bool:
        with MODEL_CHAT_RUNS_LOCK:
            run=MODEL_CHAT_RUNS.get(run_id) or {}
            return bool(run.get("cancel_requested"))

    def activity(event: dict[str,Any]) -> None:
        with MODEL_CHAT_RUNS_LOCK:
            run=MODEL_CHAT_RUNS.get(run_id)
            if not isinstance(run,dict):
                return
            event=copy.deepcopy(event)
            event["seq"]=len(run.get("activity_events") or [])+1
            run.setdefault("activity_events",[]).append(event)
            run["updated_at"]=time.time()

    try:
        result=_model_chat(payload,activity_callback=activity,cancel_check=cancelled)
        with MODEL_CHAT_RUNS_LOCK:
            run=MODEL_CHAT_RUNS.get(run_id)
            if not isinstance(run,dict):
                return
            if run.get("cancel_requested"):
                run.update({"status":"cancelled","finished":True,"updated_at":time.time()})
            else:
                run.update({"status":"completed","finished":True,"result":result,"updated_at":time.time()})
    except Exception as exc:
        with MODEL_CHAT_RUNS_LOCK:
            run=MODEL_CHAT_RUNS.get(run_id)
            if isinstance(run,dict):
                run.update({"status":"failed","finished":True,"error":str(exc),"updated_at":time.time()})


def _model_chat_start(payload: dict[str,Any]) -> dict[str,Any]:
    run_id="model-chat-"+uuid.uuid4().hex
    now=time.time()
    with MODEL_CHAT_RUNS_LOCK:
        MODEL_CHAT_RUNS[run_id]={
            "run_id":run_id,
            "status":"running",
            "finished":False,
            "cancel_requested":False,
            "activity_events":[],
            "created_at":now,
            "updated_at":now,
        }
    threading.Thread(target=_model_chat_run_worker,args=(run_id,copy.deepcopy(payload)),daemon=True).start()
    return {"status":"passed","run_id":run_id}


def _model_chat_status(payload: dict[str,Any]) -> dict[str,Any]:
    run_id=str(payload.get("run_id") or "").strip()
    after_seq=max(0,int(payload.get("after_seq") or 0))
    with MODEL_CHAT_RUNS_LOCK:
        run=copy.deepcopy(MODEL_CHAT_RUNS.get(run_id))
    if not isinstance(run,dict):
        raise ValueError("Unknown Model Chat run.")
    events=[e for e in (run.get("activity_events") or []) if int(e.get("seq") or 0)>after_seq]
    response={
        "status":"passed",
        "run_id":run_id,
        "run_status":run.get("status"),
        "finished":bool(run.get("finished")),
        "activity_events":events,
        "last_seq":max([after_seq]+[int(e.get("seq") or 0) for e in events]),
    }
    if run.get("finished"):
        if isinstance(run.get("result"),dict):
            response["result"]=run["result"]
        if run.get("error"):
            response["error"]=run["error"]
    return response


def _model_chat_cancel(payload: dict[str,Any]) -> dict[str,Any]:
    run_id=str(payload.get("run_id") or "").strip()
    with MODEL_CHAT_RUNS_LOCK:
        run=MODEL_CHAT_RUNS.get(run_id)
        if not isinstance(run,dict):
            raise ValueError("Unknown Model Chat run.")
        run["cancel_requested"]=True
        run["updated_at"]=time.time()
    return {"status":"passed","run_id":run_id,"cancel_requested":True}


def _model_chat_playbook_preview(payload: dict[str, Any]) -> dict[str, Any]:
    filename=str(payload.get("filename") or "generated.yaml")
    source=None
    if isinstance(payload.get("content_text"),str):
        try: source=yaml.safe_load(payload["content_text"])
        except yaml.YAMLError as error: raise ValueError(f"YAML preview could not be parsed: {error}") from error
    elif isinstance(payload.get("content_base64"),str):
        try: raw=base64.b64decode(payload["content_base64"],validate=True)
        except Exception as error: raise ValueError("Preview ZIP contains invalid base64 data.") from error
        if not zipfile.is_zipfile(io.BytesIO(raw)): raise ValueError("Preview file is not a valid ZIP archive.")
        with zipfile.ZipFile(io.BytesIO(raw),"r") as z:
            candidates=[]
            for name in z.namelist():
                if name.lower().endswith((".yaml",".yml")):
                    try: doc=yaml.safe_load(z.read(name).decode("utf-8"))
                    except Exception: continue
                    if _looks_like_playbook(doc): candidates.append((name,doc))
            if len(candidates)!=1: raise ValueError(f"ZIP preview requires exactly one Ordo playbook YAML; found {len(candidates)}.")
            filename,source=candidates[0]
    if not isinstance(source,dict) or not _looks_like_playbook(source): raise ValueError("This file does not contain a recognizable Ordo playbook source.")
    return {"status":"passed","filename":filename,"source":source,"graph":graph_view(source)}


def _data_lineage_assistant(payload: dict[str,Any]) -> dict[str,Any]:
    package_id=str(payload.get("package_id") or "").strip()
    package=PLAYBOOK_PACKAGES.get(package_id) or (PLAYBOOK_PACKAGE if package_id==str(PLAYBOOK_PACKAGE.get("id") or "") else None)
    credentials=_live_credentials(payload)
    if isinstance(package,dict):
        contract=(package.get("semantic_plan") or {}).get("interaction_contract") if isinstance(package.get("semantic_plan"),dict) else None
    else: contract=None
    if not isinstance(contract,dict): contract=_interaction_contract()
    locale=str(contract.get("locale") or "en-US"); language=str(contract.get("model_output_language") or "en")
    entity=payload.get("entity") if isinstance(payload.get("entity"),dict) else {}
    if not entity: raise ValueError("Select a data-flow entity before requesting an explanation.")
    context_graph=payload.get("context") if isinstance(payload.get("context"),dict) else {}
    messages=payload.get("messages") if isinstance(payload.get("messages"),list) else []
    safe=[]
    for msg in messages[-16:]:
        if not isinstance(msg,dict): continue
        role=str(msg.get("role") or "").lower(); content=str(msg.get("content") or "")
        if role in {"user","assistant"} and content: safe.append({"role":role,"content":content[:12000]})
    schema={"type":"object","additionalProperties":False,"required":["answer_markdown"],"properties":{"answer_markdown":{"type":"string"}}}
    system=(
        "You are the read-only Data Flow Assistant inside Ordo Editor. "
        f"Respond in the playbook analyst-facing language: locale={locale}, language={language}. "
        "Explain only what is supported by the supplied logical lineage. Distinguish analyst-provided values, model/deterministic/internal state, documents, and archives. "
        "Explain how the selected entity is formed, what upstream values contribute to it, what downstream entities use it, its current runtime value when supplied, and which execution nodes produce or consume it. "
        "Do not confuse this logical data-flow graph with execution control flow. Do not invent undeclared dependencies. Use Markdown. Return exactly one JSON object matching the schema."
    )
    ctx={"selected_entity":entity,"lineage_context":context_graph,"conversation":safe,"__response_json_schema":schema}
    _,_,raw,usage=_provider_api_call(credentials,system,ctx)
    parsed=_parse_model_json(raw); answer=str(parsed.get("answer_markdown") or parsed.get("explanation") or parsed.get("answer") or parsed.get("analysis") or parsed.get("message") or parsed.get("content") or "").strip()
    if not answer: raise ValueError("The model returned an empty data-flow explanation.")
    return {"status":"passed","answer_markdown":answer,"provider":credentials.get("provider"),"model":credentials.get("model"),"locale":locale,"language":language,"usage":usage}


def _playbook_settings_assistant(payload: dict[str,Any]) -> dict[str,Any]:
    package_id=str(payload.get("package_id") or "").strip()
    package=PLAYBOOK_PACKAGES.get(package_id) or (PLAYBOOK_PACKAGE if package_id==str(PLAYBOOK_PACKAGE.get("id") or "") else None)
    if not isinstance(package,dict):
        raise ValueError("Load a playbook before using the settings assistant.")
    credentials=_live_credentials(payload)
    settings=_playbook_settings_payload(package)
    contract=(package.get("semantic_plan") or {}).get("interaction_contract") if isinstance(package.get("semantic_plan"),dict) else None
    if not isinstance(contract,dict):
        source=package.get("source") if isinstance(package.get("source"),dict) else {}
        im=source.get("interaction_model") if isinstance(source.get("interaction_model"),dict) else {}
        locale=str(im.get("locale") or im.get("interaction_locale") or "uk-UA")
        language=str(im.get("model_output_language") or im.get("language") or ("uk" if locale.lower().startswith("uk") else "en"))
    else:
        locale=str(contract.get("locale") or "uk-UA")
        language=str(contract.get("model_output_language") or ("uk" if locale.lower().startswith("uk") else "en"))
    messages=payload.get("messages") if isinstance(payload.get("messages"),list) else []
    mode=str(payload.get("mode") or "chat")
    user_message=str(payload.get("message") or "").strip()
    resource_mode=mode=="resource_chat"
    resource_path=str(payload.get("resource_path") or "").strip().replace("\\","/")
    resource_text=""
    resolved_resource_path=""
    if resource_mode:
        resolved_resource_path,resolved_text=_resolve_package_text_resource(package,resource_path)
        if not resolved_resource_path or not isinstance(resolved_text,str):
            raise ValueError("Select a text package file before using the file assistant.")
        resource_text=resolved_text
    schema={
        "type":"object","additionalProperties":False,
        "required":["answer_markdown","yaml_settings_block"],
        "properties":{"answer_markdown":{"type":"string"},"yaml_settings_block":{"type":"string"}},
    }
    if resource_mode:
        system=(
            "You are the read-only Package File Assistant inside Ordo Editor. "
            f"Respond in the playbook analyst-facing language: locale={locale}, language={language}. "
            "Discuss only the selected package file and relationships that are supported by the supplied package context. "
            "Explain what the file contains, its role, relevant structure, and how it may relate to the playbook. "
            "Do not claim the file or playbook was modified, do not invent undeclared dependencies, and distinguish direct evidence from inference. "
            "Use Markdown in answer_markdown. yaml_settings_block must always be an empty string. Return one JSON object matching the schema."
        )
        context={
            "mode":"resource_chat",
            "user_message":user_message,
            "conversation":messages[-12:],
            "selected_file":{"path":resolved_resource_path,"content":resource_text[:60000]},
            "package":{"filename":package.get("filename"),"source_name":package.get("source_name")},
            "__response_json_schema":schema,
        }
    else:
        system=(
            "You are the read-only Playbook Settings Assistant inside Ordo Editor. "
            f"Respond in the playbook analyst-facing language: locale={locale}, language={language}. "
            "The settings catalog, allowed values and meanings come from the Ordo language schema/registry and are authoritative for this task. "
            "Explain the current configuration and help the user reason about desired changes. "
            "Never edit the playbook, never claim changes were applied, and never invent enum values not present in the supplied catalog. "
            "When the user asks to change settings, return a complete proposed YAML block containing only program-level/settings sections that should be inserted or replaced by the user manually. "
            "Preserve unknown package-specific settings unless the user explicitly asks to remove them. "
            "If no YAML change is warranted, yaml_settings_block must be an empty string. "
            "Use Markdown in answer_markdown. Return one JSON object matching the schema."
        )
        compact_groups=[]
        for group in settings.get("groups") or []:
            compact_fields=[]
            for field in group.get("fields") or []:
                compact_fields.append({
                    "path":field.get("path"),"specified":field.get("specified"),"current_value":field.get("current_value"),
                    "allowed_values":[{"value":o.get("value"),"meaning":o.get("meaning")} for o in field.get("options") or []],
                })
            compact_groups.append({"id":group.get("id"),"fields":compact_fields})
        context={
            "mode":mode,
            "user_message": user_message if user_message else ("Analyze the current playbook settings, explain the effective posture, important defaults that are not explicitly set, and any notable inconsistencies." if mode=="analyze" else ""),
            "conversation":messages[-12:],
            "settings_catalog":compact_groups,
            "__response_json_schema":schema,
        }
    _,_,raw,usage=_provider_api_call(credentials,system,context)
    result=_parse_model_json(raw)
    answer=str(result.get("answer_markdown") or result.get("explanation") or result.get("answer") or result.get("analysis") or result.get("message") or result.get("content") or "").strip()
    if not answer:
        raise ValueError("The model returned an empty assistant response.")
    return {
        "status":"passed","answer_markdown":answer,
        "yaml_settings_block":"" if resource_mode else str(result.get("yaml_settings_block") or ""),
        "provider":credentials.get("provider"),"model":credentials.get("model"),
        "locale":locale,"language":language,"usage":usage,
        "resource_path":resolved_resource_path if resource_mode else "",
    }


def _playbook_settings_unbound_resource_groups(package: dict[str, Any]) -> list[dict[str, Any]]:
    """Expose package-level resources that cannot be reached from any node/gate References view.

    Classification is structural and generic: it depends on package paths/roles, never domain values.
    """
    source=package.get("source") if isinstance(package.get("source"),dict) else {}
    resources=package.get("resources") if isinstance(package.get("resources"),dict) else {}
    manifest=package.get("manifest") if isinstance(package.get("manifest"),list) else []
    source_name=str(package.get("source_name") or "")
    referenced:set[str]=set()
    records=[r for r in [*(source.get("nodes") or []),*(source.get("gates") or [])] if isinstance(r,dict)]
    for record in records:
        for item in _generic_record_resource_references(record,resources):
            resolved,_=_resolve_package_resource(resources,str(item.get("path") or ""))
            if resolved:
                referenced.add(resolved)
    sizes={str(item.get("path") or ""):int(item.get("size") or 0) for item in manifest if isinstance(item,dict)}
    rows=[]
    for path,text in resources.items():
        if path==source_name or path in referenced:
            continue
        norm=str(path).replace("\\","/")
        low=norm.lower()
        parts=[p for p in low.split("/") if p]
        group=None
        # Package/startup metadata is intentionally root-scoped: nested knowledge/templates are not global metadata.
        if len(parts)==1:
            group="package"
        # Runtime infrastructure includes unbound runtime state/helpers and unbound tool executors.
        elif parts and parts[0] in {"runtime","tools"}:
            group="runtime"
        # Verification policy excludes regression/test implementation; those belong on Verify Playbook, not Settings.
        elif parts and parts[0]=="verification" and "tests" not in parts and not any(p.startswith("test") for p in parts[1:]):
            group="verification"
        if not group:
            continue
        rows.append({"path":norm,"size":sizes.get(path,0),"text":text,"extension":Path(norm).suffix.lower(),"group":group})
    defs=[
        ("runtime","Global runtime infrastructure","Unbound runtime state, runtime helpers, and tool resources used at package/execution level rather than by one node or gate."),
        ("package","Package/startup metadata","Root-level package, startup, release, manifest, and usage metadata that belongs to the playbook package but is not referenced by a tree element."),
        ("verification","Global verification policy","Unbound package-level verification policy/resources. Regression test implementations are intentionally excluded from this Settings view."),
    ]
    return [{"id":gid,"title":title,"description":desc,"files":sorted([r for r in rows if r["group"]==gid],key=lambda r:r["path"].lower())} for gid,title,desc in defs]


def _playbook_settings_payload(package: dict[str, Any]) -> dict[str, Any]:
    source=package.get("source") if isinstance(package.get("source"),dict) else {}
    groups=_language_defined_settings_catalog(source)
    total=sum(len(g.get("fields") or []) for g in groups)
    specified=sum(1 for g in groups for f in g.get("fields") or [] if f.get("specified"))
    return {
        "status":"passed",
        "package":{"id":package.get("id"),"filename":package.get("filename"),"source_name":package.get("source_name")},
        "groups":groups,
        "summary":{"total_settings":total,"specified":specified,"not_specified":max(0,total-specified)},
        "registry_source":"verification/toolkit/language/schemas + registry",
        "unbound_resource_groups":_playbook_settings_unbound_resource_groups(package),
    }


def _verification_catalog_public() -> list[dict[str, Any]]:
    rows=[]
    for item in _verification_catalog():
        rows.append({k:copy.deepcopy(item.get(k)) for k in ["id","title","group","description","order","enabled","requires","descriptor_file","registry_error"]})
    return rows

def _verification_editor_validate(source_path: Path) -> dict[str, Any]:
    if not source_path.is_file():
        return {"status":"FAIL","error":"Source file could not be resolved for current Editor validation."}
    source=parse_yaml(source_path.read_text(encoding="utf-8"))
    if not isinstance(source,dict): return {"status":"FAIL","error":"Source must be a mapping."}
    report=validate_source(source)
    errors=report.get("errors") if isinstance(report,dict) else None
    return {"status":"PASS" if not errors else "FAIL","validation":report}

def _verification_worker(run_id: str, package: dict[str, Any]) -> None:
    raw=package.get("raw_zip")
    if not isinstance(raw,(bytes,bytearray)):
        with VERIFICATION_RUNS_LOCK:
            VERIFICATION_RUNS[run_id].update({"status":"FAIL","error":"Verification requires a loaded playbook ZIP/package with source resources.","finished":True})
        return
    def progress(item:dict[str,Any])->None:
        with VERIFICATION_RUNS_LOCK:
            run=VERIFICATION_RUNS.get(run_id)
            if not run: return
            cid=str(item.get("id") or "")
            existing=next((x for x in run["checks"] if str(x.get("id"))==cid),None)
            if existing: existing.update(copy.deepcopy(item))
            else: run["checks"].append(copy.deepcopy(item))
            done=sum(1 for x in run["checks"] if x.get("status") in {"PASS","FAIL","SKIPPED","ERROR"})
            run["completed"]=done; run["total"]=len(run.get("catalog") or [])
            run["progress_percent"]=round((done/max(1,run["total"]))*100)
    try:
        result=_run_verification_catalog(raw_zip=bytes(raw),source_name=str(package.get("source_name") or ""),progress=progress,internal_editor_validate=_verification_editor_validate)
        with VERIFICATION_RUNS_LOCK:
            VERIFICATION_RUNS[run_id].update({"status":result.get("status"),"summary":result.get("summary"),"duration_ms":result.get("duration_ms"),"progress_percent":100,"finished":True})
    except Exception as exc:
        with VERIFICATION_RUNS_LOCK:
            VERIFICATION_RUNS[run_id].update({"status":"FAIL","error":str(exc),"finished":True})


def _json_response(handler: SimpleHTTPRequestHandler, payload: dict[str, Any], status: int = HTTPStatus.OK) -> None:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)




def _resolve_package_text_resource(package: dict[str, Any], requested_path: str) -> tuple[str, str] | tuple[None, None]:
    resources = package.get("resources") if isinstance(package.get("resources"), dict) else {}
    path = str(requested_path or "").strip().replace("\\", "/")
    if not path:
        return None, None
    if isinstance(resources.get(path), str):
        return path, resources[path]
    matches=[(str(k),v) for k,v in resources.items() if isinstance(v,str) and (str(k).endswith(path) or path.endswith(str(k)))]
    return matches[0] if len(matches)==1 else (None,None)



def _verification_assistant(payload: dict[str, Any]) -> dict[str, Any]:
    package_id=str(payload.get("package_id") or "").strip()
    package=PLAYBOOK_PACKAGES.get(package_id) or (PLAYBOOK_PACKAGE if package_id==str(PLAYBOOK_PACKAGE.get("id") or "") else None)
    if package is None:
        raise ValueError("Load a playbook package before discussing verification results.")
    credentials=_live_credentials(payload)
    contract=(package.get("semantic_plan") or {}).get("interaction_contract") if isinstance(package.get("semantic_plan"),dict) else None
    if not isinstance(contract,dict):
        contract=_interaction_contract()
    locale=str(contract.get("locale") or "en-US")
    language=str(contract.get("model_output_language") or "en")
    check=payload.get("verification_check") if isinstance(payload.get("verification_check"),dict) else {}
    if not check:
        raise ValueError("Verification check details are required.")
    if str(check.get("status") or "").upper()=="PASS":
        raise ValueError("Verification Assistant is intended for non-PASS checks.")
    messages=payload.get("messages") if isinstance(payload.get("messages"),list) else []
    safe_messages=[]
    for msg in messages[-16:]:
        if not isinstance(msg,dict): continue
        role=str(msg.get("role") or "").strip().lower()
        if role not in {"user","assistant"}: continue
        content=str(msg.get("content") or "")
        if content:
            safe_messages.append({"role":role,"content":content[:12000]})
    output=str(check.get("output") or "")
    if len(output)>50000:
        output=output[:50000]+"\n...[truncated]"
    evidence=check.get("evidence") if isinstance(check.get("evidence"),list) else []
    # Bound embedded evidence while preserving structured reports useful for diagnosis.
    bounded_evidence=[]
    for item in evidence[:8]:
        if not isinstance(item,dict): continue
        clone={k:v for k,v in item.items() if k not in {"content_text","content_json"}}
        if "content_json" in item:
            raw=json.dumps(item.get("content_json"),ensure_ascii=False)
            clone["content_json"]=json.loads(raw[:30000]) if len(raw)<=30000 else {"truncated_json":raw[:30000]}
        elif "content_text" in item:
            clone["content_text"]=str(item.get("content_text") or "")[:30000]
        bounded_evidence.append(clone)
    schema={
        "type":"object","additionalProperties":False,
        "required":["explanation"],
        "properties":{"explanation":{"type":"string"}},
    }
    system=(
        "You are the read-only Verification Assistant inside Ordo Editor. "
        f"Respond in the playbook's analyst-facing language: locale={locale}, language={language}. "
        "Discuss only the selected verification and supplied evidence. "
        "Do not change the playbook, do not change the verification verdict, and do not invent evidence. "
        "For SKIPPED checks, explain what context is missing, why the check was not run, whether that is expected, and how the user could make the check applicable. "
        "For FAIL/ERROR checks, distinguish playbook/graph defects from missing environment dependencies, verification-tool defects, package/release issues, and missing runtime evidence. "
        "Use Markdown. Keep technical IDs and paths unchanged. Return exactly one JSON object matching the response schema."
    )
    context={
        "task":"Continue a focused conversation about this verification result. Answer the latest user question using the selected check, its reason/output/evidence, and the prior conversation.",
        "verification_check":{**check,"output":output,"evidence":bounded_evidence},
        "conversation":safe_messages,
        "__response_json_schema":schema,
    }
    _,_,raw,usage=_provider_api_call(credentials,system,context)
    parsed=_parse_model_json(raw)
    answer=str(parsed.get("explanation") or parsed.get("answer_markdown") or parsed.get("answer") or "").strip()
    if not answer:
        raise ValueError("Model returned valid JSON but no non-empty verification-assistant explanation.")
    return {
        "status":"passed",
        "answer_markdown":answer,
        "explanation":answer,
        "provider":credentials.get("provider"),
        "model":credentials.get("model"),
        "locale":locale,
        "language":language,
        "usage":usage,
    }


def _model_explanation(payload: dict[str, Any]) -> dict[str, Any]:
    package_id=str(payload.get("package_id") or "").strip()
    package=PLAYBOOK_PACKAGES.get(package_id) or (PLAYBOOK_PACKAGE if package_id==str(PLAYBOOK_PACKAGE.get("id") or "") else None)
    if package is None:
        raise ValueError("Load a playbook package before requesting an explanation.")
    credentials=_live_credentials(payload)
    contract=(package.get("semantic_plan") or {}).get("interaction_contract") if isinstance(package.get("semantic_plan"),dict) else None
    if not isinstance(contract,dict): contract=_interaction_contract()
    locale=str(contract.get("locale") or "en-US")
    language=str(contract.get("model_output_language") or "en")
    kind=str(payload.get("kind") or "node").strip()
    if kind=="verification_check":
        schema={
            "type":"object","additionalProperties":False,
            "required":["explanation","classification"],
            "properties":{
                "explanation":{"type":"string"},
                "classification":{"type":"string","enum":["playbook_graph_defect","package_release_defect","missing_runtime_evidence_context","verification_tool_defect","inconclusive"]},
            },
        }
    else:
        schema={"type":"object","additionalProperties":False,"required":["explanation"],"properties":{"explanation":{"type":"string"}}}
    system=(
        "You are the read-only explanation layer of an Ordo playbook explorer. "
        f"Write the explanation in the playbook's analyst-facing language: locale={locale}, language={language}. "
        "Use only the supplied playbook/code material. Do not invent missing behavior, hidden business rules, inputs, outputs, routes, or implementation details. "
        "Clearly distinguish what is declared from what is not declared. Keep technical IDs, state paths, field names, enum values, and code identifiers unchanged. "
        "Return exactly one JSON object matching the supplied response schema. Use Markdown inside explanation when structure improves readability."
    )
    context={"explanation_kind":kind,"locale":locale,"language":language,"__response_json_schema":schema}
    if kind=="verification_check":
        check=payload.get("verification_check") if isinstance(payload.get("verification_check"),dict) else {}
        if not check: raise ValueError("Verification check details are required.")
        output=str(check.get("output") or "")
        if len(output)>50000: output=output[:50000]+"\n...[truncated]"
        context.update({
            "task":"Explain this verification result in human language using concise Markdown. Classify the likely source using exactly one schema classification value. Prefer embedded generated report evidence over the short command message when evidence is available. Explain what the result actually proves, what it does not prove, and the next concrete diagnostic action. Do not change the playbook, do not override PASS/FAIL, and do not invent evidence.",
            "verification_check":{**check,"output":output},
        })
    elif kind=="python_resource":
        requested=str(payload.get("resource_path") or "").strip()
        resolved,text=_resolve_package_text_resource(package,requested)
        if not resolved or text is None: raise ValueError("Python resource is not available in the loaded package.")
        if Path(resolved).suffix.lower() != ".py": raise ValueError("Model code explanation is currently enabled only for Python resources.")
        if len(text)>120000: raise ValueError("Python resource is too large for interactive explanation.")
        context.update({
            "task":"Explain what this Python script checks or validates, its inputs, success/failure conditions, important branches, and observable side effects if they are explicit in code. Do not propose changes.",
            "resource_path":resolved,"python_source":text,
        })
    else:
        node_id=str(payload.get("node_id") or "").strip()
        source=package.get("source") if isinstance(package.get("source"),dict) else {}
        records=[r for r in [*(source.get("nodes") or []),*(source.get("gates") or [])] if isinstance(r,dict)]
        record=next((r for r in records if str(r.get("id") or "")==node_id),None)
        if record is None and node_id.startswith("OUT::"):
            inspector=_template_inspector_payload(package,source,node_id)
            record=inspector.get("record") if isinstance(inspector,dict) else None
        if not isinstance(record,dict): raise ValueError("Selected node has no source definition to explain.")
        plan=package.get("semantic_plan") if isinstance(package.get("semantic_plan"),dict) else {}
        semantic=(plan.get("elements") or {}).get(node_id) if isinstance(plan.get("elements"),dict) else None
        context.update({
            "task":"Explain this playbook node in human language: its purpose, declared inputs/state dependencies, what it does, decisions/gates/routes, outputs/state writes, and referenced resources when present. Do not execute it and do not infer undeclared semantics.",
            "node_id":node_id,
            "source_definition":record,
            "compiled_semantics":semantic or {},
        })
    _,_,raw,usage=_provider_api_call(credentials,system,context)
    parsed=_parse_model_json(raw)
    explanation=str(parsed.get("explanation") or "").strip()
    if not explanation: raise ValueError("The model returned an empty explanation.")
    result={"explanation":explanation,"model":credentials.get("model"),"provider":credentials.get("provider"),"locale":locale,"language":language,"usage":usage}
    if kind=="verification_check":
        allowed_classifications={"playbook_graph_defect","package_release_defect","missing_runtime_evidence_context","verification_tool_defect","inconclusive"}
        raw_classification=str(parsed.get("classification") or "inconclusive").strip()
        result["classification"]=raw_classification if raw_classification in allowed_classifications else "inconclusive"
        if raw_classification not in allowed_classifications:
            result["classification_normalized_from"]=raw_classification
    return result


def _execute_live_step_with_revision(payload: dict[str, Any]) -> dict[str, Any]:
    """Execute the exact live runtime boundary used by the browser, with revision/artifact bookkeeping."""
    live_result = _call_openai_live(payload)
    rev_before = int(payload.get("state_revision") or 0)
    before_state = _canonicalize_runtime_state(payload.get("state") or {})
    after_state = _canonicalize_runtime_state(live_result.get("state") or before_state)
    rev_after = rev_before + (1 if after_state != before_state else 0)
    live_result["state_revision"] = rev_after
    dbg = live_result.get("debug") if isinstance(live_result.get("debug"), dict) else None
    if dbg is not None:
        runtime_dbg = dbg.get("runtime") if isinstance(dbg.get("runtime"), dict) else {}
        runtime_dbg["revision_before"] = rev_before; runtime_dbg["revision_after"] = rev_after
        dbg["runtime"] = runtime_dbg
        alpha = dbg.get("alpha20") if isinstance(dbg.get("alpha20"), dict) else None
        if alpha is not None:
            alpha["revision_before"] = rev_before; alpha["revision_after"] = rev_after
            alpha["revision_check"] = f"runtime_owned_revision_{rev_before}"
        artifact = runtime_dbg.get("artifact") if isinstance(runtime_dbg.get("artifact"), dict) else None
        if artifact is not None:
            artifact.setdefault("materialized_from_revision", rev_before)
            artifact.setdefault("producer_node", str(payload.get("current_id") or ""))
            artifact.setdefault("approval_status", _state_subtree(after_state,"document_metadata.approval_status"))
            dep_paths=_artifact_dependency_paths_for_element(str(payload.get("current_id") or ""))
            artifact["depends_on_paths"] = list(dep_paths)
            runtime_dbg["artifact_lineage"] = {
                "path": artifact.get("path"), "sha256": artifact.get("sha256"),
                "materialized_from_revision": int(artifact.get("materialized_from_revision") or rev_before), "producer_node": artifact.get("producer_node"),
                "approval_status": artifact.get("approval_status"), "depends_on_paths": list(dep_paths),
            }
        _update_artifact_registry(state_lineage=runtime_dbg.get("state_lineage") if isinstance(runtime_dbg.get("state_lineage"),list) else [], artifact_lineage=runtime_dbg.get("artifact_lineage") if isinstance(runtime_dbg.get("artifact_lineage"),dict) else None)
        runtime_dbg["artifact_registry_snapshot"]=list((RUN_ARTIFACT_REGISTRY.get(_active_run_key()) or {}).values())
    return live_result


def _managed_run_public(run: dict[str, Any], *, include_state: bool=True, include_transcript: bool=True) -> dict[str, Any]:
    artifact_rows=[]
    for row in (RUN_ARTIFACT_REGISTRY.get((run["package_id"],run.get("session_id") or "",run["run_id"])) or {}).values():
        item=copy.deepcopy(row); path=str(item.get("path") or "")
        if path:
            item["download_url"]="/api/run-artifact?"+urllib.parse.urlencode({"path":path,"package_id":run["package_id"],"session_id":run.get("session_id") or "","run_id":run["run_id"]})
        artifact_rows.append(item)
    out={
        "run_id":run["run_id"],"package_id":run["package_id"],"session_id":run.get("session_id") or "",
        "status":run.get("status"),"current_id":run.get("current_id"),"state_revision":run.get("state_revision",0),
        "awaiting_analyst":bool(run.get("awaiting_analyst")),"path":copy.deepcopy(run.get("path") or []),
        "outcome":copy.deepcopy(run.get("outcome")),"error":copy.deepcopy(run.get("error")),
        "step_count":int(run.get("step_count") or 0),"usage":copy.deepcopy(run.get("usage") or {}),"auto_answers_replay_id":run.get("auto_answers_replay_id") or None,
        "debug_available":True,"debug_endpoint":f"/api/execute-run-debug?run_id={run['run_id']}",
        "artifacts":artifact_rows,
    }
    if include_state: out["state"]=copy.deepcopy(run.get("state") or {})
    if include_transcript: out["transcript"]=copy.deepcopy(run.get("transcript") or [])
    return out


def _managed_execute_run_start(payload: dict[str, Any]) -> dict[str, Any]:
    package_id=str(payload.get("package_id") or "").strip()
    package=PLAYBOOK_PACKAGES.get(package_id) or (PLAYBOOK_PACKAGE if package_id==str(PLAYBOOK_PACKAGE.get("id") or "") else None)
    if not package: raise ValueError("Unknown playbook package. Load it with /api/playbook-package first.")
    source=package.get("source") if isinstance(package.get("source"),dict) else None
    if not source: raise ValueError("Loaded package has no executable source.")
    entry=_entry_id(source)
    if not entry: raise ValueError("Could not determine playbook entry node.")
    replay_id=str(payload.get("auto_answers_replay_id") or "").strip()
    replay=REPLAY_PACKAGES.get(replay_id) if replay_id else None
    if replay_id and replay is None: raise ValueError("Unknown auto_answers_replay_id. Load Auto Answers with /api/replay-package first.")
    run_id="execute-"+uuid.uuid4().hex
    run={
        "run_id":run_id,"package_id":package_id,"session_id":str(payload.get("session_id") or "").strip(),"source":copy.deepcopy(source),
        "status":"ready","current_id":entry,"state":{},"state_revision":0,"history":[],"path":[entry],"debug_trace":[],"transcript":[],
        "usage":{"calls":0,"input_tokens":0,"output_tokens":0,"total_tokens":0,"cached_tokens":0,"reasoning_tokens":0},
        "outcome":None,"error":None,"awaiting_analyst":False,"step_count":0,"created_at":datetime.now(timezone.utc).isoformat(),
        "updated_at":datetime.now(timezone.utc).isoformat(),"auto_answers_replay_id":replay_id or None,"stop_requested":False,
        "auto_answers":copy.deepcopy((replay or {}).get("answers_by_node") or {}),"auto_answer_cursors":{},
        "semantic_fallback_policy":str(payload.get("semantic_fallback_policy") or "automatic_safe"),
    }
    with EXECUTE_RUNS_LOCK: EXECUTE_RUNS[run_id]=run
    return _managed_run_public(run)


def _next_managed_auto_answer(run: dict[str, Any], node_id: str) -> str | None:
    values=(run.get("auto_answers") or {}).get(node_id) or []
    idx=int((run.get("auto_answer_cursors") or {}).get(node_id,0))
    if idx>=len(values): return None
    run.setdefault("auto_answer_cursors",{})[node_id]=idx+1
    return str(values[idx])


def _managed_execute_run_step(run_id: str, *, analyst_input: str | None=None, attachments: list[dict[str,Any]] | None=None) -> dict[str, Any]:
    with EXECUTE_RUNS_LOCK:
        run=EXECUTE_RUNS.get(str(run_id))
        if run is None: raise KeyError("Unknown Execute Playbook run.")
        if run.get("outcome"): return _managed_run_public(run)
        if run.get("status") == "running": raise ValueError("Execute Playbook run is busy with another runtime phase.")
        if analyst_input is not None and str(analyst_input).strip() and not run.get("awaiting_analyst"):
            raise ValueError("analyst_input is accepted only while the run is awaiting analyst input.")
        current=str(run.get("current_id") or "")
        if not current: raise ValueError("Run has no current element.")
        phase="respond" if run.get("awaiting_analyst") else "enter"
        answer=str(analyst_input or "") if analyst_input is not None else ""
        if phase=="respond" and not answer:
            answer=_next_managed_auto_answer(run,current) or ""
            if not answer:
                run["status"]="waiting_analyst"
                return _managed_run_public(run)
            run["transcript"].append({"role":"analyst","text":answer,"node_id":current,"source":"auto_answers"})
            run["history"].append({"role":"analyst","text":answer,"node_id":current})
        elif phase=="respond" and answer:
            run["transcript"].append({"role":"analyst","text":answer,"node_id":current,"source":"api_input"})
            run["history"].append({"role":"analyst","text":answer,"node_id":current})
        previous=run["path"][-2] if phase=="enter" and len(run["path"])>1 else None
        payload={
            "package_id":run["package_id"],"session_id":run.get("session_id") or "","run_id":run["run_id"],"source":run["source"],
            "state":run["state"],"state_revision":run["state_revision"],"current_id":current,"phase":phase,"history":run["history"],
            "analyst_input":answer,"attachments":attachments or [],"previous_node_id":previous,
            "entry_mode":"root" if phase=="enter" and len(run["path"])==1 else ("transition" if phase=="enter" else None),
            "semantic_fallback_policy":run.get("semantic_fallback_policy") or "automatic_safe",
        }
        run["status"]="running"
    try:
        live=_execute_live_step_with_revision(payload)
    except Exception as exc:
        with EXECUTE_RUNS_LOCK:
            run=EXECUTE_RUNS[run_id]
            run["status"]="error"; run["outcome"]={"status":"error","reason":"execution_exception","node_id":current}
            run["error"]={"type":type(exc).__name__,"message":str(exc),"node_id":current,"phase":phase,"traceback":traceback.format_exc()}
            run["updated_at"]=datetime.now(timezone.utc).isoformat()
            return _managed_run_public(run)
    with EXECUTE_RUNS_LOCK:
        run=EXECUTE_RUNS[run_id]
        run["state"]=copy.deepcopy(live.get("state") or run["state"])
        run["state_revision"]=int(live.get("state_revision") if isinstance(live.get("state_revision"),int) else run["state_revision"])
        run["step_count"]+=1; run["updated_at"]=datetime.now(timezone.utc).isoformat()
        if isinstance(live.get("debug"),dict): run["debug_trace"].append(copy.deepcopy(live["debug"]))
        if not live.get("llm_call_skipped") and isinstance(live.get("usage"),dict):
            run["usage"]["calls"]+=1
            for key in ("input_tokens","output_tokens","total_tokens","cached_tokens","reasoning_tokens"):
                run["usage"][key]+=int(live["usage"].get(key) or 0)
        if live.get("assistant_message"):
            msg={"role":"assistant","text":str(live["assistant_message"]),"node_id":current,"phase":phase,"usage":copy.deepcopy(live.get("usage"))}
            run["transcript"].append(msg); run["history"].append(copy.deepcopy(msg))
        if run.get("stop_requested"):
            run["stop_requested"]=False; run["awaiting_analyst"]=False; run["status"]="halted"
            run["outcome"]={"status":"halted","reason":"user_stop","node_id":current}
            return _managed_run_public(run)
        if live.get("await_analyst"):
            run["awaiting_analyst"]=True; run["status"]="waiting_analyst"
            return _managed_run_public(run)
        run["awaiting_analyst"]=False
        status=str(live.get("run_status") or "running")
        if status in {"completed","halted"}:
            run["status"]=status; run["outcome"]={"status":status,"reason":live.get("completion_reason"),"node_id":current}
            return _managed_run_public(run)
        next_id=str(live.get("next_id") or "")
        if not next_id:
            run["status"]="halted"; run["outcome"]={"status":"halted","reason":"missing_route","node_id":current}
            return _managed_run_public(run)
        run["transcript"].append({"role":"system","text":f"{current} -> {next_id}","node_id":current,"route_key":live.get("route_key")})
        run["current_id"]=next_id; run["path"].append(next_id)
        if next_id in _declared_terminal_ids(run["source"]):
            run["status"]="completed"; run["outcome"]={"status":"completed","reason":"terminal","node_id":next_id}
        else: run["status"]="ready"
        return _managed_run_public(run)


def _managed_execute_run_advance(run_id: str, *, max_steps: int=120) -> dict[str, Any]:
    max_steps=max(1,min(int(max_steps or 120),1000))
    for _ in range(max_steps):
        with EXECUTE_RUNS_LOCK:
            run=EXECUTE_RUNS.get(str(run_id))
            if run is None: raise KeyError("Unknown Execute Playbook run.")
            if run.get("outcome"): return _managed_run_public(run)
            if run.get("awaiting_analyst") and not _next_managed_auto_answer_preview(run,str(run.get("current_id") or "")):
                run["status"]="waiting_analyst"; return _managed_run_public(run)
        result=_managed_execute_run_step(run_id)
        if result.get("outcome") or result.get("status") in {"error","halted","completed"}: return result
        if result.get("status") == "waiting_analyst":
            with EXECUTE_RUNS_LOCK:
                current=EXECUTE_RUNS.get(str(run_id))
                if current is None or not _next_managed_auto_answer_preview(current,str(current.get("current_id") or "")):
                    return result
    with EXECUTE_RUNS_LOCK:
        run=EXECUTE_RUNS[str(run_id)]
        run["status"]="halted"; run["outcome"]={"status":"halted","reason":"max_steps_exceeded","node_id":run.get("current_id")}
        return _managed_run_public(run)


def _next_managed_auto_answer_preview(run: dict[str, Any], node_id: str) -> bool:
    values=(run.get("auto_answers") or {}).get(node_id) or []
    idx=int((run.get("auto_answer_cursors") or {}).get(node_id,0))
    return idx < len(values)


def _managed_execute_run_debug(run_id: str) -> dict[str, Any]:
    with EXECUTE_RUNS_LOCK:
        run=EXECUTE_RUNS.get(str(run_id))
        if run is None: raise KeyError("Unknown Execute Playbook run.")
        public=_managed_run_public(run,include_state=True,include_transcript=True)
        public.update({
            "history":copy.deepcopy(run.get("history") or []),"debug_trace":copy.deepcopy(run.get("debug_trace") or []),
            "auto_answer_cursors":copy.deepcopy(run.get("auto_answer_cursors") or {}),"created_at":run.get("created_at"),"updated_at":run.get("updated_at"),
        })
        return public


class EditorHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args: Any, **kwargs: Any) -> None:
        super().__init__(*args, directory=str(UTILITY_ROOT / "web"), **kwargs)

    def log_message(self, format: str, *args: Any) -> None:
        return

    def end_headers(self) -> None:
        """Keep local editor assets fresh while an extracted package is iterated."""
        self.send_header("Cache-Control", "no-store, max-age=0")
        super().end_headers()

    def do_POST(self) -> None:  # noqa: N802
        path = urlparse(self.path).path
        if path not in {"/api/parse", "/api/validate", "/api/export", "/api/update-node", "/api/update-node-sections", "/api/replay-package", "/api/playbook-package", "/api/export-playbook", "/api/live-step", "/api/execute-run-start", "/api/execute-run-step", "/api/execute-run-advance", "/api/execute-run-input", "/api/execute-run-stop", "/api/live-config", "/api/provider-models", "/api/provider-capability-probe", "/api/template-inspector", "/api/explain", "/api/recovery-diagnose", "/api/recovery-chat", "/api/verification-catalog", "/api/verification-start", "/api/verification-status", "/api/playbook-settings", "/api/playbook-settings-assistant", "/api/verification-assistant", "/api/data-lineage", "/api/embedded-data-flow", "/api/data-lineage-assistant", "/api/gitlab-playbooks", "/api/gitlab-directory", "/api/gitlab-playbook-load", "/api/gitlab-readme", "/api/model-chat", "/api/model-chat-start", "/api/model-chat-status", "/api/model-chat-cancel", "/api/model-chat-export", "/api/model-chat-playbook-preview"}:
            _json_response(self, {"status": "failed", "error": "Unknown API endpoint."}, HTTPStatus.NOT_FOUND)
            return
        try:
            length = int(self.headers.get("Content-Length", "0"))
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
            if path == "/api/gitlab-playbooks":
                root_url=str(payload.get("root_url") or EDITOR_STARTUP.get("gitlab_root") or "").strip()
                _json_response(self,_gitlab_playbook_catalog(root_url)); return
            if path == "/api/gitlab-directory":
                root_url=str(payload.get("root_url") or EDITOR_STARTUP.get("gitlab_root") or "").strip()
                directory_path=str(payload.get("path") or "").strip()
                _json_response(self,_gitlab_directory_payload(root_url,directory_path)); return
            if path == "/api/gitlab-playbook-load":
                root_url=str(payload.get("root_url") or EDITOR_STARTUP.get("gitlab_root") or "").strip()
                archive_path=str(payload.get("path") or "").strip()
                raw=_gitlab_download_archive(root_url,archive_path)
                package=parse_playbook_package(Path(archive_path).name or "playbook.zip",raw)
                _json_response(self,{"status":"passed","package":package}); return
            if path == "/api/gitlab-readme":
                root_url=str(payload.get("root_url") or EDITOR_STARTUP.get("gitlab_root") or "").strip()
                readme_path=str(payload.get("path") or "").strip()
                content=_gitlab_read_text_file(root_url,readme_path)
                _json_response(self,{"status":"passed","path":readme_path,"filename":Path(readme_path).name or "README.md","content":content}); return
            if path == "/api/provider-models":
                provider = str(payload.get("provider") or "openai").strip().lower()
                if provider not in PROVIDERS: raise ValueError("Unknown provider.")
                session_id = str(payload.get("session_id") or "").strip()
                previous = LIVE_SESSIONS.get(session_id, {}) if session_id else {}
                if provider == "openai":
                    base_url = "https://api.openai.com/v1"
                    key = LIVE_RUNTIME.get("api_key") or str(payload.get("api_key") or "").strip() or previous.get("api_key")
                    models = list(OPENAI_MODELS) if base_url == "https://api.openai.com/v1" else _provider_models(provider, base_url, key)
                elif provider == "mlx":
                    base_url = _normalize_base_url(payload.get("base_url"), DEFAULT_MLX_BASE_URL)
                    models = _provider_models(provider, base_url, str(payload.get("api_key") or "local"))
                else:
                    base_url = _normalize_base_url(payload.get("base_url"), DEFAULT_CUSTOM_BASE_URL)
                    models = _provider_models(provider, base_url, "")
                _json_response(self, {"status": "passed", "provider": provider, "base_url": base_url, "models": models})
                return
            if path == "/api/live-config":
                session_id = str(payload.get("session_id") or "").strip()
                if not session_id or len(session_id) > 160: raise ValueError("A valid browser session id is required.")
                provider = str(payload.get("provider") or "openai").strip().lower()
                if provider not in PROVIDERS: raise ValueError("Choose OpenAI, Local MLX, or Custom OpenAI-compatible provider.")
                previous = LIVE_SESSIONS.get(session_id, {})
                personal_key = str(payload.get("api_key") or "").strip()
                previous_provider = str(previous.get("provider") or "")
                effective_personal_key = personal_key or (str(previous.get("api_key") or "") if previous_provider == provider else "")
                if provider == "openai":
                    base_url = "https://api.openai.com/v1"
                    model = str(payload.get("model") or "").strip()
                    if base_url == "https://api.openai.com/v1" and model not in OPENAI_MODELS: raise ValueError(f"Choose one of: {', '.join(OPENAI_MODELS)}")
                    if not LIVE_RUNTIME.get("api_key") and not effective_personal_key: raise ValueError("Enter a personal OpenAI API key because no shared server key is configured.")
                elif provider == "mlx":
                    base_url = _normalize_base_url(payload.get("base_url"), DEFAULT_MLX_BASE_URL)
                    models = _provider_models(provider, base_url, "local")
                    model = str(payload.get("model") or "").strip()
                    if model not in models: raise ValueError("Refresh Local MLX models and select a model currently reported by /models.")
                    effective_personal_key = "local"
                else:
                    base_url = _normalize_base_url(payload.get("base_url"), DEFAULT_CUSTOM_BASE_URL)
                    effective_personal_key = ""
                    models = _provider_models(provider, base_url, "")
                    model = str(payload.get("model") or "").strip()
                    if model not in models: raise ValueError("Refresh provider models and select one reported by /models.")
                structured_output_mode=str(payload.get("structured_output_mode") or previous.get("structured_output_mode") or "auto").strip().lower()
                if structured_output_mode not in STRUCTURED_OUTPUT_MODES:
                    raise ValueError(f"Unsupported structured_output_mode: {structured_output_mode}")
                LIVE_SESSIONS[session_id] = {"provider": provider, "model": model, "base_url": base_url, "structured_output_mode": structured_output_mode}
                if effective_personal_key: LIVE_SESSIONS[session_id]["api_key"] = effective_personal_key
                _json_response(self, {"status": "passed", "live": {"enabled": True, "provider": provider, "base_url": base_url, "model": model, "structured_output_mode": structured_output_mode, "shared_key": bool(provider == "openai" and LIVE_RUNTIME.get("api_key")), "personal_key": bool(effective_personal_key and provider != "mlx"), "models": list(OPENAI_MODELS) if provider == "openai" and base_url == "https://api.openai.com/v1" else (models if 'models' in locals() else [model])}})
                return
            if path == "/api/model-chat-start":
                _json_response(self,_model_chat_start(payload)); return
            if path == "/api/model-chat-status":
                _json_response(self,_model_chat_status(payload)); return
            if path == "/api/model-chat-cancel":
                _json_response(self,_model_chat_cancel(payload)); return
            if path == "/api/model-chat":
                _json_response(self,_model_chat(payload)); return
            if path == "/api/model-chat-export":
                _json_response(self,_model_chat_export(payload)); return
            if path == "/api/model-chat-playbook-preview":
                _json_response(self,_model_chat_playbook_preview(payload)); return
            if path == "/api/data-lineage":
                package_id=str(payload.get("package_id") or "").strip()
                package=PLAYBOOK_PACKAGES.get(package_id) or (PLAYBOOK_PACKAGE if package_id==str(PLAYBOOK_PACKAGE.get("id") or "") else None)
                source=(package.get("source") if isinstance(package,dict) and isinstance(package.get("source"),dict) else (payload.get("source") if isinstance(payload.get("source"),dict) else {}))
                if not source: raise ValueError("Load a playbook before opening Show Data Flow.")
                _json_response(self,_build_data_lineage(package,source,payload.get("runtime_state") if isinstance(payload.get("runtime_state"),dict) else {})); return
            if path == "/api/embedded-data-flow":
                package_id=str(payload.get("package_id") or "").strip()
                package=PLAYBOOK_PACKAGES.get(package_id) or (PLAYBOOK_PACKAGE if package_id==str(PLAYBOOK_PACKAGE.get("id") or "") else None)
                if not isinstance(package,dict): raise ValueError("Load a playbook before opening Source Data Flow.")
                _json_response(self,_discover_embedded_authoring_data_flow(package)); return
            if path == "/api/data-lineage-assistant":
                _json_response(self,_data_lineage_assistant(payload)); return
            if path == "/api/playbook-settings":
                package_id=str(payload.get("package_id") or "")
                package=PLAYBOOK_PACKAGES.get(package_id) if package_id else _active_playbook_package()
                if not isinstance(package,dict) or not package.get("source"):
                    _json_response(self,{"status":"failed","error":"Load a playbook before opening Playbook Settings."},HTTPStatus.BAD_REQUEST); return
                _json_response(self,_playbook_settings_payload(package)); return
            if path == "/api/playbook-settings-assistant":
                _json_response(self,_playbook_settings_assistant(payload)); return
            if path == "/api/verification-assistant":
                _json_response(self,_verification_assistant(payload)); return
            if path == "/api/verification-catalog":
                _json_response(self,{"status":"passed","catalog":_verification_catalog_public()})
                return
            if path == "/api/verification-start":
                package_id=str(payload.get("package_id") or "").strip()
                package=PLAYBOOK_PACKAGES.get(package_id) or (PLAYBOOK_PACKAGE if package_id==str(PLAYBOOK_PACKAGE.get("id") or "") else None)
                if package is None:
                    _json_response(self,{"status":"failed","error":"Load a playbook package before verification."},HTTPStatus.BAD_REQUEST); return
                run_id="verify-"+uuid.uuid4().hex
                catalog=_verification_catalog_public()
                with VERIFICATION_RUNS_LOCK:
                    VERIFICATION_RUNS[run_id]={"run_id":run_id,"package_id":package_id,"status":"RUNNING","finished":False,"catalog":catalog,"checks":[{**copy.deepcopy(i),"status":"PENDING","message":"Waiting"} for i in catalog],"completed":0,"total":len(catalog),"progress_percent":0}
                threading.Thread(target=_verification_worker,args=(run_id,copy.deepcopy(package)),daemon=True).start()
                _json_response(self,{"status":"passed","run_id":run_id,"total":len(catalog)})
                return
            if path == "/api/verification-status":
                run_id=str(payload.get("run_id") or "").strip()
                with VERIFICATION_RUNS_LOCK: run=copy.deepcopy(VERIFICATION_RUNS.get(run_id))
                if run is None:
                    _json_response(self,{"status":"failed","error":"Unknown verification run."},HTTPStatus.NOT_FOUND); return
                _json_response(self,{"status":"passed","verification":run})
                return
            if path == "/api/template-inspector":
                package_id = str(payload.get("package_id") or "")
                node_id = str(payload.get("node_id") or "")
                source = payload.get("source") if isinstance(payload.get("source"), dict) else {}
                package = PLAYBOOK_PACKAGES.get(package_id) or (PLAYBOOK_PACKAGE if package_id == str(PLAYBOOK_PACKAGE.get("id") or "") else None)
                if package is None:
                    _json_response(self, {"status":"failed","error":"Unknown playbook package."}, HTTPStatus.BAD_REQUEST); return
                try:
                    result = _template_inspector_payload(package, source, node_id)
                except ValueError as exc:
                    _json_response(self, {"status":"failed","error":str(exc)}, HTTPStatus.BAD_REQUEST); return
                _json_response(self, result); return
            if path == "/api/explain":
                result=_model_explanation(payload)
                _json_response(self,{"status":"passed",**result})
                return
            if path == "/api/provider-capability-probe":
                credentials=_live_credentials(payload)
                profile=_probe_provider_json_schema_capability(credentials)
                _remember_provider_capability(profile)
                session_id=str(payload.get("session_id") or "").strip()
                if session_id and session_id in LIVE_SESSIONS:
                    LIVE_SESSIONS[session_id]["capability_profile"]=copy.deepcopy(profile)
                resolved_mode="strict_json_schema" if profile.get("supports_json_schema") else "json_object"
                _json_response(self,{"status":"passed","capability_profile":profile,"structured_output_mode":resolved_mode})
                return
            if path == "/api/live-step":
                _json_response(self, {"status": "passed", "live": _execute_live_step_with_revision(payload)})
                return
            if path == "/api/execute-run-start":
                result=_managed_execute_run_start(payload)
                if bool(payload.get("advance")):
                    result=_managed_execute_run_advance(result["run_id"],max_steps=int(payload.get("max_steps") or 120))
                _json_response(self,{"status":"passed","run":result})
                return
            if path == "/api/execute-run-step":
                run_id=str(payload.get("run_id") or "").strip()
                result=_managed_execute_run_step(run_id,analyst_input=payload.get("analyst_input"),attachments=payload.get("attachments") if isinstance(payload.get("attachments"),list) else None)
                _json_response(self,{"status":"passed","run":result})
                return
            if path == "/api/execute-run-advance":
                run_id=str(payload.get("run_id") or "").strip()
                result=_managed_execute_run_advance(run_id,max_steps=int(payload.get("max_steps") or 120))
                _json_response(self,{"status":"passed","run":result})
                return
            if path == "/api/execute-run-input":
                run_id=str(payload.get("run_id") or "").strip()
                if not str(payload.get("analyst_input") or "").strip(): raise ValueError("analyst_input is required.")
                result=_managed_execute_run_step(run_id,analyst_input=str(payload.get("analyst_input")),attachments=payload.get("attachments") if isinstance(payload.get("attachments"),list) else None)
                if bool(payload.get("advance",True)) and not result.get("outcome"):
                    result=_managed_execute_run_advance(run_id,max_steps=int(payload.get("max_steps") or 120))
                _json_response(self,{"status":"passed","run":result})
                return
            if path == "/api/execute-run-stop":
                run_id=str(payload.get("run_id") or "").strip()
                with EXECUTE_RUNS_LOCK:
                    run=EXECUTE_RUNS.get(run_id)
                    if run is None: raise KeyError("Unknown Execute Playbook run.")
                    if not run.get("outcome"):
                        if run.get("status") == "running":
                            run["stop_requested"]=True
                        else:
                            run["status"]="halted"; run["outcome"]={"status":"halted","reason":"user_stop","node_id":run.get("current_id")}
                    result=_managed_run_public(run)
                _json_response(self,{"status":"passed","run":result})
                return
            if path == "/api/recovery-diagnose":
                _json_response(self, {"status": "passed", **_recovery_diagnosis(payload)})
                return
            if path == "/api/recovery-chat":
                _json_response(self, {"status": "passed", **_recovery_conversation(payload)})
                return
            if path == "/api/export-playbook":
                package_id = str(payload.get("package_id") or "").strip()
                if not package_id or package_id != str(_active_playbook_package().get("id") or ""):
                    raise ValueError("The currently loaded playbook package is not available for export.")
                source = payload.get("source")
                if not isinstance(source, dict):
                    raise ValueError("Current playbook source is missing from the export request.")
                raw_zip = _active_playbook_package().get("raw_zip")
                source_name = str(_active_playbook_package().get("source_name") or "")
                if not isinstance(raw_zip, (bytes, bytearray)) or not source_name:
                    raise ValueError("Original playbook package bytes are not available.")
                out = io.BytesIO()
                with zipfile.ZipFile(io.BytesIO(bytes(raw_zip)), "r") as original, zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as rebuilt:
                    for info in original.infolist():
                        if info.is_dir():
                            rebuilt.writestr(info, b"")
                            continue
                        if info.filename == source_name:
                            rebuilt.writestr(info, dump_yaml(source).encode("utf-8"))
                        else:
                            rebuilt.writestr(info, original.read(info.filename))
                original_name = str(_active_playbook_package().get("filename") or "playbook.zip")
                stem = original_name[:-4] if original_name.lower().endswith(".zip") else original_name
                filename = f"{stem}.edited.zip"
                _json_response(self, {"status": "passed", "filename": filename, "data_base64": base64.b64encode(out.getvalue()).decode("ascii")})
                return
            if path == "/api/playbook-package":
                encoded = payload.get("data_base64")
                if not isinstance(encoded, str) or not encoded:
                    raise ValueError("Playbook package upload is missing file data.")
                try:
                    raw = base64.b64decode(encoded, validate=True)
                except Exception as error:
                    raise ValueError("Playbook package upload contains invalid base64 data.") from error
                package = parse_playbook_package(str(payload.get("filename", "playbook.zip")), raw)
                _json_response(self, {"status": "passed", "package": package})
                return
            if path == "/api/replay-package":
                encoded = payload.get("data_base64")
                if not isinstance(encoded, str) or not encoded:
                    raise ValueError("Replay upload is missing file data.")
                try:
                    raw = base64.b64decode(encoded, validate=True)
                except Exception as error:
                    raise ValueError("Replay upload contains invalid base64 data.") from error
                registered = _register_replay_package(str(payload.get("filename", "replay.zip")), raw)
                _json_response(self, {"status": "passed", **registered})
                return
            source = parse_yaml(payload["yaml"]) if "yaml" in payload else payload["source"]
            if not isinstance(source, dict):
                raise ValueError("source must be a mapping.")
            if path == "/api/export":
                _json_response(self, {"status": "passed", "yaml": dump_yaml(source)})
            elif path == "/api/update-node":
                replacement = parse_yaml(payload["node_yaml"])
                collection = str(payload.get("collection", "nodes"))
                replace_record(source, collection, str(payload["old_id"]), replacement)
                _json_response(self, {"status": "passed", "node_id": replacement["id"], "source": source, "graph": graph_view(source)})
            elif path == "/api/update-node-sections":
                collection = str(payload.get("collection", "nodes"))
                replace_record_sections(source, collection, str(payload["old_id"]), payload["sections"])
                _json_response(self, {"status": "passed", "node_id": yaml.safe_load(payload["sections"]["id"]), "source": source, "graph": graph_view(source)})
            else:
                response = {"status": "passed", "source": source, "graph": graph_view(source)}
                if path == "/api/validate":
                    response["validation"] = validate_source(source)
                _json_response(self, response)
        except (KeyError, TypeError, ValueError, json.JSONDecodeError, yaml.YAMLError) as error:
            _json_response(self, {"status": "failed", "error": str(error)}, HTTPStatus.BAD_REQUEST)

    def do_GET(self) -> None:  # noqa: N802
        parsed = urlparse(self.path)
        path = parsed.path
        if path == "/api/gitlab-archive":
            query=parse_qs(parsed.query)
            root_url=str((query.get("root_url") or [EDITOR_STARTUP.get("gitlab_root") or ""])[0]).strip()
            archive_path=str((query.get("path") or [""])[0]).strip()
            if not root_url or not archive_path:
                self.send_error(HTTPStatus.BAD_REQUEST,"root_url and path are required")
                return
            try:
                raw=_gitlab_download_archive(root_url,archive_path)
            except ValueError as error:
                self.send_error(HTTPStatus.BAD_REQUEST,str(error))
                return
            safe_name=re.sub(r'[^A-Za-z0-9._-]+','_',Path(archive_path).name) or "playbook.zip"
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type","application/zip")
            self.send_header("Content-Length",str(len(raw)))
            self.send_header("Content-Disposition",f'attachment; filename="{safe_name}"')
            self.send_header("Cache-Control","no-store")
            self.end_headers()
            self.wfile.write(raw)
            return
        if path == "/api/execute-run-status":
            query=parse_qs(parsed.query); run_id=str((query.get("run_id") or [""])[0]).strip()
            if not run_id: _json_response(self,{"status":"failed","error":"run_id is required"},HTTPStatus.BAD_REQUEST); return
            with EXECUTE_RUNS_LOCK: run=EXECUTE_RUNS.get(run_id)
            if run is None: _json_response(self,{"status":"failed","error":"Unknown Execute Playbook run."},HTTPStatus.NOT_FOUND); return
            _json_response(self,{"status":"passed","run":_managed_run_public(run)}); return
        if path == "/api/execute-run-debug":
            query=parse_qs(parsed.query); run_id=str((query.get("run_id") or [""])[0]).strip()
            if not run_id: _json_response(self,{"status":"failed","error":"run_id is required"},HTTPStatus.BAD_REQUEST); return
            try: result=_managed_execute_run_debug(run_id)
            except KeyError: _json_response(self,{"status":"failed","error":"Unknown Execute Playbook run."},HTTPStatus.NOT_FOUND); return
            _json_response(self,{"status":"passed","run":result}); return
        if path == "/api/runtime-config":
            session_id = str((parse_qs(parsed.query).get("session_id") or [""])[0]).strip()
            session = LIVE_SESSIONS.get(session_id, {}) if session_id else {}
            provider = str(session.get("provider") or LIVE_RUNTIME.get("provider") or "openai")
            shared_key = bool(provider == "openai" and LIVE_RUNTIME.get("api_key"))
            model = session.get("model") or LIVE_RUNTIME.get("model") or ("gpt-5.6-terra" if provider == "openai" else "")
            base_url = session.get("base_url") or (DEFAULT_MLX_BASE_URL if provider == "mlx" else DEFAULT_CUSTOM_BASE_URL if provider == "custom" else "https://api.openai.com/v1")
            personal_key = bool(provider == "openai" and session.get("api_key") and session.get("api_key") != "local")
            enabled = bool(model and (provider == "mlx" or shared_key or personal_key or provider == "custom"))
            models = list(OPENAI_MODELS) if provider == "openai" else ([model] if model else [])
            _json_response(self, {"status": "passed", "live": {"enabled": enabled, "provider": provider, "base_url": base_url, "model": model, "shared_key": shared_key, "personal_key": personal_key, "models": models}, "startup": {"gitlab_root": str(EDITOR_STARTUP.get("gitlab_root") or "")}, "package": {"id": _active_playbook_package().get("id"), "filename": _active_playbook_package().get("filename"), "source_name": _active_playbook_package().get("source_name")} if _active_playbook_package().get("id") else None})
            return
        if path == "/api/node-templates":
            _json_response(self, {"status": "passed", "templates": NODE_TEMPLATES})
            return
        if path == "/api/tree-modules":
            manifest = tree_module_manifest_path()
            _json_response(self, {"status": "passed", "library": yaml.safe_load(manifest.read_text(encoding="utf-8"))})
            return
        if path == "/api/model-chat-workspace-file":
            query=parse_qs(parsed.query)
            session_id=str((query.get("session_id") or [""])[0]).strip()
            relative=str((query.get("path") or [""])[0]).strip()
            if not session_id or not relative:
                self.send_error(HTTPStatus.BAD_REQUEST,"session_id and path are required")
                return
            workspace=_workspace_root(session_id).resolve()
            try:
                artifact=_workspace_safe_path(workspace,relative).resolve()
                artifact.relative_to(workspace)
            except Exception:
                self.send_error(HTTPStatus.FORBIDDEN,"invalid workspace path")
                return
            if not artifact.is_file():
                self.send_error(HTTPStatus.NOT_FOUND,"workspace file not found")
                return
            data=artifact.read_bytes()
            content_type=mimetypes.guess_type(artifact.name)[0] or "application/octet-stream"
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type",content_type)
            self.send_header("Content-Length",str(len(data)))
            safe_name=re.sub(r'[^A-Za-z0-9._-]+','_',artifact.name) or "download"
            self.send_header("Content-Disposition",f'attachment; filename="{safe_name}"')
            self.send_header("Cache-Control","no-store")
            self.end_headers()
            self.wfile.write(data)
            return
        if path == "/api/run-artifact":
            relative = str((parse_qs(parsed.query).get("path") or [""])[0]).strip().lstrip("/\\")
            if not relative:
                self.send_error(HTTPStatus.BAD_REQUEST, "artifact path is required")
                return
            query = parse_qs(parsed.query)
            package_id = str((query.get("package_id") or [""])[0]).strip()
            session_id = str((query.get("session_id") or [""])[0]).strip()
            run_id = str((query.get("run_id") or [""])[0]).strip()
            workspace = _runtime_workspace(package_id=package_id or None, session_id=session_id or None, run_id=run_id or None).resolve()
            artifact = (workspace / relative).resolve()
            try:
                artifact.relative_to(workspace)
            except ValueError:
                self.send_error(HTTPStatus.FORBIDDEN, "artifact path is outside runtime workspace")
                return
            if not artifact.is_file():
                self.send_error(HTTPStatus.NOT_FOUND, "artifact not found")
                return
            data = artifact.read_bytes()
            content_type = "text/markdown; charset=utf-8" if artifact.suffix.lower() in {".md", ".markdown"} else "application/octet-stream"
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Length", str(len(data)))
            safe_name = re.sub(r'[^A-Za-z0-9._-]+', '_', artifact.name) or "artifact"
            self.send_header("Content-Disposition", f'attachment; filename="{safe_name}"')
            self.send_header("Cache-Control", "no-store")
            self.end_headers()
            self.wfile.write(data)
            return
        super().do_GET()


def run_server(port: int, open_browser: bool, *, provider: str | None=None, api_key: str | None=None, model: str | None=None, base_url: str | None=None, gitlab_root: str | None=None) -> None:
    config=_resolve_startup_runtime_config(provider=provider,model=model,base_url=base_url,api_key=api_key)
    LIVE_RUNTIME.update(config)
    EDITOR_STARTUP["gitlab_root"]=str(gitlab_root or os.environ.get("ORDO_GITLAB_ROOT") or "").strip()
    if config.get("enabled"):
        print(f"Model default: {config['provider']} · {config.get('model')} · {config.get('base_url')}")
    else:
        print("Model default is not fully configured; use Model Settings in the Editor.")
    if EDITOR_STARTUP["gitlab_root"]:
        print(f"GitLab playbook root: {EDITOR_STARTUP['gitlab_root']}")
    server = ThreadingHTTPServer(("127.0.0.1", port), EditorHandler)
    url = f"http://127.0.0.1:{server.server_port}"
    print(f"Ordo Tree Editor is running at {url}")
    if open_browser:
        webbrowser.open(url)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nOrdo Tree Editor stopped.")
    finally:
        server.server_close()


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run the local Ordo Tree Editor.")
    parser.add_argument("--port", type=int, default=8765)
    parser.add_argument("--no-browser", action="store_true")
    parser.add_argument("--model-provider", choices=PROVIDERS, default=os.environ.get("ORDO_MODEL_PROVIDER"), help="Default provider: openai, mlx, or custom.")
    parser.add_argument("--model-name", default=os.environ.get("ORDO_MODEL_NAME"), help="Default model selected when the Editor opens.")
    parser.add_argument("--model-base-url", default=os.environ.get("ORDO_MODEL_BASE_URL"), help="Default endpoint for mlx/custom providers.")
    parser.add_argument("--model-api-key", default=None, help="Optional shared model API key. Prefer ORDO_MODEL_API_KEY env var.")
    parser.add_argument("--gitlab-root", default=os.environ.get("ORDO_GITLAB_ROOT"), help="Optional public GitLab repository tree URL containing playbook/version ZIPs.")
    # Backward-compatible aliases used by older launcher scripts.
    parser.add_argument("--openai-api-key", default=None, help=argparse.SUPPRESS)
    parser.add_argument("--openai-model", default=None, help=argparse.SUPPRESS)
    parser.add_argument("--openai-base-url", default=None, help=argparse.SUPPRESS)
    args = parser.parse_args(argv)
    provider=args.model_provider or ("openai" if (args.openai_api_key or args.openai_model or args.openai_base_url) else None)
    api_key=args.model_api_key or args.openai_api_key
    model=args.model_name or args.openai_model
    base_url=args.model_base_url or args.openai_base_url
    run_server(args.port,open_browser=not args.no_browser,provider=provider,api_key=api_key,model=model,base_url=base_url,gitlab_root=args.gitlab_root)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
